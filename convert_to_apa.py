#!/usr/bin/env python3
"""
Conversor ERS Markdown → DOCX con Normas APA 7ª edición.

Pipeline:
  1. Pre-procesa el Markdown: inserta saltos de página (raw openxml)
     antes de cada encabezado de nivel 1 (# Sección) y separa la
     portada como primera página.
  2. Pandoc convierte el Markdown pre-procesado a DOCX.
  3. python-docx post-procesa el DOCX para aplicar formato APA 7:
     - Times New Roman 12 pt
     - Doble espacio
     - Márgenes 1" (2.54 cm)
     - Encabezados APA: Nivel 1 centrado/negrita, Nivel 2 izq/negrita,
       Nivel 3 izq/negrita-cursiva
     - Tablas con solo líneas horizontales (top, sub-header, bottom)
     - Referencias con sangría francesa
     - Numeración de páginas en esquina superior derecha

Reglas APA implementadas:
  - Portada en UNA sola página
  - Cada sección principal (Heading 1) inicia en página nueva
  - Sub-secciones (H2, H3) NO crean nueva página
  - Contenido general fluye sin interrupciones
"""

from __future__ import annotations

import re
import subprocess
import sys
import tempfile
import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INPUT_FILE = "ERS_Shoppipai_SENA_COMPLETO.md"
DOCS_DIR = Path("DOCS")
OUTPUT_FILE = DOCS_DIR / "ERS_Shoppipai_SENA_COMPLETO_APA.docx"

# Raw OpenXML page break for Pandoc
PAGEBREAK_OPENXML = """
```{=openxml}
<w:p>
  <w:r>
    <w:br w:type="page"/>
  </w:r>
</w:p>
```

"""


# ────────────────────────── Step 1: Pre-process Markdown ──────────────────────────


def _extract_metadata(lines: list[str]) -> dict[str, str]:
    """Extract title, author, etc. from the first ~14 lines."""
    meta: dict[str, str] = {}

    # Lines 0-1: Title
    title_parts = []
    for i in range(2):
        if i < len(lines):
            title_parts.append(lines[i].strip().replace("\r", "").replace("**", ""))
    meta["title"] = " ".join(filter(None, title_parts)).strip()

    # Remaining metadata
    field_order = ["author", "program", "ficha", "institution", "center", "date"]
    idx = 0
    for i in range(3, 14):
        if i < len(lines):
            val = lines[i].strip().replace("\r", "")
            if val:
                if idx < len(field_order):
                    meta[field_order[idx]] = val
                    idx += 1
    return meta


def _build_title_page_md(meta: dict[str, str]) -> str:
    """Build a Markdown title page that Pandoc will render."""
    parts: list[str] = []

    # Empty lines for vertical centering (~6 empty lines to push down)
    parts.append("&nbsp;\n" * 6)
    parts.append("")

    # Title (centered bold)
    title = meta.get("title", "Sin Título")
    parts.append('<div style="text-align:center">\n')
    parts.append(f"**{title}**\n")
    parts.append("")  # Blank line between title and author (APA)
    parts.append("&nbsp;\n")  # Extra visual spacing
    parts.append("")

    # Metadata fields
    for field in ["author", "program", "ficha", "institution", "center", "date"]:
        val = meta.get(field, "")
        if val:
            parts.append(val + "\n")
    parts.append("</div>\n")

    # Page break after title page
    parts.append(PAGEBREAK_OPENXML)

    return "\n".join(parts)


def preprocess_markdown(text: str, meta: dict[str, str]) -> str:
    """
    Pre-process the Markdown:
      1. Replace original metadata + title lines with APA title page
      2. Insert page breaks before every # heading (level 1)
      3. Skip ## and ### headings (they stay in the same page)
    """
    lines = text.split("\n")

    # Find where content starts (# Resumen or # Contenido)
    content_start = 0
    for i, line in enumerate(lines):
        stripped = line.strip().replace("\r", "")
        if stripped.startswith("# "):
            content_start = i
            break
    if content_start == 0:
        content_start = 60  # fallback

    # Build the new markdown
    output_parts: list[str] = []

    # 1. Title page
    output_parts.append(_build_title_page_md(meta))

    # 2. Process content lines
    found_first_heading = False
    for i in range(content_start, len(lines)):
        line = lines[i]
        stripped = line.strip().replace("\r", "")

        # Check for level 1 heading
        if re.match(r"^#\s+", stripped) and not re.match(r"^##", stripped):
            if found_first_heading:
                # Insert page break BEFORE this heading (except the very first one)
                output_parts.append(PAGEBREAK_OPENXML)
            found_first_heading = True

        output_parts.append(line)

    return "\n".join(output_parts)


# ────────────────────────── Step 2: Pandoc ──────────────────────────


def run_pandoc(md_text: str, docx_path: str):
    """Use Pandoc to convert pre-processed Markdown to DOCX."""
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", encoding="utf-8", delete=False
    ) as tmp:
        tmp.write(md_text)
        tmp_path = tmp.name

    cmd = [
        "pandoc",
        tmp_path,
        "-f",
        "markdown+raw_attribute",
        "-t",
        "docx",
        "-o",
        docx_path,
        "--standalone",
    ]
    print(
        f"  ▸ Ejecutando: pandoc ... -f markdown+raw_attribute -t docx -o {docx_path}"
    )
    result = subprocess.run(cmd, capture_output=True, text=True)
    Path(tmp_path).unlink(missing_ok=True)

    if result.returncode != 0:
        print(f"  ✗ Error de Pandoc:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)
    print(f"  ✓ Pandoc generó: {docx_path}")


# ────────────────────────── Step 3: APA Post-Processing ──────────────────────────


def _clear_cell_borders(cell):
    """Remove ALL existing cell-level borders from a cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)


def _set_cell_border(cell, **kwargs):
    """Set border on a table cell (OpenXML). Clears existing first."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing cell borders
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge_name in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge_name in kwargs:
            el = OxmlElement(f"w:{edge_name}")
            for attr, val in kwargs[edge_name].items():
                el.set(qn(f"w:{attr}"), str(val))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _apply_apa_table_borders(table):
    """APA 7: only 3 horizontal lines — top, below header, bottom.

    No vertical lines, no internal horizontal lines between body rows.
    Strategy:
      1. Remove Pandoc's tblStyle (root cause of grid borders)
      2. Set table-level borders (top+bottom=single, rest=none)
      3. Explicitly set ALL borders to NONE on EVERY cell
         (clearing alone makes cells inherit from style — must override)
      4. Re-add header row bottom border
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # ── Step 1: REMOVE the Pandoc table style ──
    # Pandoc sets tblStyle="Table" which defines grid borders.
    # Removing it prevents style-level borders from being applied.
    for old_style in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(old_style)

    # Clear old table-level borders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    # Set table-level: top + bottom only, everything else = none
    borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge_name}")
        if edge_name in ("top", "bottom"):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:color"), "FFFFFF")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)

    # ── Step 2: Explicitly set ALL borders to NONE on every cell ──
    # (Just clearing tcBorders makes cells inherit from table style;
    #  we must set explicit "none" to override any inherited borders)
    none_b = {"sz": "0", "val": "none", "color": "FFFFFF"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top=none_b,
                bottom=none_b,
                start=none_b,
                end=none_b,
                insideH=none_b,
                insideV=none_b,
            )

    # ── Step 3: Set border BELOW header row (first row) ──
    if len(table.rows) > 0:
        line_b = {"sz": "8", "val": "single", "color": "000000"}
        for cell in table.rows[0].cells:
            _set_cell_border(
                cell,
                top=none_b,
                bottom=line_b,
                start=none_b,
                end=none_b,
            )


def _apply_font_to_paragraph(paragraph, font_name="Times New Roman", font_size=Pt(12)):
    """Ensure all runs have the correct font."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_page_number(section):
    """Add page number top-right in header."""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.clear()

    run = hp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def apply_apa_styles(docx_path: str, output_path: str, meta: dict = None):
    """Post-process the Pandoc DOCX to enforce APA 7 formatting."""
    if meta is None:
        meta = {}
    doc = Document(docx_path)

    # ── Page layout ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        _add_page_number(section)

    # ── Normal style ──
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Body Text style (used by Pandoc for paragraphs) ──
    for style_name in ["Body Text", "First Paragraph"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 2.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.first_line_indent = Inches(0.5)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except KeyError:
            pass

    # ── Neutralize Pandoc's "Table" style in styles.xml ──
    # The Table style may define borders. Remove all tblBorders from it.
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for style_el in doc.styles.element.findall(f"{{{ns}}}style"):
        style_id = style_el.get(f"{{{ns}}}styleId", "")
        if style_id == "Table":
            tblPr_el = style_el.find(f"{{{ns}}}tblPr")
            if tblPr_el is not None:
                for old_b in tblPr_el.findall(f"{{{ns}}}tblBorders"):
                    tblPr_el.remove(old_b)
                # Set all borders to none in the style itself
                brd = OxmlElement("w:tblBorders")
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    e = OxmlElement(f"w:{edge}")
                    e.set(qn("w:val"), "none")
                    e.set(qn("w:sz"), "0")
                    e.set(qn("w:color"), "auto")
                    e.set(qn("w:space"), "0")
                    brd.append(e)
                tblPr_el.append(brd)

    # ── Heading styles (APA Levels) ──
    heading_configs = {
        "Heading 1": {
            "bold": True,
            "italic": False,
            "align": WD_ALIGN_PARAGRAPH.CENTER,
        },
        "Heading 2": {"bold": True, "italic": False, "align": WD_ALIGN_PARAGRAPH.LEFT},
        "Heading 3": {"bold": True, "italic": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
    }
    for sn, cfg in heading_configs.items():
        try:
            h = doc.styles[sn]
            h.font.name = "Times New Roman"
            h.font.size = Pt(12)
            h.font.bold = cfg["bold"]
            h.font.italic = cfg["italic"]
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.alignment = cfg["align"]
            h.paragraph_format.line_spacing = 2.0
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(0)
            h.paragraph_format.page_break_before = False
        except KeyError:
            pass

    # ── Helper: clean text by fixing spacing issues ──
    def _fix_text_spacing(text: str) -> str:
        """Fix stuck words, phantom spaces, and comma spacing."""
        # Fix phantom spaces inside parens: ( text ) -> (text)
        text = re.sub(r"\(\s+", "(", text)
        text = re.sub(r"\s+\)", ")", text)
        # Fix space before comma: word , -> word,
        text = re.sub(r"\s+,", ",", text)
        # Fix space before period at end: word . -> word.
        text = re.sub(r"\s+\.\s", ". ", text)
        # Collapse multiple spaces into one
        text = re.sub(r"  +", " ", text)
        return text

    def _merge_and_clean_paragraph(p):
        """Merge all runs in a paragraph, fix spacing, rewrite as minimal runs.

        Preserves italic/bold formatting by creating separate runs for
        formatting transitions, but fixes stuck words and phantom spaces.
        """
        if not p.runs:
            return

        # Collect runs with their formatting
        segments = []
        for run in p.runs:
            segments.append(
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                }
            )

        # Build merged text, adding spaces between runs that need them
        merged_parts = []
        for i, seg in enumerate(segments):
            t = seg["text"]
            if not t:
                continue
            if merged_parts:
                prev_text = merged_parts[-1]["text"]
                if (
                    prev_text
                    and not prev_text[-1].isspace()
                    and not t[0].isspace()
                    and t[0] not in ".,;:!?)]}"
                    and prev_text[-1] not in "([{"
                ):
                    # Check if formatting changed - if same format, just append
                    # If different format, add space to previous or current
                    merged_parts.append(
                        {"text": " ", **{k: seg[k] for k in seg if k != "text"}}
                    )
            merged_parts.append(seg)

        if not merged_parts:
            return

        # Now consolidate consecutive segments with same formatting
        consolidated = [merged_parts[0].copy()]
        for seg in merged_parts[1:]:
            prev = consolidated[-1]
            same_fmt = prev["bold"] == seg["bold"] and prev["italic"] == seg["italic"]
            if same_fmt:
                prev["text"] += seg["text"]
            else:
                consolidated.append(seg.copy())

        # Apply text spacing fixes to each consolidated segment
        for seg in consolidated:
            seg["text"] = _fix_text_spacing(seg["text"])

        # Clear all existing runs
        p.clear()

        # Recreate runs
        for seg in consolidated:
            if not seg["text"]:
                continue
            run = p.add_run(seg["text"])
            run.bold = seg["bold"]
            run.italic = seg["italic"]
            if seg["font_name"]:
                run.font.name = seg["font_name"]
            if seg["font_size"]:
                run.font.size = seg["font_size"]

    # ── Compact / List styles ──
    for sn in ["Compact", "List Paragraph", "List Bullet", "List Number"]:
        try:
            style = doc.styles[sn]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 2.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except KeyError:
            pass

    # ── Block Text (used by Pandoc for blockquotes / references) ──
    try:
        bt = doc.styles["Block Text"]
        bt.font.name = "Times New Roman"
        bt.font.size = Pt(12)
        bt.paragraph_format.line_spacing = 2.0
        bt.paragraph_format.space_before = Pt(0)
        bt.paragraph_format.space_after = Pt(0)
    except KeyError:
        pass

    # ── Center title page paragraphs ──
    # Title page = all paragraphs before the first Heading
    first_heading_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if p.style and p.style.name.startswith("Heading"):
            first_heading_idx = idx
            break

    if first_heading_idx is not None:
        for idx in range(first_heading_idx):
            p = doc.paragraphs[idx]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0)
            _apply_font_to_paragraph(p)
            # Bold only the title line (longest text, typically > 60 chars)
            text = p.text.strip()
            if text and len(text) > 60:
                for run in p.runs:
                    run.bold = True

    # ── Process all paragraphs ──
    in_references = False
    in_toc = False
    in_abstract = False

    # Build a map: paragraph XML element -> table that follows it
    # for table numbering
    body_elements = list(doc.element.body)
    table_predecessors = {}
    for i, el in enumerate(body_elements):
        if el.tag.endswith("}tbl"):
            # Find the paragraph just before this table
            for j in range(i - 1, -1, -1):
                if body_elements[j].tag.endswith("}p"):
                    table_predecessors[id(body_elements[j])] = True
                    break

    for para_idx, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""

        # Set font on all runs
        _apply_font_to_paragraph(p)

        # Detect References section
        if style_name.startswith("Heading") and "referencia" in p.text.lower():
            in_references = True
            in_toc = False
            in_abstract = False

        # Detect Abstract section (Resumen)
        if style_name.startswith("Heading") and "resumen" in p.text.lower():
            in_abstract = True
            in_toc = False
        elif style_name.startswith("Heading") and "resumen" not in p.text.lower():
            in_abstract = False

        # Detect TOC section (Contenido)
        if style_name.startswith("Heading") and "contenido" in p.text.lower():
            in_toc = True
        elif style_name.startswith("Heading") and "contenido" not in p.text.lower():
            in_toc = False

        # Enforce double spacing everywhere
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Heading alignment
        if style_name == "Heading 1":
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name in ("Heading 2", "Heading 3"):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Compute title page flag (needed by multiple sections below)
        is_title_page = first_heading_idx is not None and para_idx < first_heading_idx

        # APA: body text left-aligned (not justified)
        # Skip title page paragraphs (they should remain centered)
        if not style_name.startswith("Heading") and not is_title_page:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # TOC entries: reformat with tab stops and dot leaders
        if in_toc and not style_name.startswith("Heading"):
            text = p.text.strip()
            if text:
                # Parse TOC entry: "Title ... N" or "   SubTitle ... N"
                # Handle both regular dots and Unicode ellipsis (…)
                import re

                normalized = text.replace("\u2026", "...")
                match = re.match(r"^(\s*)(.*?)\s*\.{3,}\s*(\d+)\s*$", normalized)
                if match:
                    indent_str, title, page_num = match.groups()
                    indent_level = len(indent_str)

                    # Clear paragraph and rebuild with tab stop
                    p.clear()
                    p.paragraph_format.first_line_indent = Inches(0)

                    # Indent sub-levels
                    if indent_level > 0:
                        p.paragraph_format.left_indent = Inches(0.5)
                    else:
                        p.paragraph_format.left_indent = Inches(0)

                    # Add right-aligned tab stop with dot leader
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(
                        Inches(
                            6.0
                        ),  # 6 inches from left margin (1" margin + 6" = right edge)
                        alignment=2,  # RIGHT
                        leader=4,  # DOTS
                    )

                    # Add title text
                    run_title = p.add_run(title.strip())
                    run_title.font.name = "Times New Roman"
                    run_title.font.size = Pt(12)

                    # Add tab character
                    run_tab = p.add_run("\t")
                    run_tab.font.name = "Times New Roman"
                    run_tab.font.size = Pt(12)

                    # Add page number
                    run_page = p.add_run(page_num)
                    run_page.font.name = "Times New Roman"
                    run_page.font.size = Pt(12)
                    continue  # Skip further processing for this paragraph

        # Body text indent (skip title page and TOC paragraphs)
        if (
            not style_name.startswith("Heading")
            and "List" not in style_name
            and not in_references
            and not is_title_page
            and not in_toc
        ):
            if style_name in ("Body Text", "First Paragraph", "Normal"):
                if p.text.strip():
                    if in_abstract:
                        p.paragraph_format.first_line_indent = Inches(0)
                    else:
                        p.paragraph_format.first_line_indent = Inches(0.5)

        # References: hanging indent
        if in_references and not style_name.startswith("Heading"):
            if p.text.strip():
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)

    # ── Merge runs & fix spacing on ALL body paragraphs ──
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading"):
            # For headings: just enforce bold/italic on each run
            cfg = heading_configs.get(style_name)
            if cfg:
                for run in p.runs:
                    run.bold = cfg["bold"]
                    run.italic = cfg["italic"]
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
            continue
        # For body text: merge runs and fix spacing
        _merge_and_clean_paragraph(p)
        # Enforce left alignment (APA 7 — no justification)
        if (
            p.paragraph_format.alignment is None
            or p.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER
        ):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Keywords: "Palabras clave:" in italic, indented, keywords not italic ──
    abstract_started = False
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading") and "resumen" in p.text.lower():
            abstract_started = True
            continue
        if abstract_started and style_name.startswith("Heading"):
            abstract_started = False
            continue
        if abstract_started and p.text.strip():
            full = p.text.strip()
            kw_match = re.match(r"^(Palabras\s+clave:\s*)(.*)", full, re.IGNORECASE)
            if kw_match:
                label = kw_match.group(1).strip()
                keywords_text = kw_match.group(2).strip()
                # Clear the paragraph and rebuild with split formatting
                p.clear()
                p.paragraph_format.first_line_indent = Inches(0.5)
                # Label in italic
                run_label = p.add_run(label + " ")
                run_label.italic = True
                run_label.bold = False
                run_label.font.name = "Times New Roman"
                run_label.font.size = Pt(12)
                # Keywords: not italic
                run_kw = p.add_run(keywords_text)
                run_kw.italic = False
                run_kw.bold = False
                run_kw.font.name = "Times New Roman"
                run_kw.font.size = Pt(12)

    # ── Page break after Palabras clave + Repeated Title (APA 7) ──
    # APA 7: The body text must start on a new page.
    # The repeated title acts as the introduction heading (no "Introducción" label).
    kw_para = None
    for i, p in enumerate(doc.paragraphs):
        if "palabras clave" in p.text.lower():
            kw_para = (i, p)
            break

    if kw_para:
        kw_idx, kw_p = kw_para
        # Find the paragraph AFTER keywords (which will be the first body text)
        next_body_idx = None
        for j in range(kw_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.strip():
                next_body_idx = j
                break

        if next_body_idx is not None:
            next_p = doc.paragraphs[next_body_idx]
            # Insert a page break RUN at the beginning of this paragraph's XML
            # We create a new paragraph BEFORE it with the page break + title
            next_p_element = next_p._element

            # 1. Create page-break paragraph
            pb_p = OxmlElement("w:p")
            pb_pPr = OxmlElement("w:pPr")
            pb_spacing = OxmlElement("w:spacing")
            pb_spacing.set(qn("w:line"), "480")
            pb_spacing.set(qn("w:lineRule"), "auto")
            pb_spacing.set(qn("w:before"), "0")
            pb_spacing.set(qn("w:after"), "0")
            pb_pPr.append(pb_spacing)
            pb_p.append(pb_pPr)
            pb_run = OxmlElement("w:r")
            pb_br = OxmlElement("w:br")
            pb_br.set(qn("w:type"), "page")
            pb_run.append(pb_br)
            pb_p.append(pb_run)
            next_p_element.addprevious(pb_p)

            # 2. Create repeated title paragraph (centered, bold, TNR 12pt)
            title_text = meta.get("title", "")
            if title_text:
                title_p = OxmlElement("w:p")
                # Paragraph properties: centered, double-spaced
                t_pPr = OxmlElement("w:pPr")
                t_jc = OxmlElement("w:jc")
                t_jc.set(qn("w:val"), "center")
                t_pPr.append(t_jc)
                t_spacing = OxmlElement("w:spacing")
                t_spacing.set(qn("w:line"), "480")
                t_spacing.set(qn("w:lineRule"), "auto")
                t_spacing.set(qn("w:before"), "0")
                t_spacing.set(qn("w:after"), "0")
                t_pPr.append(t_spacing)
                title_p.append(t_pPr)
                # Run: bold, TNR 12pt
                t_run = OxmlElement("w:r")
                t_rPr = OxmlElement("w:rPr")
                t_b = OxmlElement("w:b")
                t_rPr.append(t_b)
                t_font = OxmlElement("w:rFonts")
                t_font.set(qn("w:ascii"), "Times New Roman")
                t_font.set(qn("w:hAnsi"), "Times New Roman")
                t_rPr.append(t_font)
                t_sz = OxmlElement("w:sz")
                t_sz.set(qn("w:val"), "24")  # 12pt
                t_rPr.append(t_sz)
                t_run.append(t_rPr)
                t_text = OxmlElement("w:t")
                t_text.set(qn("xml:space"), "preserve")
                t_text.text = title_text
                t_run.append(t_text)
                title_p.append(t_run)
                next_p_element.addprevious(title_p)

    # ── Citations: replace "y" with "&" in parenthetical citations ──
    # Pattern: (Author1 y Author2, YEAR) → (Author1 & Author2, YEAR)
    citation_re = re.compile(
        r"\(([A-ZÁ-Ú][a-záéíóúñ]+(?:\s+(?:et\s+al\.))?)\s+y\s+([A-ZÁ-Ú][a-záéíóúñ]+),\s*(\d{4})\)"
    )
    for p in doc.paragraphs:
        for run in p.runs:
            if " y " in run.text and "(" in run.text:
                run.text = citation_re.sub(r"(\1 & \2, \3)", run.text)

    # ── References: merge broken paragraphs and fix Block Text indent ──
    if in_references:
        pass  # in_references was set during paragraph loop above
    # Re-scan for reference paragraphs and fix Block Text style indentation
    ref_section_started = False
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading") and "referencia" in p.text.lower():
            ref_section_started = True
            continue
        if ref_section_started and style_name.startswith("Heading"):
            break  # past references
        if ref_section_started and p.text.strip():
            # Force Block Text paragraphs to use correct indentation
            if style_name == "Block Text":
                p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    # ── Tables: APA borders + numbering + cell text cleanup ──
    def _cell_text(cell):
        """Join all runs/paragraphs in a cell with correct spacing.

        python-docx concatenates runs without separators, so
        italic 'Checkout' + plain ' Simplificado' becomes 'CheckoutSimplificado'.
        This reconstructs the text inserting spaces between runs as needed.
        """
        para_texts = []
        for p in cell.paragraphs:
            run_parts = []
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                if run_parts:
                    prev = run_parts[-1]
                    # Insert space if previous doesn't end with whitespace
                    # and current doesn't start with whitespace
                    if prev and not prev[-1].isspace() and not t[0].isspace():
                        run_parts.append(" ")
                run_parts.append(t)
            combined = "".join(run_parts).strip()
            if combined:
                para_texts.append(combined)
        return " ".join(para_texts)

    for t_idx, table in enumerate(doc.tables):
        _apply_apa_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ─── Clean cell text: merge runs, fix spacing, strip artifacts ───
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # First pass: strip Grid Table border fragments from runs
                    for run in p.runs:
                        t = run.text
                        if t:
                            t = re.sub(r"[+|][-=]{3,}[+|]?", "", t)
                            t = re.sub(r"={3,}", "", t)
                            t = t.strip("|").strip()
                            run.text = t
                    # Second pass: merge runs and fix spacing
                    _merge_and_clean_paragraph(p)

        # ─── Build structured APA table title ───
        table_num = t_idx + 1
        table_id = ""
        desc_title = ""

        for row in table.rows:
            if len(row.cells) > 1:
                header = _cell_text(row.cells[0]).replace("*", "")
                value = _cell_text(row.cells[1]).replace("*", "")
                # Clean grid artifact remnants from header/value
                header = re.sub(r"[+|][-=]{3,}[+|]?", "", header).strip("| ").strip()
                value = re.sub(r"[+|][-=]{3,}[+|]?", "", value).strip("| ").strip()
                if "ID" in header and not table_id:
                    table_id = value
                if (
                    "Nombre" in header or "Atributo de Calidad" in header
                ) and not desc_title:
                    desc_title = value
            if table_id and desc_title:
                break

        # Build a standardized title from the ID
        full_title = ""
        if table_id and desc_title:
            id_upper = table_id.upper()
            if id_upper.startswith("RF-"):
                num = id_upper.replace("RF-", "")
                full_title = f"Requisito Funcional {num}: {desc_title}"
            elif id_upper.startswith("RNF-"):
                num = id_upper.replace("RNF-", "")
                full_title = f"Requisito No Funcional {num}: {desc_title}"
            elif id_upper.startswith("HU-"):
                num = id_upper.replace("HU-", "")
                full_title = f"Historia de Usuario {num}: {desc_title}"
            else:
                full_title = desc_title
        elif desc_title:
            full_title = desc_title
        elif table_id:
            full_title = table_id

        # ─── Remove redundant ID/Nombre rows from table body ───
        # Only for 2-column requirement tables (not stakeholder or matrix)
        if table_id and desc_title and len(table.columns) <= 3:
            rows_to_remove = []
            for row in table.rows:
                if len(row.cells) > 1:
                    hdr = _cell_text(row.cells[0]).replace("*", "").strip()
                    if hdr in ("ID", "Nombre", "Atributo de Calidad"):
                        rows_to_remove.append(row)
            for row in rows_to_remove:
                tbl_xml = table._tbl
                tbl_xml.remove(row._tr)

        # Fallback for the traceability matrix (no ID/Nombre rows)
        if not full_title:
            # Check if first row header looks like a matrix
            if len(table.rows) > 0 and len(table.rows[0].cells) > 2:
                full_title = "Matriz de Trazabilidad de Requisitos"

        # Insert "Tabla N" in bold, left-aligned
        tbl_element = table._tbl
        label_p = OxmlElement("w:p")
        # Paragraph properties: left-aligned, double-spaced
        label_pPr = OxmlElement("w:pPr")
        label_jc = OxmlElement("w:jc")
        label_jc.set(qn("w:val"), "left")
        label_pPr.append(label_jc)
        label_spacing = OxmlElement("w:spacing")
        label_spacing.set(qn("w:line"), "480")  # double space = 480 twips
        label_spacing.set(qn("w:after"), "0")
        label_pPr.append(label_spacing)
        label_p.append(label_pPr)

        label_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b_elem = OxmlElement("w:b")
        rPr.append(b_elem)
        font_name = OxmlElement("w:rFonts")
        font_name.set(qn("w:ascii"), "Times New Roman")
        font_name.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(font_name)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")  # 12pt = 24 half-points
        rPr.append(sz)
        label_run.append(rPr)
        label_text = OxmlElement("w:t")
        label_text.text = f"Tabla {table_num}"
        label_run.append(label_text)
        label_p.append(label_run)

        tbl_element.addprevious(label_p)

        # Descriptive title in italic (on a separate line)
        if full_title:
            title_p = OxmlElement("w:p")
            # Left-aligned, double-spaced
            title_pPr = OxmlElement("w:pPr")
            title_jc = OxmlElement("w:jc")
            title_jc.set(qn("w:val"), "left")
            title_pPr.append(title_jc)
            title_spacing = OxmlElement("w:spacing")
            title_spacing.set(qn("w:line"), "480")
            title_spacing.set(qn("w:after"), "0")
            title_pPr.append(title_spacing)
            title_p.append(title_pPr)
            title_run = OxmlElement("w:r")
            tPr = OxmlElement("w:rPr")
            i_elem = OxmlElement("w:i")
            tPr.append(i_elem)
            font_name2 = OxmlElement("w:rFonts")
            font_name2.set(qn("w:ascii"), "Times New Roman")
            font_name2.set(qn("w:hAnsi"), "Times New Roman")
            tPr.append(font_name2)
            sz2 = OxmlElement("w:sz")
            sz2.set(qn("w:val"), "24")
            tPr.append(sz2)
            title_run.append(tPr)
            title_text = OxmlElement("w:t")
            title_text.text = full_title[:120]
            title_run.append(title_text)
            title_p.append(title_run)
            tbl_element.addprevious(title_p)

        # Apply font/spacing to all cells
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 2.0
                    p.paragraph_format.first_line_indent = Inches(0)
                    _apply_font_to_paragraph(p, font_size=Pt(10))

    doc.save(output_path)
    print(f"  ✓ Formato APA 7 aplicado → {output_path}")


# ────────────────────────── Main ──────────────────────────


def convert_to_pdf(docx_path: Path, output_dir: Path):
    """Convert DOCX to PDF using LibreOffice."""
    print("Paso 5: Generando PDF (LibreOffice)...")
    try:
        # LibreOffice requires an outdir, and it keeps the original filename
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]
        print(f"  ▸ Ejecutando: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            print(f"  ✗ Error de LibreOffice:\n{result.stderr}", file=sys.stderr)
            print(
                "  ⚠ No se pudo generar el PDF. Asegúrate de tener LibreOffice instalado."
            )
            return

        pdf_path = output_dir / docx_path.with_suffix(".pdf").name
        if pdf_path.exists():
            print(f"  ✓ PDF generado: {pdf_path}")
        else:
            print("  ⚠ El PDF no se generó por una razón desconocida.")

    except FileNotFoundError:
        print("  ⚠ LibreOffice no encontrado. Instálalo para generar PDFs.")


def main():
    parser = argparse.ArgumentParser(
        description="Conversor ERS Markdown -> DOCX (APA 7a edición)"
    )
    parser.add_argument(
        "input_file",
        nargs="?",
        default="ERS_Shoppipai_SENA_COMPLETO.md",
        help="Archivo Markdown de entrada (por defecto: ERS_Shoppipai_SENA_COMPLETO.md)",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        default="DOCS",
        help="Directorio de salida (por defecto: DOCS)",
    )
    parser.add_argument(
        "-f",
        "--format",
        choices=["docx", "pdf", "all"],
        default="docx",
        help="Formato de salida: docx, pdf o all (ambos). Por defecto: docx",
    )

    args = parser.parse_args()

    input_file = Path(args.input_file)
    output_dir = Path(args.output_dir)

    # Determine output filename based on input filename
    output_docx = output_dir / f"{input_file.stem}_APA.docx"

    print("╔══════════════════════════════════════════════════╗")
    print("║   Conversor ERS -> DOCX (APA 7a edición)         ║")
    print("║   Usando Pandoc + python-docx                    ║")
    print("╚══════════════════════════════════════════════════╝")
    print()

    if not input_file.exists():
        print(f"✗ Error: No se encontró el archivo de entrada '{input_file}'")
        sys.exit(1)

    # Create output directory if it doesn't exist
    if not output_dir.exists():
        output_dir.mkdir(parents=True, exist_ok=True)
        print(f"  ✓ Directorio creado: {output_dir}")

    # Read source
    try:
        raw_md = input_file.read_text(encoding="utf-8")
    except Exception as e:
        print(f"✗ Error leyendo archivo: {e}")
        sys.exit(1)

    lines = raw_md.split("\n")

    # Step 1: Extract metadata
    print("Paso 1: Extrayendo metadatos...")
    meta = _extract_metadata(lines)
    for k, v in meta.items():
        print(f"  ▸ {k}: {v}")

    # Step 2: Pre-process markdown
    print("Paso 2: Pre-procesando Markdown (saltos de página, portada)...")
    processed_md = preprocess_markdown(raw_md, meta)
    print(f"  ✓ Markdown pre-procesado ({len(processed_md)} caracteres)")

    # Step 3: Pandoc conversion
    temp_docx = "_pandoc_raw.docx"
    print("Paso 3: Conversión con Pandoc...")
    run_pandoc(processed_md, temp_docx)

    # Step 4: APA formatting
    print("Paso 4: Aplicando formato APA 7a edición...")
    apply_apa_styles(temp_docx, str(output_docx), meta)

    # Cleanup raw pandoc file
    Path(temp_docx).unlink(missing_ok=True)

    # Step 5: PDF Generation (if requested)
    if args.format in ["pdf", "all"]:
        convert_to_pdf(output_docx, output_dir)
        if args.format == "pdf":
            # If user ONLY wanted PDF, maybe delete the DOCX?
            # Usually users expect the intermediate DOCX if conversion fails,
            # but strict interpret might mean delete it.
            # For safety/utility, I'll keep the DOCX too or strictly follow request.
            # Let's keep DOCX as intermediate artifact is often useful, unless user explicitly complains.
            # Actually, standard behavior for "pdf" source usually implies result is PDF.
            # But here we are "converting". Let's leave DOCX as it's the source for PDF.
            pass

    # Summary
    if output_docx.exists():
        file_size = output_docx.stat().st_size
        print()
        print(f"✅ DOCX Generado: {output_docx}")
        print(f"   Tamaño: {file_size / 1024:.1f} KB")

    print()
    print("Normas APA 7 aplicadas correctamante.")
    print()


if __name__ == "__main__":
    main()
