import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..models import DocumentMetadata
from .base import DocumentFormatter


class APADocxFormatter(DocumentFormatter):
    """
    Applies APA 7th Edition formatting to a DOCX file.
    """

    def __init__(self, doc_path: str):
        super().__init__(doc_path)
        self.doc: DocumentObject = Document(doc_path)

    def process(self, meta: DocumentMetadata):
        """Run the full formatting pipeline."""
        self._setup_page_layout()
        self._create_styles()
        self._add_cover_page(meta)
        self._process_paragraphs()
        self._format_tables()
        self._format_figures()
        self._format_nota_italic()
        self._format_lists()
        self._apply_body_indent()
        self._format_keywords(meta)
        self._fix_text_spacing_global()

    def _make_figure_paragraph(self, text, bold=False, italic=False, space_after="0"):
        """Helper: create a Times New Roman 12pt paragraph for figure captions."""
        p_el = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        p_sp = OxmlElement("w:spacing")
        p_sp.set(qn("w:after"), space_after)
        p_sp.set(qn("w:line"), "480")
        p_sp.set(qn("w:lineRule"), "auto")
        p_pr.append(p_sp)
        p_el.append(p_pr)

        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        rn = OxmlElement("w:rFonts")
        rn.set(qn("w:ascii"), "Times New Roman")
        rn.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(rn)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        p_el.append(run)
        return p_el

    def _format_figures(self):
        """Add APA 7 figure captions: Label + Title ABOVE, Nota BELOW.

        APA 7 figure order:
          Figura N        (bold, left-aligned)
          Italic title    (italic, left-aligned, no trailing period)
          [image]         (centered)
          Nota. context   (italic "Nota.", then regular text)
        """
        ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        image_paragraphs = []

        for p in self.doc.paragraphs:
            drawings = p._element.findall(f".//{qn('w:drawing')}")
            if drawings:
                alt_text = ""
                for drawing in drawings:
                    for docPr in drawing.iter(f"{{{ns_wp}}}docPr"):
                        alt_text = docPr.get("descr", "") or docPr.get("name", "")
                        break
                    if not alt_text:
                        ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"
                        for cNvPr in drawing.iter(f"{{{ns_pic}}}cNvPr"):
                            alt_text = cNvPr.get("descr", "") or cNvPr.get("name", "")
                            break
                image_paragraphs.append((p, alt_text.strip()))

        for _, (p, _) in enumerate(image_paragraphs, start=1):
            # Center the image paragraph
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Scale oversized images to fit within the usable page area.
            # Inline images in DOCX CANNOT span pages — LibreOffice clips
            # them at the page boundary. We must always scale to fit.
            # Max usable: 6.5in wide x 8.5in tall (leaving room for captions).
            max_w = Inches(6.5)
            max_h = Inches(8.5)
            drawings = p._element.findall(f".//{{{ns_wp}}}inline") + p._element.findall(
                f".//{{{ns_wp}}}anchor"
            )
            for d in drawings:
                extent = d.find(f"{{{ns_wp}}}extent")
                if extent is None:
                    continue
                cx = int(extent.get("cx", 0))
                cy = int(extent.get("cy", 0))
                if cx == 0 or cy == 0:
                    continue

                scale = 1.0
                if cx > max_w:
                    scale = min(scale, max_w / cx)
                if cy > max_h:
                    scale = min(scale, max_h / cy)

                if scale < 1.0:
                    new_cx = int(cx * scale)
                    new_cy = int(cy * scale)
                    extent.set("cx", str(new_cx))
                    extent.set("cy", str(new_cy))

                    # Also update the a:ext in the spPr (shape properties)
                    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
                    for ext_el in d.iter(f"{{{ns_a}}}ext"):
                        old_cx = int(ext_el.get("cx", 0))
                        old_cy = int(ext_el.get("cy", 0))
                        if old_cx > 0:
                            ext_el.set("cx", str(int(old_cx * scale)))
                        if old_cy > 0:
                            ext_el.set("cy", str(int(old_cy * scale)))

    def _format_nota_italic(self):
        """APA 7: 'Nota.' must be italic in figure/table notes.

        Finds paragraphs starting with 'Nota.' and splits the first run
        so that 'Nota.' is italic and the rest is regular.
        """
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text.startswith("Nota."):
                continue

            # Get full text and rebuild with italic "Nota."
            full_text = p.text
            # Clear all runs
            for run in list(p.runs):
                run._element.getparent().remove(run._element)

            # Add "Nota. " as italic
            nota_run = p.add_run("Nota. ")
            nota_run.italic = True
            self._apply_font_style(nota_run, italic=True)

            # Add the rest as regular text
            rest = full_text[len("Nota.") :].strip()
            if rest:
                rest_run = p.add_run(rest)
                self._apply_font_style(rest_run)

            # Ensure left alignment and no indent for notes
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)

    def _format_lists(self):
        """Apply APA 7 list formatting.

        Regular bullet lists: bullet at 0.5in, text at 0.75in, hanging indent.
        Reference entries: no bullet, hanging indent at 0.5in (APA 7 standard).
        """
        in_references = False

        for p in self.doc.paragraphs:
            # Track sections via headings
            style_name = p.style.name if p.style else ""
            if style_name == "Heading 1":
                text_lower = p.text.lower().strip()
                in_references = "referencia" in text_lower or "reference" in text_lower
                continue

            pPr = p._element.find(qn("w:pPr"))
            if pPr is None:
                continue

            numPr = pPr.find(qn("w:numPr"))
            if numPr is None:
                continue

            # Remove Pandoc's numbering reference
            pPr.remove(numPr)

            if in_references:
                # APA 7 reference: hanging indent, NO bullet
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
            else:
                # APA 7 bullet list: bullet at 0.5in, text at 0.75in
                p.paragraph_format.left_indent = Inches(0.75)
                p.paragraph_format.first_line_indent = Inches(-0.25)

                # Tab stop so tab after bullet snaps text into position
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(0.75), alignment=0)

                # Prepend bullet character
                text = p.text
                if not text.startswith("\u2022") and not text.startswith("-"):
                    first_run = p.runs[0] if p.runs else None
                    if first_run:
                        first_run.text = "\u2022\t" + first_run.text
                    else:
                        run = p.add_run("\u2022\t")
                        self._apply_font_style(run)

    def _apply_body_indent(self):
        """Final pass: apply first-line indent to all body paragraphs.

        Runs after _format_tables so that newly-created paragraphs (table
        descriptions, etc.) are also covered.
        """
        in_references = False
        in_toc = False
        in_abstract = False
        for p in self.doc.paragraphs:
            style_name = p.style.name if p.style else ""
            text = p.text.strip()

            # Track References, TOC, and Abstract sections
            if style_name.startswith("Heading"):
                text_lower = text.lower()
                if "referencia" in text_lower or "reference" in text_lower:
                    in_references = True
                    in_toc = False
                    in_abstract = False
                elif "contenido" in text_lower or "index" in text_lower:
                    in_toc = True
                    in_references = False
                    in_abstract = False
                elif "resumen" in text_lower or "abstract" in text_lower:
                    in_abstract = True
                    in_toc = False
                    in_references = False
                else:
                    in_references = False
                    in_toc = False
                    in_abstract = False
                continue  # Skip headings

            if not text or len(text) < 5:
                continue  # Skip empty or very short labels

            # Skip TOC entries
            if in_toc:
                continue

            # Skip abstract paragraphs (block format, no indent per APA 7)
            if in_abstract:
                # End abstract section after keywords paragraph
                if text.lower().startswith("palabras clave") or text.lower().startswith("keywords"):
                    in_abstract = False
                continue

            # Skip if already indented (positive or negative)
            fli = p.paragraph_format.first_line_indent
            if fli is not None and fli != 0:
                continue

            # Skip centered paragraphs (cover page)
            align = p.paragraph_format.alignment
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                continue

            # Skip lists, captions, compact (Pandoc lists), and other special styles
            if "List" in style_name or "Caption" in style_name or "Compact" in style_name:
                continue

            # Skip table/figure caption labels and Nota paragraphs
            if (text.startswith("Tabla ") or text.startswith("Figura ")) and len(text.split()) <= 2:
                continue
            if text.startswith("Nota."):
                continue

            # Skip short all-italic paragraphs (table caption titles)
            # Long italic paragraphs are body text that should be indented
            if len(text) < 80 and p.runs and all(r.italic for r in p.runs if r.text.strip()):
                continue

            # Skip references section (handled with hanging indent)
            if in_references:
                continue

            # Apply APA first-line indent (0.5 inches / 1.27 cm)
            p.paragraph_format.first_line_indent = Inches(0.5)

    def save(self, output_path: str):
        self.doc.save(str(output_path))

    def _apply_apa_table_borders(self, table):
        """
        Apply APA-style borders:
        - Top/Bottom of table: single line
        - Bottom of header row: single line
        - No vertical lines.
        """
        # Single-pass: set correct borders for each row position

        # Clear all borders first for all cells
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(
                    cell, top={}, bottom={}, start={}, end={}, insideH={}, insideV={}
                )

        # Iterate once and set correct edges.
        num_rows = len(table.rows)
        for i, row in enumerate(table.rows):
            is_first_row = i == 0
            is_last_row = i == num_rows - 1

            for cell in row.cells:
                borders = {}
                if is_first_row:
                    borders["top"] = {"val": "single", "sz": "12", "color": "auto"}
                    borders["bottom"] = {
                        "val": "single",
                        "sz": "6",
                        "color": "auto",
                    }  # Header underline

                if is_last_row:
                    borders["bottom"] = {"val": "single", "sz": "12", "color": "auto"}

                # Apply
                self._set_cell_border(cell, **borders)

    # ────────────────────────── OXML Helpers ──────────────────────────

    def _set_cell_border(self, cell, **kwargs):
        """Set border on a table cell (OpenXML). Clears existing first."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove any existing cell borders
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)
        tcBorders = OxmlElement("w:tcBorders")
        for edge_name in ("start", "top", "end", "bottom", "insideH", "insideV"):
            if edge_name in kwargs:
                el = OxmlElement(f"w:{edge_name}")
                for attr, val in kwargs[edge_name].items():
                    el.set(qn(f"w:{attr}"), str(val))
                tcBorders.append(el)
        tcPr.append(tcBorders)

    # ────────────────────────── Formatting Steps ──────────────────────────

    def _setup_page_layout(self):
        """Set margins to 1 inch and add page numbers top-right."""
        for section in self.doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

            # Disable separate first-page header/footer
            section.different_first_page_header_footer = False

            # Clear footers completely to prevent Pandoc page numbers at bottom
            footer = section.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.clear()
                # Remove ALL child elements (runs, field codes, pPr, etc.)
                for child in list(p._element):
                    p._element.remove(child)
            # Also set footer distance to zero to suppress any residual space
            sectPr = section._sectPr
            existing_pgMar = sectPr.find(qn("w:pgMar"))
            if existing_pgMar is not None:
                existing_pgMar.set(qn("w:footer"), "0")

            # Remove the footerReference entirely to prevent LibreOffice
            # from rendering any footer content
            for fref in list(sectPr.findall(qn("w:footerReference"))):
                sectPr.remove(fref)

            self._add_page_number(section)

    def _add_page_number(self, section):
        """Add page number top-right in header."""
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing header text
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Build PAGE field with begin/separate/end sequence
        # (LibreOffice requires the 'separate' marker to render page numbers)
        run_begin = hp.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_begin)
        run_begin.font.name = "Times New Roman"
        run_begin.font.size = Pt(12)

        run_instr = hp.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)
        run_instr.font.name = "Times New Roman"
        run_instr.font.size = Pt(12)

        run_sep = hp.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)
        run_sep.font.name = "Times New Roman"
        run_sep.font.size = Pt(12)

        # Placeholder text (will be replaced by actual page number)
        run_num = hp.add_run("1")
        run_num.font.name = "Times New Roman"
        run_num.font.size = Pt(12)

        run_end = hp.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)
        run_end.font.name = "Times New Roman"
        run_end.font.size = Pt(12)

        # Strip excessive paragraphs in header
        while len(header.paragraphs) > 1:
            p_elem = header.paragraphs[-1]._element
            p_elem.getparent().remove(p_elem)

    def _create_styles(self):
        """
        Configure Normal, Headings, and other styles.
        """
        styles = self.doc.styles

        # Normal
        normal = styles["Normal"]
        self._apply_font_style(normal, size=12)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Body Text
        for style_name in ["Body Text", "First Paragraph"]:
            try:
                style = styles[style_name]
                self._apply_font_style(style, size=12)
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                style.paragraph_format.first_line_indent = Inches(0.5)
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except KeyError:
                pass

        # Headings
        configs = {
            "Heading 1": {
                "bold": True,
                "italic": False,
                "align": WD_ALIGN_PARAGRAPH.CENTER,
            },
            "Heading 2": {
                "bold": True,
                "italic": False,
                "align": WD_ALIGN_PARAGRAPH.LEFT,
            },
            "Heading 3": {
                "bold": True,
                "italic": True,
                "align": WD_ALIGN_PARAGRAPH.LEFT,
            },
        }
        for sn, cfg in configs.items():
            try:
                h = styles[sn]
                self._apply_font_style(h, size=12, bold=cfg["bold"], italic=cfg["italic"])
                h.paragraph_format.alignment = cfg["align"]
                h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                h.paragraph_format.page_break_before = False
            except KeyError:
                pass

        # Override Compact style (used by Pandoc for table cells & lists)
        # Must have single spacing + left alignment for APA tables
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for style_el in self.doc.styles.element.findall(f"{{{ns}}}style"):
            style_id = style_el.get(f"{{{ns}}}styleId", "")
            if style_id == "Compact":
                # Replace pPr with single spacing + left alignment
                old_pPr = style_el.find(f"{{{ns}}}pPr")
                if old_pPr is not None:
                    style_el.remove(old_pPr)
                pPr = OxmlElement("w:pPr")
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                spacing.set(qn("w:before"), "36")
                spacing.set(qn("w:after"), "36")
                pPr.append(spacing)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "left")
                pPr.append(jc)
                # Remove first-line indent inherited from BodyText
                ind = OxmlElement("w:ind")
                ind.set(qn("w:firstLine"), "0")
                pPr.append(ind)
                style_el.append(pPr)
                break

        # Remove Pandoc Table Style borders
        self._neutralize_table_style()

    def _neutralize_table_style(self):
        """Remove borders from 'Table' style and force left-alignment + single spacing."""
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for style_el in self.doc.styles.element.findall(f"{{{ns}}}style"):
            style_id = style_el.get(f"{{{ns}}}styleId", "")
            if style_id == "Table":
                tblPr_el = style_el.find(f"{{{ns}}}tblPr")
                if tblPr_el is not None:
                    for old_b in tblPr_el.findall(f"{{{ns}}}tblBorders"):
                        tblPr_el.remove(old_b)

                    # Set explicit no borders
                    brd = OxmlElement("w:tblBorders")
                    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        e = OxmlElement(f"w:{edge}")
                        e.set(qn("w:val"), "none")
                        e.set(qn("w:sz"), "0")
                        e.set(qn("w:space"), "0")
                        brd.append(e)
                    tblPr_el.append(brd)

                # Add paragraph properties to the Table style (left align + single spacing)
                # This is authoritative for LibreOffice's style inheritance
                old_pPr = style_el.find(f"{{{ns}}}pPr")
                if old_pPr is not None:
                    style_el.remove(old_pPr)
                pPr = OxmlElement("w:pPr")
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "left")
                pPr.append(jc)
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                spacing.set(qn("w:before"), "0")
                spacing.set(qn("w:after"), "0")
                pPr.append(spacing)
                style_el.append(pPr)

                # Add run properties (font) to the Table style
                old_rPr = style_el.find(f"{{{ns}}}rPr")
                if old_rPr is not None:
                    style_el.remove(old_rPr)
                rPr = OxmlElement("w:rPr")
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rPr.append(rFonts)
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "22")  # 11pt
                rPr.append(sz)
                style_el.append(rPr)

                # Fix the firstRow tblStylePr: change vAlign from bottom to top
                for tsp in style_el.findall(f"{{{ns}}}tblStylePr"):
                    if tsp.get(f"{{{ns}}}type") == "firstRow":
                        tcPr = tsp.find(f"{{{ns}}}tcPr")
                        if tcPr is not None:
                            old_va = tcPr.find(f"{{{ns}}}vAlign")
                            if old_va is not None:
                                tcPr.remove(old_va)
                            va = OxmlElement("w:vAlign")
                            va.set(qn("w:val"), "top")
                            tcPr.append(va)

    def _apply_font_style(
        self,
        style_or_run,
        font_name="Times New Roman",
        size=12,
        bold=None,
        italic=None,
        color_rgb=(0, 0, 0),
    ):
        """Helper to apply font settings."""
        font = style_or_run.font
        font.name = font_name
        font.size = Pt(size)
        if hasattr(font.color, "rgb"):
            font.color.rgb = RGBColor(*color_rgb)
        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic

    def _apply_font_to_paragraph(self, paragraph, font_size=None):
        size = 12 if font_size is None else font_size.pt
        for run in paragraph.runs:
            self._apply_font_style(run, size=size)

    def _add_cover_page(self, meta: DocumentMetadata):
        """
        Insert a new paragraph at index 0 for the cover page.
        """
        # Create a new paragraph for the cover content
        self.doc.add_paragraph()
        if self.doc.paragraphs:
            self.doc.paragraphs[0].insert_paragraph_before("")

        # Extract fields safely
        instructor = getattr(meta, "instructor", "") or meta.extra.get("instructor", "")
        due_date = getattr(meta, "date", "") or ""
        institution = getattr(meta, "institution", "") or ""
        program = getattr(meta, "program", "") or ""
        ficha = getattr(meta, "ficha", "") or ""
        center = getattr(meta, "center", "") or ""

        # APA 7 Student Paper: cover page with content centered vertically
        # in the upper half of the page. Calculate how many spacers to add
        # above the title: page is ~23 double-spaced lines (11in - 2in margins
        # = 9in at 24pt line height), so center ≈ ~11 lines from top.
        # Content takes ~8 lines; centering means ~(23 - 8) / 3 = ~5 spacers.
        content_lines: list[tuple[str, bool]] = [
            (meta.title, True),  # Title: centred, bold
            ("", False),  # Blank line between title and author info
            (meta.author or "", False),
        ]

        # Append remaining fields only if they exist
        for field_val in [program, ficha, institution, center, instructor, due_date]:
            if field_val:
                content_lines.append((field_val, False))

        # Calculate spacers: APA 7 wants title at about 1/3 from the top
        # of the page. With double spacing (24pt lines), the usable page
        # height is ~23 lines. 1/3 down = ~line 7-8. We use 3 blank lines
        # (header + margin already take up space above).
        n_spacers = 3
        elements: list[tuple[str, bool]] = [("", False)] * n_spacers
        elements.extend(content_lines)

        ref_p = self.doc.paragraphs[0]

        for text, is_bold in elements:
            p = ref_p.insert_paragraph_before(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = self.doc.styles["Normal"]

            # Apply bold manually if needed
            if is_bold and text.strip() and p.runs:
                p.runs[0].bold = True

        # Remove the leftover reference paragraph to avoid blank page 2.
        ref_p._element.getparent().remove(ref_p._element)

        # Remove the trailing empty paragraph added by doc.add_paragraph()
        last_p = self.doc.paragraphs[-1]
        if not last_p.text.strip():
            last_p._element.getparent().remove(last_p._element)

        # Add page break after cover page so body text starts on page 2
        # Find the last cover paragraph (centered) and set page break on the
        # first non-cover paragraph that follows
        cover_ended = False
        for p in self.doc.paragraphs:
            if cover_ended:
                # First paragraph after cover — force new page
                p.paragraph_format.page_break_before = True
                break
            # Cover paragraphs are centered; first non-centered = end of cover
            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading"):
                p.paragraph_format.page_break_before = True
                break
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER and p.text.strip():
                cover_ended = True
                p.paragraph_format.page_break_before = True
                break

    def _process_paragraphs(self):
        """Iterate through paragraphs to apply indentation, fix citations, etc."""
        in_references = False
        in_toc = False
        in_abstract = False
        just_left_abstract = False
        heading_levels = self._build_heading_level_map()

        for p in self.doc.paragraphs:
            style_name = p.style.name if p.style else ""
            self._apply_font_to_paragraph(p)

            text_lower = p.text.lower()

            # Detect sections
            if style_name.startswith("Heading"):
                if "referencia" in text_lower or "reference" in text_lower:
                    in_references = True
                    in_toc = False
                    in_abstract = False
                    p.paragraph_format.page_break_before = True
                elif "resumen" in text_lower or "abstract" in text_lower:
                    in_abstract = True
                    in_toc = False
                elif "contenido" in text_lower or "index" in text_lower:
                    in_toc = True
                    p.paragraph_format.page_break_before = True
                else:
                    # Every Level 1 heading starts a new page in APA 7
                    if style_name == "Heading 1":
                        p.paragraph_format.page_break_before = True
                    # If leaving abstract section, force page break on any heading
                    if in_abstract or just_left_abstract:
                        p.paragraph_format.page_break_before = True
                        just_left_abstract = False
                    in_abstract = False
                    in_toc = False

                # Explicit APA heading alignment and bold on every heading
                if style_name == "Heading 1":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                elif style_name in ("Heading 2", "Heading 3"):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.bold = True

                # Strip numbering property from headings (APA 7 doesn't use numbered headings)
                try:
                    p_pr = p._element.find(qn("w:pPr"))
                    if p_pr is not None:
                        num_pr = p_pr.find(qn("w:numPr"))
                        if num_pr is not None:
                            p_pr.remove(num_pr)
                except Exception:
                    pass
                p.paragraph_format.first_line_indent = Inches(0)

            # Line spacing
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Special section handling
            if in_toc and not style_name.startswith("Heading"):
                self._format_toc_entry(p, heading_levels)
                continue

            # Indentation logic
            if (
                not style_name.startswith("Heading")
                and "List" not in style_name
                and "Caption" not in style_name
            ):
                text_strip = p.text.strip()

                # Remove purely numeric paragraphs (Pandoc page numbers injected as text)
                if text_strip.isdigit():
                    p._element.getparent().remove(p._element)
                    continue

                if in_references:
                    if text_strip:
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(-0.5)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Force XML tag to prevent tag dropping
                        p_pr = p._element.get_or_add_pPr()
                        jc = p_pr.get_or_add_jc()
                        jc.set(qn("w:val"), "left")
                elif in_abstract:
                    if p.text.strip():
                        # Abstract block format (no indent per APA 7)
                        p.paragraph_format.first_line_indent = Inches(0)
                    # End abstract after keywords paragraph
                    if text_strip.lower().startswith(
                        "palabras clave"
                    ) or text_strip.lower().startswith("keywords"):
                        in_abstract = False
                        just_left_abstract = True
                elif (
                    style_name in ("Body Text", "Normal", "First Paragraph", "Compact")
                    and text_strip
                    and p.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER
                ):
                    # Force page break on first paragraph after abstract/keywords
                    if just_left_abstract:
                        p.paragraph_format.page_break_before = True
                        just_left_abstract = False
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Force XML tag to prevent tag dropping
                    p_pr = p._element.get_or_add_pPr()
                    jc = p_pr.get_or_add_jc()
                    jc.set(qn("w:val"), "left")

            # Fix citations (y -> &)
            self._fix_citations(p)

    def _format_toc_entry(self, p, heading_levels):
        """Format Table of Contents entries with correct indentation."""
        text = p.text.strip()
        if not text:
            return

        # Parse TOC entry: "Title\tPageNum" or "Title ... PageNum"
        title = None
        page_num = None

        # Try tab-separated first (Pandoc default)
        if "\t" in text:
            parts = text.rsplit("\t", 1)
            if len(parts) == 2 and parts[1].strip().isdigit():
                title = parts[0].strip()
                page_num = parts[1].strip()

        # Fallback: dot-separated
        if title is None:
            normalized = text.replace("\u2026", "...")
            match = re.match(r"^(.*?)\s*\.{3,}\s*(\d+)\s*$", normalized)
            if match:
                title = match.group(1).strip()
                page_num = match.group(2)

        # Fallback: space + trailing digits
        if title is None:
            match = re.match(r"^(.*?)\s{2,}(\d+)\s*$", text)
            if match:
                title = match.group(1).strip()
                page_num = match.group(2)

        if not title or not page_num:
            return

        # Determine heading level from the map
        title_lower = title.lower().strip()
        level = heading_levels.get(title_lower, 1)

        # Indentation: H1=0in, H2=0.5in, H3=1.0in
        indent_map = {1: 0, 2: 0.5, 3: 1.0}
        left_indent = indent_map.get(level, 0)

        p.clear()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(left_indent)

        # Tab stop at fixed 6.5in (right margin) for all entries,
        # ensuring page numbers are perfectly vertically aligned.
        # leader=1 = DOTS (......), alignment=2 = RIGHT
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), alignment=2, leader=1)

        run = p.add_run(title)
        self._apply_font_style(run)

        # Prevent Word/LibreOffice from auto-formatting "1. " as a numbered list
        # which breaks the tab leader layout during PDF conversion
        if title and title[:1].isdigit():
            run.text = "\u200b" + run.text

        run = p.add_run("\t")
        self._apply_font_style(run)

        run = p.add_run(page_num)
        self._apply_font_style(run)

    def _build_heading_level_map(self):
        """Build a map of heading text -> heading level from the document."""
        levels = {}
        for p in self.doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading"):
                parts = style_name.split()
                if len(parts) >= 2 and parts[-1].isdigit():
                    level = int(parts[-1])
                    text = p.text.strip().lower()
                    if text:
                        levels[text] = level
        return levels

    def _fix_citations(self, p):
        """Replace ' (Author y Author, YEAR)' with '&'."""
        citation_re = re.compile(
            r"\(([A-ZÁ-Ú][a-záéíóúñ]+(?:\s+(?:et\s+al\.))?)\s+y\s+([A-ZÁ-Ú][a-záéíóúñ]+),\s*(\d{4})\)"
        )
        for run in p.runs:
            if " y " in run.text and "(" in run.text:
                run.text = citation_re.sub(r"(\1 & \2, \3)", run.text)

    def _format_tables(self):
        """Apply APA borders and formatting to all tables."""
        for i, table in enumerate(self.doc.tables):
            self._apply_apa_table_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set table to full page width with FIXED layout
            tbl_pr = table._tbl.tblPr
            if tbl_pr is not None:
                # Use FIXED layout so LibreOffice respects explicit column widths
                existing_layout = tbl_pr.find(qn("w:tblLayout"))
                if existing_layout is not None:
                    tbl_pr.remove(existing_layout)
                layout = OxmlElement("w:tblLayout")
                layout.set(qn("w:type"), "fixed")
                tbl_pr.append(layout)

                # Set table width to 100% of page
                tbl_w = tbl_pr.find(qn("w:tblW"))
                if tbl_w is None:
                    tbl_w = OxmlElement("w:tblW")
                    tbl_pr.append(tbl_w)
                tbl_w.set(qn("w:type"), "pct")
                tbl_w.set(qn("w:w"), "5000")  # 100% in fifths of a percent

                # Set table-level cell margins
                existing_tblcm = tbl_pr.find(qn("w:tblCellMar"))
                if existing_tblcm is not None:
                    tbl_pr.remove(existing_tblcm)
                tbl_cell_mar = OxmlElement("w:tblCellMar")
                for side in ("top", "bottom", "start", "end"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:w"), "57")  # ~1mm padding
                    el.set(qn("w:type"), "dxa")
                    tbl_cell_mar.append(el)
                tbl_pr.append(tbl_cell_mar)

            # Reduce cell margins, set vertical top-alignment and left-alignment
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._element.find(qn("w:tcPr"))
                    if tc_pr is None:
                        tc_pr = OxmlElement("w:tcPr")
                        cell._element.insert(0, tc_pr)

                    # Vertical alignment: top (APA 7 requirement)
                    existing_valign = tc_pr.find(qn("w:vAlign"))
                    if existing_valign is not None:
                        tc_pr.remove(existing_valign)
                    v_align = OxmlElement("w:vAlign")
                    v_align.set(qn("w:val"), "top")
                    tc_pr.append(v_align)

                    # Remove existing margins
                    existing_mar = tc_pr.find(qn("w:tcMar"))
                    if existing_mar is not None:
                        tc_pr.remove(existing_mar)
                    tc_mar = OxmlElement("w:tcMar")
                    for side in ("top", "bottom", "start", "end"):
                        el = OxmlElement(f"w:{side}")
                        el.set(qn("w:w"), "28")  # small margin (~0.5mm)
                        el.set(qn("w:type"), "dxa")
                        tc_mar.append(el)
                    tc_pr.append(tc_mar)

                    # Left-align all cell paragraphs (APA 7 for text content)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_pr = p._element.get_or_add_pPr()
                        jc = p_pr.get_or_add_jc()
                        jc.set(qn("w:val"), "left")

            # Scale font size based on column count
            num_cols = len(table.columns)
            if num_cols >= 6:
                font_size = 8
            elif num_cols >= 5:
                font_size = 9
            elif num_cols >= 4:
                font_size = 10
            else:
                font_size = 11

            # Set proportional column widths based on content
            if num_cols >= 2:
                # Calculate max content length per column (across all rows)
                max_content_len = [0] * num_cols
                max_word_len = [0] * num_cols
                for row in table.rows:
                    for ci, cell in enumerate(row.cells):
                        if ci < num_cols:
                            text = cell.text.strip()
                            max_content_len[ci] = max(max_content_len[ci], len(text))
                            for word in text.split():
                                max_word_len[ci] = max(max_word_len[ci], len(word))

                # Char width estimate per font size (includes ~0.1in padding)
                cw_map = {8: 0.10, 9: 0.10, 10: 0.12, 11: 0.12}
                cw = cw_map.get(font_size, 0.10)

                # Minimum column width: at least wide enough for longest word
                # plus enforce a minimum of 1.5in for all columns
                min_col = 1.5
                min_widths = [max(w * cw + 0.12, min_col) for w in max_word_len]
                total_min = sum(min_widths)

                avail = 6.5
                col_widths_inches = []
                if total_min <= avail:
                    remaining = avail - total_min
                    total_content = sum(max_content_len) or 1
                    for ci in range(num_cols):
                        extra = remaining * (max_content_len[ci] / total_content)
                        col_widths_inches.append(min_widths[ci] + extra)
                else:
                    # Proportional fallback with floor
                    total_content = sum(max_content_len) or 1
                    floor_w = max(avail / num_cols * 0.4, 0.85)
                    for ci in range(num_cols):
                        proportion = max_content_len[ci] / total_content
                        col_widths_inches.append(max(avail * proportion, floor_w))

                # CRITICAL: Normalize total to exactly 6.5in so LibreOffice
                # doesn't proportionally downscale all columns
                current_total = sum(col_widths_inches)
                if current_total > 0 and abs(current_total - avail) > 0.01:
                    scale = avail / current_total
                    col_widths_inches = [w * scale for w in col_widths_inches]

                # Apply widths to columns, cells, AND gridCol elements
                tbl_grid = table._tbl.find(qn("w:tblGrid"))
                grid_cols = tbl_grid.findall(qn("w:gridCol")) if tbl_grid is not None else []
                for ci, col in enumerate(table.columns):
                    col_width = Inches(col_widths_inches[ci])
                    col.width = col_width
                    # Update gridCol (authoritative for LibreOffice)
                    if ci < len(grid_cols):
                        grid_cols[ci].set(qn("w:w"), str(int(col_widths_inches[ci] * 1440)))
                    # Update each cell width
                    for row in table.rows:
                        if ci < len(row.cells):
                            row.cells[ci].width = col_width

            # Repeat table headers across pages
            if len(table.rows) > 0:
                tr = table.rows[0]._tr
                tblHeader = OxmlElement("w:tblHeader")
                tr.get_or_add_trPr().append(tblHeader)

            # Clean and merge cell text
            for row in table.rows:
                for cell in row.cells:
                    # Collect ALL text from ALL paragraphs in this cell
                    cell_texts = []
                    for p in cell.paragraphs:
                        para_text = ""
                        for run in p.runs:
                            t = run.text or ""
                            t = re.sub(r"[+|][-=]{3,}[+|]?", "", t)
                            t = re.sub(r"={3,}", "", t)
                            t = t.strip("|").strip()
                            para_text += " " + t if para_text else t
                        para_text = para_text.strip()
                        if para_text:
                            cell_texts.append(para_text)

                    # Join all paragraphs into one continuous text
                    merged = " ".join(cell_texts)
                    # Strip ** bold markers left by Pandoc
                    merged = merged.replace("**", "")
                    # Strip * italic markers
                    merged = merged.replace("*", "")
                    # Fix missing spaces between words (camelCase joins from line merging)
                    # E.g: "deStock" -> "de Stock", "ProductOwner)" -> "Product Owner)"
                    merged = re.sub(r"([a-záéíóúñ])([A-ZÁÉÍÓÚÑ])", r"\1 \2", merged)
                    # Collapse multiple spaces
                    merged = re.sub(r"\s{2,}", " ", merged).strip()

                    # Clear all paragraphs except the first, set merged text
                    if cell.paragraphs:
                        first_p = cell.paragraphs[0]
                        # Detect if header row (first row of table)
                        is_header = row == table.rows[0]

                        # Clear first paragraph
                        first_p.clear()
                        new_run = first_p.add_run(merged)
                        self._apply_font_style(new_run, size=font_size)
                        if is_header:
                            new_run.bold = True

                        # Remove extra paragraphs from the cell
                        for extra_p in list(cell.paragraphs[1:]):
                            extra_p._element.getparent().remove(extra_p._element)

            # FINAL PASS: Force left alignment + single spacing on ALL cell
            # paragraphs. Must happen AFTER the merge/clear cycle above,
            # because .clear() strips paragraph-level formatting.
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        # --- Left alignment (remove old jc, add new) ---
                        p_pr = p._element.get_or_add_pPr()
                        old_jc = p_pr.find(qn("w:jc"))
                        if old_jc is not None:
                            p_pr.remove(old_jc)
                        new_jc = OxmlElement("w:jc")
                        new_jc.set(qn("w:val"), "left")
                        p_pr.append(new_jc)

                        # --- Single line spacing (APA 7 exception for tables) ---
                        old_spacing = p_pr.find(qn("w:spacing"))
                        if old_spacing is not None:
                            p_pr.remove(old_spacing)
                        spacing_el = OxmlElement("w:spacing")
                        spacing_el.set(qn("w:line"), "240")  # single spacing
                        spacing_el.set(qn("w:lineRule"), "auto")
                        spacing_el.set(qn("w:before"), "0")
                        spacing_el.set(qn("w:after"), "40")  # tiny gap between rows
                        p_pr.append(spacing_el)

            # Add spacing paragraph after table (APA 7: double-space gap)
            table_element = table._tbl
            spacing_p = OxmlElement("w:p")
            spacing_pPr = OxmlElement("w:pPr")
            spacing_spacing = OxmlElement("w:spacing")
            spacing_spacing.set(qn("w:line"), "480")
            spacing_spacing.set(qn("w:lineRule"), "auto")
            spacing_pPr.append(spacing_spacing)
            spacing_p.append(spacing_pPr)
            table_element.addnext(spacing_p)

            self._add_table_caption(table, i + 1)

    def _add_table_caption(self, table, num):
        """
        Add 'Tabla N' (Bold) and Table Title (Italic).
        Extracts the title from the preceding paragraph, or from the table's
        own 'Nombre' / 'Atributo de Calidad' row if no preceding title exists.
        """
        parent = table._tbl.getparent()
        if parent is None:
            return

        idx = parent.index(table._tbl)
        table_title = None

        # Try to retrieve the title from the immediately preceding paragraph
        if idx > 0:
            prev_el = parent[idx - 1]
            if prev_el.tag.endswith("p"):
                from docx.text.paragraph import Paragraph

                prev_p = Paragraph(prev_el, self.doc)
                ptxt = prev_p.text.strip()
                if ptxt and not ptxt.isdigit() and "[Título" not in ptxt:
                    table_title = ptxt
                    parent.remove(prev_el)

        # If no title found, try extracting from the table's own content
        if not table_title:
            table_title = self._extract_table_title(table)

        if not table_title:
            table_title = "[Título de la Tabla]"

        # Create paragraph for "Tabla N"
        p_label = OxmlElement("w:p")
        p_label_pr = OxmlElement("w:pPr")
        p_spacing = OxmlElement("w:spacing")
        p_spacing.set(qn("w:after"), "0")  # No space after label
        p_spacing.set(qn("w:line"), "480")  # Double spacing
        p_spacing.set(qn("w:lineRule"), "auto")
        p_label_pr.append(p_spacing)
        p_label.append(p_label_pr)

        run_label = OxmlElement("w:r")
        rPr_label = OxmlElement("w:rPr")
        b_label = OxmlElement("w:b")
        rPr_label.append(b_label)
        run_label.append(rPr_label)
        t_label = OxmlElement("w:t")
        t_label.text = f"Tabla {num}"
        run_label.append(t_label)
        p_label.append(run_label)

        # Create paragraph for Title placeholder
        p_title = OxmlElement("w:p")
        p_title_pr = OxmlElement("w:pPr")
        p_spacing_title = OxmlElement("w:spacing")
        p_spacing_title.set(qn("w:after"), "0")
        p_spacing_title.set(qn("w:line"), "480")  # Double spacing
        p_spacing_title.set(qn("w:lineRule"), "auto")
        p_title_pr.append(p_spacing_title)
        p_title.append(p_title_pr)

        run_title = OxmlElement("w:r")
        rPr_title = OxmlElement("w:rPr")
        i_title = OxmlElement("w:i")
        rPr_title.append(i_title)
        run_title.append(rPr_title)
        t_title = OxmlElement("w:t")
        t_title.text = table_title
        run_title.append(t_title)
        p_title.append(run_title)

        # Insert before table
        parent.insert(parent.index(table._tbl), p_label)
        parent.insert(parent.index(table._tbl), p_title)

    @staticmethod
    def _extract_table_title(table):
        """
        Extract a descriptive title from a grid table's own content.
        Looks for 'Nombre' or 'Atributo de Calidad' rows and combines with the ID.
        Returns None if no suitable data found.
        """
        if len(table.rows) < 2 or len(table.columns) < 2:
            return None

        # Build a dict of label → value from first column pairs
        data = {}
        for row in table.rows[:4]:  # Only check first 4 rows
            cells = row.cells
            if len(cells) >= 2:
                label = cells[0].text.strip().lower()
                value = cells[1].text.strip()
                if label and value:
                    data[label] = value

        # Extract ID and name/attribute
        table_id = data.get("id", "")
        name = data.get("nombre", "") or data.get("atributo de calidad", "")

        if name and table_id:
            return f"{table_id}: {name}"
        if name:
            return name

        return None

    def _format_keywords(self, meta: DocumentMetadata):
        """Format 'Palabras clave', and add repeated title + page break."""
        found_kw = False
        kw_idx = -1

        for i, p in enumerate(self.doc.paragraphs):
            if "palabras clave" in p.text.lower() or "keywords" in p.text.lower():
                found_kw = True
                kw_idx = i

                full = p.text.strip()
                # Regex for "Palabras clave: ..." or "Keywords: ..."
                match = re.search(r"((?:Palabras\s+clave|Keywords):)(.*)", full, re.IGNORECASE)
                if match:
                    label, content = match.groups()
                    p.clear()
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    r1 = p.add_run(label + " ")
                    r1.italic = True
                    self._apply_font_style(r1)
                    r2 = p.add_run(content.strip())
                    self._apply_font_style(r2)
                break

        if found_kw and kw_idx != -1:
            # Logic to insert page break after keywords would go here
            pass

    def _fix_text_spacing_global(self):
        """Run merge_and_clean on all paragraphs."""
        for p in self.doc.paragraphs:
            if not p.style or not p.style.name.startswith("Heading"):
                self._merge_and_clean_paragraph(p)
                # Enforce left align
                if p.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _merge_and_clean_paragraph(self, p):
        """Consolidate runs while preserving inline formatting boundaries.

        Groups consecutive runs by their (bold, italic) attributes and merges
        only within each group, so italic titles in references and the
        "Nota." italic split are kept intact.

        Skips paragraphs containing embedded images (w:drawing) to avoid
        destroying their XML structure.
        """
        if not p.runs:
            return

        # Skip paragraphs with embedded images — clearing runs would destroy them
        if p._element.findall(f".//{qn('w:drawing')}"):
            return

        # Build groups of consecutive runs with the same formatting
        groups: list[tuple[bool, bool, str, str | None, object | None]] = []
        for run in p.runs:
            t = run.text or ""
            t = t.replace("\n", " ").replace("\r", " ")
            # Clean grid-table artifacts
            t = re.sub(r"[+|][-=]{3,}[+|]?", "", t)
            t = re.sub(r"={3,}", "", t)
            t = t.strip("|")

            is_bold = bool(run.bold)
            is_italic = bool(run.italic)
            font_name = run.font.name
            font_size = run.font.size

            if groups and groups[-1][0] == is_bold and groups[-1][1] == is_italic:
                # Same formatting — append text to the current group
                groups[-1] = (
                    is_bold,
                    is_italic,
                    groups[-1][2] + t,
                    font_name or groups[-1][3],
                    font_size or groups[-1][4],
                )
            else:
                groups.append((is_bold, is_italic, t, font_name, font_size))

        # Clear all existing runs
        for run in list(p.runs):
            run._r.getparent().remove(run._r)

        # Re-create one clean run per formatting group, preserving boundary spaces
        for idx, (is_bold, is_italic, text, font_name, font_size) in enumerate(groups):
            # Collapse multiple internal spaces but keep single boundary spaces
            text = re.sub(r"\s{2,}", " ", text)

            # Only strip leading space on the very first group
            if idx == 0:
                text = text.lstrip()
            # Only strip trailing space on the very last group
            if idx == len(groups) - 1:
                text = text.rstrip()

            if not text:
                continue

            new_run = p.add_run(text)
            new_run.font.name = font_name or "Times New Roman"
            new_run.font.size = font_size or Pt(12)
            if is_bold:
                new_run.bold = True
            if is_italic:
                new_run.italic = True
