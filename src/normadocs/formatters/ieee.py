"""
IEEE 8th Edition formatter for academic documents.

Applies IEEE formatting standards to DOCX files.
"""

from typing import Any, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

from ..models import DocumentMetadata
from .base import DocumentFormatter


class IEEEDocxFormatter(DocumentFormatter):
    """
    Applies IEEE 8th Edition formatting to a DOCX file.

    IEEE 8th Edition requirements:
    - Times New Roman 10pt for body text
    - Single spacing throughout
    - 1 inch margins on all sides
    - Page numbers in header, top right
    - Centered, bold headings (Level 1)
    - Tables with full borders
    - Figures: "Fig." caption prefix
    """

    def __init__(self, doc_path: str, config: dict[str, Any] | None = None):
        super().__init__(doc_path, config)
        self.doc: DocumentObject = Document(doc_path)

    def _get_margins(self) -> dict[str, float | str]:
        """Get margins from config with IEEE defaults (1 inch all sides)."""
        margins = self.config.get("margins", {})
        return {
            "top": cast(float, margins.get("top", 1.0)),
            "bottom": cast(float, margins.get("bottom", 1.0)),
            "left": cast(float, margins.get("left", 1.0)),
            "right": cast(float, margins.get("right", 1.0)),
            "unit": cast(str, margins.get("unit", "inches")),
        }

    def _get_font_name(self, key: str = "body") -> str:
        """Get font name from config."""
        return cast(str, self.config.get("fonts", {}).get(key, {}).get("name", "Times New Roman"))

    def _get_font_size(self, key: str = "body") -> int:
        """Get font size from config (IEEE default is 10pt)."""
        return cast(int, self.config.get("fonts", {}).get(key, {}).get("size", 10))

    def _get_spacing_line(self) -> str:
        """Get line spacing from config (IEEE default is single)."""
        return cast(str, self.config.get("spacing", {}).get("line", "single"))

    def process(self, meta: DocumentMetadata) -> None:
        """Run the IEEE formatting pipeline."""
        self._setup_page_layout()
        self._setup_headers()
        self._create_styles()
        self._format_paragraphs()
        self._format_tables()
        self._format_figures()

    def save(self, output_path: str) -> None:
        """Save the formatted document."""
        self.doc.save(str(output_path))

    def _setup_page_layout(self) -> None:
        """Set 1 inch margins on all sides."""
        margins = self._get_margins()
        unit = margins["unit"]

        for section in self.doc.sections:
            top_val = float(margins["top"])
            bottom_val = float(margins["bottom"])
            left_val = float(margins["left"])
            right_val = float(margins["right"])

            if unit == "inches":
                section.top_margin = Inches(top_val)
                section.bottom_margin = Inches(bottom_val)
                section.left_margin = Inches(left_val)
                section.right_margin = Inches(right_val)
            elif unit == "cm":
                section.top_margin = Cm(top_val)
                section.bottom_margin = Cm(bottom_val)
                section.left_margin = Cm(left_val)
                section.right_margin = Cm(right_val)
            else:
                section.top_margin = Emu(int(top_val * 914400 / 1.0))
                section.bottom_margin = Emu(int(bottom_val * 914400 / 1.0))
                section.left_margin = Emu(int(left_val * 914400 / 1.0))
                section.right_margin = Emu(int(right_val * 914400 / 1.0))

    def _setup_headers(self) -> None:
        """Add page numbers in header, top right."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = header_para.add_run()

            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            run._r.append(fldChar1)

            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = "PAGE"
            run._r.append(instrText)

            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar2)

            run.font.name = self._get_font_name()
            run.font.size = Pt(self._get_font_size())

    def _create_styles(self) -> None:
        """Apply IEEE text styles: Times New Roman 10pt, single spacing."""
        styles = self.doc.styles
        body_font = self._get_font_name()
        body_size = self._get_font_size()
        line_spacing = WD_LINE_SPACING.SINGLE

        normal = styles["Normal"]
        normal.font.name = body_font
        normal.font.size = Pt(body_size)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        normal.paragraph_format.line_spacing_rule = line_spacing
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)

        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.name = body_font
            h1.font.size = Pt(body_size)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(12)
            h1.paragraph_format.line_spacing_rule = line_spacing

        if "Heading 2" in styles:
            h2 = styles["Heading 2"]
            h2.font.name = body_font
            h2.font.size = Pt(body_size)
            h2.font.bold = True
            h2.font.italic = True
            h2.paragraph_format.space_before = Pt(12)
            h2.paragraph_format.space_after = Pt(6)
            h2.paragraph_format.line_spacing_rule = line_spacing

    def _format_paragraphs(self) -> None:
        """Format paragraphs: justify text, single spacing."""
        body_font = self._get_font_name()
        body_size = self._get_font_size()

        for p in self.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue

            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            for run in p.runs:
                run.font.name = body_font
                run.font.size = Pt(body_size)

    def _format_tables(self) -> None:
        """Format tables with full borders."""
        for table in self.doc.tables:
            if "Table Grid" in self.doc.styles:
                table.style = "Table Grid"

            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _format_figures(self) -> None:
        """Format figure captions: centered, italic."""
        figure_config = self.config.get("figures", {})
        caption_prefix = figure_config.get("caption_prefix", "Fig")

        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text.startswith(caption_prefix) and "." in text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = self._get_font_name()
                    run.font.size = Pt(self._get_font_size())
                    run.italic = True
