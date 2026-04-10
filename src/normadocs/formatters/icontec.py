from typing import Any, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from ..models import DocumentMetadata
from .base import DocumentFormatter


class IcontecFormatter(DocumentFormatter):
    """
    Applies ICONTEC (NTC 1486) formatting to a DOCX file.
    """

    def __init__(self, doc_path: str, config: dict[str, Any] | None = None):
        super().__init__(doc_path, config)
        self.doc: DocumentObject = Document(doc_path)

    def _get_margins(self) -> dict[str, float]:
        """Get margins from config with defaults."""
        margins = self.config.get("margins", {})
        return {
            "top": margins.get("top", 3.0),
            "bottom": margins.get("bottom", 2.0),
            "left": margins.get("left", 3.0),
            "right": margins.get("right", 2.0),
            "unit": margins.get("unit", "cm"),
        }

    def _get_font_name(self, key: str = "body") -> str:
        """Get font name from config."""
        result: str = cast(str, self.config.get("fonts", {}).get(key, {}).get("name", "Arial"))
        return result

    def _get_font_size(self, key: str = "body") -> int:
        """Get font size from config."""
        result: int = cast(int, self.config.get("fonts", {}).get(key, {}).get("size", 12))
        return result

    def _get_spacing_line(self) -> float:
        """Get line spacing from config with default."""
        result: float = self.config.get("spacing", {}).get("line", 1.5)
        return result

    def _margin_to_unit(self, value: float, unit: str) -> Cm | Inches:
        """Convert margin value based on unit."""
        if unit == "inches":
            return Inches(value)
        return Cm(value)

    def process(self, meta: DocumentMetadata) -> None:
        """Run the ICONTEC formatting pipeline."""
        self._setup_page_layout()
        self._create_styles()
        self._add_cover_page(meta)
        self._process_paragraphs()
        # Tables and Citations logic can be added later/adapted
        # For now, we focus on layout and text style.

    def save(self, output_path: str) -> None:
        self.doc.save(str(output_path))

    def _setup_page_layout(self) -> None:
        """
        NTC 1486 Margins from config.
        """
        margins = self._get_margins()
        unit = str(margins["unit"])

        for section in self.doc.sections:
            section.top_margin = self._margin_to_unit(float(margins["top"]), unit)
            section.bottom_margin = self._margin_to_unit(float(margins["bottom"]), unit)
            section.left_margin = self._margin_to_unit(float(margins["left"]), unit)
            section.right_margin = self._margin_to_unit(float(margins["right"]), unit)

    def _create_styles(self) -> None:
        """
        Font and spacing from config.
        """
        styles = self.doc.styles
        body_font = self._get_font_name("body")
        body_size = self._get_font_size("body")
        spacing_line = self._get_spacing_line()
        line_spacing = (
            WD_LINE_SPACING.ONE_POINT_FIVE if spacing_line == 1.5 else WD_LINE_SPACING.SINGLE
        )

        # Normal
        normal = styles["Normal"]
        font = normal.font
        font.name = body_font
        font.size = Pt(body_size)
        font.color.rgb = RGBColor(0, 0, 0)

        pf = normal.paragraph_format
        pf.line_spacing_rule = line_spacing
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # Headings
        # Heading 1: Centered, Bold, Uppercase (handled in text)
        body_font = self._get_font_name("body")
        body_size = self._get_font_size("body")
        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.name = body_font
            h1.font.size = Pt(body_size)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)  # Space before title
            h1.paragraph_format.space_after = Pt(12)

    def _add_cover_page(self, meta: DocumentMetadata) -> None:
        """
        ICONTEC Cover Page:
        - Title (Centered, Vertical align top approx)
        - Author (Centered, vertical align middle)
        - Legend (Institution, Faculty, etc. at bottom)
        - City, Year (Bottom)
        """
        # Simplistic implementation: Insert at top
        self.doc.add_paragraph()
        if self.doc.paragraphs:
            self.doc.paragraphs[0].insert_paragraph_before("")

        ref_p = self.doc.paragraphs[0]

        # Helper to add centered line
        def add_line(text: str, bold: bool = False, space_after: int = 0) -> Paragraph:
            p = ref_p.insert_paragraph_before(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = self.doc.styles["Normal"]
            if bold and p.runs:
                p.runs[0].bold = True
            p.paragraph_format.space_after = Pt(space_after)
            return p

        # Title formatting
        add_line(meta.title.upper(), bold=True, space_after=120)  # Approx space

        # Author
        add_line(meta.author or "Author Name", space_after=200)

        # Legend/Footer block
        # Institution
        institution = meta.institution or "Institution Name"
        add_line(institution.upper())

        # Program/Faculty
        if meta.program:
            add_line(meta.program.upper())

        # City, Year
        city = meta.extra.get("city", "City")
        year = meta.date or "2024"
        add_line(f"{city.upper()}, {year}")

        # Page Break
        pb_p = ref_p.insert_paragraph_before()
        pb_p.add_run().add_break()

    def _process_paragraphs(self) -> None:
        """Justify text and ensure config font."""
        body_font = self._get_font_name("body")
        body_size = self._get_font_size("body")
        for p in self.doc.paragraphs:
            # Skip if cover page created just now...
            # Ideally we skip based on section or style
            if p.style and p.style.name.startswith("Heading"):
                continue

            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)  # No indent in ICONTEC usually

            # Ensure font
            for run in p.runs:
                run.font.name = body_font
                run.font.size = Pt(body_size)
