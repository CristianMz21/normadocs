"""APA page layout and page number handling."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, cast

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from ...config import DEFAULT_BODY_FONT, W_TYPE, W_VAL
from ...utils.docx_helpers import paragraph_style_name

if TYPE_CHECKING:
    from docx.document import Document as DocType
    from docx.section import Section as SectionType


_FLD_CHAR = "w:fldChar"
_FLD_CHAR_TYPE = "w:fldCharType"
_W_VAL = W_VAL
_W_TYPE = W_TYPE


class APAPageHandler:
    """Handles page layout and page numbers per APA 7th Edition.

    Sets up page margins, headers with page numbers, and section
    page breaks according to APA 7th Edition requirements.

    Args:
        doc: The python-docx Document object.
        config: Optional configuration dictionary.
    """

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        """Initialize APAPageHandler.

        Args:
            doc: The python-docx Document object.
            config: Optional configuration dictionary.
        """
        self.doc = doc
        self.config = config if config is not None else {}

    def get_config(self, *keys: str, default: Any = None) -> Any:
        """Get nested config value via dot-notation keys."""
        value: Any = self.config
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key)
            else:
                return default
            if value is None:
                return default
        return value

    def _get_margins(self) -> dict[str, float]:
        """Get margins from config with defaults."""
        margins = cast(dict[str, Any], self.get_config("margins", default={}))
        return {
            "top": margins.get("top", 1.0),
            "bottom": margins.get("bottom", 1.0),
            "left": margins.get("left", 1.0),
            "right": margins.get("right", 1.0),
            "unit": margins.get("unit", "inches"),
        }

    def _margin_to_inches(self, value: float, unit: str) -> Inches | Cm:
        """Convert margin value to inches based on unit.

        Args:
            value: The margin value.
            unit: The unit type ("inches", "cm", or "Emu").

        Returns:
            The value converted to inches.
        """
        if unit == "cm":
            return Cm(value)
        return Inches(value)

    def setup_page_layout(self) -> None:
        """Set margins and add page numbers top-right."""
        margins = self._get_margins()
        unit = str(margins["unit"])

        for section in self.doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = self._margin_to_inches(margins["left"], unit)
            section.right_margin = self._margin_to_inches(margins["right"], unit)
            section.top_margin = self._margin_to_inches(margins["top"], unit)
            section.bottom_margin = self._margin_to_inches(margins["bottom"], unit)

            # Enable separate first-page header/footer so the cover can show
            # only its page number while later pages may use a running head.
            section.different_first_page_header_footer = True

            # Clear footers completely to prevent Pandoc page numbers at bottom
            footer = section.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.clear()
                # Remove ALL child elements (runs, field codes, pPr, etc.)
                del p._element[:]
            # Also set footer distance to zero to suppress any residual space
            sect_pr = section._sectPr
            existing_pg_mar = sect_pr.find(qn("w:pgMar"))
            if existing_pg_mar is not None:
                existing_pg_mar.set(qn("w:footer"), "0")

            # Remove the footerReference entirely to prevent LibreOffice
            # from rendering any footer content
            for fref in sect_pr.findall(qn("w:footerReference")):
                sect_pr.remove(fref)

            # Student APA papers show the page number on the cover too.
            self._add_page_number(section)
            self._add_page_number(section, section.first_page_header)

    def _add_page_number(self, section: SectionType, header: Any | None = None) -> None:
        """Add only the page number to a selected header, default or first-page."""
        if header is None:
            header = section.header
        header.is_linked_to_previous = False

        # Clear existing header text
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Build PAGE field with begin/separate/end sequence
        # (LibreOffice requires the 'separate' marker to render page numbers)
        run_begin = hp.add_run()
        fld_begin = OxmlElement(_FLD_CHAR)
        fld_begin.set(qn(_FLD_CHAR_TYPE), "begin")
        run_begin._r.append(fld_begin)
        run_begin.font.name = DEFAULT_BODY_FONT
        run_begin.font.size = Pt(12)

        run_instr = hp.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)
        run_instr.font.name = DEFAULT_BODY_FONT
        run_instr.font.size = Pt(12)

        run_sep = hp.add_run()
        fld_sep = OxmlElement(_FLD_CHAR)
        fld_sep.set(qn(_FLD_CHAR_TYPE), "separate")
        run_sep._r.append(fld_sep)
        run_sep.font.name = DEFAULT_BODY_FONT
        run_sep.font.size = Pt(12)

        # Placeholder text (will be replaced by actual page number)
        run_num = hp.add_run("1")
        run_num.font.name = DEFAULT_BODY_FONT
        run_num.font.size = Pt(12)

        run_end = hp.add_run()
        fld_end = OxmlElement(_FLD_CHAR)
        fld_end.set(qn(_FLD_CHAR_TYPE), "end")
        run_end._r.append(fld_end)
        run_end.font.name = DEFAULT_BODY_FONT
        run_end.font.size = Pt(12)

        # Strip excessive paragraphs in header
        while len(header.paragraphs) > 1:
            p_elem = header.paragraphs[-1]._element
            p_elem.getparent().remove(p_elem)

    @staticmethod
    def _has_page_break_before(paragraph: Any) -> bool:
        """Return whether a paragraph already starts on a new page."""
        if bool(paragraph.paragraph_format.page_break_before):
            return True
        previous = paragraph._element.getprevious()
        while previous is not None:
            if any(br.get(qn("w:type")) == "page" for br in previous.iter(qn("w:br"))):
                return True
            if previous.tag == qn("w:p") and not "".join(previous.itertext()).strip():
                previous = previous.getprevious()
                continue
            return False
        return False

    def add_section_page_breaks(self) -> None:
        """Add section breaks only when one is not already before the heading.

        Markdown preprocessing already inserts explicit page breaks before
        level-1 headings. Repeating those breaks here creates blank pages in
        LibreOffice. Direct callers that provide an unprocessed DOCX retain
        the legacy fallback for Conclusions/References.
        """
        # Sections that need to start on a new page
        new_page_sections = [
            "Conclusiones",
            "Referencias",
            "References",
            "Appendix A",
            "Appendices",
        ]

        # Find headings and add page breaks before them
        for _i, p in enumerate(self.doc.paragraphs):
            if paragraph_style_name(p).startswith("Heading"):
                # Check if this heading is one of our target sections
                heading_text = p.text.strip()
                for section in new_page_sections:
                    if heading_text.lower() == section.lower():
                        if self._has_page_break_before(p):
                            break
                        # Create a valid WordprocessingML paragraph containing
                        # the page-break run. A standalone w:br sibling is not
                        # valid WordprocessingML and may be dropped by converters.
                        break_paragraph = OxmlElement("w:p")
                        break_run = OxmlElement("w:r")
                        br = OxmlElement("w:br")
                        br.set(qn("w:type"), "page")
                        break_run.append(br)
                        break_paragraph.append(break_run)
                        p._element.addprevious(break_paragraph)
                        break

    def setup_running_head(self, short_title: str | None = None) -> None:
        """Add running head to headers on pages after the cover page.

        The running head consists of:
        - Left side: Short title in ALL CAPS (Times New Roman 12pt)
        - Right side: Page number

        APA 7th Edition requires the running head only on pages after
        the cover page (not on the cover page itself).

        This is implemented using different_first_page_header_footer = True:
        - first_page_header (cover page): page number only
        - header (pages 2+): running head with short title and page number

        Args:
            short_title: The short title to display. If None or empty,
                        the running head is not added.
        """
        display_title = self._resolve_display_title(short_title)
        if display_title is None:
            return
        for section in self.doc.sections:
            self._apply_running_head_to_section(section, display_title)

    def _resolve_display_title(self, short_title: str | None) -> str | None:
        """Resolve and truncate the running head title, or None if disabled."""
        if not short_title:
            return None
        enabled = cast(bool, self.get_config("running_head", "enabled", default=True))
        if not enabled:
            return None
        max_length = int(cast(int, self.get_config("running_head", "max_length", default=50)))
        display_title = short_title.upper()
        if len(display_title) > max_length:
            display_title = display_title[:max_length]
        return display_title

    def _apply_running_head_to_section(self, section: SectionType, display_title: str) -> None:
        """Apply running head configuration to a single section."""
        self._reset_first_page_header(section)
        self._setup_header_with_running_head(section.header, display_title)

    def _reset_first_page_header(self, section: SectionType) -> None:
        """Clear first-page header and restore page number only."""
        first_page_header = section.first_page_header
        first_page_header.is_linked_to_previous = False
        for p in first_page_header.paragraphs:
            p.clear()
            del p._element[:]
        self._add_page_number(section, first_page_header)

    def _setup_header_with_running_head(self, header: Any, display_title: str) -> None:
        """Configure default header with running head and page number."""
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        self._configure_header_tabs(hp)
        self._add_running_head_content(hp, display_title)
        self._strip_excess_header_paragraphs(header)

    def _configure_header_tabs(self, hp: Any) -> None:
        """Set center and right tab stops on a header paragraph."""
        p_pr = hp._p.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            hp._p.insert(0, p_pr)
        tabs = OxmlElement("w:tabs")
        tab1 = OxmlElement("w:tab")
        tab1.set(qn(_W_VAL), "center")
        tab1.set(qn("w:pos"), "4680")
        tabs.append(tab1)
        tab2 = OxmlElement("w:tab")
        tab2.set(qn(_W_VAL), "right")
        tab2.set(qn("w:pos"), "9360")
        tabs.append(tab2)
        existing_tabs = p_pr.find(qn("w:tabs"))
        if existing_tabs is not None:
            p_pr.remove(existing_tabs)
        p_pr.append(tabs)

    def _add_running_head_content(self, hp: Any, display_title: str) -> None:
        """Add title, tabs, and PAGE field to header paragraph."""
        title_run = hp.add_run(display_title)
        title_run.font.name = DEFAULT_BODY_FONT
        title_run.font.size = Pt(12)
        hp.add_run("\t")
        hp.add_run("\t")
        self._add_page_field_runs(hp)

    def _add_page_field_runs(self, hp: Any) -> None:
        """Build PAGE field with begin/separate/end sequence."""
        run_begin = hp.add_run()
        fld_begin = OxmlElement(_FLD_CHAR)
        fld_begin.set(qn(_FLD_CHAR_TYPE), "begin")
        run_begin._r.append(fld_begin)
        run_begin.font.name = DEFAULT_BODY_FONT
        run_begin.font.size = Pt(12)
        run_instr = hp.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)
        run_instr.font.name = DEFAULT_BODY_FONT
        run_instr.font.size = Pt(12)
        run_sep = hp.add_run()
        fld_sep = OxmlElement(_FLD_CHAR)
        fld_sep.set(qn(_FLD_CHAR_TYPE), "separate")
        run_sep._r.append(fld_sep)
        run_sep.font.name = DEFAULT_BODY_FONT
        run_sep.font.size = Pt(12)
        run_num = hp.add_run("1")
        run_num.font.name = DEFAULT_BODY_FONT
        run_num.font.size = Pt(12)
        run_end = hp.add_run()
        fld_end = OxmlElement(_FLD_CHAR)
        fld_end.set(qn(_FLD_CHAR_TYPE), "end")
        run_end._r.append(fld_end)
        run_end.font.name = DEFAULT_BODY_FONT
        run_end.font.size = Pt(12)

    def _strip_excess_header_paragraphs(self, header: Any) -> None:
        """Remove extra paragraphs from a header."""
        while len(header.paragraphs) > 1:
            p_elem = header.paragraphs[-1]._element
            p_elem.getparent().remove(p_elem)
