"""APA cover page handling."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, cast

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from ...config import HEADING_1_STYLE, NORMAL_STYLE
from ...models import DocumentMetadata
from ...utils.docx_helpers import paragraph_style, paragraph_style_name

if TYPE_CHECKING:
    from docx.document import Document as DocType


class APACoverHandler:
    """Handles creation of APA 7th Edition cover page."""

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        """Initialize APACoverHandler.

        Args:
            doc: The python-docx Document object.
            config: Optional configuration dictionary.
        """
        self.doc = doc
        self.config = config if config is not None else {}

    def get_config(self, *keys: str, default: Any = None) -> Any:
        """Get nested config value via dot-notation keys."""
        value: Any = self.config
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key)
            else:
                return default
            if value is None:
                return default
        return value

    def _get_cover_config(self) -> dict[str, Any]:
        """Get cover configuration from config with defaults."""
        default: dict[str, Any] = {"title_align": "center", "author_align": "center"}
        return cast(dict[str, Any], self.get_config("cover", default=default))

    def _get_spacing_line(self) -> str:
        """Get line spacing from config with default."""
        return cast(str, self.get_config("spacing", "line", default="double"))

    def add_cover_page(self, meta: DocumentMetadata) -> None:
        """
        Insert a cover page at the beginning of the document.

        APA 7 Cover Page format:
        - Title (bold, centered) in upper half of page
        - Subtitle (centered) below title
        - Author name (centered)
        - Institutional affiliation (centered) - company/department
        - Date (centered)
        """
        self._prepare_initial_paragraph()
        content_lines = self._build_content_lines(meta)
        elements = self._build_elements_with_spacers(content_lines)
        ref_p = self.doc.paragraphs[0]
        self._insert_cover_elements(ref_p, elements)
        self._cleanup_cover_ref(ref_p)
        self._ensure_cover_title(meta)
        self._add_page_break_after_cover()

    def _prepare_initial_paragraph(self) -> None:
        """Ensure a leading empty paragraph exists for insertion anchor."""
        self.doc.add_paragraph()
        if self.doc.paragraphs:
            self.doc.paragraphs[0].insert_paragraph_before("")

    def _build_content_lines(self, meta: DocumentMetadata) -> list[tuple[str, bool]]:
        """Build the ordered cover content lines."""
        lines: list[tuple[str, bool]] = [(meta.title, True)]
        self._append_subtitle(lines, meta)
        lines.append(("", False))
        lines.append((meta.author or "", False))
        self._append_affiliation(lines, meta)
        self._append_institution(lines, meta)
        self._append_program_ficha(lines, meta)
        self._append_subject_instructor(lines, meta)
        self._append_location(lines, meta)
        self._append_date(lines, meta)
        return lines

    def _append_subtitle(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append subtitle with preceding blank line when present."""
        subtitle = getattr(meta, "subtitle", None) or meta.extra.get("subtitle", "")
        if subtitle:
            lines.append(("", False))
            lines.append((subtitle, False))

    def _append_affiliation(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append affiliation combining center when present."""
        center = getattr(meta, "center", None) or ""
        affiliation = getattr(meta, "affiliation", None) or ""
        if center and affiliation:
            affiliation = f"{affiliation}\n{center}"
        elif center:
            affiliation = center
        if affiliation:
            lines.append((affiliation, False))

    def _append_institution(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append institution when different from affiliation."""
        # affiliation already resolved; recompute for comparison without duplication
        center = getattr(meta, "center", None) or ""
        affiliation = getattr(meta, "affiliation", None) or ""
        if center and affiliation:
            affiliation = f"{affiliation}\n{center}"
        elif center:
            affiliation = center
        institution = getattr(meta, "institution", None) or ""
        if institution and institution != affiliation:
            lines.append((institution, False))

    def _append_program_ficha(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append program and ficha lines for SENA format."""
        program = getattr(meta, "program", None) or ""
        ficha = getattr(meta, "ficha", None) or ""
        if program:
            lines.append(("", False))
            lines.append((program, False))
        if ficha:
            lines.append((ficha, False))

    def _append_subject_instructor(
        self, lines: list[tuple[str, bool]], meta: DocumentMetadata
    ) -> None:
        """Append subject and instructor lines."""
        subject = meta.subject or ""
        instructor = meta.instructor or ""
        if subject:
            lines.append((subject, False))
        if instructor:
            lines.append((f"Instructor: {instructor}", False))

    def _append_location(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append location (city/country) when present."""
        location = meta.location or ""
        if location:
            lines.append((location, False))

    def _append_date(self, lines: list[tuple[str, bool]], meta: DocumentMetadata) -> None:
        """Append date with preceding blank line."""
        date = meta.date or ""
        if date:
            lines.append(("", False))
            lines.append((date, False))

    def _build_elements_with_spacers(
        self, content_lines: list[tuple[str, bool]]
    ) -> list[tuple[str, bool]]:
        """Prefix spacers to position title in upper third."""
        n_spacers = 6
        elements: list[tuple[str, bool]] = [("", False)] * n_spacers
        elements.extend(content_lines)
        return elements

    def _insert_cover_elements(self, ref_p: Any, elements: list[tuple[str, bool]]) -> None:
        """Insert cover elements before the reference paragraph."""
        spacing_line = self._get_spacing_line()
        for text, is_bold in elements:
            p = ref_p.insert_paragraph_before(text)
            self._configure_cover_paragraph(p, text, is_bold, spacing_line)

    def _configure_cover_paragraph(
        self, p: Any, text: str, is_bold: bool, spacing_line: str
    ) -> None:
        """Apply APA cover paragraph formatting."""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = paragraph_style(self.doc.styles, NORMAL_STYLE)
        self._apply_spacing_rule(p, spacing_line)
        if is_bold and text.strip() and p.runs:
            p.runs[0].bold = True

    def _apply_spacing_rule(self, p: Any, spacing_line: str) -> None:
        """Set line spacing rule based on config."""
        if spacing_line == "double":
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif spacing_line == "1.5":
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _cleanup_cover_ref(self, ref_p: Any) -> None:
        """Remove leftover reference and trailing empty paragraphs."""
        ref_p._element.getparent().remove(ref_p._element)
        last_p = self.doc.paragraphs[-1]
        if not last_p.text.strip():
            last_p._element.getparent().remove(last_p._element)

    def _ensure_cover_title(self, meta: DocumentMetadata) -> None:
        """Repeat title as centered bold heading on first text page when needed."""
        first_heading = self._find_first_heading()
        if first_heading is None:
            return
        if self._has_title_heading(meta):
            return
        title_heading = first_heading.insert_paragraph_before(meta.title)
        title_heading.style = first_heading.style
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_heading.paragraph_format.page_break_before = True
        for run in title_heading.runs:
            run.bold = True

    def _find_first_heading(self) -> Any | None:
        """Return the first Heading 1 paragraph with text, if any."""
        return next(
            (
                p
                for p in self.doc.paragraphs
                if paragraph_style_name(p) == HEADING_1_STYLE and p.text.strip()
            ),
            None,
        )

    def _has_title_heading(self, meta: DocumentMetadata) -> bool:
        """Check if a title heading already exists."""
        return any(
            paragraph_style_name(p) == HEADING_1_STYLE
            and p.text.strip().casefold() == meta.title.strip().casefold()
            for p in self.doc.paragraphs
        )

    def _add_page_break_after_cover(self) -> None:
        """Add page break after cover before first non-center heading/text."""
        for p in self.doc.paragraphs:
            style_name = paragraph_style_name(p)
            if style_name.startswith("Heading"):
                p.paragraph_format.page_break_before = True
                break
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER and p.text.strip():
                p.paragraph_format.page_break_before = True
                break
