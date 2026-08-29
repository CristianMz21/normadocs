"""APA figure formatting and captions."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, cast

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml.etree import Element

from ...config import DEFAULT_BODY_FONT, W_VAL
from ...utils.docx_helpers import paragraph_style_name

if TYPE_CHECKING:
    from docx.document import Document as DocType
    from docx.text.run import Run as RunType


_W_RPR = "w:rPr"
_W_VAL = W_VAL
_XML_SPACE = "xml:space"
_W_P = "w:p"
_W_DRAWING_QN = "w:drawing"
_DESCR = "descr"
_NAME = "name"
_CX = "cx"
_CY = "cy"
_W_T = "w:t"
_W_R = "w:r"
_W_JC = "w:jc"
_W_SPACING = "w:spacing"
_W_LINE = "w:line"
_W_LINE_RULE = "w:lineRule"
_W_AFTER = "w:after"
_W_PR = "w:pPr"
_NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PRESERVE = "preserve"


def _extract_prefix(text: str) -> tuple[str, str] | None:
    stripped = text.strip()
    if stripped.startswith("Figura") and len(stripped) > 6 and stripped[6].isspace():
        return ("Figura", stripped[6:].lstrip())
    if stripped.startswith("Figure") and len(stripped) > 6 and stripped[6].isspace():
        return ("Figure", stripped[6:].lstrip())
    return None


def _extract_number(rest: str) -> tuple[str, int] | None:
    i = 0
    num = ""
    while i < len(rest) and rest[i].isdigit():
        num += rest[i]
        i += 1
    if not num:
        return None
    return (num, i)


def _skip_punct_and_spaces(rest: str, idx: int) -> int:
    while idx < len(rest) and rest[idx].isspace():
        idx += 1
    if idx < len(rest) and rest[idx] in ".:":
        idx += 1
    while idx < len(rest) and rest[idx].isspace():
        idx += 1
    return idx


def _parse_caption(text: str) -> tuple[str, str, str] | None:
    """Manual linear parse for 'Figura N. Title' (replaces S8786 _CAPTION_RE)."""
    pref = _extract_prefix(text)
    if pref is None:
        return None
    prefix, rest = pref
    num_res = _extract_number(rest)
    if num_res is None:
        return None
    num, idx = num_res
    idx = _skip_punct_and_spaces(rest, idx)
    title = rest[idx:] if idx < len(rest) else ""
    return (prefix, num, title)


class _CaptionMatch:
    """Mimics re.Match for caption parsing."""

    def __init__(self, prefix: str, num: str, title: str) -> None:
        self._groups = (prefix, num, title)

    def group(self, n: int) -> str:
        return self._groups[n - 1]


class _CaptionPattern:
    """Minimal replacement for compiled _CAPTION_RE (S8786 linear)."""

    def match(self, text: str) -> _CaptionMatch | None:
        parsed = _parse_caption(text)
        if parsed is None:
            return None
        return _CaptionMatch(parsed[0], parsed[1], parsed[2])


_CAPTION_RE = _CaptionPattern()


class APAFiguresHandler:
    """Handles figure formatting and captions per APA 7th Edition."""

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        self.doc = doc
        self.config = config if config is not None else {}

    def get_config(self, *keys: str, default: Any = None) -> Any:
        """Get nested config value via dot-notation keys."""
        value: Any = self.config
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key)
            else:
                return default
            if value is None:
                return default
        return value

    def _get_figure_config(self) -> dict[str, Any]:
        """Get figure configuration from config with defaults."""
        default_config: dict[str, Any] = {
            "caption_prefix": "Figure",
            "title_above": True,
            "nota_prefix": "Nota.",
        }
        return cast(dict[str, Any], self.get_config("figures", default=default_config))

    def _get_body_font(self) -> str:
        """Get body font name from config."""
        return cast(str, self.get_config("fonts", "body", "name", default=DEFAULT_BODY_FONT))

    def _apply_font_style(self, run: RunType, bold: bool = False, italic: bool = False) -> None:
        """Apply font style to a run (helper for this handler)."""
        from .apa_styles import APAStylesHandler

        handler = APAStylesHandler(self.doc)
        handler._apply_font_style(run, bold=bold, italic=italic)

    def _make_figure_paragraph(
        self, text: str, bold: bool = False, italic: bool = False, space_after: str = "0"
    ) -> Element:
        """Helper: create a Times New Roman 12pt paragraph for figure captions."""
        p_el = OxmlElement(_W_P)
        p_pr = OxmlElement(_W_PR)
        p_sp = OxmlElement(_W_SPACING)
        p_sp.set(qn(_W_AFTER), space_after)
        p_sp.set(qn(_W_LINE), "480")
        p_sp.set(qn(_W_LINE_RULE), "auto")
        p_pr.append(p_sp)
        p_el.append(p_pr)

        run = OxmlElement(_W_R)
        r_pr = OxmlElement(_W_RPR)
        if bold:
            r_pr.append(OxmlElement("w:b"))
        if italic:
            r_pr.append(OxmlElement("w:i"))
        rn = OxmlElement("w:rFonts")
        rn.set(qn("w:ascii"), DEFAULT_BODY_FONT)
        rn.set(qn("w:hAnsi"), DEFAULT_BODY_FONT)
        r_pr.append(rn)
        sz = OxmlElement("w:sz")
        sz.set(qn(_W_VAL), "24")
        r_pr.append(sz)
        run.append(r_pr)
        t = OxmlElement(_W_T)
        t.set(qn(_XML_SPACE), _PRESERVE)
        t.text = text
        run.append(t)
        p_el.append(run)
        return p_el

    def format_figures(self) -> None:
        """Add APA 7 figure captions: Label + Title ABOVE, Nota BELOW."""
        image_paragraphs = self._collect_image_paragraphs()
        for p, _ in image_paragraphs:
            self._center_image_paragraph(p)
            self._scale_image_drawings(p)

    def _collect_image_paragraphs(self) -> list[tuple[Any, str]]:
        """Collect paragraphs containing images and their alt text."""
        result: list[tuple[Any, str]] = []
        for p in self.doc.paragraphs:
            drawings = p._element.findall(f".//{qn(_W_DRAWING_QN)}")
            if not drawings:
                continue
            alt = self._extract_alt_text_for_drawings(drawings, _NS_WP)
            result.append((p, alt.strip()))
        return result

    def _extract_alt_text_for_drawings(self, drawings: list[Any], ns_wp: str) -> str:
        """Extract alt text from drawings."""
        for drawing in drawings:
            for doc_pr in drawing.iter(f"{{{ns_wp}}}docPr"):
                return str((doc_pr.get(_DESCR, "") or doc_pr.get(_NAME, "")).strip())
            for c_nv_pr in drawing.iter(f"{{{_NS_PIC}}}cNvPr"):
                return str((c_nv_pr.get(_DESCR, "") or c_nv_pr.get(_NAME, "")).strip())
        return ""

    def _center_image_paragraph(self, p: Any) -> None:
        """Center image paragraph and clear spacing."""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    def _scale_image_drawings(self, p: Any) -> None:
        """Scale oversized drawings to fit page."""
        max_w = Inches(6.5)
        max_h = Inches(8.5)
        drawings = p._element.findall(f".//{{{_NS_WP}}}inline") + p._element.findall(
            f".//{{{_NS_WP}}}anchor"
        )
        for d in drawings:
            self._scale_single_drawing(d, max_w, max_h, _NS_WP)

    def _scale_single_drawing(self, drawing: Any, max_w: Any, max_h: Any, ns_wp: str) -> None:
        """Scale one drawing element if oversized."""
        extent = drawing.find(f"{{{ns_wp}}}extent")
        if extent is None:
            return
        cx = int(extent.get(_CX, 0))
        cy = int(extent.get(_CY, 0))
        if cx == 0 or cy == 0:
            return
        scale = self._compute_scale(cx, cy, max_w, max_h)
        if scale >= 1.0:
            return
        self._apply_scale(extent, drawing, scale, cx, cy)

    def _compute_scale(self, cx: int, cy: int, max_w: Any, max_h: Any) -> float:
        """Compute scale factor for drawing."""
        scale = 1.0
        if cx > max_w:
            scale = min(scale, max_w / cx)
        if cy > max_h:
            scale = min(scale, max_h / cy)
        return scale

    def _apply_scale(self, extent: Any, drawing: Any, scale: float, cx: int, cy: int) -> None:
        """Apply scale to extent and shape extents."""
        new_cx = int(cx * scale)
        new_cy = int(cy * scale)
        extent.set(_CX, str(new_cx))
        extent.set(_CY, str(new_cy))
        for ext_el in drawing.iter(f"{{{_NS_A}}}ext"):
            self._scale_ext(ext_el, scale)

    def _scale_ext(self, ext_el: Any, scale: float) -> None:
        """Scale a single a:ext element."""
        old_cx = int(ext_el.get(_CX, 0))
        old_cy = int(ext_el.get(_CY, 0))
        if old_cx > 0:
            ext_el.set(_CX, str(int(old_cx * scale)))
        if old_cy > 0:
            ext_el.set(_CY, str(int(old_cy * scale)))

    def add_figure_captions(self) -> None:
        """Ensure every figure has an APA 7 caption: 'Figura N.' bold + title italic."""
        image_paragraphs = [
            p for p in self.doc.paragraphs if p._element.findall(f".//{qn(_W_DRAWING_QN)}")
        ]
        self._move_existing_captions(image_paragraphs, _CAPTION_RE)
        max_used = self._find_max_caption_number(_CAPTION_RE)
        self._insert_missing_captions(image_paragraphs, _CAPTION_RE, max_used)
        self._normalize_caption_runs(_CAPTION_RE)

    def _move_existing_captions(self, images: list[Any], caption_re: Any) -> None:
        """Move captions sitting directly below image to above it."""
        from docx.text.paragraph import Paragraph

        for img_p in images:
            nxt = img_p._element.getnext()
            if nxt is None or nxt.tag != qn(_W_P):
                continue
            next_p = Paragraph(nxt, img_p._parent)
            if caption_re.match(next_p.text.strip()):
                img_p._element.addprevious(nxt)

    def _find_max_caption_number(self, caption_re: Any) -> int:
        """Find highest existing figure number."""
        max_used = 0
        for p in self.doc.paragraphs:
            match = caption_re.match(p.text.strip())
            if match:
                max_used = max(max_used, int(match.group(2)))
        return max_used

    def _insert_missing_captions(self, images: list[Any], caption_re: Any, max_used: int) -> None:
        """Insert missing captions for images without adjacent captions."""
        current_max = max_used
        for img_p in images:
            if self._has_adjacent_caption(img_p, caption_re):
                continue
            if self._has_manual_title(img_p):
                continue
            current_max += 1
            alt_text = self._extract_alt_text(img_p)
            caption_el = self._build_caption_element(current_max, alt_text)
            img_p._element.addprevious(caption_el)

    def _normalize_caption_runs(self, caption_re: Any) -> None:
        """Normalize caption paragraphs into bold label + italic title."""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            m = caption_re.match(text)
            if not m:
                continue
            label, num, title = m.group(1), m.group(2), m.group(3).strip()
            runs = para.runs
            if self._is_caption_formatted(runs, title):
                continue
            self._reformat_caption(para, label, num, title)

    def _is_caption_formatted(self, runs: Any, title: str) -> bool:
        """Check if caption already has correct formatting."""
        label_ok = bool(runs) and bool(runs[0].bold)
        italic_ok = not title or any(r.italic for r in runs)
        return label_ok and italic_ok

    def _reformat_caption(self, para: Any, label: str, num: str, title: str) -> None:
        """Reformat caption runs to bold label and italic title."""
        for r in tuple(para.runs):
            parent = r._element.getparent()
            if parent is not None:
                parent.remove(r._element)
        label_run = para.add_run(f"{label} {num}. ")
        label_run.bold = True
        self._apply_font_style(label_run, bold=True)
        if title:
            title_run = para.add_run(title)
            title_run.italic = True
            self._apply_font_style(title_run, italic=True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _has_adjacent_caption(self, img_p: Any, caption_re: Any) -> bool:
        """Return whether the paragraphs around an image contain its caption."""
        from docx.text.paragraph import Paragraph

        for el in (img_p._element.getprevious(), img_p._element.getnext()):
            if el is None or el.tag != qn(_W_P):
                continue
            neighbor = Paragraph(el, img_p._parent)
            if caption_re.match(neighbor.text.strip()):
                return True
        return False

    def _has_manual_title(self, img_p: Any) -> bool:
        """Return whether a short untitled line next to an image is its caption.

        Authors may write a plain caption line below an image ("My Figure
        Title") instead of a numbered "Figura N" label; auto-numbering over
        it would duplicate the caption, so it is left untouched.
        """
        from docx.text.paragraph import Paragraph

        for el in (img_p._element.getprevious(), img_p._element.getnext()):
            if el is None or el.tag != qn(_W_P):
                continue
            neighbor = Paragraph(el, img_p._parent)
            if self._is_skip_manual_title(neighbor):
                continue
            text = neighbor.text.strip()
            if len(text) < 80 and not text.endswith((".", ",", ";")):
                return True
        return False

    def _is_skip_manual_title(self, neighbor: Any) -> bool:
        """Check if neighbor should be skipped for manual title detection."""
        text = neighbor.text.strip()
        if not text:
            return True
        if text.startswith(("Nota.", "Note.")):
            return True
        style_name = paragraph_style_name(neighbor)
        return style_name.startswith("Heading")

    def _build_caption_element(self, number: int, title: str) -> Element:
        """Build a 'Figura N' (bold) + title (italic) caption paragraph."""
        prefix = cast(str, self._get_figure_config().get("caption_prefix", "Figure"))
        p_el = OxmlElement(_W_P)
        p_pr = OxmlElement(_W_PR)
        jc = OxmlElement(_W_JC)
        jc.set(qn(_W_VAL), "left")
        p_pr.append(jc)
        spacing = OxmlElement(_W_SPACING)
        spacing.set(qn(_W_AFTER), "0")
        spacing.set(qn(_W_LINE), "480")
        spacing.set(qn(_W_LINE_RULE), "auto")
        p_pr.append(spacing)
        p_el.append(p_pr)
        p_el.append(self._build_label_run(prefix, number))
        if title:
            p_el.append(self._build_title_run(title))
        return p_el

    def _build_label_run(self, prefix: str, number: int) -> Element:
        """Build bold label run for caption."""
        label_run = OxmlElement(_W_R)
        label_rpr = OxmlElement(_W_RPR)
        label_rpr.append(OxmlElement("w:b"))
        self._append_font_props(label_rpr)
        label_run.append(label_rpr)
        label_t = OxmlElement(_W_T)
        label_t.set(qn(_XML_SPACE), _PRESERVE)
        label_t.text = f"{prefix} {number}. "
        label_run.append(label_t)
        return label_run

    def _build_title_run(self, title: str) -> Element:
        """Build italic title run for caption."""
        title_run = OxmlElement(_W_R)
        title_rpr = OxmlElement(_W_RPR)
        title_rpr.append(OxmlElement("w:i"))
        self._append_font_props(title_rpr)
        title_run.append(title_rpr)
        title_t = OxmlElement(_W_T)
        title_t.set(qn(_XML_SPACE), _PRESERVE)
        title_t.text = title
        title_run.append(title_t)
        return title_run

    def _append_font_props(self, rpr: Element) -> None:
        """Append Times New Roman 12pt run properties to a w:rPr element."""
        font = OxmlElement("w:rFonts")
        font.set(qn("w:ascii"), DEFAULT_BODY_FONT)
        font.set(qn("w:hAnsi"), DEFAULT_BODY_FONT)
        rpr.append(font)
        sz = OxmlElement("w:sz")
        sz.set(qn(_W_VAL), "24")
        rpr.append(sz)

    @staticmethod
    def _extract_alt_text(p: Any) -> str:
        """Extract alt text from the first drawing of a paragraph."""
        for drawing in p._element.findall(f".//{qn(_W_DRAWING_QN)}"):
            for doc_pr in drawing.iter(f"{{{_NS_WP}}}docPr"):
                return str((doc_pr.get(_DESCR, "") or doc_pr.get(_NAME, "")).strip())
            for c_nv_pr in drawing.iter(f"{{{_NS_PIC}}}cNvPr"):
                return str((c_nv_pr.get(_DESCR, "") or c_nv_pr.get(_NAME, "")).strip())
        return ""
