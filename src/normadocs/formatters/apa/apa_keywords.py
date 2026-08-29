"""APA keywords and foreign word formatting."""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Any, cast

from docx.shared import Inches

from ...utils.docx_helpers import paragraph_style_name

_NOTA_PREFIX = "Nota."

if TYPE_CHECKING:
    from docx.document import Document as DocType
    from docx.text.paragraph import Paragraph as ParagraphType
    from docx.text.run import Run as RunType


_FOREIGN_WORDS = frozenset(
    (
        "Backend",
        "Frontend",
        "backend",
        "frontend",
        "PostgreSQL",
        "Redis",
        "Django",
        "React",
        "Next.js",
        "JavaScript",
        "Python",
        "Celery",
        "Docker",
        "Wompi",
        "WhatsApp",
        "iPhone",
        "iOS",
        "DDoS",
        "SSL",
        "PCI DSS",
        "RESTful",
        "API",
        "APIs",
        "SQL",
        "ORM",
        "CDN",
        "CEO",
    )
)

_KEYWORDS_RE = re.compile(r"((?:Palabras\s+clave|Keywords):)(.*)", re.IGNORECASE)


def _clear_paragraph(p: ParagraphType) -> ParagraphType:
    """Clear a paragraph's content while preserving formatting."""
    cast(Any, p._p).clear_content()
    return p


class APAKeywordsHandler:
    """Handles keywords and foreign word formatting per APA 7th Edition.

    Processes the keywords section and applies italics to foreign words
    according to APA 7th Edition standards.

    Args:
        doc: The python-docx Document object.
        config: Optional configuration dictionary.
    """

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        """Initialize APAKeywordsHandler.

        Args:
            doc: The python-docx Document object.
            config: Optional configuration dictionary.
        """
        self.doc = doc
        self.config = config if config is not None else {}

    def get_config(self, *keys: str, default: Any = None) -> Any:
        """Get nested config value via dot-notation keys."""
        value: Any = self.config
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key)
            else:
                return default
            if value is None:
                return default
        return value

    def _get_keywords_config(self) -> dict[str, Any]:
        """Get keywords configuration from config with defaults.

        Returns:
            Keywords configuration dictionary.
        """
        default: dict[str, Any] = {
            "keywords_label": "Keywords",
            "keywords_prefix": _NOTA_PREFIX,
        }
        return cast(dict[str, Any], self.get_config("keywords", default=default))

    def _apply_font_style(self, run: RunType, italic: bool | None = None) -> None:
        """Apply font style to a run (helper for this handler)."""
        from .apa_styles import APAStylesHandler

        handler = APAStylesHandler(self.doc)
        handler._apply_font_style(run, italic=italic)

    def format_keywords(self) -> None:
        """Format 'Palabras clave' per APA 7.

        APA 7 requires:
        - Keywords on its own line after abstract
        - Left indent of 0.5 inches
        - "Keywords:" label in italics
        """
        found_kw = self._format_keywords_paragraphs()
        if found_kw:
            self._add_page_break_before_introduction()

    def _format_keywords_paragraphs(self) -> bool:
        """Search and format the keywords paragraph, return whether found."""
        for p in self.doc.paragraphs:
            if self._is_keywords_paragraph(p):
                self._format_single_keywords_paragraph(p)
                return True
        return False

    def _is_keywords_paragraph(self, p: ParagraphType) -> bool:
        """Return True if paragraph contains keywords marker."""
        text_lower = p.text.lower()
        return "palabras clave" in text_lower or "keywords" in text_lower

    def _format_single_keywords_paragraph(self, p: ParagraphType) -> None:
        """Format a single keywords paragraph."""
        full = p.text.strip()
        match = _KEYWORDS_RE.search(full)
        if match is None:
            return
        label, content = match.groups()
        _clear_paragraph(p)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(0)
        r1 = p.add_run(label + " ")
        r1.italic = True
        self._apply_font_style(r1)
        r2 = p.add_run(content.strip())
        self._apply_font_style(r2)

    def _add_page_break_before_introduction(self) -> None:
        """Add page break before the Introduction section.

        APA 7: After keywords, the introduction starts on a new page.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        from .apa_page import APAPageHandler

        for p in self.doc.paragraphs:
            if paragraph_style_name(p).startswith("Heading"):
                text = p.text.strip().lower()
                if "introducción" in text or "introduction" in text:
                    if APAPageHandler._has_page_break_before(p):
                        break
                    break_paragraph = OxmlElement("w:p")
                    break_run = OxmlElement("w:r")
                    page_break = OxmlElement("w:br")
                    page_break.set(qn("w:type"), "page")
                    break_run.append(page_break)
                    break_paragraph.append(break_run)
                    p._element.addprevious(break_paragraph)
                    break

    def format_nota_italic(self) -> None:
        """APA 7: 'Nota.' must be italic in figure/table notes.

        Finds paragraphs starting with 'Nota.' and splits the first run
        so that 'Nota.' is italic and the rest is regular.
        """
        for p in self.doc.paragraphs:
            if not p.text.strip().startswith(_NOTA_PREFIX):
                continue
            self._format_single_nota_paragraph(p)

    def _format_single_nota_paragraph(self, p: ParagraphType) -> None:
        """Rebuild a single Nota paragraph with italic prefix."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        full_text = p.text
        self._clear_runs(p)
        nota_run = p.add_run(f"{_NOTA_PREFIX} ")
        nota_run.italic = True
        self._apply_font_style(nota_run, italic=True)
        rest = full_text[len(_NOTA_PREFIX) :].strip()
        if rest:
            rest_run = p.add_run(rest)
            self._apply_font_style(rest_run)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0)

    def _clear_runs(self, p: ParagraphType) -> None:
        """Remove all runs from a paragraph."""
        for run in tuple(p.runs):
            parent = run._element.getparent()
            if parent is not None:
                parent.remove(run._element)

    def apply_foreign_word_italics(self) -> None:
        """Apply italics to foreign words per APA 7 (Backend, Frontend, etc.).

        APA 7 requires that foreign words used as nouns be italicized on first use.
        This method searches for these terms in table cells and applies italics.
        """
        for table in self.doc.tables:
            self._process_table_for_foreign(table)

    def _process_table_for_foreign(self, table: Any) -> None:
        """Process a single table for foreign words."""
        for row in table.rows:
            self._process_row_for_foreign(row)

    def _process_row_for_foreign(self, row: Any) -> None:
        """Process a single row."""
        for cell in row.cells:
            self._process_cell_for_foreign(cell)

    def _process_cell_for_foreign(self, cell: Any) -> None:
        """Process a single cell."""
        for p in cell.paragraphs:
            self._process_paragraph_for_foreign(p)

    def _process_paragraph_for_foreign(self, p: ParagraphType) -> None:
        """Process paragraph runs for foreign words."""
        for run in tuple(p.runs):
            self._process_run_for_foreign(run)

    def _process_run_for_foreign(self, run: RunType) -> None:
        """Apply italic to a run if it contains foreign words."""
        if not run.text:
            return
        if run.italic:
            return
        if not self._contains_foreign_word(run.text):
            return
        run.italic = True

    def _contains_foreign_word(self, text: str) -> bool:
        """Return True if text contains any foreign word."""
        return any(fw in text for fw in _FOREIGN_WORDS)
