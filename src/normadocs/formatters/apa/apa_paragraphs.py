"""APA paragraph processing, formatting, and cleanup."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any, cast

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ...config import (
    BODY_TEXT_STYLE,
    COMPACT_STYLE,
    DEFAULT_BODY_FONT,
    HEADING_1_STYLE,
    HEADING_4_STYLE,
    HEADING_5_STYLE,
    NORMAL_STYLE,
    W_VAL,
)
from ...utils.docx_helpers import paragraph_style_name
from .apa_citations import REFERENCE_HEADINGS

if TYPE_CHECKING:
    from docx.document import Document as DocType
    from docx.text.paragraph import Paragraph as ParagraphType

_HEADING_1 = HEADING_1_STYLE
_HEADING_5 = HEADING_5_STYLE
_RUN_IN_HEADINGS = (HEADING_4_STYLE, _HEADING_5)

_TOC_DOTS_RE = re.compile(r"\.{3,}\s*\d+\s*$")
_TOC_SPACES_RE = re.compile(r"\s{2,}\d+\s*$")
_BLOCK_CLOSING_RE = re.compile(r'["\u201c\u201d\u00ab\u00bb]\s*(\([^()]*\))?\.?\s*$')
_CITATION_RE = re.compile(
    r"\(([A-ZÁ-Ú][a-záéíóúñ]+(?:\s+et\s+al\.)?)\s+y\s+([A-ZÁ-Ú][a-záéíóúñ]+),\s*(\d{4})\)"
)
_GRID_TABLE_RE = re.compile(r"[+|][-=]{3,}[+|]?")
_EQUALS_RE = re.compile(r"={3,}")
_MULTI_SPACE_RE = re.compile(r"\s{2,}")
_HEADING_NUMBER_RE = re.compile(r"^\d+(\.\d+)*\s*")


@dataclass
class ParagraphState:
    """Mutable state tracked while iterating paragraphs."""

    in_references: bool = False
    in_toc: bool = False
    in_abstract: bool = False
    just_left_abstract: bool = False
    first_paragraph_after_heading: bool = False
    first_heading_seen: bool = False
    heading_levels: dict[str, int] = field(default_factory=dict)


def _is_references_heading(stripped_lower: str) -> bool:
    """Return whether a stripped, lowercased heading title starts the references."""
    return stripped_lower in REFERENCE_HEADINGS or stripped_lower.startswith(
        ("referencias ", "references ", "lista de referencias")
    )


def _clear_paragraph(p: ParagraphType) -> ParagraphType:
    """Clear a paragraph's content while preserving formatting."""
    cast(Any, p._p).clear_content()
    return p


class APAParagraphsHandler:
    """Handles paragraph processing, formatting, and cleanup per APA 7th Edition."""

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        """Initialize APAParagraphsHandler.

        Args:
            doc: The python-docx Document object.
            config: Optional configuration dictionary.
        """
        self.doc = doc
        self.config = config if config is not None else {}

    def _get_spacing_line(self) -> str:
        """Get line spacing from config with default."""
        spacing: dict[str, str] = {"line": "double"}
        return cast(str, self.config.get("spacing", spacing).get("line", "double"))

    def _get_body_font(self) -> str:
        """Get body font name from config."""
        fonts: dict[str, Any] = {}
        return cast(
            str, self.config.get("fonts", fonts).get("body", {}).get("name", DEFAULT_BODY_FONT)
        )

    @staticmethod
    def _has_page_break_before(paragraph: ParagraphType) -> bool:
        """Return whether the nearest preceding body content is a page break."""
        previous = paragraph._element.getprevious()
        while previous is not None:
            if any(br.get(qn("w:type")) == "page" for br in previous.iter(qn("w:br"))):
                return True
            if previous.tag == qn("w:p") and not "".join(previous.itertext()).strip():
                previous = previous.getprevious()
                continue
            return False
        return False

    def _set_page_break_before(self, paragraph: ParagraphType) -> None:
        """Add a page break only when one is not already immediately before it.

        The Markdown preprocessor emits explicit OpenXML page breaks before
        level-1 headings. Adding ``page_break_before`` again creates blank
        pages in LibreOffice, so the formatter deduplicates the two paths per
        heading rather than using a document-wide flag.
        """
        if not self._has_page_break_before(paragraph):
            paragraph.paragraph_format.page_break_before = True

    def process(self) -> None:
        """Iterate through paragraphs to apply APA 7 formatting.

        APA 7 requires:
        - RESUMEN (Abstract): Title centered, bold, text without indent
        - Body text paragraphs: First line indent of 0.5 inches
        - First paragraph after any heading: No indent
        - References: Hanging indent (0.5 inches)
        """
        from .apa_styles import APAStylesHandler

        styles_handler = APAStylesHandler(self.doc)
        state = ParagraphState(heading_levels=self._build_heading_level_map())

        for p in self.doc.paragraphs:
            style_name = paragraph_style_name(p)
            styles_handler._apply_font_to_paragraph(p)
            text_lower = p.text.lower()
            if style_name.startswith("Heading"):
                self._process_heading(p, state, style_name, text_lower)
            self._apply_spacing(p, style_name)
            if self._handle_toc_entry(p, state, style_name):
                continue
            self._apply_body_formatting(p, state, style_name)
            self._fix_citations(p)

    def _process_heading(
        self, p: ParagraphType, state: ParagraphState, style_name: str, text_lower: str
    ) -> None:
        """Dispatch heading handling and apply heading style."""
        heading_stripped = text_lower.strip().rstrip(".")
        if _is_references_heading(heading_stripped):
            self._handle_references_heading(p, state)
        elif "resumen" in text_lower or "abstract" in text_lower:
            self._handle_abstract_heading(p, state)
        elif "contenido" in text_lower or "index" in text_lower:
            self._handle_toc_heading(p, state)
        else:
            self._handle_generic_heading(p, state, style_name, text_lower)
        self._apply_heading_alignment(p, style_name)
        self._strip_heading_numbering(p)
        self._apply_heading_indent(p, style_name)

    def _handle_references_heading(self, p: ParagraphType, state: ParagraphState) -> None:
        """Update state for a references heading."""
        state.in_references = True
        state.in_toc = False
        state.in_abstract = False
        self._set_page_break_before(p)
        state.first_paragraph_after_heading = True

    def _handle_abstract_heading(self, p: ParagraphType, state: ParagraphState) -> None:
        """Update state and formatting for abstract heading."""
        state.in_abstract = True
        state.in_toc = False
        state.in_references = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
        p.paragraph_format.first_line_indent = Inches(0)

    def _handle_toc_heading(self, p: ParagraphType, state: ParagraphState) -> None:
        """Update state for table-of-contents heading."""
        state.in_toc = True
        state.in_references = False
        self._set_page_break_before(p)

    def _handle_generic_heading(
        self, p: ParagraphType, state: ParagraphState, style_name: str, text_lower: str
    ) -> None:
        """Handle non-special headings and reset section flags."""
        del text_lower
        if style_name == _HEADING_1:
            self._handle_h1_page_break(p, state)
            state.first_paragraph_after_heading = True
        if state.in_abstract or state.just_left_abstract:
            self._set_page_break_before(p)
            state.just_left_abstract = False
        state.in_references = False
        state.in_abstract = False
        state.in_toc = False

    def _handle_h1_page_break(self, p: ParagraphType, state: ParagraphState) -> None:
        """Clear stale breaks or mark first heading seen for H1."""
        if state.first_heading_seen:
            p.paragraph_format.page_break_before = False
        else:
            state.first_heading_seen = True

    def _apply_heading_alignment(self, p: ParagraphType, style_name: str) -> None:
        """Apply APA alignment and bold/italic per heading level."""
        if style_name == _HEADING_1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
            return
        config_map: dict[str, tuple[WD_ALIGN_PARAGRAPH, bool, bool]] = {
            "Heading 2": (WD_ALIGN_PARAGRAPH.LEFT, True, False),
            "Heading 3": (WD_ALIGN_PARAGRAPH.LEFT, True, True),
        }
        if style_name in config_map:
            align, bold, italic = config_map[style_name]
            p.alignment = align
            for run in p.runs:
                run.bold = bold
                run.italic = italic
            return
        if style_name in _RUN_IN_HEADINGS:
            self._format_run_in_heading(p, italic=(style_name == _HEADING_5))

    def _strip_heading_numbering(self, p: ParagraphType) -> None:
        """Remove numbering property from headings."""
        p_pr = p._element.find(qn("w:pPr"))
        if p_pr is None:
            return
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)

    def _apply_heading_indent(self, p: ParagraphType, style_name: str) -> None:
        """Ensure non run-in headings have no first-line indent."""
        if style_name not in _RUN_IN_HEADINGS:
            p.paragraph_format.first_line_indent = Inches(0)

    def _apply_spacing(self, p: ParagraphType, style_name: str) -> None:
        """Apply line spacing and widow/orphan controls."""
        spacing_line = self._get_spacing_line()
        self._set_line_spacing(p, spacing_line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        self._apply_paragraph_spacing_control(p)
        if style_name.startswith("Heading"):
            self._apply_keep_with_next(p)

    def _set_line_spacing(self, p: ParagraphType, spacing_line: str) -> None:
        """Set line spacing rule based on config."""
        if spacing_line == "double":
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            return
        if spacing_line == "1.5":
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            return
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _handle_toc_entry(self, p: ParagraphType, state: ParagraphState, style_name: str) -> bool:
        """Format TOC entry if in TOC section; return True if handled."""
        if not state.in_toc or style_name.startswith("Heading"):
            return False
        self._format_toc_entry(p, state.heading_levels)
        return True

    def _apply_body_formatting(
        self, p: ParagraphType, state: ParagraphState, style_name: str
    ) -> None:
        """Dispatch body indentation after heading/TOC checks."""
        if style_name.startswith("Heading"):
            return
        if "List" in style_name or "Caption" in style_name:
            return
        text_strip = p.text.strip()
        if self._remove_numeric_paragraph(p, text_strip):
            return
        if state.in_references:
            self._format_reference_body(p, text_strip)
            return
        if state.in_abstract:
            self._format_abstract_body(p, state, text_strip)
            return
        if self._is_body_style(style_name):
            self._format_normal_body(p, state, text_strip)

    def _remove_numeric_paragraph(self, p: ParagraphType, text_strip: str) -> bool:
        """Remove purely numeric paragraphs; return True if removed."""
        if not text_strip.isdigit():
            return False
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
        return True

    def _is_body_style(self, style_name: str) -> bool:
        """Return whether style is a body paragraph style."""
        return style_name in (BODY_TEXT_STYLE, NORMAL_STYLE, "First Paragraph", COMPACT_STYLE)

    def _format_reference_body(self, p: ParagraphType, text_strip: str) -> None:
        """Apply hanging indent to reference entries."""
        if not text_strip:
            return
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pr = p._element.get_or_add_pPr()
        jc = p_pr.get_or_add_jc()
        jc.set(qn(W_VAL), "left")

    def _format_abstract_body(
        self, p: ParagraphType, state: ParagraphState, text_strip: str
    ) -> None:
        """Format abstract block and detect keywords end."""
        if p.text.strip():
            p.paragraph_format.first_line_indent = Inches(0)
        if text_strip.lower().startswith(("palabras clave", "keywords")):
            state.in_abstract = False
            state.just_left_abstract = True

    def _format_normal_body(self, p: ParagraphType, state: ParagraphState, text_strip: str) -> None:
        """Apply normal body indentation and block-quote handling."""
        if not text_strip:
            return
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return
        if state.just_left_abstract:
            self._set_page_break_before(p)
            state.just_left_abstract = False
        if self._is_block_quote(text_strip):
            self._convert_block_quote(p, text_strip)
            return
        self._apply_first_line_rule(p, state)

    def _apply_first_line_rule(self, p: ParagraphType, state: ParagraphState) -> None:
        """Apply first-line indent rule based on heading proximity."""
        if state.first_paragraph_after_heading:
            p.paragraph_format.first_line_indent = Inches(0)
            state.first_paragraph_after_heading = False
        else:
            p.paragraph_format.first_line_indent = Inches(0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pr = p._element.get_or_add_pPr()
        jc = p_pr.get_or_add_jc()
        jc.set(qn(W_VAL), "left")

    def _apply_paragraph_spacing_control(self, p: ParagraphType) -> None:
        """Apply widow/orphan control per APA 7.

        APA 7 requires widow/orphan control to prevent single lines
        at the top or bottom of a page.
        """
        p_pr = p._element.get_or_add_pPr()
        widow_orphan = OxmlElement("w:widowControl")
        p_pr.append(widow_orphan)
        doc_grid = OxmlElement("w:docGrid")
        doc_grid.set(qn("w:type"), "lines")
        doc_grid.set(qn("w:linePitch"), "360")
        p_pr.append(doc_grid)

    def _apply_keep_with_next(self, p: ParagraphType) -> None:
        """Apply keep-with-next to heading paragraphs per APA 7.

        Headings should stay on the same page as the following paragraph.
        """
        p_pr = p._element.get_or_add_pPr()
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)

    def _format_run_in_heading(self, p: ParagraphType, italic: bool) -> None:
        """Format a Level 4/5 heading as an APA 7 run-in heading.

        Level 4: indented, bold, ends with a period, text begins on the
        same line. Level 5 adds italics to the heading portion. The first
        body paragraph after the heading is merged onto the heading line so
        the text truly runs in, with explicit non-bold runs so the bold
        heading style does not bleed into the body text.
        """
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0.5)
        for run in p.runs:
            run.bold = True
            run.italic = italic
        self._ensure_heading_period(p)
        self._merge_next_paragraph_into_heading(p, italic)

    def _ensure_heading_period(self, p: ParagraphType) -> None:
        """Ensure run-in heading ends with a period."""
        if p.text.strip().endswith("."):
            return
        if p.runs:
            p.runs[-1].text = f"{p.runs[-1].text.rstrip()}."
        else:
            p.add_run(".")

    def _merge_next_paragraph_into_heading(self, p: ParagraphType, italic: bool) -> None:
        """Merge the following body paragraph into a run-in heading."""
        from docx.text.paragraph import Paragraph

        next_el = p._element.getnext()
        if next_el is None or next_el.tag != qn("w:p"):
            return
        if next_el.findall(f".//{qn('w:drawing')}") or next_el.findall(f".//{qn('w:hyperlink')}"):
            return
        next_p = Paragraph(next_el, p._parent)
        next_style = next_p.style.name if next_p.style else ""
        if next_style not in (BODY_TEXT_STYLE, NORMAL_STYLE, "First Paragraph"):
            return
        if not next_p.text.strip():
            return
        runs = list(next_p.runs)
        if runs:
            runs[0].text = f" {runs[0].text.lstrip()}"
        for run in runs:
            run.bold = False
            if italic and run.italic is None:
                run.italic = False
            p._element.append(run._element)
        parent = next_el.getparent()
        if parent is not None:
            parent.remove(next_el)

    def _get_block_quote_config(self) -> dict[str, Any]:
        """Get block-quote configuration with APA defaults."""
        default_config: dict[str, Any] = {"min_words": 40, "indent_inches": 0.5}
        return cast(dict[str, Any], self.config.get("block_quote", default_config))

    def _is_block_quote(self, text: str) -> bool:
        """Return whether a quoted paragraph qualifies as an APA block quote.

        APA 8.27: quotations of 40+ words are formatted as a freestanding
        block without quotation marks.
        """
        if not text or text[:1] not in ('"', "\u201c", "\u00ab"):
            return False
        min_words = cast(int, self._get_block_quote_config().get("min_words", 40))
        return len(text.split()) >= min_words

    def _convert_block_quote(self, p: ParagraphType, text: str) -> None:
        """Convert a long quotation into an APA 8.27 block quote.

        Removes the surrounding quotation marks, indents the whole block
        0.5" from the left (no first-line indent), and moves the final
        punctuation before the closing citation parenthesis.
        """
        indent = cast(float, self._get_block_quote_config().get("indent_inches", 0.5))
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.first_line_indent = Inches(0)
        stripped = text.strip()
        quote_chars = ('"', "\u201c", "\u201d", "\u00ab", "\u00bb")
        if stripped[:1] in quote_chars:
            stripped = stripped[1:].lstrip()
        stripped = self._fix_block_quote_closing(stripped)
        if stripped and stripped != text.strip():
            self._replace_paragraph_text(p, stripped)

    def _fix_block_quote_closing(self, stripped: str) -> str:
        """Move closing quote and punctuation before citation."""
        closing = _BLOCK_CLOSING_RE.search(stripped)
        if closing is None:
            return stripped
        citation = closing.group(1)
        body = stripped[: closing.start()].rstrip()
        if body.endswith((",", ";", ":")):
            body = f"{body[:-1]}."
        if not body.endswith((".", "!", "?")):
            body = f"{body}."
        return f"{body} {citation}" if citation else body

    def _replace_paragraph_text(self, p: ParagraphType, text: str) -> None:
        """Replace paragraph text with a single run, keeping the first run's font."""
        keep_font = None
        keep_size = None
        for run in p.runs:
            if run.text.strip():
                keep_font = run.font.name
                keep_size = run.font.size
                break
        for run in p.runs:
            parent = run._element.getparent()
            if parent is not None:
                parent.remove(run._element)
        new_run = p.add_run(text)
        if keep_font is not None:
            new_run.font.name = keep_font
        if keep_size is not None:
            new_run.font.size = keep_size

    def _fix_citations(self, p: ParagraphType) -> None:
        """Replace 'y' citations with '&' per APA 7.

        Converts Spanish-style 'y' conjunctions in citations to '&'.

        Args:
            p: The paragraph to process.
        """
        for run in p.runs:
            if " y " in run.text and "(" in run.text:
                run.text = _CITATION_RE.sub(r"(\1 & \2, \3)", run.text)

    def _format_toc_entry(self, p: ParagraphType, heading_levels: dict[str, int]) -> None:
        """Format Table of Contents entries with correct indentation."""
        from .apa_styles import APAStylesHandler

        styles_handler = APAStylesHandler(self.doc)
        text = p.text.strip()
        if not text:
            return
        title, page_num = self._parse_toc_entry(text)
        if not title or not page_num:
            return
        level = heading_levels.get(title.lower().strip(), 1)
        left_indent = {1: 0, 2: 0.5, 3: 1.0}.get(level, 0)
        _clear_paragraph(p)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(left_indent)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(
            Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS
        )
        self._populate_toc_runs(p, styles_handler, title, page_num)

    def _parse_toc_entry(self, text: str) -> tuple[str | None, str | None]:
        """Parse TOC entry into title and page number using linear checks."""
        title, page = self._parse_toc_tab(text)
        if title is not None:
            return title, page
        normalized = text.replace("\u2026", "...")
        title, page = self._parse_toc_dots(normalized)
        if title is not None:
            return title, page
        return self._parse_toc_spaces(text)

    def _parse_toc_tab(self, text: str) -> tuple[str | None, str | None]:
        """Try tab-separated TOC entry."""
        if "\t" not in text:
            return None, None
        parts = text.rsplit("\t", 1)
        if len(parts) == 2 and parts[1].strip().isdigit():
            return parts[0].strip(), parts[1].strip()
        return None, None

    def _parse_toc_dots(self, normalized: str) -> tuple[str | None, str | None]:
        """Try dot-separated TOC entry with linear regex."""
        if "..." not in normalized:
            return None, None
        match = _TOC_DOTS_RE.search(normalized)
        if match is None:
            return None, None
        # Extract page digits at end
        page_match = re.search(r"\d+\s*$", match.group(0))
        page = page_match.group(0).strip() if page_match else ""
        title = normalized[: match.start()].strip()
        return (title, page) if title and page else (None, None)

    def _parse_toc_spaces(self, text: str) -> tuple[str | None, str | None]:
        """Try space-separated TOC entry with linear regex."""
        match = _TOC_SPACES_RE.search(text)
        if match is None:
            return None, None
        page_match = re.search(r"\d+\s*$", match.group(0))
        page = page_match.group(0).strip() if page_match else ""
        title = text[: match.start()].strip()
        return (title, page) if title and page else (None, None)

    def _populate_toc_runs(
        self, p: ParagraphType, styles_handler: Any, title: str, page_num: str
    ) -> None:
        """Create TOC runs with leaders and formatting."""
        run = p.add_run(title)
        styles_handler._apply_font_style(run)
        if title and title[:1].isdigit():
            run.text = "\u200b" + run.text
        run = p.add_run("\t")
        styles_handler._apply_font_style(run)
        run = p.add_run(page_num)
        styles_handler._apply_font_style(run)

    def _build_heading_level_map(self) -> dict[str, int]:
        """Build a map of heading text -> heading level from the document."""
        levels = {}
        for p in self.doc.paragraphs:
            style_name = paragraph_style_name(p)
            if style_name.startswith("Heading"):
                parts = style_name.split()
                if len(parts) >= 2 and parts[-1].isdigit():
                    level = int(parts[-1])
                    text = p.text.strip().lower()
                    if text:
                        levels[text] = level
        return levels

    def format_lists(self) -> None:
        """Apply APA 7 list formatting.

        Regular bullet lists: bullet at 0.5in, text at 0.75in, hanging indent.
        Reference entries: no bullet, hanging indent at 0.5in (APA 7 standard).
        """
        from .apa_styles import APAStylesHandler

        styles_handler = APAStylesHandler(self.doc)
        in_references = False
        for p in self.doc.paragraphs:
            style_name = paragraph_style_name(p)
            if style_name == _HEADING_1:
                text_lower = p.text.lower().strip().rstrip(".")
                in_references = _is_references_heading(text_lower)
                continue
            if self._should_skip_list_paragraph(p):
                continue
            if in_references:
                self._format_reference_list(p)
            else:
                self._format_bullet_list(p, styles_handler)

    def _should_skip_list_paragraph(self, p: ParagraphType) -> bool:
        """Return whether paragraph lacks list numbering."""
        p_pr = p._element.find(qn("w:pPr"))
        if p_pr is None:
            return True
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return True
        p_pr.remove(num_pr)
        return False

    def _format_reference_list(self, p: ParagraphType) -> None:
        """Apply hanging indent without bullet for references."""
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    def _format_bullet_list(self, p: ParagraphType, styles_handler: Any) -> None:
        """Apply bullet formatting with hanging indent and tab."""
        p.paragraph_format.left_indent = Inches(0.75)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(0.75), alignment=WD_TAB_ALIGNMENT.LEFT)
        text = p.text
        if text.startswith("\u2022") or text.startswith("-"):
            return
        first_run = p.runs[0] if p.runs else None
        if first_run:
            first_run.text = "\u2022\t" + first_run.text
        else:
            run = p.add_run("\u2022\t")
            styles_handler._apply_font_style(run)

    def apply_body_indent(self) -> None:
        """Final pass: apply first-line indent to all body paragraphs.

        Runs after _format_tables so that newly-created paragraphs (table
        descriptions, etc.) are also covered.
        """
        state = ParagraphState()
        for p in self.doc.paragraphs:
            style_name = paragraph_style_name(p)
            text = p.text.strip()
            if style_name.startswith("Heading"):
                self._update_indent_state(p, state, text)
                continue
            if self._should_skip_body_indent(p, state, text, style_name):
                continue
            p.paragraph_format.first_line_indent = Inches(0.5)

    def _update_indent_state(self, _p: ParagraphType, state: ParagraphState, text: str) -> None:
        """Update indent-tracking state for heading paragraphs."""
        text_lower = text.lower().strip().rstrip(".")
        if _is_references_heading(text_lower):
            state.in_references = True
            state.in_toc = False
            state.in_abstract = False
            return
        if "contenido" in text_lower or "index" in text_lower:
            state.in_toc = True
            state.in_references = False
            state.in_abstract = False
            return
        if "resumen" in text_lower or "abstract" in text_lower:
            state.in_abstract = True
            state.in_toc = False
            state.in_references = False
            return
        state.in_references = False
        state.in_toc = False
        state.in_abstract = False

    def _should_skip_body_indent(
        self, p: ParagraphType, state: ParagraphState, text: str, style_name: str
    ) -> bool:
        """Return whether body indent should be skipped for this paragraph."""
        if not text or len(text) < 5:
            return True
        if self._has_context_skip(state, text):
            return True
        return self._has_format_skip(p, style_name, text)

    def _has_context_skip(self, state: ParagraphState, text: str) -> bool:
        """Check context-based skips (TOC, abstract, references)."""
        if state.in_toc:
            return True
        if state.in_abstract:
            if text.lower().startswith(("palabras clave", "keywords")):
                state.in_abstract = False
            return True
        return state.in_references

    def _has_format_skip(self, p: ParagraphType, style_name: str, text: str) -> bool:
        """Check formatting-based skips (indent, alignment, style)."""
        if self._has_existing_indent(p):
            return True
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        if self._has_block_left_indent(p):
            return True
        if "List" in style_name or "Caption" in style_name or COMPACT_STYLE in style_name:
            return True
        if self._is_caption_or_note(text):
            return True
        return self._is_short_italic_caption(p, text)

    def _has_existing_indent(self, p: ParagraphType) -> bool:
        """Check if paragraph already has first-line indent."""
        fli = p.paragraph_format.first_line_indent
        return fli is not None and fli != 0

    def _has_block_left_indent(self, p: ParagraphType) -> bool:
        """Check if paragraph has block left indent (block quote)."""
        left_indent = p.paragraph_format.left_indent
        return left_indent is not None and left_indent >= Inches(0.25)

    def _is_caption_or_note(self, text: str) -> bool:
        """Check if text is a short caption or nota paragraph."""
        if text.startswith(("Tabla ", "Figura ")) and len(text.split()) <= 2:
            return True
        return text.startswith("Nota.")

    def _is_short_italic_caption(self, p: ParagraphType, text: str) -> bool:
        """Check if short all-italic paragraph is a caption title."""
        if len(text) >= 80:
            return False
        if not p.runs:
            return False
        return all(r.italic for r in p.runs if r.text.strip())

    def fix_text_spacing_global(self) -> None:
        """Run merge_and_clean on all paragraphs."""
        for p in self.doc.paragraphs:
            if not paragraph_style_name(p).startswith("Heading"):
                self._merge_and_clean_paragraph(p)
                if p.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _merge_and_clean_paragraph(self, p: ParagraphType) -> None:
        """Consolidate runs while preserving inline formatting boundaries.

        Groups consecutive runs by their (bold, italic) attributes and merges
        only within each group, so italic titles in references and the
        "Nota." italic split are kept intact.

        Skips paragraphs containing embedded images (w:drawing) to avoid
        destroying their XML structure.
        Skips paragraphs containing OMML math (m:oMath) — re-created runs are
        appended at the end of the paragraph, which would reorder text around
        the equation.
        Skips Source Code paragraphs to preserve ASCII art and code formatting.
        """
        if self._should_skip_merge(p):
            return
        groups = self._build_format_groups(p)
        self._clear_runs(p)
        self._recreate_runs(p, groups)

    def _should_skip_merge(self, p: ParagraphType) -> bool:
        """Return whether paragraph should skip merge and clean."""
        if not p.runs:
            return True
        if p._element.findall(f".//{qn('w:drawing')}"):
            return True
        if p._element.findall(f".//{qn('m:oMath')}"):
            return True
        if p._element.findall(f".//{qn('w:br')}"):
            return True
        style_name = p.style.name if p.style else ""
        return style_name in ("Source Code", "Source", "Code", "Preformatted", "HTMLPre")

    def _build_format_groups(
        self, p: ParagraphType
    ) -> list[tuple[bool, bool, str, str | None, Any]]:
        """Group consecutive runs by bold/italic attributes."""
        groups: list[tuple[bool, bool, str, str | None, Any]] = []
        for run in p.runs:
            t = run.text or ""
            t = t.replace("\r", " ")
            t = _GRID_TABLE_RE.sub("", t)
            t = _EQUALS_RE.sub("", t)
            t = t.strip("|")
            is_bold = bool(run.bold)
            is_italic = bool(run.italic)
            font_name = run.font.name
            font_size = run.font.size
            if groups and groups[-1][0] == is_bold and groups[-1][1] == is_italic:
                groups[-1] = (
                    is_bold,
                    is_italic,
                    groups[-1][2] + t,
                    font_name or groups[-1][3],
                    font_size or groups[-1][4],
                )
            else:
                groups.append((is_bold, is_italic, t, font_name, font_size))
        return groups

    def _clear_runs(self, p: ParagraphType) -> None:
        """Clear all runs from paragraph."""
        for run in list(p.runs):
            parent = run._r.getparent()
            if parent is not None:
                parent.remove(run._r)

    def _recreate_runs(
        self, p: ParagraphType, groups: list[tuple[bool, bool, str, str | None, Any]]
    ) -> None:
        """Re-create runs from format groups with cleaned text."""
        from docx.shared import Pt

        for idx, (is_bold, is_italic, text, font_name, font_size) in enumerate(groups):
            text = _MULTI_SPACE_RE.sub(" ", text)
            if idx == 0:
                text = text.lstrip()
            if idx == len(groups) - 1:
                text = text.rstrip()
            if not text:
                continue
            new_run = p.add_run(text)
            new_run.font.name = font_name or DEFAULT_BODY_FONT
            new_run.font.size = font_size or Pt(12)
            if is_bold:
                new_run.bold = True
            if is_italic:
                new_run.italic = True
