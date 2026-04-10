"""APA 7th Edition formatter - main class that delegates to handlers."""

from typing import Any, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tc
from docx.section import Section
from docx.table import Table as TableObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ...models import DocumentMetadata
from .apa_cover import APACoverHandler
from .apa_figures import APAFiguresHandler
from .apa_keywords import APAKeywordsHandler
from .apa_page import APAPageHandler
from .apa_paragraphs import APAParagraphsHandler
from .apa_styles import APAStylesHandler
from .apa_tables import APATablesHandler


class APADocxFormatter:
    """Applies APA 7th Edition formatting to a DOCX file.

    This class orchestrates multiple handler classes to apply comprehensive
    APA 7th Edition formatting including cover page, styles, page layout,
    paragraphs, tables, figures, and keywords.

    Args:
        doc_path: Path to the DOCX file to format.
        config: Optional configuration dictionary to override defaults.
    """

    def __init__(self, doc_path: str, config: dict[str, Any] | None = None) -> None:
        """Initialize the formatter with document path and optional config.

        Args:
            doc_path: Path to the DOCX file to format.
            config: Optional configuration dictionary to override defaults.
        """
        self._doc: DocumentObject = Document(doc_path)
        self.config = config if config is not None else {}

        # Initialize handlers with config
        self._styles = APAStylesHandler(self._doc, self.config)
        self._page = APAPageHandler(self._doc, self.config)
        self._cover = APACoverHandler(self._doc, self.config)
        self._paragraphs = APAParagraphsHandler(self._doc, self.config)
        self._tables = APATablesHandler(self._doc, self.config)
        self._figures = APAFiguresHandler(self._doc, self.config)
        self._keywords = APAKeywordsHandler(self._doc, self.config)

    def process(self, meta: DocumentMetadata) -> None:
        """Run the full formatting pipeline.

        Processes the document through all APA 7th Edition formatting stages
        including cover page, styles, page layout, paragraphs, tables,
        figures, and keywords.

        Args:
            meta: DocumentMetadata containing title, author, institution, etc.
        """
        self._page.setup_page_layout()
        self._styles.create_styles()
        self._cover.add_cover_page(meta)
        self._paragraphs.process()

        self._tables.add_table_captions()
        self._tables.add_table_notes()
        self._tables.format_tables()

        self._figures.add_figure_captions()
        self._figures.format_figures()
        self._keywords.format_nota_italic()
        self._paragraphs.format_lists()
        self._paragraphs.apply_body_indent()
        self._keywords.format_keywords(meta)
        self._page.add_section_page_breaks()
        self._paragraphs.fix_text_spacing_global()

        self._tables.add_table_header_bold()
        self._keywords.apply_foreign_word_italics()

    def save(self, output_path: str) -> None:
        """Save the formatted document to output_path.

        Args:
            output_path: Path where the formatted DOCX will be saved.
        """
        self._doc.save(str(output_path))

    @property
    def doc(self) -> DocumentObject:
        """The underlying python-docx Document object."""
        return self._doc

    @doc.setter
    def doc(self, value: DocumentObject) -> None:
        """Set the underlying document."""
        self._doc = value
        # Initialize config if not already set (for tests using __new__)
        if not hasattr(self, "config"):
            self.config = {}
        # Reinitialize handlers with the new document and same config
        self._styles = APAStylesHandler(self._doc, self.config)
        self._page = APAPageHandler(self._doc, self.config)
        self._cover = APACoverHandler(self._doc, self.config)
        self._paragraphs = APAParagraphsHandler(self._doc, self.config)
        self._tables = APATablesHandler(self._doc, self.config)
        self._figures = APAFiguresHandler(self._doc, self.config)
        self._keywords = APAKeywordsHandler(self._doc, self.config)

    # ─────────────────── Delegate methods for backward compatibility ───────────────────

    # Cover handler delegation
    def _add_cover_page(self, meta: DocumentMetadata) -> None:
        """Delegate to APACoverHandler to add the cover page."""
        self._cover.add_cover_page(meta)

    # Page handler delegation
    def _setup_page_layout(self) -> None:
        """Delegate to APAPageHandler to set up page margins."""
        self._page.setup_page_layout()

    def _add_page_number(self, section: Section) -> None:
        """Delegate to APAPageHandler to add page numbers."""
        self._page._add_page_number(section)

    def _add_section_page_breaks(self) -> None:
        """Delegate to APAPageHandler to add section page breaks."""
        self._page.add_section_page_breaks()

    # Styles handler delegation
    def _create_styles(self) -> None:
        """Delegate to APAStylesHandler to create APA styles."""
        self._styles.create_styles()

    def _apply_font_style(
        self,
        style_or_run: Paragraph | Run,
        font_name: str = "Times New Roman",
        size: int = 12,
        bold: bool | None = None,
        italic: bool | None = None,
        color_rgb: tuple[int, int, int] = (0, 0, 0),
    ) -> None:
        """Apply font style settings to a style or run element."""
        self._styles._apply_font_style(style_or_run, font_name, size, bold, italic, color_rgb)

    def _apply_font_to_paragraph(self, paragraph: Paragraph, font_size: int | None = None) -> None:
        """Apply font style to all runs in a paragraph."""
        self._styles._apply_font_to_paragraph(paragraph, font_size)

    def _neutralize_table_style(self) -> None:
        """Remove all custom table styling to reset to default."""
        self._styles._neutralize_table_style()

    # Paragraphs handler delegation
    def _process_paragraphs(self) -> None:
        """Delegate to APAParagraphsHandler to process paragraphs."""
        self._paragraphs.process()

    def _apply_paragraph_spacing_control(self, p: Paragraph) -> None:
        """Apply widow/orphan control per APA 7."""
        self._paragraphs._apply_paragraph_spacing_control(p)

    def _apply_keep_with_next(self, p: Paragraph) -> None:
        """Apply keep-with-next to heading paragraphs per APA 7."""
        self._paragraphs._apply_keep_with_next(p)

    def _fix_citations(self, p: Paragraph) -> None:
        """Replace 'y' citations with '&' per APA 7."""
        self._paragraphs._fix_citations(p)

    def _format_toc_entry(self, p: Paragraph, heading_levels: dict[str, int]) -> None:
        """Format Table of Contents entries with correct indentation."""
        self._paragraphs._format_toc_entry(p, heading_levels)

    def _build_heading_level_map(self) -> dict[str, Any]:
        """Build a map of heading text to heading level."""
        return self._paragraphs._build_heading_level_map()

    def _format_lists(self) -> None:
        """Format lists with proper APA styling."""
        self._paragraphs.format_lists()

    def _apply_body_indent(self) -> None:
        """Apply first-line indent to body paragraphs per APA 7."""
        self._paragraphs.apply_body_indent()

    def _fix_text_spacing_global(self) -> None:
        """Fix text spacing issues across all paragraphs."""
        self._paragraphs.fix_text_spacing_global()

    def _merge_and_clean_paragraph(self, p: Paragraph) -> None:
        """Merge and clean paragraph formatting."""
        self._paragraphs._merge_and_clean_paragraph(p)

    # Tables handler delegation
    def _format_tables(self) -> None:
        """Delegate to APATablesHandler to format tables."""
        self._tables.format_tables()

    def _apply_apa_table_borders(self, table: TableObject) -> None:
        """Apply APA 7th Edition table borders (horizontal only)."""
        self._tables._apply_apa_table_borders(table)

    def _set_cell_border(self, cell: CT_Tc, **kwargs: Any) -> None:
        """Set cell border properties."""
        self._tables._set_cell_border(cell, **kwargs)

    def _add_table_captions(self) -> None:
        """Delegate to APATablesHandler to add table captions."""
        self._tables.add_table_captions()

    def _extract_table_title(self, table: TableObject) -> str:
        """Extract the table title from preceding paragraph."""
        return self._tables._extract_table_title(table)

    def _get_nearest_section_heading(self, table_pos: int) -> str:
        """Find the nearest section heading above the table."""
        return self._tables._get_nearest_section_heading(table_pos)

    def _add_table_notes(self) -> None:
        """Add notes below table per APA 7 format."""
        self._tables.add_table_notes()

    def _add_table_header_bold(self) -> None:
        """Make table header row bold."""
        self._tables.add_table_header_bold()

    # Figures handler delegation
    def _make_figure_paragraph(
        self, text: str, bold: bool = False, italic: bool = False, space_after: str = "0"
    ) -> Paragraph:
        """Create or modify a paragraph for figure placement."""
        result = self._figures._make_figure_paragraph(text, bold, italic, space_after)
        return cast(Paragraph, result)

    def _format_figures(self) -> None:
        """Delegate to APAFiguresHandler to format figures."""
        self._figures.format_figures()

    def _add_figure_captions(self) -> None:
        """Delegate to APAFiguresHandler to add figure captions."""
        self._figures.add_figure_captions()

    # Keywords handler delegation
    def _format_keywords(self, meta: DocumentMetadata) -> None:
        """Delegate to APAKeywordsHandler to format keywords section."""
        self._keywords.format_keywords(meta)

    def _add_page_break_before_introduction(self) -> None:
        """Add page break before the introduction section."""
        self._keywords._add_page_break_before_introduction()

    def _format_nota_italic(self) -> None:
        """Format 'Nota' paragraphs in italics per APA 7."""
        self._keywords.format_nota_italic()

    def _apply_foreign_word_italics(self) -> None:
        """Apply italics to foreign words per APA 7."""
        self._keywords.apply_foreign_word_italics()
