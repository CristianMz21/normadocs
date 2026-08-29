"""APA table formatting, borders, captions, and notes."""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Any, cast

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from ...config import DEFAULT_BODY_FONT, W_AFTER, W_LINE, W_LINE_RULE, W_SPACING, W_TYPE, W_VAL
from ...utils.docx_helpers import paragraph_style_name

if TYPE_CHECKING:
    from docx.document import Document as DocType
    from docx.table import Table as TableType
    from docx.table import _Cell as CellType


PAGE_CONTENT_WIDTH = 6.5

_W_VAL = W_VAL
_W_TYPE = W_TYPE
_W_SPACING = W_SPACING
_W_LINE = W_LINE
_W_LINE_RULE = W_LINE_RULE
_W_AFTER = W_AFTER
_W_PPR = "w:pPr"
_W_RPR = "w:rPr"

COMPANY_KEYWORDS = frozenset(["mackroph", "tecnoshop", "devsoft"])

_SOURCE_CAPTION_RE = re.compile(r"^(?:Tabla|Table|Cuadro)\s+(\d+)\s*[.:\u2014\u2013-]?\s*(.*)$")
_GRID_TABLE_RE = re.compile(r"[+|][-=]{3,}[+|]?")
_EQUALS_RE = re.compile(r"={3,}")
_CAMEL_SPLIT_RE = re.compile(r"([a-záéíóúñ])([A-ZÁÉÍÓÚÑ])")
_MONEY_SPLIT_RE = re.compile(r"(\$\d+,\d+,\d+)\s+(\d{1,2})")
_MULTI_SPACE_RE = re.compile(r"\s{2,}")


class APATablesHandler:
    """Handles table formatting, borders, captions, and notes per APA 7th Edition."""

    def __init__(self, doc: DocType, config: dict[str, Any] | None = None) -> None:
        """Initialize APATablesHandler.

        Args:
            doc: The python-docx Document object.
            config: Optional configuration dictionary.
        """
        self.doc = doc
        self.config = config if config is not None else {}

    def _get_table_config(self) -> dict[str, Any]:
        """Get table configuration from config with defaults."""
        default_config: dict[str, Any] = {
            "borders": "horizontal_only",
            "caption_prefix": "Table",
            "caption_above": True,
            "note_suffix": "Elaboración propia.",
            "vertical_align": "top",
        }
        return cast(dict[str, Any], self.config.get("tables", default_config))

    def _get_body_font(self) -> str:
        """Get body font name from config."""
        fonts: dict[str, Any] = {}
        return cast(
            str, self.config.get("fonts", fonts).get("body", {}).get("name", DEFAULT_BODY_FONT)
        )

    def _apply_font_style(
        self,
        element: Any,
        font_name: str | None = None,
        size: int = 12,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Apply font style to table elements.

        Args:
            element: The table or cell element to apply formatting to.
            font_name: Font name to apply (optional).
            size: Font size to apply (optional).
            bold: Whether to apply bold (default False).
            italic: Whether to apply italic (default False).
        """
        from .apa_styles import APAStylesHandler

        handler = APAStylesHandler(self.doc)
        handler._apply_font_style(element, font_name=font_name, size=size, bold=bold, italic=italic)

    def format_tables(self) -> None:
        """Apply APA borders and formatting to all tables."""
        for table in self.doc.tables:
            self._format_single_table(table)

    def _format_single_table(self, table: TableType) -> None:
        """Format a single table through the APA pipeline."""
        self._apply_apa_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._apply_table_layout(table)
        self._apply_cell_geometry(table)
        font_size = self._determine_font_size(table)
        col_widths = self._calc_col_widths(table, font_size)
        self._apply_col_widths(table, col_widths)
        self._configure_table_header(table)
        self._prevent_row_split(table)
        self._configure_table_look(table)
        self._clean_cells(table, font_size)
        self._apply_final_cell_formatting(table)
        self._add_table_spacing(table)

    def _apply_table_layout(self, table: TableType) -> None:
        """Set FIXED layout, width and cell margins for table."""
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            return
        self._set_fixed_layout(tbl_pr)
        self._set_table_width(tbl_pr)
        self._set_table_cell_margins(tbl_pr)

    def _set_fixed_layout(self, tbl_pr: Any) -> None:
        """Set tblLayout to fixed."""
        existing = tbl_pr.find(qn("w:tblLayout"))
        if existing is not None:
            tbl_pr.remove(existing)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn(_W_TYPE), "fixed")
        tbl_pr.append(layout)

    def _set_table_width(self, tbl_pr: Any) -> None:
        """Set table width to 100 percent."""
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn(_W_TYPE), "pct")
        tbl_w.set(qn("w:w"), "5000")

    def _set_table_cell_margins(self, tbl_pr: Any) -> None:
        """Set table-level cell margins."""
        existing = tbl_pr.find(qn("w:tblCellMar"))
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        for side in ("top", "bottom", "start", "end"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "57")
            el.set(qn(_W_TYPE), "dxa")
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)

    def _apply_cell_geometry(self, table: TableType) -> None:
        """Apply vertical alignment, margins and paragraph alignment to cells."""
        for row in table.rows:
            for cell in row.cells:
                self._format_single_cell_geometry(cell)

    def _format_single_cell_geometry(self, cell: CellType) -> None:
        """Format geometry for one cell."""
        tc_pr = self._ensure_tc_pr(cell)
        self._set_valign_top(tc_pr)
        self._set_cell_margins(tc_pr)
        self._set_cell_paragraphs_left(cell)
        self._ensure_no_wrap(tc_pr)

    def _ensure_tc_pr(self, cell: CellType) -> Any:
        """Ensure tcPr exists for cell."""
        tc_pr = cell._element.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = OxmlElement("w:tcPr")
            cell._element.insert(0, tc_pr)
        return tc_pr

    def _set_valign_top(self, tc_pr: Any) -> None:
        """Set vertical alignment to top."""
        existing = tc_pr.find(qn("w:vAlign"))
        if existing is not None:
            tc_pr.remove(existing)
        v_align = OxmlElement("w:vAlign")
        v_align.set(qn(_W_VAL), "top")
        tc_pr.append(v_align)

    def _set_cell_margins(self, tc_pr: Any) -> None:
        """Set small cell margins."""
        existing = tc_pr.find(qn("w:tcMar"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "bottom", "start", "end"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "28")
            el.set(qn(_W_TYPE), "dxa")
            tc_mar.append(el)
        tc_pr.append(tc_mar)

    def _set_cell_paragraphs_left(self, cell: CellType) -> None:
        """Left-align cell paragraphs and set overflow."""
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_pr = p._element.get_or_add_pPr()
            jc = p_pr.get_or_add_jc()
            jc.set(qn(_W_VAL), "left")
            overflow = OxmlElement("w:overflow")
            overflow.set(qn(_W_VAL), "continue")
            p_pr.append(overflow)

    def _ensure_no_wrap(self, tc_pr: Any) -> None:
        """Ensure noWrap is present."""
        if tc_pr.find(qn("w:noWrap")) is None:
            tc_pr.append(OxmlElement("w:noWrap"))

    def _determine_font_size(self, table: TableType) -> int:
        """Determine font size based on column count."""
        num_cols = len(table.columns)
        if num_cols >= 8:
            return 9
        if num_cols >= 6:
            return 10
        return 12

    def _calc_col_widths(self, table: TableType, font_size: int) -> list[float]:
        """Calculate column widths in inches."""
        num_cols = len(table.columns)
        if num_cols < 2:
            return []
        max_content_len, max_word_len = self._collect_column_metrics(table, num_cols)
        cw = {9: 0.08, 10: 0.09, 11: 0.10, 12: 0.10}.get(font_size, 0.09)
        min_col = max(0.8, 6.0 / num_cols)
        min_widths = [max(w * cw + 0.08, min_col) for w in max_word_len]
        return self._distribute_widths(max_content_len, min_widths, min_col)

    def _collect_column_metrics(
        self, table: TableType, num_cols: int
    ) -> tuple[list[int], list[int]]:
        """Collect max content and word lengths per column."""
        max_content_len = [0] * num_cols
        max_word_len = [0] * num_cols
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= num_cols:
                    continue
                text = cell.text.strip()
                max_content_len[ci] = max(max_content_len[ci], len(text))
                for word in text.split():
                    max_word_len[ci] = max(max_word_len[ci], len(word))
        return max_content_len, max_word_len

    def _distribute_widths(
        self, max_content_len: list[int], min_widths: list[float], min_col: float
    ) -> list[float]:
        """Distribute remaining width proportionally and normalize."""
        total_min = sum(min_widths)
        avail = PAGE_CONTENT_WIDTH
        num_cols = len(max_content_len)
        col_widths = self._initial_widths(max_content_len, min_widths, total_min, avail, min_col)
        return self._normalize_widths(col_widths, avail, num_cols)

    def _initial_widths(
        self,
        max_content_len: list[int],
        min_widths: list[float],
        total_min: float,
        avail: float,
        min_col: float,
    ) -> list[float]:
        """Compute initial widths before normalization."""
        if total_min <= avail:
            return self._widths_with_remaining(max_content_len, min_widths, avail, total_min)
        return self._widths_proportional(max_content_len, avail, min_col)

    def _widths_with_remaining(
        self, max_content_len: list[int], min_widths: list[float], avail: float, total_min: float
    ) -> list[float]:
        """Distribute remaining space proportionally."""
        remaining = avail - total_min
        total_content = sum(max_content_len) or 1
        widths: list[float] = []
        for ci in range(len(max_content_len)):
            extra = remaining * (max_content_len[ci] / total_content)
            widths.append(min_widths[ci] + extra)
        return widths

    def _widths_proportional(
        self, max_content_len: list[int], avail: float, min_col: float
    ) -> list[float]:
        """Proportional fallback with floor."""
        total_content = sum(max_content_len) or 1
        widths: list[float] = []
        for ci in range(len(max_content_len)):
            proportion = max_content_len[ci] / total_content
            widths.append(max(avail * proportion, min_col))
        return widths

    def _normalize_widths(
        self, col_widths: list[float], avail: float, _num_cols: int
    ) -> list[float]:
        """Normalize total to exactly PAGE_CONTENT_WIDTH."""
        current_total = sum(col_widths)
        if current_total > 0 and abs(current_total - avail) > 0.01:
            scale = avail / current_total
            return [w * scale for w in col_widths]
        return col_widths

    def _apply_col_widths(self, table: TableType, col_widths: list[float]) -> None:
        """Apply widths to columns, gridCols and cells."""
        if not col_widths:
            return
        tbl_grid = table._tbl.find(qn("w:tblGrid"))
        grid_cols = tbl_grid.findall(qn("w:gridCol")) if tbl_grid is not None else []
        for ci, col in enumerate(table.columns):
            width = Inches(col_widths[ci])
            col.width = width
            if ci < len(grid_cols):
                grid_cols[ci].set(qn("w:w"), str(int(col_widths[ci] * 1440)))
            for row in table.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = width

    def _configure_table_header(self, table: TableType) -> None:
        """Repeat header row across pages."""
        if not table.rows:
            return
        tr = table.rows[0]._tr
        tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    def _prevent_row_split(self, table: TableType) -> None:
        """Prevent rows from splitting across pages."""
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            existing = tr_pr.find(qn("w:cantSplit"))
            if existing is not None:
                tr_pr.remove(existing)
            cant_split = OxmlElement("w:cantSplit")
            cant_split.set(qn(_W_VAL), "1")
            tr_pr.append(cant_split)

    def _configure_table_look(self, table: TableType) -> None:
        """Add table look and split properties."""
        tbl_pr_elem = table._tbl.tblPr
        if tbl_pr_elem is None:
            tbl_pr_elem = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr_elem)
        self._set_tbl_look(tbl_pr_elem)
        self._set_tbl_split(tbl_pr_elem)

    def _set_tbl_look(self, tbl_pr_elem: Any) -> None:
        """Set tblLook element."""
        tbl_look = OxmlElement("w:tblLook")
        tbl_look.set(qn(_W_VAL), "04A0")
        tbl_look.set(qn("w:first"), "1")
        tbl_look.set(qn("w:last"), "1")
        tbl_look.set(qn("w:hBand"), "1")
        tbl_look.set(qn("w:vBand"), "1")
        tbl_pr_elem.append(tbl_look)

    def _set_tbl_split(self, tbl_pr_elem: Any) -> None:
        """Set tblSplit to prevent row splitting."""
        existing = tbl_pr_elem.find(qn("w:tblSplit"))
        if existing is not None:
            tbl_pr_elem.remove(existing)
        tbl_split = OxmlElement("w:tblSplit")
        tbl_split.set(qn(_W_VAL), "0")
        tbl_pr_elem.append(tbl_split)

    def _clean_cells(self, table: TableType, font_size: int) -> None:
        """Clean and merge cell text for all cells."""
        for row in table.rows:
            for cell in row.cells:
                self._clean_single_cell(cell, row, table, font_size)

    def _clean_single_cell(
        self, cell: CellType, row: Any, table: TableType, font_size: int
    ) -> None:
        """Clean merged text for a single cell."""
        merged = self._collect_merged_text(cell)
        merged = self._normalize_merged_text(merged)
        self._apply_cleaned_text(cell, row, table, merged, font_size)

    def _collect_merged_text(self, cell: CellType) -> str:
        """Collect all paragraph texts in cell into single string."""
        cell_texts: list[str] = []
        for p in cell.paragraphs:
            para_text = ""
            for run in p.runs:
                t = run.text or ""
                t = _GRID_TABLE_RE.sub("", t)
                t = _EQUALS_RE.sub("", t)
                t = t.strip("|").strip()
                para_text += f" {t}" if para_text else t
            para_text = para_text.strip()
            if para_text:
                cell_texts.append(para_text)
        return " ".join(cell_texts)

    def _normalize_merged_text(self, merged: str) -> str:
        """Apply text normalizations to merged cell text."""
        merged = merged.replace("**", "").replace("*", "")
        merged = _CAMEL_SPLIT_RE.sub(r"\1 \2", merged)
        merged = self._fix_technology_names(merged)
        merged = self._fix_split_words(merged)
        merged = _MONEY_SPLIT_RE.sub(r"\1\2", merged)
        return _MULTI_SPACE_RE.sub(" ", merged).strip()

    def _fix_technology_names(self, text: str) -> str:
        """Fix incorrectly split technology names."""
        text = text.replace("i Phone", "iPhone").replace("i OS", "iOS")
        return text.replace("Whats App", "WhatsApp").replace("DDo S", "DDoS")

    def _fix_split_words(self, text: str) -> str:
        """Fix specific split words and phrases."""
        text = text.replace("Java Script", "JavaScript").replace("Postgre SQL", "PostgreSQL")
        text = text.replace("Dedicació n", "Dedicación")
        text = text.replace("REQUERIMIENTO S", "REQUERIMIENTOS")
        return text.replace("REQUERIMIENT S", "REQUERIMIENTOS")

    def _apply_cleaned_text(
        self, cell: CellType, row: Any, table: TableType, merged: str, font_size: int
    ) -> None:
        """Clear cell and apply cleaned text with formatting."""
        if not cell.paragraphs:
            return
        first_p = cell.paragraphs[0]
        is_header = row == table.rows[0]
        first_p.clear()
        new_run = first_p.add_run(merged)
        self._apply_font_style(new_run, size=font_size)
        if is_header:
            new_run.bold = True
        for extra_p in list(cell.paragraphs[1:]):
            parent = extra_p._element.getparent()
            if parent is not None:
                parent.remove(extra_p._element)

    def _apply_final_cell_formatting(self, table: TableType) -> None:
        """Apply left alignment and spacing to all cell paragraphs."""
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._format_cell_paragraph(p)

    def _format_cell_paragraph(self, p: Any) -> None:
        """Apply left, single spacing and widow controls to paragraph."""
        p_pr = p._element.get_or_add_pPr()
        self._set_paragraph_alignment(p_pr)
        self._set_paragraph_spacing(p_pr)
        self._set_widow_keep(p_pr)

    def _set_paragraph_alignment(self, p_pr: Any) -> None:
        """Set left alignment."""
        old_jc = p_pr.find(qn("w:jc"))
        if old_jc is not None:
            p_pr.remove(old_jc)
        jc = OxmlElement("w:jc")
        jc.set(qn(_W_VAL), "left")
        p_pr.append(jc)

    def _set_paragraph_spacing(self, p_pr: Any) -> None:
        """Set single spacing."""
        old = p_pr.find(qn("w:spacing"))
        if old is not None:
            p_pr.remove(old)
        spacing = OxmlElement(_W_SPACING)
        spacing.set(qn(_W_LINE), "240")
        spacing.set(qn(_W_LINE_RULE), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn(_W_AFTER), "40")
        p_pr.append(spacing)

    def _set_widow_keep(self, p_pr: Any) -> None:
        """Set widow and keepLines controls."""
        old = p_pr.find(qn("w:widowControl"))
        if old is not None:
            p_pr.remove(old)
        p_pr.append(OxmlElement("w:widowControl"))
        old_keep = p_pr.find(qn("w:keepLines"))
        if old_keep is not None:
            p_pr.remove(old_keep)
        p_pr.append(OxmlElement("w:keepLines"))

    def _add_table_spacing(self, table: TableType) -> None:
        """Add spacing paragraph after table."""
        spacing_p = OxmlElement("w:p")
        spacing_p_pr = OxmlElement(_W_PPR)
        spacing = OxmlElement(_W_SPACING)
        spacing.set(qn(_W_LINE), "480")
        spacing.set(qn(_W_LINE_RULE), "auto")
        spacing_p_pr.append(spacing)
        spacing_p.append(spacing_p_pr)
        table._tbl.addnext(spacing_p)

    def _apply_apa_table_borders(self, table: TableType) -> None:
        """
        Apply APA-style borders:
        - Top/Bottom of table: single line
        - Bottom of header row: single line
        - No vertical lines.
        """
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(
                    cell, top={}, bottom={}, start={}, end={}, insideH={}, insideV={}
                )
        num_rows = len(table.rows)
        for i, row in enumerate(table.rows):
            is_first = i == 0
            is_last = i == num_rows - 1
            for cell in row.cells:
                borders: dict[str, Any] = {}
                if is_first:
                    borders["top"] = {"val": "single", "sz": "12", "color": "auto"}
                    borders["bottom"] = {"val": "single", "sz": "6", "color": "auto"}
                if is_last:
                    borders["bottom"] = {"val": "single", "sz": "12", "color": "auto"}
                self._set_cell_border(cell, **borders)

    def _set_cell_border(self, cell: CellType, **kwargs: Any) -> None:
        """Set border on a table cell (OpenXML). Clears existing first."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(old)
        tc_borders = OxmlElement("w:tcBorders")
        for edge_name in ("start", "top", "end", "bottom", "insideH", "insideV"):
            if edge_name in kwargs:
                el = OxmlElement(f"w:{edge_name}")
                for attr, val in kwargs[edge_name].items():
                    el.set(qn(f"w:{attr}"), str(val))
                tc_borders.append(el)
        tc_pr.append(tc_borders)

    def add_table_captions(self) -> None:
        """Add APA 7 captions to tables: 'Tabla N' (bold) + title (italic)."""
        body = self.doc._body._element
        table_positions = [
            (pos, child) for pos, child in enumerate(list(body)) if child.tag == qn("w:tbl")
        ]
        offset = 0
        max_used = 0
        for _idx, (orig_pos, tbl) in enumerate(table_positions):
            current_pos = orig_pos + offset
            source = self._extract_source_caption(current_pos)
            source_title = ""
            if source is not None:
                source_el, source_num, source_title = source
                body.remove(source_el)
                offset -= 1
                current_pos -= 1
                max_used = max(max_used, source_num)
                table_num = source_num
            else:
                max_used += 1
                table_num = max_used
            title_text = self._resolve_table_title(current_pos, tbl, source_title)
            caption_p = self._build_caption_paragraph(table_num)
            body.insert(current_pos, caption_p)
            offset += 1
            if title_text:
                title_p = self._build_title_paragraph(title_text)
                body.insert(current_pos + 1, title_p)
                offset += 1

    def _resolve_table_title(self, current_pos: int, tbl: Any, source_title: str) -> str:
        """Resolve caption title from source, heading or table content."""
        if source_title:
            return source_title
        title = self._get_nearest_section_heading(current_pos)
        if title:
            return title
        from docx.table import Table

        return self._extract_table_title(Table(tbl, self.doc))

    def _build_caption_paragraph(self, table_num: int) -> Any:
        """Build 'Tabla N' bold caption paragraph."""
        caption_p = OxmlElement("w:p")
        caption_p_pr = OxmlElement(_W_PPR)
        jc = OxmlElement("w:jc")
        jc.set(qn(_W_VAL), "left")
        caption_p_pr.append(jc)
        spacing = OxmlElement(_W_SPACING)
        spacing.set(qn(_W_AFTER), "0")
        spacing.set(qn(_W_LINE), "240")
        spacing.set(qn(_W_LINE_RULE), "auto")
        caption_p_pr.append(spacing)
        caption_p.append(caption_p_pr)
        run = OxmlElement("w:r")
        r_pr = OxmlElement(_W_RPR)
        r_pr.append(OxmlElement("w:b"))
        font = OxmlElement("w:rFonts")
        font.set(qn("w:ascii"), DEFAULT_BODY_FONT)
        font.set(qn("w:hAnsi"), DEFAULT_BODY_FONT)
        r_pr.append(font)
        sz = OxmlElement("w:sz")
        sz.set(qn(_W_VAL), "24")
        r_pr.append(sz)
        run.append(r_pr)
        t = OxmlElement("w:t")
        prefix = cast(str, self._get_table_config().get("caption_prefix", "Table"))
        t.text = f"{prefix} {table_num}"
        run.append(t)
        caption_p.append(run)
        return caption_p

    def _build_title_paragraph(self, title_text: str) -> Any:
        """Build italic title paragraph for table."""
        title_p = OxmlElement("w:p")
        title_p_pr = OxmlElement(_W_PPR)
        jc2 = OxmlElement("w:jc")
        jc2.set(qn(_W_VAL), "left")
        title_p_pr.append(jc2)
        spacing2 = OxmlElement(_W_SPACING)
        spacing2.set(qn(_W_AFTER), "120")
        spacing2.set(qn(_W_LINE), "240")
        spacing2.set(qn(_W_LINE_RULE), "auto")
        title_p_pr.append(spacing2)
        title_p.append(title_p_pr)
        title_run = OxmlElement("w:r")
        title_r_pr = OxmlElement(_W_RPR)
        title_r_pr.append(OxmlElement("w:i"))
        font2 = OxmlElement("w:rFonts")
        font2.set(qn("w:ascii"), DEFAULT_BODY_FONT)
        font2.set(qn("w:hAnsi"), DEFAULT_BODY_FONT)
        title_r_pr.append(font2)
        sz2 = OxmlElement("w:sz")
        sz2.set(qn(_W_VAL), "24")
        title_r_pr.append(sz2)
        title_run.append(title_r_pr)
        title_t = OxmlElement("w:t")
        title_t.text = title_text
        title_t.set(qn("xml:space"), "preserve")
        title_run.append(title_t)
        title_p.append(title_run)
        return title_p

    def _extract_source_caption(self, pos: int) -> tuple[Any, int, str] | None:
        """Find a Markdown-source caption immediately before a table."""
        from docx.text.paragraph import Paragraph

        body = self.doc._body._element
        children = list(body)
        for i in (pos - 1, pos - 2):
            if i < 0 or i >= len(children):
                continue
            el = children[i]
            if el.tag != qn("w:p"):
                continue
            p = Paragraph(el, self.doc)
            text = p.text.strip()
            if not text:
                continue
            match = _SOURCE_CAPTION_RE.match(text)
            if match is None:
                return None
            style_name = paragraph_style_name(p)
            if style_name.startswith("Heading"):
                return None
            return (el, int(match.group(1)), match.group(2).strip())
        return None

    def _extract_table_title(self, table: TableType) -> str:
        """Extract a descriptive title from the table content."""
        try:
            if len(table.rows) < 2 or len(table.columns) < 2:
                return ""
            first_row = table.rows[0]
            first_cell = first_row.cells[0].text.strip()
            if self._is_short_header(first_cell):
                return self._title_from_short_header(first_row)
            if first_cell and 3 < len(first_cell) < 80:
                return first_cell
            return self._title_from_combined_cells(first_row, first_cell)
        except Exception:
            return ""

    def _is_short_header(self, first_cell: str) -> bool:
        """Check if first cell is a short header indicator."""
        short_headers = ["n°", "no.", "campo", "nombre", "característica", "concepto", "rubro"]
        return first_cell.lower() in short_headers or first_cell.lower().startswith("tabla")

    def _title_from_short_header(self, first_row: Any) -> str:
        """Get title from second cell when first is short header."""
        if len(first_row.cells) <= 1:
            return ""
        second = cast(str, first_row.cells[1].text.strip())
        if second and len(second) < 80:
            return second
        return ""

    def _title_from_combined_cells(self, first_row: Any, first_cell: str) -> str:
        """Combine first two cells for title."""
        if len(first_row.cells) < 2:
            return ""
        second_text = cast(str, first_row.cells[1].text.strip())
        if first_cell and second_text:
            combined = f"{first_cell} - {second_text}"
            if len(combined) < 100:
                return combined
        return ""

    def _get_nearest_section_heading(self, table_pos: int) -> str:
        """Find the nearest section heading before a table."""
        body = self.doc._body._element
        body_children = list(body)
        for i in range(table_pos - 1, -1, -1):
            if i >= len(body_children):
                continue
            elem = body_children[i]
            if elem.tag == qn("w:p"):
                heading = self._heading_from_element(elem)
                if heading is not None:
                    return heading
            if elem.tag == qn("w:tbl"):
                break
        return ""

    def _heading_from_element(self, elem: Any) -> str | None:
        """Extract heading text from element if valid."""
        p_idx_map = {p._element: idx for idx, p in enumerate(self.doc.paragraphs)}
        if elem not in p_idx_map:
            return None
        p = self.doc.paragraphs[p_idx_map[elem]]
        text = p.text.strip()
        if re.match(r"^(Tabla|Figura)\s+\d+", text):
            return None
        style_name = paragraph_style_name(p)
        if not style_name.startswith("Heading"):
            return None
        cleaned = self._strip_leading_numbers(text)
        return cleaned if cleaned else None

    def _strip_leading_numbers(self, text: str) -> str:
        """Strip leading chapter numbers like '2.2 ' without super-linear regex."""
        if not text or not text[0].isdigit():
            return text
        idx = 0
        while idx < len(text) and (text[idx].isdigit() or text[idx] == "."):
            idx += 1
        return text[idx:].lstrip() if idx > 0 else text

    def add_table_notes(self) -> None:
        """Add table notes after each table (APA 7 requirement)."""
        tables = list(self.doc.tables)
        descriptions = self._build_all_descriptions(tables)
        self._insert_notes(tables, descriptions)

    def _build_all_descriptions(self, tables: list[TableType]) -> list[str]:
        """Build note descriptions for all tables."""
        return [self._describe_table(t) for t in tables]

    def _describe_table(self, table: TableType) -> str:
        """Describe table based on its header content."""
        if not table.rows:
            return ""
        first_row_text = " ".join(cell.text.strip().lower() for cell in table.rows[0].cells)
        all_text = " ".join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
        if self._is_caracteristica(first_row_text):
            return self._desc_caracteristica(table, first_row_text)
        if self._is_software_version(first_row_text):
            return self._desc_software(first_row_text)
        if self._is_servicio_proveedor(first_row_text):
            return "Infraestructura en la nube y servicios externos planificados."
        if self._is_rubro_costo(first_row_text):
            return self._desc_rubro(all_text)
        if self._is_componente_porcentaje(first_row_text):
            return "Componentes del modelo de costos AIU (Administración, Imprevistos, Utilidad)."
        if self._is_rol_dedicacion(first_row_text):
            return "Estructura del equipo de desarrollo con roles y dedicación temporal."
        if self._is_concepto_valor(first_row_text):
            return self._desc_concepto(all_text)
        if self._is_numerado(first_row_text):
            return self._desc_numerado(all_text)
        return ""

    def _is_caracteristica(self, text: str) -> bool:
        return "característica" in text and "especificación" in text

    def _is_software_version(self, text: str) -> bool:
        return "software" in text and "versión" in text

    def _is_servicio_proveedor(self, text: str) -> bool:
        return "servicio" in text and "proveedor" in text

    def _is_rubro_costo(self, text: str) -> bool:
        return "rubro" in text and "costo" in text

    def _is_componente_porcentaje(self, text: str) -> bool:
        return "componente" in text and "porcentaje" in text

    def _is_rol_dedicacion(self, text: str) -> bool:
        return "rol" in text and "dedicación" in text

    def _is_concepto_valor(self, text: str) -> bool:
        return "concepto" in text and "valor" in text

    def _is_numerado(self, text: str) -> bool:
        return "n°" in text or "º" in text

    def _desc_caracteristica(self, table: TableType, first_row_text: str) -> str:
        """Description for característica/especificación tables."""
        if "procesador" not in first_row_text or "memoria ram" not in first_row_text:
            return "Características técnicas del componente especificado."
        return self._desc_hardware_detail(table)

    def _desc_hardware_detail(self, table: TableType) -> str:
        """Detail for hardware spec tables with procesador/memoria."""
        if not table.rows or len(table.rows) <= 1:
            return "Especificaciones técnicas del equipo de hardware."
        first_cell = table.rows[1].cells[0].text.strip().lower() if table.rows[1].cells else ""
        return self._map_hardware_second_cell(first_cell)

    def _map_hardware_second_cell(self, second_cell: str) -> str:
        """Map second cell keyword to description."""
        if "desarrollador" in second_cell or "computador" in second_cell:
            return "Especificaciones técnicas del computador para desarrollo."
        if "analista" in second_cell or "qa" in second_cell:
            return "Especificaciones técnicas del computador para pruebas de calidad."
        if "servidor" in second_cell or "vps" in second_cell:
            return "Configuración técnica del servidor virtual privado."
        if "móvil" in second_cell or "smartphone" in second_cell or "iphone" in second_cell:
            return "Dispositivos móviles seleccionados para pruebas de compatibilidad."
        return "Especificaciones técnicas del equipo de hardware."

    def _desc_software(self, first_row_text: str) -> str:
        """Description for software tables."""
        if "función" in first_row_text or "licencia" in first_row_text:
            return self._desc_software_funcion(first_row_text)
        if "tipo" in first_row_text and "costo" in first_row_text:
            return self._desc_software_costo(first_row_text)
        return "Software de desarrollo requerido con versiones compatibles."

    def _desc_software_funcion(self, text: str) -> str:
        """Software with función/licencia."""
        if "django" in text or "react" in text or "framework" in text:
            return "Frameworks y bibliotecas de desarrollo seleccionados para el proyecto."
        if "python" in text or "javascript" in text:
            return "Lenguajes de programación utilizados en el desarrollo."
        return "Software de desarrollo requerido con versiones compatibles."

    def _desc_software_costo(self, text: str) -> str:
        """Software with tipo/costo."""
        if "postgresql" in text or "mysql" in text or "redis" in text:
            return "Sistemas de gestión de bases de datos seleccionados."
        return "Software y servicios con costos asociados."

    def _desc_rubro(self, all_text: str) -> str:
        """Description for rubro tables."""
        if "herramientas" in all_text or "licencias" in all_text:
            return "Resumen de inversiones en software y servicios del proyecto."
        if "hardware" in all_text or "servidor" in all_text:
            return "Resumen de inversión en equipos y servicios de infraestructura."
        return "Resumen de inversiones en software y servicios del proyecto."

    def _desc_concepto(self, all_text: str) -> str:
        """Description for concepto tables."""
        if "iva" in all_text:
            return "Resumen del presupuesto total del proyecto con impuestos aplicables."
        if "total costos directos" in all_text and "aiu" not in all_text:
            return "Desglose de costos directos del proyecto de desarrollo."
        if "aiu" in all_text or "subtotal" in all_text:
            return "Componentes del presupuesto con aplicación del modelo AIU."
        return "Conceptos y valores del presupuesto del proyecto."

    def _desc_numerado(self, all_text: str) -> str:
        """Description for n° numbered tables."""
        if COMPANY_KEYWORDS.intersection(all_text.split()):
            return self._desc_company_table(all_text)
        if "backend" in all_text or "frontend" in all_text or "panel" in all_text:
            return "Desglose de costos por módulo de desarrollo y servicio."
        if "hardware" in all_text:
            return "Resumen de inversión en equipos de hardware."
        return "Detalle de costos y Rubros del presupuesto."

    def _desc_company_table(self, all_text: str) -> str:
        """Description for company comparison tables."""
        if "criterio" in all_text or "peso" in all_text or "funcional" in all_text:
            return "Matriz de evaluación técnica y ponderación de criterios por proveedor."
        return "Comparación de costos directos entre proveedores evaluados."

    def _insert_notes(self, tables: list[TableType], descriptions: list[str]) -> None:
        """Insert nota paragraphs after each table."""
        for idx, table in enumerate(tables):
            parent = table._tbl.getparent()
            if parent is None:
                continue
            table_idx = parent.index(table._tbl)
            desc = descriptions[idx] if idx < len(descriptions) else ""
            nota_p = self._build_nota_paragraph(desc)
            parent.insert(table_idx + 1, nota_p)

    def _build_nota_paragraph(self, desc: str) -> Any:
        """Build Nota paragraph element."""
        nota_p = OxmlElement("w:p")
        nota_p_pr = OxmlElement(_W_PPR)
        nota_spacing = OxmlElement(_W_SPACING)
        nota_spacing.set(qn(_W_AFTER), "0")
        nota_p_pr.append(nota_spacing)
        nota_p.append(nota_p_pr)
        nota_p.append(self._build_nota_run1())
        nota_p.append(self._build_nota_run2(desc))
        nota_p.append(self._build_nota_run3())
        return nota_p

    def _build_nota_run1(self) -> Any:
        """Build italic 'Nota. ' run."""
        r1 = OxmlElement("w:r")
        r_pr1 = OxmlElement(_W_RPR)
        r_pr1.append(OxmlElement("w:i"))
        sz1 = OxmlElement("w:sz")
        sz1.set(qn(_W_VAL), "24")
        r_pr1.append(sz1)
        r1.append(r_pr1)
        t1 = OxmlElement("w:t")
        t1.text = "Nota. "
        r1.append(t1)
        return r1

    def _build_nota_run2(self, desc: str) -> Any:
        """Build description run."""
        r2 = OxmlElement("w:r")
        r_pr2 = OxmlElement(_W_RPR)
        sz2 = OxmlElement("w:sz")
        sz2.set(qn(_W_VAL), "24")
        r_pr2.append(sz2)
        r2.append(r_pr2)
        t2 = OxmlElement("w:t")
        t2.text = desc
        r2.append(t2)
        return r2

    def _build_nota_run3(self) -> Any:
        """Build suffix run."""
        r3 = OxmlElement("w:r")
        r_pr3 = OxmlElement(_W_RPR)
        sz3 = OxmlElement("w:sz")
        sz3.set(qn(_W_VAL), "24")
        r_pr3.append(sz3)
        r3.append(r_pr3)
        t3 = OxmlElement("w:t")
        suffix = cast(str, self._get_table_config().get("note_suffix", " Author's elaboration."))
        t3.text = f" {suffix}"
        r3.append(t3)
        return r3

    def add_table_header_bold(self) -> None:
        """Set bold on all table header rows (APA 7 requirement)."""
        for table in self.doc.tables:
            for row in table.rows[:1]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
