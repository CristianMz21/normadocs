"""Tables verification for APA 7th Edition.

Verifies table formatting meets APA 7th Edition requirements:
- Table caption: "Table N" bold + title italic, positioned ABOVE table
- Table borders: Horizontal only (no vertical borders)
- Table note: "Nota." italic, positioned BELOW table
- Vertical alignment: Top
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Any, TypedDict

from docx.oxml.ns import qn

from ...utils.docx_helpers import paragraph_style_name
from .. import CheckCategory, VerificationIssue
from ..docx_analyzer import DOCXParagraphInfo

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

    from ..apa_verifier import VerificationContext


class TableCaption(TypedDict):
    """Typed dict for table caption data."""

    text: str
    index: int
    paragraph_info: DOCXParagraphInfo


_CAPTION_RE = re.compile(r"^(Tabla|Table|Figura|Figure)\s+\d+")
_NOTA_RE = re.compile(r"^Not[ae]\b")


class TablesCheck:
    """Check table formatting against APA 7th Edition requirements."""

    def run(self, ctx: VerificationContext) -> list[VerificationIssue]:
        """Run tables verification.

        Args:
            ctx: Verification context with access to PDF and DOCX analyzers.

        Returns:
            List of verification issues found.
        """
        issues: list[VerificationIssue] = []
        paragraphs_info = ctx.docx.get_paragraphs_info()
        tables_info = ctx.docx.get_tables_info()
        table_numbers = self._collect_table_captions(paragraphs_info)
        self._check_all_tables(tables_info, table_numbers, paragraphs_info, ctx, issues)
        self._check_numbering_sequence(table_numbers, issues)
        self._check_table_notes(ctx, issues)
        return issues

    def _collect_table_captions(
        self, paragraphs_info: list[DOCXParagraphInfo]
    ) -> list[TableCaption]:
        table_numbers: list[TableCaption] = []
        for i, p_info in enumerate(paragraphs_info):
            text = p_info.text.strip()
            if not text.startswith(("Table ", "Tabla ")):
                continue
            parts = text.split()
            if len(parts) >= 2 and parts[1].replace(".", "").isdigit():
                table_numbers.append(
                    {
                        "text": text,
                        "index": i,
                        "paragraph_info": p_info,
                    }
                )
        return table_numbers

    def _check_all_tables(
        self,
        tables_info: list[Any],
        table_numbers: list[TableCaption],
        paragraphs_info: list[DOCXParagraphInfo],
        ctx: VerificationContext,
        issues: list[VerificationIssue],
    ) -> None:
        for idx in range(len(tables_info)):
            table_has_caption = idx < len(table_numbers)
            if table_has_caption:
                caption_data = table_numbers[idx]
                self._check_caption_bold(caption_data, idx, ctx, issues)
                self._check_caption_italic(caption_data, paragraphs_info, idx, ctx, issues)
            if not table_has_caption:
                self._check_caption_present(idx, ctx, issues)
            if ctx.strict and table_has_caption:
                self._check_caption_position(idx, table_numbers, ctx, issues)
            if ctx.strict:
                self._check_vertical_borders(idx, ctx, issues)

    def _check_caption_bold(
        self,
        caption_data: TableCaption,
        idx: int,
        ctx: VerificationContext,
        issues: list[VerificationIssue],
    ) -> None:
        runs = caption_data["paragraph_info"].runs
        has_bold = any(run.get("bold") for run in runs)
        all_bold = all(run.get("bold") for run in runs) if runs else False
        if has_bold and (not ctx.strict or all_bold):
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.caption_bold",
                severity="error",
                expected="'Table N' in bold",
                actual="Caption not bold",
                evidence=f"Table {idx + 1} caption lacks bold formatting",
            )
        )

    def _check_caption_italic(
        self,
        caption_data: TableCaption,
        paragraphs_info: list[DOCXParagraphInfo],
        idx: int,
        ctx: VerificationContext,
        issues: list[VerificationIssue],
    ) -> None:
        caption_idx = caption_data["index"]
        has_italic = self._has_italic_title(caption_idx, paragraphs_info)
        if has_italic:
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.caption_italic",
                severity="error" if ctx.strict else "warning",
                expected="Title should be italic (in paragraph after 'Tabla N')",
                actual="Title not italic",
                evidence=f"Table {idx + 1} caption title should be italic",
            )
        )

    def _has_italic_title(self, caption_idx: int, paragraphs_info: list[DOCXParagraphInfo]) -> bool:
        if caption_idx + 1 >= len(paragraphs_info):
            return False
        next_para = paragraphs_info[caption_idx + 1]
        next_text = next_para.text.strip()
        if not next_text:
            return False
        if _CAPTION_RE.match(next_text):
            return False
        if next_text.startswith("Nota."):
            return False
        return any(run.get("italic") for run in next_para.runs)

    def _check_caption_present(
        self, idx: int, ctx: VerificationContext, issues: list[VerificationIssue]
    ) -> None:
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.caption_present",
                severity="error" if ctx.strict else "warning",
                expected="Table caption above table",
                actual="No caption found",
                evidence=f"Table {idx + 1} lacks a proper table caption",
            )
        )

    def _check_caption_position(
        self,
        idx: int,
        table_numbers: list[TableCaption],
        ctx: VerificationContext,
        issues: list[VerificationIssue],
    ) -> None:
        table_element = ctx.docx.tables[idx]._tbl
        caption_element = ctx.docx.paragraphs[table_numbers[idx]["index"]]._element
        body_children = list(ctx.docx.doc._body._element)
        if body_children.index(caption_element) <= body_children.index(table_element):
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.caption_position",
                severity="error",
                expected="Caption before table",
                actual="Caption appears after table",
                evidence=f"Table {idx + 1} caption is not above the table",
            )
        )

    def _check_vertical_borders(
        self, idx: int, ctx: VerificationContext, issues: list[VerificationIssue]
    ) -> None:
        table_element = ctx.docx.tables[idx]._tbl
        properties = table_element.tblPr
        borders = properties.find(qn("w:tblBorders")) if properties is not None else None
        if borders is None:
            return
        vertical_edges = ("left", "right", "insideV")
        has_vertical = any(
            (edge := borders.find(qn(f"w:{name}"))) is not None
            and edge.get(qn("w:val")) not in {None, "nil", "none"}
            for name in vertical_edges
        )
        if not has_vertical:
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.vertical_borders",
                severity="error",
                expected="Horizontal borders only",
                actual="Vertical table border detected",
                evidence=f"Table {idx + 1} contains a vertical border",
            )
        )

    def _check_numbering_sequence(
        self, table_numbers: list[TableCaption], issues: list[VerificationIssue]
    ) -> None:
        """Verify table numbers run 1..N without gaps or duplicates."""
        numbers: list[int] = []
        for caption_data in table_numbers:
            parts = caption_data["text"].split()
            if len(parts) >= 2 and parts[1].rstrip(".").isdigit():
                numbers.append(int(parts[1].rstrip(".")))
        if not numbers:
            return
        if sorted(numbers) == list(range(1, len(numbers) + 1)):
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.numbering_sequence",
                severity="error",
                expected=f"Sequential numbering 1..{len(numbers)}",
                actual=f"Table numbers found: {numbers}",
                evidence="Table numbers must be consecutive starting at 1",
            )
        )

    def _check_table_notes(self, ctx: VerificationContext, issues: list[VerificationIssue]) -> None:
        """Verify each table has an APA-formatted 'Nota.' below it."""
        body_children = list(ctx.docx.doc._body._element)
        for t_idx, table in enumerate(ctx.docx.tables):
            note_para = self._find_note_for_table(t_idx, table, body_children, ctx)
            if note_para is None:
                self._report_missing_note(t_idx, issues)
                continue
            self._validate_note_format(note_para, t_idx, issues)

    def _find_note_for_table(
        self,
        t_idx: int,
        table: Any,
        body_children: list[Any],
        ctx: VerificationContext,
    ) -> Paragraph | None:
        from docx.text.paragraph import Paragraph

        tbl = getattr(table, "_tbl", None)
        try:
            t_pos = body_children.index(tbl)
        except ValueError:
            return None
        for el in body_children[t_pos + 1 :]:
            tag = getattr(el, "tag", "")
            if tag == qn("w:tbl"):
                break
            if tag != qn("w:p"):
                continue
            para = Paragraph(el, ctx.docx.doc)
            text = para.text.strip()
            if not text:
                continue
            if _CAPTION_RE.match(text):
                break
            if paragraph_style_name(para).startswith("Heading"):
                break
            if _NOTA_RE.match(text):
                return para
            break
        return None

    def _report_missing_note(self, t_idx: int, issues: list[VerificationIssue]) -> None:
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.note_present",
                severity="info",
                expected="'Nota.' paragraph below the table",
                actual="No table note found",
                evidence=(
                    f"Table {t_idx + 1} has no note; APA 7 only requires one when the"
                    " table needs explanation"
                ),
            )
        )

    def _validate_note_format(
        self, note_para: Paragraph, t_idx: int, issues: list[VerificationIssue]
    ) -> None:
        note_text = note_para.text.strip()
        if not note_text.startswith(("Nota.", "Note.")):
            issues.append(
                VerificationIssue(
                    check=f"{CheckCategory.TABLES}.note_format",
                    severity="warning",
                    expected="'Nota.' label ending with a period",
                    actual=note_text[:40],
                    evidence=f"Table {t_idx + 1} note label must be 'Nota.'",
                )
            )
        nota_run = next(
            (r for r in note_para.runs if r.text.strip().lower().startswith("nota")), None
        )
        if nota_run is None or nota_run.italic:
            return
        issues.append(
            VerificationIssue(
                check=f"{CheckCategory.TABLES}.note_italic",
                severity="warning",
                expected="'Nota.' label in italics",
                actual="Note label not italic",
                evidence=f"Table {t_idx + 1} note label should be italic",
            )
        )
