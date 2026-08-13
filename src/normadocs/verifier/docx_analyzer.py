"""DOCX analysis utilities using python-docx.

Provides deep analysis of DOCX documents for APA verification,
including styles, paragraphs, tables, headers, and footers.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.text.paragraph import Paragraph

PageMargins = tuple[float, float, float, float]


@dataclass
class DOCXStyleInfo:
    """Information about a document style."""

    name: str
    font_name: str | None
    font_size: float | None
    bold: bool
    italic: bool
    alignment: str | None
    first_line_indent: float | None


@dataclass
class DOCXPageInfo:
    """Information about page layout and headers."""

    page_width: float
    page_height: float
    margins: PageMargins
    has_different_first_page_header_footer: bool = False
    header_text_pages: dict[int, str] = field(default_factory=dict)


@dataclass
class DOCXParagraphInfo:
    """Detailed paragraph information."""

    text: str
    style_name: str | None
    alignment: str | None
    first_line_indent: float | None
    space_before: float | None
    space_after: float | None
    line_spacing: float | None
    runs: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class DOCXTableInfo:
    """Table information."""

    rows: int
    cols: int
    caption: str | None
    caption_position: str | None
    has_horizontal_borders_only: bool = True


class DOCXAnalyzer:
    """Analyzes DOCX documents using python-docx.

    Extracts comprehensive information about styles, formatting, content structure,
    and all APA-relevant elements.
    """

    def __init__(self, docx_path: str | Path) -> None:
        """Initialize the DOCX analyzer.

        Args:
            docx_path: Path to the DOCX file to analyze.
        """
        self.doc_path = Path(docx_path)
        self.doc = Document(str(docx_path))

    @property
    def paragraphs(self) -> list[Paragraph]:
        """Get all paragraphs from the document."""
        return list(self.doc.paragraphs)

    @property
    def tables(self) -> list[Any]:
        """Get all tables from the document."""
        return list(self.doc.tables)

    @property
    def styles(self) -> dict[str, Any]:
        """Get all styles from the document."""
        return {s.name: s for s in self.doc.styles}

    def get_page_info(self) -> DOCXPageInfo:
        """Extract page layout information.

        Returns:
            DOCXPageInfo with all layout details.
        """
        section = self.doc.sections[0]

        def emu_to_inches(emu: Any) -> float:
            if emu is None:
                return 1.0
            emu_val = int(emu)
            return emu_val / 914400.0

        def emu_to_inches_page(emu: Any) -> float:
            if emu is None:
                return 8.5
            emu_val = int(emu)
            return emu_val / 914400.0

        margins = (
            emu_to_inches(section.top_margin),
            emu_to_inches(section.right_margin),
            emu_to_inches(section.bottom_margin),
            emu_to_inches(section.left_margin),
        )

        return DOCXPageInfo(
            page_width=emu_to_inches_page(section.page_width),
            page_height=emu_to_inches_page(section.page_height),
            margins=margins,
            has_different_first_page_header_footer=section.different_first_page_header_footer,
        )

    @staticmethod
    def _effective_line_spacing(paragraph: Paragraph) -> float | None:
        """Return direct or inherited line spacing as a numeric multiplier."""
        paragraph_format = paragraph.paragraph_format
        direct = paragraph_format.line_spacing
        if direct is not None:
            return float(direct)

        style = paragraph.style
        if style is None:
            return None
        style_format = style.paragraph_format
        inherited = style_format.line_spacing
        if inherited is not None:
            return float(inherited)

        rule = style_format.line_spacing_rule
        if rule == WD_LINE_SPACING.DOUBLE:
            return 2.0
        if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            return 1.5
        if rule == WD_LINE_SPACING.SINGLE:
            return 1.0
        return None

    @staticmethod
    def _effective_indent(paragraph: Paragraph) -> Any:
        """Return direct or inherited first-line indentation."""
        direct = paragraph.paragraph_format.first_line_indent
        if direct is not None:
            return direct
        if paragraph.style is not None:
            return paragraph.style.paragraph_format.first_line_indent
        return None

    @staticmethod
    def _effective_spacing(paragraph: Paragraph, attribute: str) -> Any:
        """Return direct or inherited paragraph spacing for an attribute."""
        direct = getattr(paragraph.paragraph_format, attribute)
        if direct is not None:
            return direct
        if paragraph.style is not None:
            return getattr(paragraph.style.paragraph_format, attribute)
        return None

    def get_paragraphs_info(self) -> list[DOCXParagraphInfo]:
        """Extract detailed paragraph information, including inherited styles.

        Returns:
            List of DOCXParagraphInfo for each paragraph.
        """
        paragraphs_info: list[DOCXParagraphInfo] = []

        for p in self.paragraphs:
            runs_data: list[dict[str, Any]] = []
            style_font = p.style.font if p.style is not None else None
            for run in p.runs:
                runs_data.append(
                    {
                        "text": run.text,
                        "font_name": run.font.name or (style_font.name if style_font else None),
                        "font_size": run.font.size or (style_font.size if style_font else None),
                        "bold": run.bold,
                        "italic": run.italic,
                    }
                )

            alignment_value = p.alignment
            if alignment_value is None and p.style is not None:
                alignment_value = p.style.paragraph_format.alignment

            alignment = None
            if alignment_value == WD_ALIGN_PARAGRAPH.LEFT:
                alignment = "left"
            elif alignment_value == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = "center"
            elif alignment_value == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = "right"
            elif alignment_value == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = "justify"

            paragraphs_info.append(
                DOCXParagraphInfo(
                    text=p.text,
                    style_name=p.style.name if p.style else None,
                    alignment=alignment,
                    first_line_indent=self._effective_indent(p),
                    space_before=self._effective_spacing(p, "space_before"),
                    space_after=self._effective_spacing(p, "space_after"),
                    line_spacing=self._effective_line_spacing(p),
                    runs=runs_data,
                )
            )

        return paragraphs_info

    def get_style_info(self, style_name: str) -> DOCXStyleInfo | None:
        """Get detailed information about a specific style.

        Args:
            style_name: Name of the style to inspect.

        Returns:
            DOCXStyleInfo or None if style not found.
        """
        has_get_by_name = hasattr(self.doc.styles, "get_by_name")
        style = self.doc.styles.get_by_name(style_name) if has_get_by_name else None
        if style is None:
            for s in self.doc.styles:
                if s.name == style_name:
                    style = s
                    break

        if style is None:
            return None

        font = style.font
        return DOCXStyleInfo(
            name=style.name,
            font_name=font.name,
            font_size=font.size.pt if font.size else None,
            bold=font.bold or False,
            italic=font.italic or False,
            alignment=None,
            first_line_indent=None,
        )

    def get_tables_info(self) -> list[DOCXTableInfo]:
        """Extract table information.

        Returns:
            List of DOCXTableInfo for each table.
        """
        tables_info: list[DOCXTableInfo] = []

        for table in self.tables:
            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0

            caption = None
            caption_position = None

            tables_info.append(
                DOCXTableInfo(
                    rows=rows,
                    cols=cols,
                    caption=caption,
                    caption_position=caption_position,
                )
            )

        return tables_info

    def get_footer_text(self, footer_type: str = "default") -> str:
        """Get text content from a section footer."""
        section = self.doc.sections[0]
        if footer_type == "first":
            footer = section.first_page_footer
        elif footer_type == "even":
            footer = section.even_page_footer
        else:
            footer = section.footer
        return "\n".join(p.text for p in footer.paragraphs if p.text)

    def get_header_text(self, header_type: str = "default") -> str:
        """Get text content from headers.

        Args:
            header_type: Type of header - "first", "default", or "even".

        Returns:
            All text content concatenated from the header.
        """
        section = self.doc.sections[0]

        if header_type == "first":
            header = section.first_page_header
        elif header_type == "even":
            header = section.even_page_header
        else:
            header = section.header

        return "\n".join(p.text for p in header.paragraphs if p.text)

    def has_heading(self, level: int) -> bool:
        """Check if document has a heading of specified level.

        Args:
            level: Heading level (1-9).

        Returns:
            True if at least one heading of that level exists.
        """
        return any(p.style and p.style.name == f"Heading {level}" for p in self.paragraphs)

    def count_headings(self, level: int) -> int:
        """Count headings of a specific level.

        Args:
            level: Heading level (1-9).

        Returns:
            Count of headings at that level.
        """
        count = 0
        for p in self.paragraphs:
            if p.style and p.style.name == f"Heading {level}":
                count += 1
        return count
