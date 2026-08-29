#!/usr/bin/env python3
"""
Script para verificar TODOS los cálculos del INFORME TÉCNICO ECONÓMICO.
"""

import re
import sys
from pathlib import Path
from typing import Any

_SEP = "   ─────────────────────────────────"


def extract_currency(value: str) -> int:
    """Extrae valor numérico de formato colombiano $3.500.000."""
    if not value:
        return 0
    cleaned = value.strip().replace("$", "").replace(" ", "")
    if "." in cleaned:
        cleaned = re.sub(r"[^\d]", "", cleaned)
        return int(cleaned) if cleaned else 0
    return 0


def _strip_cells(line: str) -> list[str]:
    cells = [c.strip() for c in line.split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def _is_separator(cells: list[str]) -> bool:
    return all(re.match(r"^[:\- ]+$", c) or c == "" for c in cells)


def parse_table_lines(section: str) -> list[list[str]]:
    """Extrae filas de datos de una sección markdown (que contiene tablas)."""
    rows: list[list[str]] = []
    in_table = False
    for line in section.split("\n"):
        if not line.startswith("|"):
            in_table = False
            continue
        cells = _strip_cells(line)
        if not cells:
            continue
        if _is_separator(cells):
            in_table = True
            continue
        if in_table and cells:
            rows.append(cells)
    return rows


def _section_slice(content: str, start: str, end: str) -> str:
    s = content.find(start)
    e = content.find(end)
    if s == -1 or e == -1:
        return ""
    return content[s:e]


def _first_currency_reversed(cells: list[str]) -> int:
    for cell in reversed(cells):
        val = extract_currency(cell)
        if val > 0:
            return val
    return 0


def _is_personal_subtotal_row(row: list[str]) -> bool:
    return len(row) >= 2 and row[0] == "" and "subtotal" in " ".join(row).lower()


def _extract_personal_item(row: list[str]) -> tuple[str, int] | None:
    if len(row) >= 5:
        val = extract_currency(row[4])
        if val > 0:
            return (row[0], val)
    return None


def _collect_personal_rows(
    rows: list[list[str]],
) -> tuple[list[tuple[str, int]], int]:
    items: list[tuple[str, int]] = []
    total = 0
    for row in rows:
        if _is_personal_subtotal_row(row):
            found = _first_currency_reversed(row)
            if found:
                total = found
            continue
        maybe = _extract_personal_item(row)
        if maybe is not None:
            items.append(maybe)
    return items, total


def _summarize_personal(items: list[tuple[str, int]], total: int) -> tuple[int, int]:
    p_sum = sum(t for _, t in items)
    for name, val in items:
        print(f"   - {name}: ${val:,}")
    print(_SEP)
    print(f"   TOTAL NÓMINA: ${p_sum:,}")
    if total > 0 and p_sum == total:
        print("   ✓ Suma coincide con subtotal")
    elif total > 0:
        print(f"   ⚠️ Diferencia - usando suma: ${p_sum:,}")
        total = p_sum
    else:
        total = p_sum
    return total, p_sum


def _is_total_row(row: list[str]) -> bool:
    return "total" in " ".join(row).lower()


def _extract_hardware_item(row: list[str]) -> tuple[str, int] | None:
    if len(row) >= 2:
        val = extract_currency(row[-1])
        if val > 0:
            name = row[-2] if len(row) > 2 else row[0]
            return (name, val)
    return None


def _collect_hardware_rows(
    rows: list[list[str]],
) -> tuple[list[tuple[str, int]], int]:
    items: list[tuple[str, int]] = []
    total = 0
    for row in rows:
        if _is_total_row(row):
            found = _first_currency_reversed(row)
            if found:
                total = found
            continue
        maybe = _extract_hardware_item(row)
        if maybe is not None:
            items.append(maybe)
    return items, total


def _summarize_hardware(items: list[tuple[str, int]], total: int) -> tuple[int, int]:
    hw_sum = sum(t for _, t in items)
    for name, val in items:
        print(f"   - {name}: ${val:,}")
    print(_SEP)
    print(f"   TOTAL HARDWARE: ${hw_sum:,}")
    if total > 0 and hw_sum == total:
        print("   ✓ Suma coincide con total")
    elif total > 0:
        print("   ⚠️ Diferencia - usando suma")
    return hw_sum, hw_sum


def _extract_software_item(row: list[str]) -> tuple[str, int] | None:
    if len(row) >= 2:
        val = extract_currency(row[-1])
        name = row[0]
        if val > 0 and "rubro" not in name.lower():
            return (name, val)
    return None


def _collect_software_rows(
    rows: list[list[str]],
) -> tuple[list[tuple[str, int]], int]:
    items: list[tuple[str, int]] = []
    total = 0
    for row in rows:
        if _is_total_row(row):
            found = _first_currency_reversed(row)
            if found:
                total = found
            continue
        maybe = _extract_software_item(row)
        if maybe is not None:
            items.append(maybe)
    return items, total


def _summarize_software(items: list[tuple[str, int]], total: int) -> tuple[int, int]:
    sw_sum = sum(t for _, t in items)
    for name, val in items:
        print(f"   - {name}: ${val:,}")
    print(_SEP)
    print(f"   TOTAL SOFTWARE: ${sw_sum:,}")
    if total > 0 and sw_sum == total:
        print("   ✓ Suma coincide con total")
    elif total > 0:
        print("   ⚠️ Diferencia - usando suma")
    return sw_sum, sw_sum


def _extract_direct_item(row: list[str]) -> tuple[str, int] | None:
    if len(row) >= 2:
        val = extract_currency(row[-1])
        if val > 0:
            return (row[0], val)
    return None


def _collect_direct_rows(
    rows: list[list[str]],
) -> tuple[list[tuple[str, int]], int]:
    items: list[tuple[str, int]] = []
    direct = 0
    for row in rows:
        if _is_total_row(row):
            found = _first_currency_reversed(row)
            if found:
                direct = found
            continue
        maybe = _extract_direct_item(row)
        if maybe is not None:
            items.append(maybe)
    return items, direct


def _classify_aiu_row(text: str, val: int, vals: dict[str, int]) -> None:
    if "administración" in text:
        vals["a"] = val
    elif "imprevistos" in text:
        vals["i"] = val
    elif "utilidad" in text:
        vals["u"] = val
    elif "total aiu" in text:
        vals["total"] = val


def _parse_aiu_vals(rows: list[list[str]]) -> dict[str, int]:
    vals: dict[str, int] = {}
    for row in rows:
        text = " ".join(row).lower()
        if len(row) >= 4:
            val = extract_currency(row[3])
            if val > 0:
                _classify_aiu_row(text, val, vals)
    return vals


def _report_aiu_row(label: str, key: str, calc: int, vals: dict[str, int]) -> bool:
    rep = vals.get(key, 0)
    ok = calc == rep
    st = "✓" if ok else "✗"
    print(f"      {label}: Calc=${calc:,} vs Reportado=${rep:,} {st}")
    return ok


def _verify_provider_sums(
    providers: list[tuple[str, int, int, int]],
) -> bool:
    all_ok = True
    for pname, sub, aiu, sub_aiu in providers:
        expected = sub + aiu
        if sub_aiu == expected:
            print(f"   ✓ {pname}: Subtotal=${sub:,} + AIU=${aiu:,} = ${sub_aiu:,} ✓")
        else:
            print(
                f"   ✗ {pname}: Subtotal=${sub:,} + AIU=${aiu:,} = "
                f"${expected:,} pero tabla dice ${sub_aiu:,}"
            )
            all_ok = False
    return all_ok


def _check_swapped_columns(
    mack: dict, tecno: dict, dev: dict, providers: list[tuple[str, int, int, int]]
) -> bool:
    all_ok = True
    for pname, sub, aiu, sub_aiu in providers:
        expected = sub + aiu
        if sub_aiu != expected:
            print(
                f"   ✗ ERROR: {pname} tiene valor incorrecto "
                f"${sub_aiu:,} (debería ser ${expected:,})"
            )
            all_ok = False
    actual_tecno = tecno["sub"] + tecno["aiu"]
    actual_dev = dev["sub"] + dev["aiu"]
    if actual_tecno == dev["sub_aiu"] and actual_dev == tecno["sub_aiu"]:
        print("   ✗ CRÍTICO: ¡Valores de TecnoShop y DevSoft están INTERCAMBIADOS!")
        all_ok = False
    elif actual_tecno == mack["sub_aiu"] or actual_dev == mack["sub_aiu"]:
        print("   ✗ CRÍTICO: ¡Valores están en columna equivocada!")
        all_ok = False
    else:
        print("   ✓ Columnas correctas - no hay valores intercambiados")
    return all_ok


def _handle_personal(content: str) -> tuple[int, int]:
    print("\n" + "=" * 70)
    print("1. COSTOS DE PERSONAL (Sección 4.2)")
    print("=" * 70)
    sec = _section_slice(content, "## 4.2", "## 4.3")
    rows = parse_table_lines(sec)
    items, total = _collect_personal_rows(rows)
    return _summarize_personal(items, total)


def _handle_hardware(content: str) -> tuple[int, int]:
    print("\n" + "=" * 70)
    print("2. COSTOS DE HARDWARE (Sección 2.6)")
    print("=" * 70)
    sec = _section_slice(content, "## 2.6", "## 3.")
    rows = parse_table_lines(sec)
    items, total = _collect_hardware_rows(rows)
    return _summarize_hardware(items, total)


def _handle_software(content: str) -> tuple[int, int]:
    print("\n" + "=" * 70)
    print("3. COSTOS DE SOFTWARE (Sección 3.7)")
    print("=" * 70)
    sec = _section_slice(content, "## 3.7", "## 4.")
    rows = parse_table_lines(sec)
    items, total = _collect_software_rows(rows)
    return _summarize_software(items, total)


def _print_direct_summary(
    items: list[tuple[str, int]],
    direct: int,
    personal: int,
    hw: int,
    sw: int,
) -> int:
    cost_sum = sum(t for _, t in items)
    for name, val in items:
        print(f"   - {name}: ${val:,}")
    print(_SEP)
    print(f"   TOTAL COSTOS DIRECTOS: ${cost_sum:,}")
    sum_parts = personal + hw + sw
    status = "✓" if sum_parts == direct else "⚠️"
    print(f"\n   Verificación: ${personal:,} + ${hw:,} + ${sw:,} = ${sum_parts:,}")
    print(f"   vs Costos Directos: ${direct:,} {status}")
    return cost_sum


def _handle_direct(content: str, personal: int, hw: int, sw: int) -> int:
    print("\n" + "=" * 70)
    print("4. ESTRUCTURA DE COSTOS DIRECTOS (Sección 4.3)")
    print("=" * 70)
    sec = _section_slice(content, "## 4.3", "## 4.4")
    rows = parse_table_lines(sec)
    items, direct = _collect_direct_rows(rows)
    return _print_direct_summary(items, direct, personal, hw, sw)


def _calc_aiu_values(direct: int) -> tuple[int, int, int, int, int, int]:
    admin = int(direct * 0.05)
    imp = int(direct * 0.10)
    util = int(direct * 0.20)
    aiu = admin + imp + util
    iva = int(util * 0.19)
    total_c = direct + aiu + iva
    return admin, imp, util, aiu, iva, total_c


def _report_aiu_comparison(
    vals: dict[str, int],
    admin: int,
    imp: int,
    util: int,
    aiu: int,
    all_ok: bool,
) -> bool:
    for label, key, calc in [
        ("Administración (5 %)", "a", admin),
        ("Imprevistos (10 %)", "i", imp),
        ("Utilidad (20 %)", "u", util),
        ("Total AIU (35 %)", "total", aiu),
    ]:
        if not _report_aiu_row(label, key, calc, vals):
            all_ok = False
    return all_ok


def _handle_aiu(content: str, direct: int, all_ok: bool) -> tuple[int, bool]:
    print("\n" + "=" * 70)
    print("5. CÁLCULO AIU (Sección 4.4)")
    print("=" * 70)
    sec = _section_slice(content, "## 4.4", "## 4.5")
    rows = parse_table_lines(sec)
    vals = _parse_aiu_vals(rows)
    admin, imp, util, aiu, iva, total_c = _calc_aiu_values(direct)
    print(f"   Base Imponible:     ${direct:,}")
    print(_SEP)
    all_ok = _report_aiu_comparison(vals, admin, imp, util, aiu, all_ok)
    print(f"   IVA (19 %):  ${iva:,}")
    print("   ════════════════════════════════════════")
    print(f"   TOTAL PROYECTO:   ${total_c:,}")
    return total_c, all_ok


def _handle_budget(content: str) -> None:
    print("\n" + "=" * 70)
    print("6. PRESUPUESTO TOTAL (Sección 4.5)")
    print("=" * 70)
    sec = _section_slice(content, "## 4.5", "## 5.")
    rows = parse_table_lines(sec)
    for row in rows:
        if len(row) >= 2:
            name = row[0]
            for cell in reversed(row):
                val = extract_currency(cell)
                if val > 0:
                    print(f"   - {name}: ${val:,}")
                    break


def _assign_fin(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if "total final" in text and len(nums) >= 3:
        mack["fin"], tecno["fin"], dev["fin"] = nums[0], nums[1], nums[2]
        return True
    return False


def _assign_sub_aiu(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if "subtotal" in text and "antes de iva" in text and len(nums) >= 3:
        mack["sub_aiu"], tecno["sub_aiu"], dev["sub_aiu"] = (
            nums[0],
            nums[1],
            nums[2],
        )
        return True
    return False


def _assign_iva(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if "iva" in text and "19" in text and len(nums) >= 3:
        mack["iva"], tecno["iva"], dev["iva"] = nums[0], nums[1], nums[2]
        return True
    return False


def _assign_aiu(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if "aiu" in text and "35" in text and len(nums) >= 3:
        mack["aiu"], tecno["aiu"], dev["aiu"] = nums[0], nums[1], nums[2]
        return True
    return False


def _assign_sub(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if "subtotal" in text and "directos" in text and len(nums) >= 3:
        mack["sub"], tecno["sub"], dev["sub"] = nums[0], nums[1], nums[2]
        return True
    return False


def _assign_totals(text: str, nums: list[int], mack: dict, tecno: dict, dev: dict) -> bool:
    if _assign_fin(text, nums, mack, tecno, dev):
        return True
    if _assign_sub_aiu(text, nums, mack, tecno, dev):
        return True
    if _assign_iva(text, nums, mack, tecno, dev):
        return True
    if _assign_aiu(text, nums, mack, tecno, dev):
        return True
    return _assign_sub(text, nums, mack, tecno, dev)


def _handle_comparative(content: str) -> tuple[dict, dict, dict]:
    print("\n" + "=" * 70)
    print("7. TABLA COMPARATIVA (Sección 6.2)")
    print("=" * 70)
    sec = _section_slice(content, "## 6.2", "# 7.")
    rows = parse_table_lines(sec)
    mack: dict[str, Any] = {
        "items": [],
        "sub": 0,
        "aiu": 0,
        "sub_aiu": 0,
        "iva": 0,
        "fin": 0,
    }
    tecno: dict[str, Any] = {
        "items": [],
        "sub": 0,
        "aiu": 0,
        "sub_aiu": 0,
        "iva": 0,
        "fin": 0,
    }
    dev: dict[str, Any] = {
        "items": [],
        "sub": 0,
        "aiu": 0,
        "sub_aiu": 0,
        "iva": 0,
        "fin": 0,
    }
    for row in rows:
        if len(row) < 4:
            continue
        text = " ".join(row).lower()
        vals = [extract_currency(c) for c in row]
        nums = [v for v in vals if v > 0]
        if _assign_totals(text, nums, mack, tecno, dev):
            continue
        if row[0].strip().isdigit() and len(nums) >= 3:
            mack["items"].append(nums[0])
            tecno["items"].append(nums[1])
            dev["items"].append(nums[2])
    return mack, tecno, dev


def _check_unique_columns(mack: dict, tecno: dict, dev: dict) -> bool:
    all_subs = {mack["sub"], tecno["sub"], dev["sub"]}
    all_aius = {mack["aiu"], tecno["aiu"], dev["aiu"]}
    all_sub_aius = {mack["sub_aiu"], tecno["sub_aiu"], dev["sub_aiu"]}
    return len(all_subs) == 3 and len(all_aius) == 3 and len(all_sub_aius) == 3


def _handle_swap_check(mack: dict, tecno: dict, dev: dict, all_ok: bool) -> bool:
    print("\n" + "-" * 70)
    print("   ⚡ VERIFICACIÓN DE COLUMNAS (Detección de valores intercambiados)")
    print("-" * 70)
    providers = [
        ("Mackroph", mack["sub"], mack["aiu"], mack["sub_aiu"]),
        ("TecnoShop", tecno["sub"], tecno["aiu"], tecno["sub_aiu"]),
        ("DevSoft", dev["sub"], dev["aiu"], dev["sub_aiu"]),
    ]
    if not _verify_provider_sums(providers):
        all_ok = False
    print("\n   🔍 Verificando que valores NO estén intercambiados entre columnas...")
    if _check_unique_columns(mack, tecno, dev):
        if not _check_swapped_columns(mack, tecno, dev, providers):
            all_ok = False
    else:
        print("   ⚠️ No se puede verificar intercambio (hay valores duplicados)")
    return all_ok


def _check_items_subtotal(data: dict, all_ok: bool) -> bool:
    items_sum = sum(data["items"])
    if items_sum == data["sub"]:
        print(f"      ✓ Items ({len(data['items'])}): ${items_sum:,} = subtotal ${data['sub']:,}")
    else:
        print(f"      ✗ Items ({len(data['items'])}): ${items_sum:,} ≠ subtotal ${data['sub']:,}")
        all_ok = False
    return all_ok


def _check_provider_aiu_iva(data: dict, all_ok: bool) -> bool:
    base = data["sub"]
    for label, key, calc in [
        ("AIU", "aiu", int(base * 0.35)),
        ("IVA", "iva", int(base * 0.20 * 0.19)),
    ]:
        ok = data[key] == calc
        if ok:
            print(f"      ✓ {label}: ${data[key]:,} = calc ${calc:,}")
        else:
            print(f"      ✗ {label}: ${data[key]:,} ≠ calc ${calc:,}")
            all_ok = False
    return all_ok


def _check_provider_sub_aiu(data: dict, all_ok: bool) -> bool:
    base = data["sub"]
    sub_aiu_calc = base + int(base * 0.35)
    if data["sub_aiu"] > 0:
        if data["sub_aiu"] == sub_aiu_calc:
            print(
                f"      ✓ Subtotal antes de IVA: ${data['sub_aiu']:,} = "
                f"${base:,} + AIU ${int(base * 0.35):,}"
            )
        else:
            print(
                f"      ✗ Subtotal antes de IVA: ${data['sub_aiu']:,} ≠ "
                f"${base:,} + AIU ${int(base * 0.35):,} = ${sub_aiu_calc:,}"
            )
            all_ok = False
    return all_ok


def _handle_provider_details(
    mack: dict, tecno: dict, dev: dict, total_c: int, all_ok: bool
) -> bool:
    for name, data in [("MACKROPH", mack), ("TECNO SHOP", tecno), ("DEV SOFT", dev)]:
        print(f"\n   {name}:")
        all_ok = _check_items_subtotal(data, all_ok)
        all_ok = _check_provider_aiu_iva(data, all_ok)
        all_ok = _check_provider_sub_aiu(data, all_ok)
        base = data["sub"]
        fin_calc = base + int(base * 0.35) + int(base * 0.20 * 0.19)
        if data["fin"] == fin_calc:
            print(f"      ✓ TOTAL: ${data['fin']:,}")
        else:
            print(f"      ✗ TOTAL: ${data['fin']:,} ≠ calc ${fin_calc:,}")
            all_ok = False
    print("\n   Mackroph vs Proyecto:")
    if mack["fin"] == total_c:
        print(f"      ✓ TOTAL Mackroph (${mack['fin']:,}) = TOTAL Proyecto (${total_c:,})")
    else:
        print(f"      ✗ TOTAL Mackroph (${mack['fin']:,}) ≠ TOTAL Proyecto (${total_c:,})")
        all_ok = False
    return all_ok


def _handle_tables(docx_path: str) -> None:
    print("\n" + "=" * 70)
    print("8. VERIFICACIÓN DE TABLAS (No cortadas)")
    print("=" * 70)
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        print("   python-docx no instalado")
        return
    if not Path(docx_path).exists():
        print("   DOCX no encontrado")
        return
    doc = Document(docx_path)
    ok_count = 0
    for i, table in enumerate(doc.tables):

        def _row_protected(row: Any) -> bool:
            tr_pr = row._tr.find(qn("w:trPr"))
            if tr_pr is None:
                return False
            cant = tr_pr.find(qn("w:cantSplit"))
            return cant is not None and cant.get(qn("w:val")) == "1"

        rows_ok = sum(1 for row in table.rows if _row_protected(row))
        if rows_ok == len(table.rows):
            ok_count += 1
            print(f"   Tabla {i + 1}: {len(table.rows)} filas ✓")
    print(f"\n   Tablas protegidas: {ok_count}/{len(doc.tables)}")


def _print_summary(content: str, total_c: int, all_ok: bool) -> bool:
    print("\n" + "=" * 70)
    print("RESUMEN FINAL")
    print("=" * 70)
    if "margen operativo" in content.lower():
        print("   ⚠️ Posible doble margen detectado")
        all_ok = False
    else:
        print("   ✓ Sin problema de doble margen")
    print(f"\n   📊 TOTAL PROYECTO: ${total_c:,} COP")
    if all_ok:
        print("\n   ✅ TODOS LOS CÁLCULOS VERIFICADOS")
    else:
        print("\n   ❌ HAY ERRORES")
    return all_ok


def main() -> None:
    md_file = "IDocs/Evidencia_AA1_EV02/02_Informes/INFORME_TECNICO_ECONOMICO.md"
    docx_path = "ExportDocs/INFORME_TECNICO_ECONOMICO_APA.docx"
    if not Path(md_file).exists():
        print(f"Error: No se encontró {md_file}")
        sys.exit(1)
    print("=" * 70)
    print("VERIFICACIÓN COMPLETA DE CÁLCULOS")
    print("INFORME TÉCNICO ECONÓMICO - SHOPPIPAI")
    print("=" * 70)
    content = Path(md_file).read_text(encoding="utf-8")
    all_ok = True
    personal_total, _ = _handle_personal(content)
    hw_total, _ = _handle_hardware(content)
    sw_total, _ = _handle_software(content)
    direct = _handle_direct(content, personal_total, hw_total, sw_total)
    total_c, all_ok = _handle_aiu(content, direct, all_ok)
    _handle_budget(content)
    mack, tecno, dev = _handle_comparative(content)
    all_ok = _handle_swap_check(mack, tecno, dev, all_ok)
    all_ok = _handle_provider_details(mack, tecno, dev, total_c, all_ok)
    _handle_tables(docx_path)
    all_ok = _print_summary(content, total_c, all_ok)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
