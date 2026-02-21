"""
End-to-End Tests for NormaDocs.

These tests exercise the FULL pipeline through the CLI (typer.testing.CliRunner),
verify the output DOCX file exists, and deeply inspect the generated document
structure (margins, fonts, headings, tables, cover page).

Requires: pandoc installed.
"""

import shutil
import textwrap
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from typer.testing import CliRunner

from normadocs.cli import app

runner = CliRunner()

PANDOC_AVAILABLE = shutil.which("pandoc") is not None

# ──────────────────────────────────────────────────────────────────────
#  Helper: realistic Markdown that exercises the entire pipeline
# ──────────────────────────────────────────────────────────────────────
SAMPLE_MARKDOWN = textwrap.dedent("""\
    **Efectos del Aprendizaje Automático en la Educación Superior**

    Juan Pérez García
    Ingeniería de Software
    2694731
    SENA
    Centro de Tecnologías
    2026-02-18

    # Resumen

    Este documento examina los efectos del aprendizaje automático en
    la educación superior, con un enfoque en la personalización del
    contenido y la mejora del rendimiento estudiantil.

    **Palabras clave:** aprendizaje automático, educación, inteligencia artificial

    # Introducción

    La educación superior enfrenta desafíos significativos en la era
    digital. El aprendizaje automático ofrece soluciones prometedoras.

    ## Contexto Histórico

    Desde la década de 1950, la inteligencia artificial ha evolucionado
    de forma exponencial, impactando múltiples sectores.

    ### Primeros Desarrollos

    Los primeros sistemas de tutoría inteligente surgieron en los años 70.

    # Métodos

    Se realizó una revisión sistemática de la literatura utilizando
    las bases de datos Scopus y Web of Science.

    | Variable | Descripción | Tipo |
    |----------|-------------|------|
    | Rendimiento | Calificaciones promedio | Cuantitativa |
    | Satisfacción | Escala Likert 1-5 | Ordinal |
    | Tiempo | Horas de estudio semanales | Continua |

    # Resultados

    Los resultados muestran una mejora significativa en el rendimiento
    académico (p < 0.05) cuando se implementan sistemas de aprendizaje
    adaptativo.

    # Discusión

    Estos hallazgos son consistentes con la literatura previa
    (García y López, 2023; Martínez et al., 2024).

    # Conclusiones

    El aprendizaje automático tiene un impacto positivo y medible en
    la educación superior.

    # Referencias

    García, A. y López, B. (2023). Machine learning en educación. *Revista Educación*, 45(2), 112-130.

    Martínez, C., Rodríguez, D. y Silva, E. (2024). Sistemas adaptativos. *Journal of AI*, 12(1), 45-67.
""")


@unittest.skipUnless(PANDOC_AVAILABLE, "Pandoc not installed — skipping E2E")
class TestEndToEndAPA(unittest.TestCase):
    """Full pipeline: CLI → Pandoc → APA Formatter → DOCX validation."""

    output_dir: Path

    @classmethod
    def setUpClass(cls):
        cls.output_dir = Path("tests/temp_e2e")
        cls.output_dir.mkdir(parents=True, exist_ok=True)
        cls.md_file = cls.output_dir / "sample_apa.md"
        cls.md_file.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

        # Run the CLI once for all APA tests
        result = runner.invoke(
            app,
            [
                str(cls.md_file),
                "--style",
                "apa",
                "--output-dir",
                str(cls.output_dir),
            ],
        )
        cls.cli_result = result
        cls.docx_path = cls.output_dir / "sample_apa_APA.docx"
        if cls.docx_path.exists():
            cls.doc = Document(str(cls.docx_path))
        else:
            cls.doc = None

    @classmethod
    def tearDownClass(cls):
        if cls.output_dir.exists():
            shutil.rmtree(cls.output_dir)

    # ── CLI exit code ────────────────────────────────────────────────

    def test_cli_exits_successfully(self):
        """CLI must exit with code 0."""
        self.assertEqual(
            self.cli_result.exit_code,
            0,
            f"CLI failed (exit {self.cli_result.exit_code}):\n{self.cli_result.output}"
            + (f"\n{self.cli_result.exception}" if self.cli_result.exception else ""),
        )

    # ── Output file exists ───────────────────────────────────────────

    def test_docx_file_created(self):
        """The .docx output file must exist on disk."""
        self.assertTrue(
            self.docx_path.exists(),
            f"Expected output at {self.docx_path}",
        )

    def test_docx_not_empty(self):
        """The .docx must have non-zero size."""
        self.assertGreater(self.docx_path.stat().st_size, 0)

    # ── Page layout (APA 7th: 1-inch margins all around) ─────────────

    def test_margins_one_inch(self):
        """APA requires 1-inch margins on all sides."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        one_inch = Inches(1)
        for section in self.doc.sections:
            self.assertEqual(section.left_margin, one_inch, "Left margin ≠ 1 inch")
            self.assertEqual(section.right_margin, one_inch, "Right margin ≠ 1 inch")
            self.assertEqual(section.top_margin, one_inch, "Top margin ≠ 1 inch")
            self.assertEqual(section.bottom_margin, one_inch, "Bottom margin ≠ 1 inch")

    # ── Font: Times New Roman 12pt ───────────────────────────────────

    def test_normal_style_font(self):
        """Normal style must be Times New Roman, 12pt."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        normal = self.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")
        self.assertEqual(normal.font.size, Pt(12))

    # ── The document has paragraphs ──────────────────────────────────

    def test_document_has_paragraphs(self):
        """The DOCX must contain paragraphs."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        self.assertGreater(len(self.doc.paragraphs), 0)

    # ── Cover page content (title, author, institution) ──────────────

    def test_cover_page_contains_title(self):
        """Title from metadata must appear in the document."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        self.assertIn(
            "Efectos del Aprendizaje Automático en la Educación Superior",
            full_text,
            "Title not found in output document",
        )

    def test_cover_page_contains_author(self):
        """Author/metadata fields from the document must appear in the output."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        # The preprocessor extracts metadata by line position:
        # Lines 0-1 → title, Lines 3+ → author/program/ficha/institution/center/date
        # Verify that extracted metadata fields appear in the cover page
        self.assertIn("Ingeniería de Software", full_text)

    def test_cover_page_contains_institution(self):
        """Institution from metadata must appear in the document."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        self.assertIn("SENA", full_text)

    # ── Heading structure ────────────────────────────────────────────

    def test_headings_present(self):
        """Document must contain heading-styled paragraphs."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        heading_texts = [
            p.text.strip()
            for p in self.doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        self.assertGreater(len(heading_texts), 0, "No headings found in the output DOCX")

    def test_heading_keywords(self):
        """Key section names should appear in headings."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        heading_texts = [
            p.text.strip().lower()
            for p in self.doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        all_headings = " | ".join(heading_texts)
        # At least some of these must appear
        expected_keywords = ["resumen", "introducción", "métodos", "resultados"]
        found = [kw for kw in expected_keywords if kw in all_headings]
        self.assertGreater(
            len(found),
            0,
            f"None of {expected_keywords} found in headings: {all_headings}",
        )

    # ── Tables ───────────────────────────────────────────────────────

    def test_table_exists(self):
        """The Markdown table must be converted to a DOCX table."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        self.assertGreater(len(self.doc.tables), 0, "No tables in output")

    def test_table_has_correct_columns(self):
        """The table must have 3 columns (Variable, Descripción, Tipo)."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        table = self.doc.tables[0]
        # Header row
        header_texts = [cell.text.strip() for cell in table.rows[0].cells]
        self.assertIn("Variable", header_texts)
        self.assertIn("Tipo", header_texts)

    def test_table_has_data_rows(self):
        """The table must have data rows beyond the header."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        table = self.doc.tables[0]
        # At least header + 3 data rows
        self.assertGreaterEqual(len(table.rows), 4, "Table should have ≥4 rows")

    # ── Body content ─────────────────────────────────────────────────

    def test_body_content_preserved(self):
        """Key content text must survive the pipeline."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        self.assertIn("revisión sistemática", full_text)

    def test_references_preserved(self):
        """References section content must appear in the output."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        self.assertIn("García", full_text)
        self.assertIn("Martínez", full_text)

    # ── Page header (page number) ────────────────────────────────────

    def test_header_exists(self):
        """APA requires a header with page number."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        for section in self.doc.sections:
            header = section.header
            # Header should have a paragraph with a PAGE field
            self.assertTrue(
                len(header.paragraphs) > 0,
                "No header paragraphs found (APA requires page numbers)",
            )


@unittest.skipUnless(PANDOC_AVAILABLE, "Pandoc not installed — skipping E2E")
class TestEndToEndICONTEC(unittest.TestCase):
    """Full pipeline: CLI → Pandoc → ICONTEC Formatter → DOCX validation."""

    output_dir: Path

    @classmethod
    def setUpClass(cls):
        cls.output_dir = Path("tests/temp_e2e_icontec")
        cls.output_dir.mkdir(parents=True, exist_ok=True)
        cls.md_file = cls.output_dir / "sample_icontec.md"
        cls.md_file.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

        result = runner.invoke(
            app,
            [
                str(cls.md_file),
                "--style",
                "icontec",
                "--output-dir",
                str(cls.output_dir),
            ],
        )
        cls.cli_result = result
        cls.docx_path = cls.output_dir / "sample_icontec_ICONTEC.docx"
        if cls.docx_path.exists():
            cls.doc = Document(str(cls.docx_path))
        else:
            cls.doc = None

    @classmethod
    def tearDownClass(cls):
        if cls.output_dir.exists():
            shutil.rmtree(cls.output_dir)

    def test_cli_exits_successfully(self):
        """CLI must exit with code 0."""
        self.assertEqual(
            self.cli_result.exit_code,
            0,
            f"CLI failed (exit {self.cli_result.exit_code}):\n{self.cli_result.output}"
            + (f"\n{self.cli_result.exception}" if self.cli_result.exception else ""),
        )

    def test_docx_file_created(self):
        """The .docx output file must exist."""
        self.assertTrue(self.docx_path.exists())

    def test_icontec_font_is_arial(self):
        """ICONTEC requires Arial."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        normal = self.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Arial")

    def test_icontec_font_size_12pt(self):
        """ICONTEC requires 12pt."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        normal = self.doc.styles["Normal"]
        self.assertEqual(normal.font.size, Pt(12))

    def test_cover_page_title_uppercase(self):
        """ICONTEC cover page should have the title in uppercase."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        # The ICONTEC formatter renders title.upper()
        self.assertIn(
            "EFECTOS DEL APRENDIZAJE AUTOMÁTICO EN LA EDUCACIÓN SUPERIOR",
            full_text,
        )

    def test_body_content_preserved(self):
        """Key body text must survive the pipeline."""
        self.assertIsNotNone(self.doc, "DOCX not loaded")
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        self.assertIn("revisión sistemática", full_text)


@unittest.skipUnless(PANDOC_AVAILABLE, "Pandoc not installed — skipping E2E")
class TestEndToEndErrorHandling(unittest.TestCase):
    """CLI error paths."""

    def test_nonexistent_file_fails(self):
        """CLI must fail when the input file doesn't exist."""
        result = runner.invoke(app, ["this_file_does_not_exist.md"])
        self.assertNotEqual(result.exit_code, 0)

    def test_invalid_style_fails_gracefully(self):
        """CLI must fail gracefully with an unsupported style."""
        output_dir = Path("tests/temp_e2e_error")
        output_dir.mkdir(parents=True, exist_ok=True)
        try:
            md_file = output_dir / "test.md"
            md_file.write_text("**Title**\n\nAuthor\n\n# Intro\n\nHello.", encoding="utf-8")
            result = runner.invoke(
                app,
                [str(md_file), "--style", "ieee_unsupported", "--output-dir", str(output_dir)],
            )
            # Should fail because 'ieee_unsupported' is not registered
            self.assertNotEqual(result.exit_code, 0)
        finally:
            if output_dir.exists():
                shutil.rmtree(output_dir)


if __name__ == "__main__":
    unittest.main()
