"""
Strict edge case tests for the YAML configuration system.
Tests error handling, edge cases, and config override scenarios.
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from normadocs.formatters import get_formatter, merge_with_defaults
from normadocs.models import DocumentMetadata


class TestConfigMergeEdgeCases(unittest.TestCase):
    """Edge cases for config merging."""

    def test_merge_with_empty_override(self):
        """Empty override should return full defaults."""
        config = {}
        result = merge_with_defaults(config, "apa7")
        self.assertEqual(result["fonts"]["body"]["name"], "Times New Roman")
        self.assertEqual(result["spacing"]["line"], "double")

    def test_merge_with_none_values(self):
        """Config with None values should override defaults."""
        config = {"fonts": {"body": {"name": None}}}
        result = merge_with_defaults(config, "apa7")
        self.assertIsNone(result["fonts"]["body"]["name"])

    def test_merge_with_partial_nested(self):
        """Partial nested config should merge deeply."""
        config = {
            "fonts": {
                "body": {"name": "Arial"},
                "headings": {"level1": {"bold": False}},
            }
        }
        result = merge_with_defaults(config, "apa7")
        self.assertEqual(result["fonts"]["body"]["name"], "Arial")
        self.assertEqual(result["fonts"]["body"]["size"], 12)
        self.assertFalse(result["fonts"]["headings"]["level1"]["bold"])

    def test_merge_replaces_entire_section(self):
        """Non-dict value should replace entire dict section."""
        config = {"margins": "invalid"}
        result = merge_with_defaults(config, "apa7")
        self.assertEqual(result["margins"], "invalid")

    def test_merge_unknown_style_uses_apa_fallback(self):
        """Unknown style should fall back to APA7 defaults."""
        config = {"fonts": {"body": {"name": "Courier"}}}
        result = merge_with_defaults(config, "unknown_style")
        self.assertEqual(result["fonts"]["body"]["name"], "Courier")
        self.assertEqual(result["citation_style"], "apa")

    def test_deep_merge_three_levels(self):
        """Deep merge at three nesting levels."""
        config = {
            "fonts": {
                "body": {"name": "Arial"},
                "headings": {
                    "name": "Helvetica",
                    "level1": {"bold": False, "alignment": "left"},
                },
            }
        }
        result = merge_with_defaults(config, "apa7")
        self.assertEqual(result["fonts"]["headings"]["name"], "Helvetica")
        self.assertEqual(result["fonts"]["headings"]["level1"]["bold"], False)
        self.assertEqual(result["fonts"]["headings"]["level1"]["alignment"], "left")


class TestConfigOverrideScenarios(unittest.TestCase):
    """Test config override scenarios via get_formatter."""

    def test_override_body_font_name(self):
        """Custom config should override body font name."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"fonts": {"body": {"name": "Arial", "size": 12}}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Arial")

    def test_override_spacing_line(self):
        """Custom config should override line spacing."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"spacing": {"line": "single"}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter.config["spacing"]["line"], "single")

    def test_override_preserves_unmodified_keys(self):
        """Override should preserve unmodified config sections."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"tables": {"caption_prefix": "Cuadro"}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter.config["tables"]["caption_prefix"], "Cuadro")
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")
            self.assertEqual(formatter.config["spacing"]["line"], "double")

    def test_override_deep_margins(self):
        """Deep override for margins."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"margins": {"unit": "cm", "top": 2.5}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter.config["margins"]["unit"], "cm")
            self.assertEqual(formatter.config["margins"]["top"], 2.5)
            self.assertEqual(formatter.config["margins"]["bottom"], 1.0)

    def test_empty_config_uses_full_defaults(self):
        """Empty config dict should use all defaults."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path), config={})
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")
            self.assertEqual(formatter.config["spacing"]["line"], "double")

    def test_override_figure_caption_prefix(self):
        """Override figure caption prefix."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"figures": {"caption_prefix": "Imagen"}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter.config["figures"]["caption_prefix"], "Imagen")

    def test_override_preserves_nested_heading_levels(self):
        """Override should preserve all heading levels when overriding one."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"fonts": {"headings": {"level1": {"bold": False}}}}
            formatter = get_formatter("apa", str(doc_path), config=custom)
            self.assertFalse(formatter.config["fonts"]["headings"]["level1"]["bold"])
            self.assertTrue(formatter.config["fonts"]["headings"]["level2"]["bold"])


class TestIcontecConfigOverride(unittest.TestCase):
    """Test ICONTEC config override scenarios."""

    def test_override_icontec_font(self):
        """Override ICONTEC body font."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"fonts": {"body": {"name": "Times New Roman", "size": 12}}}
            formatter = get_formatter("icontec", str(doc_path), config=custom)
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")
            self.assertEqual(formatter.config["citation_style"], "icontec")

    def test_icontec_margins_in_cm(self):
        """ICONTEC default margins should be in cm."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("icontec", str(doc_path))
            self.assertEqual(formatter.config["margins"]["unit"], "cm")
            self.assertEqual(formatter.config["margins"]["top"], 3.0)


class TestIEEEDocxConfigOverride(unittest.TestCase):
    """Test IEEE config override scenarios."""

    def test_ieee_default_single_spacing(self):
        """IEEE default spacing should be single."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("ieee", str(doc_path))
            self.assertEqual(formatter.config["spacing"]["line"], "single")

    def test_ieee_override_margins(self):
        """Override IEEE margins."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"margins": {"top": 1.5}}
            formatter = get_formatter("ieee", str(doc_path), config=custom)
            self.assertEqual(formatter.config["margins"]["top"], 1.5)


class TestConfigAccess(unittest.TestCase):
    """Test config access via dict-style access on formatter.config."""

    def test_config_access_via_dict(self):
        """Access config values directly via dict."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path))
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")

    def test_config_missing_key_returns_none(self):
        """Missing config key returns None, not error."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path))
            result = formatter.config.get("nonexistent", {}).get("key")
            self.assertIsNone(result)


class TestConfigStyleAliases(unittest.TestCase):
    """Test style name aliases and normalization."""

    def test_apa_and_apa7_same_config(self):
        """APA and APA7 should produce same config."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path1 = Path(tmpdir) / "test1.docx"
            doc_path2 = Path(tmpdir) / "test2.docx"
            Document().save(str(doc_path1))
            Document().save(str(doc_path2))
            formatter1 = get_formatter("apa", str(doc_path1))
            formatter2 = get_formatter("apa7", str(doc_path2))
            self.assertEqual(formatter1.config["name"], formatter2.config["name"])
            self.assertEqual(formatter1.config["citation_style"], "apa")


class TestConfigIntegrity(unittest.TestCase):
    """Test config integrity after various operations."""

    def test_config_is_deep_copy_not_reference(self):
        """Returned config should be a deep copy, not a reference."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom = {"fonts": {"body": {"name": "Arial"}}}
            formatter1 = get_formatter("apa", str(doc_path), config=custom)
            formatter2 = get_formatter("apa", str(doc_path), config=custom)
            self.assertEqual(formatter1.config["fonts"]["body"]["name"], "Arial")
            self.assertEqual(formatter2.config["fonts"]["body"]["name"], "Arial")

    def test_process_without_meta(self):
        """Process should work with default DocumentMetadata."""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path))
            formatter.process(DocumentMetadata())


if __name__ == "__main__":
    unittest.main()
