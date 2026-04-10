"""
Tests for PDFGenerator - DOCX to PDF conversion.
"""

import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from normadocs.pdf_generator import PDFGenerator


class TestPDFGeneratorLibreOffice(unittest.TestCase):
    """Tests for convert_with_libreoffice method."""

    @patch("builtins.print")
    @patch("subprocess.run")
    def test_libreoffice_success(self, mock_run, mock_print):
        """LibreOffice conversion should return True on success."""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_dir = tmpdir

            mock_result = MagicMock()
            mock_result.returncode = 0
            mock_run.return_value = mock_result

            result = PDFGenerator.convert_with_libreoffice(str(docx_path), output_dir)
            self.assertTrue(result)

    @patch("builtins.print")
    @patch("subprocess.run")
    def test_libreoffice_failure(self, mock_run, mock_print):
        """LibreOffice should return False on non-zero returncode."""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_dir = tmpdir

            mock_result = MagicMock()
            mock_result.returncode = 1
            mock_result.stderr = "Error message"
            mock_run.return_value = mock_result

            result = PDFGenerator.convert_with_libreoffice(str(docx_path), output_dir)
            self.assertFalse(result)

    @patch("builtins.print")
    @patch("subprocess.run")
    def test_libreoffice_not_found(self, mock_run, mock_print):
        """LibreOffice FileNotFoundError should return False."""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_dir = tmpdir

            mock_run.side_effect = FileNotFoundError()

            result = PDFGenerator.convert_with_libreoffice(str(docx_path), output_dir)
            self.assertFalse(result)

    @patch("builtins.print")
    @patch("subprocess.run")
    def test_libreoffice_cmd_structure(self, mock_run, mock_print):
        """LibreOffice command should have correct structure."""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_dir = tmpdir

            mock_result = MagicMock()
            mock_result.returncode = 0
            mock_run.return_value = mock_result

            PDFGenerator.convert_with_libreoffice(str(docx_path), output_dir)

            call_args = mock_run.call_args
            cmd = call_args[0][0]
            self.assertTrue(cmd[0].endswith("libreoffice") or cmd[0] == "libreoffice")
            self.assertIn("--headless", cmd)
            self.assertIn("--convert-to", cmd)
            self.assertIn("pdf", cmd)
            self.assertIn("--outdir", cmd)


class TestPDFGeneratorWeasyPrint(unittest.TestCase):
    """Tests for convert_with_weasyprint method."""

    def test_weasyprint_success(self):
        """WeasyPrint conversion should return True on success."""
        md_content = "# Title\n\nContent paragraph."
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.pdf"

            with patch("subprocess.run") as mock_run:
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = "<html><body>Content</body></html>"
                mock_run.return_value = mock_result

                mock_weasyprint = MagicMock()
                mock_html_cls = MagicMock()
                mock_css_cls = MagicMock()
                mock_html_instance = MagicMock()
                mock_html_cls.return_value = mock_html_instance
                mock_weasyprint.HTML = mock_html_cls
                mock_weasyprint.CSS = mock_css_cls

                with patch.dict("sys.modules", {"weasyprint": mock_weasyprint}):
                    result = PDFGenerator.convert_with_weasyprint(md_content, str(output_path))
                    self.assertTrue(result)

    def test_weasyprint_import_error(self):
        """WeasyPrint ImportError should return False."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.pdf"

            with patch("subprocess.run") as mock_run:
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = "<html></html>"
                mock_run.return_value = mock_result

                with (
                    patch.dict("sys.modules", {"weasyprint": None}),
                    patch("builtins.__import__", side_effect=ImportError("No module")),
                ):
                    result = PDFGenerator.convert_with_weasyprint(md_content, str(output_path))
                    self.assertFalse(result)

    def test_weasyprint_pandoc_failure(self):
        """Pandoc failure should return False."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.pdf"

            with patch("subprocess.run") as mock_run:
                mock_result = MagicMock()
                mock_result.returncode = 1
                mock_run.return_value = mock_result

                result = PDFGenerator.convert_with_weasyprint(md_content, str(output_path))
                self.assertFalse(result)

    def test_weasyprint_general_exception(self):
        """General exception in WeasyPrint should return False."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.pdf"

            with patch("subprocess.run") as mock_run:
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = "<html></html>"
                mock_run.return_value = mock_result

                mock_weasyprint = MagicMock()
                mock_html_cls = MagicMock()
                mock_css_cls = MagicMock()
                mock_html_cls.side_effect = Exception("PDF write error")
                mock_weasyprint.HTML = mock_html_cls
                mock_weasyprint.CSS = mock_css_cls

                with patch.dict("sys.modules", {"weasyprint": mock_weasyprint}):
                    result = PDFGenerator.convert_with_weasyprint(md_content, str(output_path))
                    self.assertFalse(result)


class TestPDFGeneratorConvert(unittest.TestCase):
    """Tests for the convert method (primary flow)."""

    def test_convert_falls_back_to_weasyprint(self):
        """Should fallback to WeasyPrint when LibreOffice fails."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_path = Path(tmpdir) / "output.pdf"

            with (
                patch.object(PDFGenerator, "convert_with_libreoffice", return_value=False),
                patch.object(PDFGenerator, "convert_with_weasyprint", return_value=True),
            ):
                result = PDFGenerator.convert(str(docx_path), tmpdir, md_content, str(output_path))
                self.assertTrue(result)

    def test_convert_succeeds_with_libreoffice(self):
        """Should return True when LibreOffice succeeds (no WeasyPrint)."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_path = Path(tmpdir) / "output.pdf"

            with (
                patch.object(PDFGenerator, "convert_with_libreoffice", return_value=True),
                patch.object(PDFGenerator, "convert_with_weasyprint", return_value=False),
            ):
                result = PDFGenerator.convert(str(docx_path), tmpdir, md_content, str(output_path))
                self.assertTrue(result)

    def test_convert_fails_when_both_fail(self):
        """Should return False when both LibreOffice and WeasyPrint fail."""
        md_content = "# Title"
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "test.docx"
            Document().save(str(docx_path))
            output_path = Path(tmpdir) / "output.pdf"

            with (
                patch.object(PDFGenerator, "convert_with_libreoffice", return_value=False),
                patch.object(PDFGenerator, "convert_with_weasyprint", return_value=False),
            ):
                result = PDFGenerator.convert(str(docx_path), tmpdir, md_content, str(output_path))
                self.assertFalse(result)


if __name__ == "__main__":
    unittest.main()
