"""
Unit tests for the IEEE document formatter.

These tests exercise the IEEEDocxFormatter class directly,
verifying page layout, styles, headers, paragraph formatting,
table formatting, and figure caption formatting.
"""

import shutil
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

from normadocs.formatters.ieee import IEEEDocxFormatter
from normadocs.models import DocumentMetadata


class TestGetMargins(unittest.TestCase):
    """Tests for _get_margins method."""

    def test_default_margins_are_1_inch(self):
        """Default margins should be 1 inch on all sides."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_margins") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            formatter = IEEEDocxFormatter(str(temp_path))
            margins = formatter._get_margins()

            self.assertEqual(margins["top"], 1.0)
            self.assertEqual(margins["bottom"], 1.0)
            self.assertEqual(margins["left"], 1.0)
            self.assertEqual(margins["right"], 1.0)
            self.assertEqual(margins["unit"], "inches")
        finally:
            shutil.rmtree("tests/temp_ieee_margins", ignore_errors=True)

    def test_custom_margins_from_config(self):
        """Custom margins should be read from config."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_custom_margins") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            config = {
                "margins": {
                    "top": 2.0,
                    "bottom": 2.0,
                    "left": 2.5,
                    "right": 2.5,
                    "unit": "inches",
                }
            }
            formatter = IEEEDocxFormatter(str(temp_path), config=config)
            margins = formatter._get_margins()

            self.assertEqual(margins["top"], 2.0)
            self.assertEqual(margins["bottom"], 2.0)
            self.assertEqual(margins["left"], 2.5)
            self.assertEqual(margins["right"], 2.5)
            self.assertEqual(margins["unit"], "inches")
        finally:
            shutil.rmtree("tests/temp_ieee_custom_margins", ignore_errors=True)


class TestGetFontName(unittest.TestCase):
    """Tests for _get_font_name method."""

    def test_default_font_is_times_new_roman(self):
        """Default font should be Times New Roman."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_font") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            formatter = IEEEDocxFormatter(str(temp_path))
            font_name = formatter._get_font_name()
            self.assertEqual(font_name, "Times New Roman")
        finally:
            shutil.rmtree("tests/temp_ieee_font", ignore_errors=True)

    def test_custom_font_from_config(self):
        """Custom font should be read from config."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_custom_font") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            config = {"fonts": {"body": {"name": "Arial", "size": 11}}}
            formatter = IEEEDocxFormatter(str(temp_path), config=config)
            font_name = formatter._get_font_name("body")
            self.assertEqual(font_name, "Arial")
        finally:
            shutil.rmtree("tests/temp_ieee_custom_font", ignore_errors=True)


class TestGetFontSize(unittest.TestCase):
    """Tests for _get_font_size method."""

    def test_default_font_size_is_10pt(self):
        """Default font size should be 10pt."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_size") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            formatter = IEEEDocxFormatter(str(temp_path))
            font_size = formatter._get_font_size()
            self.assertEqual(font_size, 10)
        finally:
            shutil.rmtree("tests/temp_ieee_size", ignore_errors=True)

    def test_custom_font_size_from_config(self):
        """Custom font size should be read from config."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_custom_size") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            config = {"fonts": {"body": {"name": "Times New Roman", "size": 12}}}
            formatter = IEEEDocxFormatter(str(temp_path), config=config)
            font_size = formatter._get_font_size("body")
            self.assertEqual(font_size, 12)
        finally:
            shutil.rmtree("tests/temp_ieee_custom_size", ignore_errors=True)


class TestGetSpacingLine(unittest.TestCase):
    """Tests for _get_spacing_line method."""

    def test_default_line_spacing_is_single(self):
        """Default line spacing should be single."""
        doc = Document()
        temp_path = Path("tests/temp_ieee_spacing") / "blank.docx"
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        try:
            formatter = IEEEDocxFormatter(str(temp_path))
            line_spacing = formatter._get_spacing_line()
            self.assertEqual(line_spacing, "single")
        finally:
            shutil.rmtree("tests/temp_ieee_spacing", ignore_errors=True)


class TestSetupPageLayout(unittest.TestCase):
    """Tests for _setup_page_layout method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_layout")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_top_margin_is_1_inch_by_default(self):
        """Top margin must be 1 inch per IEEE defaults."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_page_layout()
        actual_inches = formatter.doc.sections[0].top_margin.inches
        self.assertAlmostEqual(actual_inches, 1.0, places=2)

    def test_bottom_margin_is_1_inch_by_default(self):
        """Bottom margin must be 1 inch per IEEE defaults."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_page_layout()
        actual_inches = formatter.doc.sections[0].bottom_margin.inches
        self.assertAlmostEqual(actual_inches, 1.0, places=2)

    def test_left_margin_is_1_inch_by_default(self):
        """Left margin must be 1 inch per IEEE defaults."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_page_layout()
        actual_inches = formatter.doc.sections[0].left_margin.inches
        self.assertAlmostEqual(actual_inches, 1.0, places=2)

    def test_right_margin_is_1_inch_by_default(self):
        """Right margin must be 1 inch per IEEE defaults."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_page_layout()
        actual_inches = formatter.doc.sections[0].right_margin.inches
        self.assertAlmostEqual(actual_inches, 1.0, places=2)

    def test_cm_unit_conversion(self):
        """Page layout should handle cm unit conversion correctly."""
        config = {
            "margins": {
                "top": 2.54,
                "bottom": 2.54,
                "left": 2.54,
                "right": 2.54,
                "unit": "cm",
            }
        }
        formatter = IEEEDocxFormatter(str(self.docx_path), config=config)
        formatter._setup_page_layout()
        section = formatter.doc.sections[0]
        self.assertAlmostEqual(section.top_margin.cm, 2.54, places=1)
        self.assertAlmostEqual(section.bottom_margin.cm, 2.54, places=1)


class TestSetupHeaders(unittest.TestCase):
    """Tests for _setup_headers method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_headers")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_header_paragraph_is_right_aligned(self):
        """Header paragraph should be right aligned."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_headers()
        header = formatter.doc.sections[0].header
        header_para = header.paragraphs[0]
        self.assertEqual(header_para.alignment, WD_ALIGN_PARAGRAPH.RIGHT)

    def test_header_contains_page_field(self):
        """Header should contain PAGE field element."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_headers()
        header = formatter.doc.sections[0].header
        header_xml = header._element.xml
        self.assertIn("fldChar", header_xml)
        self.assertIn("PAGE", header_xml)

    def test_header_run_uses_times_new_roman(self):
        """Header run should use Times New Roman font."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_headers()
        header = formatter.doc.sections[0].header
        header_para = header.paragraphs[0]
        for run in header_para.runs:
            if run.text or len(run._r):
                self.assertEqual(run.font.name, "Times New Roman")
                break

    def test_header_run_uses_10pt(self):
        """Header run should use 10pt font size."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._setup_headers()
        header = formatter.doc.sections[0].header
        header_para = header.paragraphs[0]
        for run in header_para.runs:
            if run.text or len(run._r):
                self.assertEqual(run.font.size, Pt(10))
                break


class TestCreateStyles(unittest.TestCase):
    """Tests for _create_styles method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_styles")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_normal_style_is_times_new_roman(self):
        """Normal style must use Times New Roman."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")

    def test_normal_style_is_10pt(self):
        """Normal style must be 10pt."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(normal.font.size, Pt(10))

    def test_normal_style_is_black(self):
        """Normal style must be black (RGB 0,0,0)."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        from docx.shared import RGBColor

        self.assertEqual(normal.font.color.rgb, RGBColor(0, 0, 0))

    def test_normal_style_line_spacing_is_single(self):
        """Normal style must have single line spacing."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(
            normal.paragraph_format.line_spacing_rule,
            WD_LINE_SPACING.SINGLE,
        )

    def test_normal_style_is_justified(self):
        """Normal style paragraph alignment must be justified."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(
            normal.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    def test_normal_style_space_after_is_zero(self):
        """Normal style space_after must be 0pt."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(normal.paragraph_format.space_after, Pt(0))

    def test_normal_style_space_before_is_zero(self):
        """Normal style space_before must be 0pt."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        normal = formatter.doc.styles["Normal"]
        self.assertEqual(normal.paragraph_format.space_before, Pt(0))

    def test_heading_1_is_centered_and_bold(self):
        """Heading 1 style must be centered and bold."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        h1 = formatter.doc.styles["Heading 1"]
        self.assertEqual(
            h1.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        self.assertTrue(h1.font.bold)

    def test_heading_1_has_12pt_spacing_before_and_after(self):
        """Heading 1 must have 12pt space before and after."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        h1 = formatter.doc.styles["Heading 1"]
        self.assertEqual(h1.paragraph_format.space_before, Pt(12))
        self.assertEqual(h1.paragraph_format.space_after, Pt(12))

    def test_heading_2_is_bold_and_italic(self):
        """Heading 2 style must be bold and italic."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._create_styles()
        h2 = formatter.doc.styles["Heading 2"]
        self.assertTrue(h2.font.bold)
        self.assertTrue(h2.font.italic)


class TestFormatParagraphs(unittest.TestCase):
    """Tests for _format_paragraphs method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_paras")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "content.docx"
        doc = Document()
        body_p = doc.add_paragraph("Este es un párrafo de contenido.")
        body_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_heading("Introducción", level=1)
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_body_text_is_justified(self):
        """Body paragraphs must be justified (not headings)."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_paragraphs()
        for p in formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue
            self.assertEqual(
                p.paragraph_format.alignment,
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                f"Paragraph '{p.text}' is not justified",
            )

    def test_body_font_is_times_new_roman(self):
        """Body paragraph runs must have Times New Roman font set."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_paragraphs()
        for p in formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue
            for run in p.runs:
                if run.text.strip():
                    self.assertEqual(
                        run.font.name,
                        "Times New Roman",
                        f"Run '{run.text}' is not Times New Roman",
                    )

    def test_body_font_size_is_10pt(self):
        """Body paragraph runs must have 10pt font size."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_paragraphs()
        for p in formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue
            for run in p.runs:
                if run.text.strip():
                    self.assertEqual(
                        run.font.size,
                        Pt(10),
                        f"Run '{run.text}' is not 10pt",
                    )

    def test_heading_paragraphs_are_not_modified(self):
        """Heading paragraphs should not be modified by _format_paragraphs."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_paragraphs()
        for p in formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                self.assertNotEqual(
                    p.paragraph_format.alignment,
                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                )


class TestFormatTables(unittest.TestCase):
    """Tests for _format_tables method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_tables")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "table.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        cell.text = "Cell 1"
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_table_uses_table_grid_style(self):
        """Tables should use Table Grid style."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_tables()
        table = formatter.doc.tables[0]
        self.assertEqual(table.style.name, "Table Grid")

    def test_table_cell_paragraphs_have_zero_spacing(self):
        """Table cell paragraphs should have zero spacing."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_tables()
        table = formatter.doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self.assertEqual(p.paragraph_format.space_before, Pt(0))
                    self.assertEqual(p.paragraph_format.space_after, Pt(0))


class TestFormatFigures(unittest.TestCase):
    """Tests for _format_figures method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_figures")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "figure.docx"
        doc = Document()
        doc.add_paragraph("Fig. 1: This is a figure caption")
        doc.add_paragraph("Regular paragraph text")
        doc.add_paragraph("Fig. 2. Another figure caption")
        doc.save(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_figure_caption_is_centered(self):
        """Figure captions must be centered."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_figures()
        for p in formatter.doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Fig"):
                self.assertEqual(
                    p.alignment,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    f"Figure caption '{text}' is not centered",
                )

    def test_figure_caption_run_is_italic(self):
        """Figure caption runs must be italic."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_figures()
        for p in formatter.doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Fig"):
                for run in p.runs:
                    self.assertTrue(
                        run.italic,
                        f"Figure caption run '{run.text}' is not italic",
                    )

    def test_figure_caption_uses_times_new_roman(self):
        """Figure caption runs must use Times New Roman."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_figures()
        for p in formatter.doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Fig"):
                for run in p.runs:
                    if run.text.strip():
                        self.assertEqual(
                            run.font.name,
                            "Times New Roman",
                            f"Figure caption run '{run.text}' is not Times New Roman",
                        )

    def test_figure_caption_uses_10pt(self):
        """Figure caption runs must use 10pt font size."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_figures()
        for p in formatter.doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Fig"):
                for run in p.runs:
                    if run.text.strip():
                        self.assertEqual(
                            run.font.size,
                            Pt(10),
                            f"Figure caption run '{run.text}' is not 10pt",
                        )

    def test_non_figure_paragraphs_not_centered(self):
        """Non-figure paragraphs should not be centered."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter._format_figures()
        for p in formatter.doc.paragraphs:
            text = p.text.strip()
            if not text.startswith("Fig"):
                self.assertNotEqual(
                    p.alignment,
                    WD_ALIGN_PARAGRAPH.CENTER,
                )


class TestProcessMethod(unittest.TestCase):
    """Integration tests for the full process method."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_ieee_process")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "full.docx"
        doc = Document()
        doc.add_paragraph("This is body text for the IEEE document.")
        doc.add_heading("Introduction", level=1)
        doc.add_heading("Related Work", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        doc.add_paragraph("Fig. 1: Sample figure")
        doc.save(str(cls.docx_path))
        cls.meta = DocumentMetadata(
            title="IEEE Document Title",
            author="John Doe",
        )

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_process_completes_without_error(self):
        """process() should complete without raising exceptions."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        try:
            formatter.process(self.meta)
        except Exception as e:
            self.fail(f"process() raised an exception: {e}")

    def test_process_applies_all_formatting(self):
        """process() should apply margins, styles, headers, paragraphs, tables, figures."""
        formatter = IEEEDocxFormatter(str(self.docx_path))
        formatter.process(self.meta)

        section = formatter.doc.sections[0]
        self.assertAlmostEqual(section.top_margin.inches, 1.0, places=1)
        self.assertAlmostEqual(section.bottom_margin.inches, 1.0, places=1)

        normal = formatter.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")


class TestSaveMethod(unittest.TestCase):
    """Tests for the save method."""

    def test_save_creates_file(self):
        """save() should create the output file."""
        doc = Document()
        temp_input = Path("tests/temp_ieee_save") / "input.docx"
        temp_input.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_input))

        temp_output = Path("tests/temp_ieee_save") / "output.docx"

        try:
            formatter = IEEEDocxFormatter(str(temp_input))
            formatter._setup_page_layout()
            formatter.save(str(temp_output))

            self.assertTrue(temp_output.exists())
        finally:
            shutil.rmtree("tests/temp_ieee_save", ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
