"""
Tests for formatters/__init__.py - formatter factory and config loading.
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from normadocs.formatters import (
    deep_merge,
    get_formatter,
    list_available_standards,
    load_standard_config,
)
from normadocs.formatters.apa import APADocxFormatter
from normadocs.formatters.icontec import IcontecFormatter
from normadocs.formatters.ieee import IEEEDocxFormatter


class TestGetFormatter(unittest.TestCase):
    def test_get_formatter_apa_returns_apa_formatter(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path))
            self.assertIsInstance(formatter, APADocxFormatter)

    def test_get_formatter_apa7_returns_apa_formatter(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa7", str(doc_path))
            self.assertIsInstance(formatter, APADocxFormatter)

    def test_get_formatter_icontec_returns_icontec_formatter(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("icontec", str(doc_path))
            self.assertIsInstance(formatter, IcontecFormatter)

    def test_get_formatter_ieee_returns_ieee_formatter(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("ieee", str(doc_path))
            self.assertIsInstance(formatter, IEEEDocxFormatter)

    def test_get_formatter_invalid_raises(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            with self.assertRaises(ValueError):
                get_formatter("invalid_style", str(doc_path))

    def test_get_formatter_passes_config(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom_config = {"fonts": {"body": {"name": "Arial", "size": 12}}}
            formatter = get_formatter("apa", str(doc_path), config=custom_config)
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Arial")

    def test_get_formatter_default_config_loaded(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            formatter = get_formatter("apa", str(doc_path))
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")
            self.assertEqual(formatter.config["spacing"]["line"], "double")

    def test_get_formatter_custom_config_merged(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "test.docx"
            Document().save(str(doc_path))
            custom_config = {"spacing": {"line": "single"}}
            formatter = get_formatter("apa", str(doc_path), config=custom_config)
            self.assertEqual(formatter.config["spacing"]["line"], "single")
            self.assertEqual(formatter.config["fonts"]["body"]["name"], "Times New Roman")


class TestLoadStandardConfig(unittest.TestCase):
    def test_load_apa7_config(self):
        config = load_standard_config("apa7")
        self.assertEqual(config["name"], "APA 7th Edition")
        self.assertEqual(config["citation_style"], "apa")

    def test_load_icontec_config(self):
        config = load_standard_config("icontec")
        self.assertEqual(config["name"], "ICONTEC (NTC 1486)")

    def test_load_nonexistent_raises(self):
        with self.assertRaises(FileNotFoundError):
            load_standard_config("nonexistent")


class TestListAvailableStandards(unittest.TestCase):
    def test_returns_list(self):
        standards = list_available_standards()
        self.assertIsInstance(standards, list)

    def test_contains_apa7(self):
        standards = list_available_standards()
        self.assertIn("apa7", standards)

    def test_contains_icontec(self):
        standards = list_available_standards()
        self.assertIn("icontec", standards)


class TestDeepMerge(unittest.TestCase):
    def test_deep_merge_in_place(self):
        base = {"a": {"b": 1}}
        override = {"a": {"c": 2}}
        deep_merge(base, override)
        self.assertEqual(base["a"]["b"], 1)
        self.assertEqual(base["a"]["c"], 2)

    def test_deep_merge_replaces_value(self):
        base = {"x": 1}
        override = {"x": {"y": 2}}
        deep_merge(base, override)
        self.assertEqual(base["x"], {"y": 2})


if __name__ == "__main__":
    unittest.main()
