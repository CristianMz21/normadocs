"""Unit tests for DOCXAnalyzer."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.shared import Inches, Pt


class TestDOCXAnalyzer(unittest.TestCase):
    """Tests for DOCXAnalyzer with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_docx_with_margins(
        self,
        top: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        right: float = 1.0,
    ) -> Path:
        """Create a DOCX with specified margins."""
        path = self.temp_dir / f"test_margins_{top}_{bottom}_{left}_{right}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Test content for margin verification.")
        doc.save(str(path))
        return path

    def _create_docx_with_fonts(
        self,
        body_font: str = "Times New Roman",
        body_size: int = 12,
        heading_font: str = "Times New Roman",
    ) -> Path:
        """Create a DOCX with specified fonts."""
        path = self.temp_dir / f"test_fonts_{body_font.replace(' ', '_')}_{body_size}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        body_para = doc.add_paragraph("This is body text.")
        for run in body_para.runs:
            run.font.name = body_font
            run.font.size = Pt(body_size)

        heading = doc.add_heading("Test Heading", level=1)
        for run in heading.runs:
            run.font.name = heading_font
            run.font.size = Pt(14)

        doc.save(str(path))
        return path

    def _create_docx_with_spacing(
        self,
        line_spacing: float = 2.0,
    ) -> Path:
        """Create a DOCX with specified line spacing."""
        path = self.temp_dir / f"test_spacing_{line_spacing}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        para = doc.add_paragraph("Double spaced text for testing.")
        para.paragraph_format.line_spacing = line_spacing

        doc.save(str(path))
        return path

    def _create_simple_docx(self) -> Path:
        """Create a minimal DOCX for general testing."""
        path = self.temp_dir / "test_simple.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        doc.save(str(path))
        return path

    def test_get_page_info_standard_margins(self) -> None:
        """Test that standard 1-inch margins are extracted correctly."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 1.0)
        try:
            analyzer = DOCXAnalyzer(path)
            info = analyzer.get_page_info()

            assert abs(info.margins[0] - 1.0) < 0.01
            assert abs(info.margins[1] - 1.0) < 0.01
            assert abs(info.margins[2] - 1.0) < 0.01
            assert abs(info.margins[3] - 1.0) < 0.01
            assert abs(info.page_width - 8.5) < 0.01
            assert abs(info.page_height - 11.0) < 0.01
        finally:
            path.unlink(missing_ok=True)

    def test_get_page_info_wrong_margins(self) -> None:
        """Test that non-standard margins are detected."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_docx_with_margins(1.5, 1.5, 1.5, 1.5)
        try:
            analyzer = DOCXAnalyzer(path)
            info = analyzer.get_page_info()

            assert abs(info.margins[0] - 1.5) < 0.01
            assert abs(info.margins[1] - 1.5) < 0.01
            assert abs(info.margins[2] - 1.5) < 0.01
            assert abs(info.margins[3] - 1.5) < 0.01
        finally:
            path.unlink(missing_ok=True)

    def test_get_page_info_letter_size(self) -> None:
        """Test that Letter size is correctly identified."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 1.0)
        try:
            analyzer = DOCXAnalyzer(path)
            info = analyzer.get_page_info()

            assert abs(info.page_width - 8.5) < 0.1
            assert abs(info.page_height - 11.0) < 0.1
        finally:
            path.unlink(missing_ok=True)

    def test_get_paragraphs_info_extracts_text(self) -> None:
        """Test that paragraph text is correctly extracted."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_simple_docx()
        try:
            analyzer = DOCXAnalyzer(path)
            paragraphs = analyzer.get_paragraphs_info()

            texts = [p.text for p in paragraphs if p.text.strip()]
            assert "First paragraph." in texts
            assert "Second paragraph." in texts
        finally:
            path.unlink(missing_ok=True)

    def test_get_paragraphs_info_extracts_runs(self) -> None:
        """Test that run properties (fonts, bold, italic) are extracted."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_docx_with_fonts("Arial", 11)
        try:
            analyzer = DOCXAnalyzer(path)
            paragraphs = analyzer.get_paragraphs_info()

            body_para = next((p for p in paragraphs if "body" in p.text.lower()), None)
            assert body_para is not None
            assert len(body_para.runs) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_paragraphs_count(self) -> None:
        """Test that all paragraphs are returned."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_simple_docx()
        try:
            analyzer = DOCXAnalyzer(path)
            paragraphs = analyzer.get_paragraphs_info()

            assert len(paragraphs) >= 2
        finally:
            path.unlink(missing_ok=True)

    def test_headings_detection(self) -> None:
        """Test that heading styles are detected."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_simple_docx()
        try:
            doc = Document(str(path))
            doc.add_heading("Test Heading", level=1)
            doc.add_heading("Another Heading", level=2)
            doc.save(str(path))

            analyzer = DOCXAnalyzer(path)
            assert analyzer.has_heading(1) is True
            assert analyzer.has_heading(2) is True
            assert analyzer.count_headings(1) >= 1
            assert analyzer.count_headings(2) >= 1
        finally:
            path.unlink(missing_ok=True)

    def test_tables_property(self) -> None:
        """Test that tables are extracted."""
        from normadocs.verifier.docx_analyzer import DOCXAnalyzer

        path = self._create_simple_docx()
        try:
            doc = Document(str(path))
            doc.add_table(rows=3, cols=3)
            doc.save(str(path))

            analyzer = DOCXAnalyzer(path)
            tables = analyzer.tables
            assert len(tables) >= 1
        finally:
            path.unlink(missing_ok=True)


class TestDOCXAnalyzerUtilities(unittest.TestCase):
    """Tests for DOCXAnalyzer utility methods."""

    def test_normalize_font_times(self) -> None:
        """Test font name normalization."""
        from normadocs.verifier.checks.fonts import FontsCheck

        check = FontsCheck()
        assert "times new roman" in check._normalize_font("Times New Roman")
        assert "times new roman" in check._normalize_font("Times New Roman Italic")
        assert "arial" in check._normalize_font("Arial")

    def test_pt_from_emu_conversion(self) -> None:
        """Test EMU to point conversion."""
        from normadocs.verifier.checks.fonts import FontsCheck

        check = FontsCheck()
        assert abs(check._pt_from_emu(12700) - 1.0) < 0.01
        assert abs(check._pt_from_emu(12700 * 12) - 12.0) < 0.01


if __name__ == "__main__":
    unittest.main()
