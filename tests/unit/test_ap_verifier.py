"""Unit tests for APAVerifier orchestrator."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory
from unittest.mock import MagicMock, patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier import CheckCategory, VerificationIssue, VerificationResult
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.paragraphs import ParagraphsCheck


class TestAPAVerifier(unittest.TestCase):
    """Tests for APAVerifier orchestrator."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_simple_docx(self) -> Path:
        """Create a minimal DOCX."""
        from docx import Document
        from docx.shared import Inches

        path = self.temp_dir / "test_simple.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        doc.save(str(path))
        return path

    def _create_mock_pdf_path(self) -> Path:
        """Create a mock PDF path that doesn't exist."""
        return self.temp_dir / "mock_output.pdf"

    def test_verification_context_creation(self) -> None:
        """Test VerificationContext can be created."""
        mock_pdf = MagicMock()
        mock_docx = MagicMock()
        meta = DocumentMetadata(title="Test Document")

        ctx = VerificationContext(pdf=mock_pdf, docx=mock_docx, meta=meta, strict=False)

        assert ctx.pdf is mock_pdf
        assert ctx.docx is mock_docx
        assert ctx.meta is meta
        assert ctx.strict is False

    def test_verification_context_strict_mode(self) -> None:
        """Test VerificationContext strict mode."""
        mock_pdf = MagicMock()
        mock_docx = MagicMock()
        meta = DocumentMetadata(title="Test Document")

        ctx = VerificationContext(pdf=mock_pdf, docx=mock_docx, meta=meta, strict=True)

        assert ctx.strict is True

    def test_verifier_initialization(self) -> None:
        """Test APAVerifier can be initialized."""
        docx_path = self._create_simple_docx()
        pdf_path = self._create_mock_pdf_path()

        try:
            verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path)
            assert verifier.pdf_path == pdf_path
            assert verifier.docx_path == docx_path
        finally:
            docx_path.unlink(missing_ok=True)

    def test_verifier_defaults_to_strict_validation(self) -> None:
        """The public verifier defaults to rejecting all detected warnings."""
        docx_path = self._create_simple_docx()
        pdf_path = self._create_mock_pdf_path()

        try:
            verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path)
            assert verifier.strict is True
        finally:
            docx_path.unlink(missing_ok=True)

    def test_verifier_initialization_with_metadata(self) -> None:
        """Test APAVerifier with metadata."""
        docx_path = self._create_simple_docx()
        pdf_path = self._create_mock_pdf_path()
        meta = DocumentMetadata(title="Test Document")

        try:
            verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
            assert verifier.meta is meta
        finally:
            docx_path.unlink(missing_ok=True)

    def test_verifier_initialization_strict_mode(self) -> None:
        """Test APAVerifier with strict mode."""
        docx_path = self._create_simple_docx()
        pdf_path = self._create_mock_pdf_path()

        try:
            verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, strict=True)
            assert verifier.strict is True
        finally:
            docx_path.unlink(missing_ok=True)

    def test_init_checks_returns_checks(self) -> None:
        """Test that _init_checks returns all expected checks."""
        docx_path = self._create_simple_docx()
        pdf_path = self._create_mock_pdf_path()

        try:
            verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path)
            checks = verifier._init_checks()

            assert len(checks) == 14
            check_categories = [c[0] for c in checks]
            assert "margins" in check_categories
            assert "fonts" in check_categories
            assert "spacing" in check_categories
            assert "headings" in check_categories
            assert "paragraphs" in check_categories
            assert "citations" in check_categories
            assert "equations" in check_categories
        finally:
            docx_path.unlink(missing_ok=True)


class TestVerificationResult(unittest.TestCase):
    """Tests for VerificationResult structure."""

    def test_result_structure(self) -> None:
        """Test that verification result has expected structure."""
        from normadocs.verifier import VerificationResult

        result = VerificationResult(
            passed=True,
            score=100.0,
            issues=[],
        )

        assert result.passed is True
        assert result.score == 100.0

    def test_result_with_issues(self) -> None:
        """Test verification result with issues."""
        from normadocs.verifier import CheckCategory, VerificationIssue, VerificationResult

        issue = VerificationIssue(
            check=f"{CheckCategory.MARGINS}.top",
            severity="error",
            expected="1.0 inches",
            actual="2.0 inches",
        )

        result = VerificationResult(
            passed=False,
            score=80.0,
            issues=[issue],
        )

        assert result.passed is False
        assert len(result.issues) == 1
        assert result.issues[0].check == "margins.top"


class TestAPAVerifierEndToEnd(unittest.TestCase):
    """End-to-end tests for APAVerifier.verify_all on a compliant DOCX."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_compliant_docx(self, justified: bool = False) -> Path:
        path = self.temp_path / "compliant.docx"
        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        section.header.paragraphs[0].add_run("1")
        section.first_page_header.paragraphs[0].add_run("1")

        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Author"

        def style_run(run, bold=False, italic=False) -> None:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(cover_title.add_run("Test Document"), bold=True)
        cover_author = doc.add_paragraph()
        cover_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(cover_author.add_run("Author"))
        cover_date = doc.add_paragraph()
        cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(cover_date.add_run("2026"))

        content_title = doc.add_paragraph("Test Document", style="Heading 1")
        content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        content_title.paragraph_format.line_spacing = 2.0
        content_title.paragraph_format.space_before = Inches(0)
        content_title.paragraph_format.space_after = Inches(0)
        style_run(content_title.runs[0], bold=True)

        intro_heading = doc.add_paragraph("Introducción", style="Heading 1")
        intro_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro_heading.paragraph_format.line_spacing = 2.0
        intro_heading.paragraph_format.space_before = Inches(0)
        intro_heading.paragraph_format.space_after = Inches(0)
        style_run(intro_heading.runs[0], bold=True)
        intro_body = doc.add_paragraph("Introduction content for the report structure.")
        intro_body.paragraph_format.first_line_indent = Inches(0.5)
        intro_body.paragraph_format.line_spacing = 2.0
        style_run(intro_body.runs[0])

        development_heading = doc.add_paragraph("Desarrollo", style="Heading 1")
        development_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        development_heading.paragraph_format.line_spacing = 2.0
        development_heading.paragraph_format.space_before = Inches(0)
        development_heading.paragraph_format.space_after = Inches(0)
        style_run(development_heading.runs[0], bold=True)
        development_body = doc.add_paragraph("Development content for the report structure.")
        development_body.paragraph_format.first_line_indent = Inches(0.5)
        development_body.paragraph_format.line_spacing = 2.0
        style_run(development_body.runs[0])

        conclusion_heading = doc.add_paragraph("Conclusiones", style="Heading 1")
        conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conclusion_heading.paragraph_format.line_spacing = 2.0
        conclusion_heading.paragraph_format.space_before = Inches(0)
        conclusion_heading.paragraph_format.space_after = Inches(0)
        style_run(conclusion_heading.runs[0], bold=True)
        conclusion_body = doc.add_paragraph("Conclusions content for the report structure.")
        conclusion_body.paragraph_format.first_line_indent = Inches(0.5)
        conclusion_body.paragraph_format.line_spacing = 2.0
        style_run(conclusion_body.runs[0])

        for i in range(3):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 2.0
            para.alignment = alignment
            style_run(para.add_run(f"Body paragraph {i + 1} of the document text."))

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(caption.add_run("Tabla 1"), bold=True)

        title = doc.add_paragraph()
        title.paragraph_format.first_line_indent = Inches(0.5)
        title.paragraph_format.line_spacing = 2.0
        title.alignment = alignment
        style_run(title.add_run("Title of the Table in Italic"), italic=True)

        doc.add_table(rows=2, cols=2)

        figura = doc.add_paragraph()
        figura.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(figura.add_run("Figura 1"), bold=True)

        ref_heading = doc.add_paragraph(style="Heading 1")
        ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_heading.paragraph_format.line_spacing = 2.0
        ref_heading.paragraph_format.space_before = Inches(0)
        ref_heading.paragraph_format.space_after = Inches(0)
        style_run(ref_heading.add_run("Referencias"), bold=True)

        for text in (
            "Alpha, A. (2023). First work. Publisher.",
            "Bravo, B. (2024). Second work. Journal.",
        ):
            ref = doc.add_paragraph()
            ref.paragraph_format.left_indent = Inches(0.5)
            ref.paragraph_format.first_line_indent = Inches(-0.5)
            ref.paragraph_format.line_spacing = 2.0
            ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_run(ref.add_run(text))

        doc.save(str(path))
        return path

    def _pdf_path(self) -> Path:
        path = self.temp_path / "compliant.pdf"
        path.touch()
        return path

    def test_verify_all_on_compliant_doc_passes(self) -> None:
        """A fully APA-compliant DOCX should pass verify_all with no errors."""
        docx_path = self._create_compliant_docx()
        verifier = APAVerifier(
            pdf_path=self._pdf_path(),
            docx_path=docx_path,
            meta=DocumentMetadata(title="Test Document"),
        )
        try:
            result = verifier.verify_all()
            self.assertTrue(result.passed, f"Expected passed but got issues: {result.errors}")
            self.assertEqual(result.errors, [], f"Expected no errors but got: {result.errors}")
        finally:
            verifier.close()

    def test_verify_all_strict_mode_fails_on_warning(self) -> None:
        """Strict mode should fail when a justification warning is present."""
        docx_path = self._create_compliant_docx(justified=True)
        verifier = APAVerifier(
            pdf_path=self._pdf_path(),
            docx_path=docx_path,
            meta=DocumentMetadata(title="Test Document"),
            strict=True,
        )
        try:
            result = verifier.verify_all()
            self.assertFalse(
                result.passed, f"Expected strict mode to fail on warning but got: {result.passed}"
            )
            self.assertEqual(result.warnings, [], "Strict mode must promote warnings to errors")
            self.assertTrue(
                any(
                    "justification" in issue.check or "strict_alignment" in issue.check
                    for issue in result.errors
                ),
                f"Expected strict justification error, got: {result.errors}",
            )
        finally:
            verifier.close()

    def test_verify_all_extracts_meta_from_docx(self) -> None:
        """verify_all with meta=None should extract metadata from the DOCX (no crash)."""
        docx_path = self._create_compliant_docx()
        verifier = APAVerifier(pdf_path=self._pdf_path(), docx_path=docx_path)
        try:
            result = verifier.verify_all()
            self.assertIsNotNone(result, "verify_all should return a result")
            self.assertEqual(result.errors, [], f"Expected no errors but got: {result.errors}")
        finally:
            verifier.close()

    def test_verify_all_swallows_check_exception_as_error(self) -> None:
        """An exception inside a check should become a `<category>.check_failed` error."""
        docx_path = self._create_compliant_docx()
        verifier = APAVerifier(
            pdf_path=self._pdf_path(),
            docx_path=docx_path,
            meta=DocumentMetadata(title="Test Document"),
        )
        try:
            with patch.object(ParagraphsCheck, "run", side_effect=RuntimeError("boom")):
                result = verifier.verify_all()
            failed = [i for i in result.errors if i.check.endswith(".check_failed")]
            self.assertGreater(
                len(failed), 0, f"Expected a check_failed error but got: {result.errors}"
            )
            self.assertIn("boom", failed[0].actual, f"Exception text should be in actual: {failed}")
            self.assertFalse(result.passed, "A check_failed error must fail the run")
        finally:
            verifier.close()


class TestAPAVerifierReport(unittest.TestCase):
    """Tests for APAVerifier.generate_report across text/markdown/html formats."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _verifier(self) -> APAVerifier:
        return APAVerifier(
            pdf_path=self.temp_path / "report.pdf",
            docx_path=self.temp_path / "report.docx",
        )

    def _result_with_one_error_and_one_warning(self) -> VerificationResult:
        error = VerificationIssue(
            check=f"{CheckCategory.MARGINS}.top",
            severity="error",
            expected="1.0 inches",
            actual="2.0 inches",
            evidence="top margin too big",
        )
        warning = VerificationIssue(
            check=f"{CheckCategory.SPACING}.line_spacing",
            severity="warning",
            expected="Double spacing",
            actual="Some paragraphs single-spaced",
        )
        return VerificationResult(
            passed=False,
            score=70.0,
            issues=[error, warning],
            errors=[error],
            warnings=[warning],
            pdf_path=self.temp_path / "report.pdf",
        )

    def test_generate_report_text(self) -> None:
        """Text report should contain Errors, Warnings and Score sections."""
        verifier = self._verifier()
        report = verifier.generate_report(self._result_with_one_error_and_one_warning(), "text")
        self.assertIn("Errors", report, f"Text report missing Errors: {report}")
        self.assertIn("Warnings", report, f"Text report missing Warnings: {report}")
        self.assertIn("Score", report, f"Text report missing Score: {report}")

    def test_generate_report_markdown(self) -> None:
        """Markdown report should contain title, File and Errors headers."""
        verifier = self._verifier()
        report = verifier.generate_report(self._result_with_one_error_and_one_warning(), "markdown")
        self.assertIn("# APA 7th", report, f"Markdown report missing title: {report}")
        self.assertIn("**File**", report, f"Markdown report missing File: {report}")
        self.assertIn("## Errors", report, f"Markdown report missing Errors header: {report}")

    def test_generate_report_html(self) -> None:
        """HTML report should contain doctype plus error/warning CSS classes."""
        verifier = self._verifier()
        report = verifier.generate_report(self._result_with_one_error_and_one_warning(), "html")
        self.assertIn("<!DOCTYPE html>", report, f"HTML report missing doctype: {report}")
        self.assertIn('class="err"', report, f"HTML report missing error class: {report}")
        self.assertIn('class="warn"', report, f"HTML report missing warn class: {report}")


class TestAPAVerifierDocxDiscovery(unittest.TestCase):
    """Tests for APAVerifier DOCX auto-discovery (suffix-stripping) and missing-DOCX errors."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def test_find_docx_by_suffix_stripped(self) -> None:
        """PDF 'Foo_APA.pdf' should discover DOCX 'Foo.docx' (suffix stripping L92-98)."""
        docx_path = self.temp_path / "Foo.docx"
        docx_path.touch()
        pdf_path = self.temp_path / "Foo_APA.pdf"

        verifier = APAVerifier(pdf_path=pdf_path, docx_path=None)
        self.assertEqual(verifier.docx_path, docx_path, "Suffix stripping should find Foo.docx")

    def test_docx_property_raises_when_missing(self) -> None:
        """Accessing .docx when the DOCX file does not exist should raise FileNotFoundError."""
        pdf_path = self.temp_path / "missing.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=self.temp_path / "absent.docx")
        with self.assertRaises(FileNotFoundError):
            _ = verifier.docx


class TestAPAVerifierClose(unittest.TestCase):
    """Tests for APAVerifier.close resetting analyzer handles."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def test_close_resets_analyzers(self) -> None:
        """close() should set both cached analyzer handles back to None."""
        pdf_path = self.temp_path / "close.pdf"
        pdf_path.touch()
        docx_path = self.temp_path / "close.docx"
        Document().save(str(docx_path))

        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path)
        _ = verifier.pdf
        _ = verifier.docx
        self.assertIsNotNone(verifier._pdf_analyzer, "PDF analyzer should be cached before close")
        self.assertIsNotNone(verifier._docx_analyzer, "DOCX analyzer should be cached before close")
        verifier.close()
        self.assertIsNone(verifier._pdf_analyzer, "close should reset PDF analyzer to None")
        self.assertIsNone(verifier._docx_analyzer, "close should reset DOCX analyzer to None")


if __name__ == "__main__":
    unittest.main()
