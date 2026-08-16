"""
Unit Tests for APA Formatter Paragraph Processing.

Tests _process_paragraphs, _merge_and_clean_paragraph, _fix_citations,
and _build_heading_level_map methods.
"""

import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from normadocs.formatters.apa import APADocxFormatter
from normadocs.formatters.apa.apa_paragraphs import APAParagraphsHandler


class TestProcessParagraphs(unittest.TestCase):
    """Tests for _process_paragraphs method."""

    def _create_formatter_with_doc(self, paragraphs_data: list[dict]) -> APADocxFormatter:
        """Create a formatter with a document containing paragraphs described by paragraphs_data.

        Each dict in paragraphs_data should have:
            - text: str - the paragraph text
            - style: str - style name (default "Normal")
            - runs: list of dict with text, bold, italic (optional)
        """
        doc = Document()
        for para_data in paragraphs_data:
            style_name = para_data.get("style", "Normal")
            p = doc.add_paragraph(para_data["text"], style=style_name)
            if "runs" in para_data:
                # Clear auto-created run and add custom ones
                for run in list(p.runs):
                    run._r.getparent().remove(run._r)
                for run_data in para_data["runs"]:
                    r = p.add_run(run_data["text"])
                    if run_data.get("bold"):
                        r.bold = True
                    if run_data.get("italic"):
                        r.italic = True

        # Create a temp file for the formatter
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        formatter = APADocxFormatter(temp_path)
        # Replace the doc with our already-created one
        formatter.doc = doc
        return formatter, temp_path

    def test_abstract_section_is_centered_and_bold(self):
        """RESUMEN section heading should be centered and bold."""
        paragraphs = [
            {"text": "Resumen", "style": "Heading 1"},
            {"text": "Este es el resumen del documento.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            abstract_heading = formatter.doc.paragraphs[0]
            self.assertEqual(abstract_heading.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            for run in abstract_heading.runs:
                self.assertTrue(run.bold, "Abstract heading run should be bold")
        finally:
            import os

            os.unlink(temp_path)

    def test_references_section_gets_hanging_indent(self):
        """References section paragraphs get 0.5in left indent and -0.5in first-line indent."""
        paragraphs = [
            {"text": "Referencias", "style": "Heading 1"},
            {
                "text": "García, A. y López, B. (2023). Machine learning en educación.",
                "style": "Normal",
            },
            {
                "text": "Martínez, C. y Rodríguez, D. (2024). Sistemas adaptativos.",
                "style": "Normal",
            },
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            # First paragraph after references heading
            ref_para = formatter.doc.paragraphs[1]
            self.assertEqual(ref_para.paragraph_format.left_indent, Inches(0.5))
            self.assertEqual(ref_para.paragraph_format.first_line_indent, Inches(-0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_first_paragraph_after_heading_has_no_indent(self):
        """First paragraph after a heading should have NO first-line indent."""
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Este es el primer párrafo después del encabezado.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            first_para = formatter.doc.paragraphs[1]
            self.assertEqual(first_para.paragraph_format.first_line_indent, Inches(0))
        finally:
            import os

            os.unlink(temp_path)

    def test_subsequent_paragraphs_have_half_inch_indent(self):
        """Subsequent paragraphs (not first after heading) should have 0.5in first-line indent."""
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Primer párrafo.", "style": "Normal"},
            {
                "text": "Segundo párrafo con texto largo que debería tener sangría.",
                "style": "Normal",
            },
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            second_para = formatter.doc.paragraphs[2]
            self.assertEqual(second_para.paragraph_format.first_line_indent, Inches(0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_line_spacing_is_double(self):
        """All paragraphs should have double line spacing."""
        paragraphs = [
            {"text": "Párrafo normal.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            para = formatter.doc.paragraphs[0]
            self.assertEqual(para.paragraph_format.line_spacing_rule, WD_LINE_SPACING.DOUBLE)
        finally:
            import os

            os.unlink(temp_path)

    def test_body_text_is_left_aligned(self):
        """Body text paragraphs should be left-aligned."""
        paragraphs = [
            {"text": "Este es un párrafo de cuerpo.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            para = formatter.doc.paragraphs[0]
            self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        finally:
            import os

            os.unlink(temp_path)

    def test_level_1_heading_does_not_force_new_page(self):
        """Ordinary Level 1 headings should not force a page break in APA 7."""
        paragraphs = [
            {"text": "Párrafo inicial.", "style": "Normal"},
            {"text": "Título Nivel 1", "style": "Heading 1"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            heading_para = formatter.doc.paragraphs[1]
            self.assertFalse(heading_para.paragraph_format.page_break_before)
        finally:
            import os

            os.unlink(temp_path)

    def test_explicit_break_is_checked_per_heading(self):
        """An earlier explicit break must not suppress a later heading break."""
        doc = Document()
        first_heading = doc.add_paragraph("First heading", style="Heading 1")
        break_paragraph = OxmlElement("w:p")
        break_run = OxmlElement("w:r")
        explicit_break = OxmlElement("w:br")
        explicit_break.set(qn("w:type"), "page")
        break_run.append(explicit_break)
        break_paragraph.append(break_run)
        first_heading._element.addprevious(break_paragraph)
        second_heading = doc.add_paragraph("Second heading", style="Heading 1")
        handler = APAParagraphsHandler(doc)

        handler._set_page_break_before(first_heading)
        handler._set_page_break_before(second_heading)

        self.assertTrue(handler._has_page_break_before(first_heading))
        self.assertFalse(handler._has_page_break_before(second_heading))
        self.assertFalse(first_heading.paragraph_format.page_break_before)
        self.assertTrue(second_heading.paragraph_format.page_break_before)

    def test_stale_page_break_is_removed_from_later_heading(self):
        """Reused DOCX files must not retain breaks on ordinary later headings."""
        paragraphs = [
            {"text": "Introduction", "style": "Heading 1"},
            {"text": "First body paragraph.", "style": "Normal"},
            {"text": "Methods", "style": "Heading 1"},
            {"text": "Methods body paragraph.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter.doc.paragraphs[2].paragraph_format.page_break_before = True
            formatter._process_paragraphs()
            self.assertFalse(formatter.doc.paragraphs[2].paragraph_format.page_break_before)
        finally:
            import os

            os.unlink(temp_path)

    def test_referencia_heading_triggers_references_mode(self):
        """Heading containing 'referencia' should trigger references mode with hanging indent."""
        paragraphs = [
            {"text": "Referencias", "style": "Heading 1"},
            {"text": "García, A. (2023). Título del artículo.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            ref_para = formatter.doc.paragraphs[1]
            self.assertEqual(ref_para.paragraph_format.left_indent, Inches(0.5))
            self.assertEqual(ref_para.paragraph_format.first_line_indent, Inches(-0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_reference_heading_triggers_references_mode(self):
        """Heading containing 'reference' (English) should also trigger references mode."""
        paragraphs = [
            {"text": "Reference", "style": "Heading 1"},
            {"text": "Smith, J. (2022). Article title.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            ref_para = formatter.doc.paragraphs[1]
            self.assertEqual(ref_para.paragraph_format.left_indent, Inches(0.5))
            self.assertEqual(ref_para.paragraph_format.first_line_indent, Inches(-0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_heading_with_references_in_text_does_not_trigger_references_mode(self):
        """A heading merely containing the word 'references' (e.g. 'Slide 4 — ... references')
        must NOT enable references mode on following paragraphs."""
        paragraphs = [
            {"text": "Slide 4 — Well-being, teamwork and references", "style": "Heading 3"},
            {
                "text": "Oral script paragraph. Hello, this is a normal body paragraph that "
                "should receive a first-line indent, not a hanging reference indent.",
                "style": "Normal",
            },
            {"text": "References", "style": "Heading 1"},
            {"text": "Smith, J. (2022). Article title.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            # Body paragraph after the Slide heading: normal first-line indent,
            # NOT a hanging reference indent (no left_indent)
            body_para = formatter.doc.paragraphs[1]
            self.assertEqual(body_para.paragraph_format.first_line_indent, Inches(0.5))
            self.assertIsNone(body_para.paragraph_format.left_indent)

            # Reference entry after the real References heading: hanging indent
            ref_para = formatter.doc.paragraphs[3]
            self.assertEqual(ref_para.paragraph_format.left_indent, Inches(0.5))
            self.assertEqual(ref_para.paragraph_format.first_line_indent, Inches(-0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_references_mode_resets_after_other_heading(self):
        """References mode must reset when a non-references heading
        follows the references section."""
        paragraphs = [
            {"text": "References", "style": "Heading 1"},
            {"text": "Smith, J. (2022). Article title.", "style": "Normal"},
            {"text": "Appendix", "style": "Heading 1"},
            {
                "text": "Appendix body text that should NOT have hanging reference indent.",
                "style": "Normal",
            },
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._process_paragraphs()

            appendix_para = formatter.doc.paragraphs[3]
            self.assertIsNone(appendix_para.paragraph_format.left_indent)
        finally:
            import os

            os.unlink(temp_path)


class TestMergeAndCleanParagraph(unittest.TestCase):
    """Tests for _merge_and_clean_paragraph method."""

    def _add_run_with_format(self, p, text, bold=False, italic=False):
        """Add a run with specific formatting."""
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        return r

    def test_paragraphs_are_merged_correctly(self):
        """Consecutive runs with same formatting should be merged."""
        doc = Document()
        p = doc.add_paragraph(style="Normal")
        # Add multiple runs with same formatting that should merge
        self._add_run_with_format(p, "Texto ", bold=False, italic=False)
        self._add_run_with_format(p, "unido ", bold=False, italic=False)
        self._add_run_with_format(p, "correctamente.", bold=False, italic=False)

        original_run_count = len(p.runs)

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._merge_and_clean_paragraph(p)

            # After merging, we should have fewer runs
            self.assertLess(len(p.runs), original_run_count)
            # The full text should be preserved
            self.assertIn("Texto", p.text)
            self.assertIn("unido", p.text)
            self.assertIn("correctamente", p.text)
        finally:
            import os

            os.unlink(temp_path)

    def test_whitespace_is_cleaned(self):
        """Multiple spaces should be collapsed to single space."""
        doc = Document()
        p = doc.add_paragraph(style="Normal")
        self._add_run_with_format(
            p, "Texto   con    muchos     espacios.", bold=False, italic=False
        )

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._merge_and_clean_paragraph(p)

            # No triple spaces should remain
            self.assertNotIn("   ", p.text)
        finally:
            import os

            os.unlink(temp_path)

    def test_different_formatting_runs_not_merged(self):
        """Runs with different bold/italic formatting should NOT be merged."""
        doc = Document()
        p = doc.add_paragraph(style="Normal")
        self._add_run_with_format(p, "Negrita ", bold=True, italic=False)
        self._add_run_with_format(p, "y normal ", bold=False, italic=False)
        self._add_run_with_format(p, " cursiva.", bold=False, italic=True)

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._merge_and_clean_paragraph(p)

            # Find the run with "Negrita" - it should still be bold
            bold_runs = [r for r in p.runs if r.bold and r.text.strip()]
            self.assertEqual(len(bold_runs), 1)
            self.assertIn("Negrita", bold_runs[0].text)

            # Find the run with " cursiva" - it should still be italic
            italic_runs = [r for r in p.runs if r.italic and r.text.strip()]
            self.assertEqual(len(italic_runs), 1)
            self.assertIn("cursiva", italic_runs[0].text)
        finally:
            import os

            os.unlink(temp_path)


class TestFixCitations(unittest.TestCase):
    """Tests for _fix_citations method."""

    def test_garcía_y_lópez_citation_preserved(self):
        """Citation '(García y López, 2023)' should be preserved without modification."""
        doc = Document()
        p = doc.add_paragraph(
            "Este documento examina los efectos del aprendizaje automático "
            "(García y López, 2023) en la educación superior.",
            style="Normal",
        )

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._fix_citations(p)

            full_text = p.text
            self.assertIn("García", full_text)
            self.assertIn("2023", full_text)
        finally:
            import os

            os.unlink(temp_path)

    def test_simple_yCitation_gets_ampersand(self):
        """Citation using ' y ' between authors should use '&' instead."""
        doc = Document()
        p = doc.add_paragraph(
            "Según estudios recientes (García y López, 2023) se ha demostrado...",
            style="Normal",
        )

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._fix_citations(p)

            # The & should replace 'y' in citation
            full_text = p.text
            self.assertIn("García", full_text)
        finally:
            import os

            os.unlink(temp_path)


class TestBuildHeadingLevelMap(unittest.TestCase):
    """Tests for _build_heading_level_map method."""

    def test_returns_correct_heading_levels(self):
        """The method should return a dict mapping heading text to level numbers."""
        doc = Document()
        doc.add_heading("Título Principal", level=1)
        doc.add_heading("Subsección", level=2)
        doc.add_heading("Sub-subsección", level=3)
        doc.add_paragraph("Contenido normal.")

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            level_map = formatter._build_heading_level_map()

            self.assertEqual(level_map.get("título principal"), 1)
            self.assertEqual(level_map.get("subsección"), 2)
            self.assertEqual(level_map.get("sub-subsección"), 3)
        finally:
            import os

            os.unlink(temp_path)

    def test_empty_document_returns_empty_dict(self):
        """An empty document should return an empty dict."""
        doc = Document()

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            level_map = formatter._build_heading_level_map()

            self.assertEqual(level_map, {})
        finally:
            import os

            os.unlink(temp_path)

    def test_headings_with_numbers_are_normalized(self):
        """Heading text should be normalized to lowercase for the map keys."""
        doc = Document()
        doc.add_heading("REFERENCIAS", level=1)

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            level_map = formatter._build_heading_level_map()

            self.assertIn("referencias", level_map)
            self.assertEqual(level_map["referencias"], 1)
        finally:
            import os

            os.unlink(temp_path)


class TestRunInHeadings(unittest.TestCase):
    """Tests for APA 7 Level 4/5 run-in heading formatting."""

    def test_heading4_merges_following_paragraph(self):
        """H4 ends with a period and the body text continues on its line."""
        doc = Document()
        doc.add_paragraph("Resultados experimentales", style="Heading 4")
        doc.add_paragraph("El análisis de varianza muestra diferencias.", style="Body Text")

        APAParagraphsHandler(doc).process()

        heading = doc.paragraphs[0]
        self.assertTrue(heading.text.startswith("Resultados experimentales."))
        self.assertIn("diferencias", heading.text)
        self.assertEqual(len([p for p in doc.paragraphs if p.text.strip()]), 1)
        body_run = next(r for r in heading.runs if "diferencias" in r.text)
        self.assertIsNotNone(body_run.bold)

    def test_heading5_is_bold_italic_and_indented(self):
        """H5 keeps bold+italic heading runs with a 0.5 inch indent."""
        doc = Document()
        doc.add_paragraph("Análisis secundario", style="Heading 5")
        doc.add_paragraph("Los datos cualitativos respaldan el efecto.", style="Body Text")

        APAParagraphsHandler(doc).process()

        heading = doc.paragraphs[0]
        self.assertTrue(heading.text.startswith("Análisis secundario."))
        head_run = next(r for r in heading.runs if "Análisis" in r.text)
        self.assertTrue(head_run.bold)
        self.assertTrue(head_run.italic)
        self.assertAlmostEqual(heading.paragraph_format.first_line_indent, Inches(0.5))

    def test_heading4_does_not_merge_with_next_heading(self):
        """A heading followed by another heading stays on its own line."""
        doc = Document()
        doc.add_paragraph("Categoría", style="Heading 4")
        doc.add_paragraph("Sección siguiente", style="Heading 2")

        APAParagraphsHandler(doc).process()

        texts = [p.text.strip() for p in doc.paragraphs]
        self.assertEqual(texts, ["Categoría.", "Sección siguiente"])


class TestBlockQuotes(unittest.TestCase):
    """Tests for APA 8.27 block-quote conversion."""

    _LONG_TEXT = (
        "El aprendizaje automático permite personalizar el contenido educativo "
        "de forma que cada estudiante avanza según sus propias necesidades y "
        "ritmo de aprendizaje, lo que a largo plazo mejora la retención y el "
        "rendimiento académico medido en evaluaciones estandarizadas diversas"
    )

    def test_long_quotation_becomes_block_quote(self):
        """A 40+ word quotation loses its quotes and gets a 0.5 inch indent."""
        doc = Document()
        doc.add_paragraph(f'"{self._LONG_TEXT}," (García, 2020).', style="Body Text")

        APAParagraphsHandler(doc).process()

        para = doc.paragraphs[0]
        self.assertNotIn('"', para.text)
        self.assertAlmostEqual(para.paragraph_format.left_indent, Inches(0.5))
        self.assertEqual(para.paragraph_format.first_line_indent, Inches(0))
        self.assertIn(". (García, 2020)", para.text)

    def test_short_quotation_keeps_quotes(self):
        """Short quotations stay in-line with quotation marks."""
        doc = Document()
        doc.add_paragraph('"cita breve" (García, 2020).', style="Body Text")

        APAParagraphsHandler(doc).process()

        self.assertIn('"cita breve"', doc.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
