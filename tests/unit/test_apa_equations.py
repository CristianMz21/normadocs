"""Unit tests for the APA equations handler (display math numbering/layout)."""

import unittest

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from normadocs.formatters.apa.apa_equations import APAEquationsHandler
from normadocs.formatters.apa.apa_paragraphs import APAParagraphsHandler


def _add_display_math(doc: Document, latex: str = "x") -> None:
    """Add a paragraph holding a pandoc-style display equation (m:oMathPara)."""
    p = doc.add_paragraph()
    o_math_para = OxmlElement("m:oMathPara")
    o_math = OxmlElement("m:oMath")
    m_r = OxmlElement("m:r")
    m_t = OxmlElement("m:t")
    m_t.text = latex
    m_r.append(m_t)
    o_math.append(m_r)
    o_math_para.append(o_math)
    p._element.append(o_math_para)


class TestFormatEquations(unittest.TestCase):
    """Tests for APA 7 display equation numbering and tab layout."""

    def test_display_math_gets_numbered_with_tabs(self):
        """A pure display equation becomes [tab] oMath [tab] (1)."""
        doc = Document()
        _add_display_math(doc, "E=mc^2")

        APAEquationsHandler(doc).format_equations()

        p = doc.paragraphs[0]
        self.assertIn("(1)", p.text)
        self.assertEqual(p.paragraph_format.first_line_indent, Inches(0))
        self.assertIsNone(p._element.find(f".//{qn('m:oMathPara')}"))
        self.assertIsNotNone(p._element.find(f"./{qn('m:oMath')}"))

        alignments = [(tab.position, tab.alignment) for tab in p.paragraph_format.tab_stops]
        positions = [pos for pos, _ in alignments]
        self.assertTrue(any(abs(pos - Inches(3.25)) < 1000 for pos in positions))
        right_tabs = [align for pos, align in alignments if abs(pos - Inches(6.5)) < 1000]
        self.assertEqual(right_tabs, [WD_TAB_ALIGNMENT.RIGHT])

        children = list(p._element)
        tags = [child.tag.split("}")[-1] for child in children]
        self.assertEqual(tags, ["pPr", "r", "oMath", "r"])

    def test_equations_are_numbered_consecutively(self):
        """The second display equation is numbered (2)."""
        doc = Document()
        _add_display_math(doc, "a+b")
        _add_display_math(doc, "c=d")

        APAEquationsHandler(doc).format_equations()

        self.assertIn("(1)", doc.paragraphs[0].text)
        self.assertIn("(2)", doc.paragraphs[1].text)

    def test_mixed_paragraph_with_math_is_left_alone(self):
        """Display math mixed with text keeps pandoc's layout (no renumbering)."""
        doc = Document()
        p = doc.add_paragraph("texto alrededor ")
        o_math_para = OxmlElement("m:oMathPara")
        o_math = OxmlElement("m:oMath")
        o_math_para.append(o_math)
        p._element.append(o_math_para)

        APAEquationsHandler(doc).format_equations()

        self.assertNotIn("(1)", doc.paragraphs[0].text)
        self.assertIsNotNone(doc.paragraphs[0]._element.find(f".//{qn('m:oMathPara')}"))

    def test_numbering_can_be_disabled(self):
        """config equations.numbering=false keeps the layout without numbers."""
        doc = Document()
        _add_display_math(doc, "z=9")
        handler = APAEquationsHandler(doc, {"equations": {"numbering": False}})

        handler.format_equations()

        self.assertNotIn("(1)", doc.paragraphs[0].text)
        children = list(doc.paragraphs[0]._element)
        tags = [child.tag.split("}")[-1] for child in children]
        self.assertEqual(tags, ["pPr", "r", "oMath"])

    def test_merge_and_clean_skips_math_paragraphs(self):
        """Run consolidation must not reorder text around an equation."""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("antes ")
        r1.font.name = "Times New Roman"
        o_math = OxmlElement("m:oMath")
        p._element.append(o_math)
        r2 = p.add_run(" después")
        r2.font.name = "Times New Roman"

        APAParagraphsHandler(doc).fix_text_spacing_global()

        children = list(doc.paragraphs[0]._element)
        tags = [child.tag.split("}")[-1] for child in children]
        self.assertEqual(tags, ["pPr", "r", "oMath", "r"])
        self.assertIn("antes", p.runs[0].text)
        self.assertIn("después", p.runs[-1].text)


if __name__ == "__main__":
    unittest.main()
