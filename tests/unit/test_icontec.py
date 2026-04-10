"""
Unit tests for the ICONTEC document formatter.

These tests exercise the IcontecFormatter class directly,
verifying page layout, styles, cover page content, and paragraph
formatting without going through the full CLI pipeline.
"""

import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

from normadocs.formatters.icontec import IcontecFormatter
from normadocs.models import DocumentMetadata


class TestSetupPageLayout(unittest.TestCase):
    """Tests for _setup_page_layout method."""

    @classmethod
    def setUpClass(cls):
        """Create a blank DOCX to test with."""
        cls.temp_dir = Path("tests/temp_icontec_unit")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))
        cls.formatter = IcontecFormatter(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        """Clean up temporary files."""
        import shutil

        if cls.temp_dir.exists():
            shutil.rmtree(cls.temp_dir)

    def test_top_margin_is_3cm(self):
        """Top margin must be 3 cm per NTC 1486."""
        self.formatter._setup_page_layout()
        actual_cm = self.formatter.doc.sections[0].top_margin.cm
        self.assertAlmostEqual(actual_cm, 3.0, places=1)

    def test_bottom_margin_is_2cm(self):
        """Bottom margin must be 2 cm per NTC 1486."""
        self.formatter._setup_page_layout()
        actual_cm = self.formatter.doc.sections[0].bottom_margin.cm
        self.assertAlmostEqual(actual_cm, 2.0, places=1)

    def test_left_margin_is_3cm(self):
        """Left margin must be 3 cm per NTC 1486."""
        self.formatter._setup_page_layout()
        actual_cm = self.formatter.doc.sections[0].left_margin.cm
        self.assertAlmostEqual(actual_cm, 3.0, places=1)

    def test_right_margin_is_2cm(self):
        """Right margin must be 2 cm per NTC 1486."""
        self.formatter._setup_page_layout()
        actual_cm = self.formatter.doc.sections[0].right_margin.cm
        self.assertAlmostEqual(actual_cm, 2.0, places=1)


class TestCreateStyles(unittest.TestCase):
    """Tests for _create_styles method."""

    @classmethod
    def setUpClass(cls):
        """Create a blank DOCX to test with."""
        cls.temp_dir = Path("tests/temp_icontec_styles")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))
        cls.formatter = IcontecFormatter(str(cls.docx_path))

    @classmethod
    def tearDownClass(cls):
        """Clean up temporary files."""
        import shutil

        if cls.temp_dir.exists():
            shutil.rmtree(cls.temp_dir)

    def test_normal_style_is_arial_12pt(self):
        """Normal style must use Arial 12pt."""
        self.formatter._create_styles()
        normal = self.formatter.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Arial")
        self.assertEqual(normal.font.size, Pt(12))

    def test_normal_style_line_spacing_is_1_5_lines(self):
        """Normal style must have 1.5 line spacing."""
        self.formatter._create_styles()
        normal = self.formatter.doc.styles["Normal"]
        self.assertEqual(
            normal.paragraph_format.line_spacing_rule,
            WD_LINE_SPACING.ONE_POINT_FIVE,
        )

    def test_normal_style_paragraph_is_justified(self):
        """Normal style paragraph alignment must be justified."""
        self.formatter._create_styles()
        normal = self.formatter.doc.styles["Normal"]
        self.assertEqual(
            normal.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    def test_heading_1_is_centered_and_bold(self):
        """Heading 1 style must be centered and bold."""
        self.formatter._create_styles()
        h1 = self.formatter.doc.styles["Heading 1"]
        self.assertEqual(
            h1.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        self.assertTrue(h1.font.bold)


class TestAddCoverPage(unittest.TestCase):
    """Tests for _add_cover_page method."""

    @classmethod
    def setUpClass(cls):
        """Create a blank DOCX to test with."""
        cls.temp_dir = Path("tests/temp_icontec_cover")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "blank.docx"
        doc = Document()
        doc.save(str(cls.docx_path))
        cls.meta = DocumentMetadata(
            title="Mi Documento de Prueba",
            author="Juan Pérez García",
            institution="Universidad Nacional",
            program="Ingeniería de Software",
            date="2025",
            extra={"city": "Bogotá"},
        )
        cls.formatter = IcontecFormatter(str(cls.docx_path))
        cls.formatter._add_cover_page(cls.meta)

    @classmethod
    def tearDownClass(cls):
        """Clean up temporary files."""
        import shutil

        if cls.temp_dir.exists():
            shutil.rmtree(cls.temp_dir)

    def test_title_is_uppercase(self):
        """Title on cover page must appear in uppercase."""
        full_text = "\n".join(p.text for p in self.formatter.doc.paragraphs)
        self.assertIn("MI DOCUMENTO DE PRUEBA", full_text)

    def test_author_appears(self):
        """Author name must appear in the document."""
        full_text = "\n".join(p.text for p in self.formatter.doc.paragraphs)
        self.assertIn("Juan Pérez García", full_text)

    def test_institution_appears_uppercase(self):
        """Institution must appear in uppercase."""
        full_text = "\n".join(p.text for p in self.formatter.doc.paragraphs)
        self.assertIn("UNIVERSIDAD NACIONAL", full_text)

    def test_program_appears_uppercase(self):
        """Program must appear in uppercase."""
        full_text = "\n".join(p.text for p in self.formatter.doc.paragraphs)
        self.assertIn("INGENIERÍA DE SOFTWARE", full_text)

    def test_city_and_year_at_bottom(self):
        """City and year must appear at the bottom of cover page."""
        full_text = "\n".join(p.text for p in self.formatter.doc.paragraphs)
        self.assertIn("BOGOTÁ, 2025", full_text)

    def test_page_break_inserted(self):
        """A page break must be inserted after the cover page."""

        found_break = False
        for p in self.formatter.doc.paragraphs:
            if not p.text.strip():
                for child in p._element.iter():
                    if child.tag == qn("w:br"):
                        found_break = True
                        break
            if found_break:
                break
        self.assertTrue(found_break, "No page break found in document")


class TestProcessParagraphs(unittest.TestCase):
    """Tests for _process_paragraphs method."""

    @classmethod
    def setUpClass(cls):
        """Create a DOCX with body content to test with."""
        cls.temp_dir = Path("tests/temp_icontec_paras")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "content.docx"
        doc = Document()
        body_p = doc.add_paragraph("Este es un párrafo de contenido.")
        body_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_heading("Introducción", level=1)
        doc.save(str(cls.docx_path))
        cls.formatter = IcontecFormatter(str(cls.docx_path))
        cls.formatter._process_paragraphs()

    @classmethod
    def tearDownClass(cls):
        """Clean up temporary files."""
        import shutil

        if cls.temp_dir.exists():
            shutil.rmtree(cls.temp_dir)

    def test_body_text_is_justified(self):
        """Body paragraphs must be justified (not headings)."""
        for p in self.formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue
            self.assertEqual(
                p.paragraph_format.alignment,
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                f"Paragraph '{p.text}' is not justified",
            )

    def test_body_font_is_arial(self):
        """Body paragraph runs must have Arial font set."""
        for p in self.formatter.doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                continue
            for run in p.runs:
                if run.text.strip():
                    self.assertEqual(
                        run.font.name,
                        "Arial",
                        f"Run '{run.text}' is not Arial",
                    )


if __name__ == "__main__":
    unittest.main()
