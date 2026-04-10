"""
Unit tests for APA formatter figure methods.

Tests cover:
- _make_figure_paragraph helper
- format_figures (image scaling)
- add_figure_captions (caption creation)
- Config getters
"""

import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from normadocs.formatters.apa.apa_figures import APAFiguresHandler


def _create_drawing_element(cx: int = 9144000, cy: int = 6858000, descr: str = "") -> OxmlElement:
    """Create a w:drawing element with inline image for testing.

    Args:
        cx: Width in EMUs (9144000 = 1 inch)
        cy: Height in EMUs
        descr: Alt text / description for the image
    """
    ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"

    drawing = OxmlElement("w:drawing")
    inline = OxmlElement("wp:inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
    inline.append(extent)

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")
    inline.append(effect_extent)

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", "1")
    doc_pr.set("name", "Image 1")
    if descr:
        doc_pr.set("descr", descr)
    inline.append(doc_pr)

    c_nv_graphic_frame_locks = OxmlElement("wp:c_nvGraphicFrameLocks")
    no_change_aspect = OxmlElement("wp:noChangeAspect")
    c_nv_graphic_frame_locks.append(no_change_aspect)
    inline.append(c_nv_graphic_frame_locks)

    graphic = OxmlElement("a:graphic")
    graphic.set("uri", "")

    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set("uri", f"{{{ns_pic}}}")

    pic = OxmlElement("pic:pic")
    c_nv_pic_properties = OxmlElement("pic:c_nvPicPr")

    c_nv_document_properties = OxmlElement("pic:cNvPr")
    c_nv_document_properties.set("id", "1")
    c_nv_document_properties.set("name", "Image 1")
    if descr:
        c_nv_document_properties.set("descr", descr)
    c_nv_pic_properties.append(c_nv_document_properties)

    c_nv_fill = OxmlElement("pic:cNvC")
    c_nv_pic_properties.append(c_nv_fill)

    blip_fill = OxmlElement("pic:blipFill")
    blip = OxmlElement("pic:blip")
    blip.set(qn("r:embed"), "rId1")
    blip_fill.append(blip)

    stretch = OxmlElement("pic:stretch")
    fill_rect = OxmlElement("pic:fillRect")
    stretch.append(fill_rect)
    blip_fill.append(stretch)
    pic.append(blip_fill)

    sp_pr = OxmlElement("pic:spPr")
    xfrm = OxmlElement("a:xfrm")
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    xfrm.append(off)
    ext = OxmlElement("a:ext")
    ext.set("cx", str(cx))
    ext.set("cy", str(cy))
    xfrm.append(ext)
    sp_pr.append(xfrm)

    prst_geom = OxmlElement("a:prstGeom")
    prst_geom.set("prst", "rect")
    av_lst = OxmlElement("a:avLst")
    prst_geom.append(av_lst)
    sp_pr.append(prst_geom)

    pic.append(sp_pr)
    graphic_data.append(pic)
    graphic.append(graphic_data)
    inline.append(graphic)
    drawing.append(inline)

    return drawing


def _add_paragraph_with_drawing(
    doc: Document, text: str = "", cx: int = 9144000, cy: int = 6858000, descr: str = ""
) -> None:
    """Add a paragraph with a drawing element to the document."""
    p = doc.add_paragraph()
    if text:
        p.add_run(text)
    drawing = _create_drawing_element(cx=cx, cy=cy, descr=descr)
    p._element.append(drawing)


class TestGetFigureConfig(unittest.TestCase):
    """Tests for _get_figure_config method."""

    def test_default_config(self):
        """Default config should have expected keys."""
        doc = Document()
        handler = APAFiguresHandler(doc)
        config = handler._get_figure_config()
        self.assertEqual(config["caption_prefix"], "Figura")
        self.assertTrue(config["title_above"])
        self.assertEqual(config["nota_prefix"], "Nota.")

    def test_custom_config(self):
        """Custom config should override defaults."""
        doc = Document()
        handler = APAFiguresHandler(doc, config={"figures": {"caption_prefix": "Image"}})
        config = handler._get_figure_config()
        self.assertEqual(config["caption_prefix"], "Image")


class TestGetBodyFont(unittest.TestCase):
    """Tests for _get_body_font method."""

    def test_default_font(self):
        """Default body font should be Times New Roman."""
        doc = Document()
        handler = APAFiguresHandler(doc)
        font = handler._get_body_font()
        self.assertEqual(font, "Times New Roman")

    def test_custom_font(self):
        """Custom config should override font."""
        doc = Document()
        handler = APAFiguresHandler(doc, config={"fonts": {"body": {"name": "Arial"}}})
        font = handler._get_body_font()
        self.assertEqual(font, "Arial")


class TestMakeFigureParagraph(unittest.TestCase):
    """Tests for _make_figure_paragraph helper."""

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = Path("tests/temp_figures")
        cls.temp_dir.mkdir(parents=True, exist_ok=True)
        cls.docx_path = cls.temp_dir / "test.docx"
        doc = Document()
        doc.save(str(cls.docx_path))
        cls.handler = APAFiguresHandler(doc)

    @classmethod
    def tearDownClass(cls):
        import shutil

        if cls.temp_dir.exists():
            shutil.rmtree(cls.temp_dir)

    def test_figure_paragraph_returns_element(self):
        """_make_figure_paragraph returns an OxmlElement."""
        result = self.handler._make_figure_paragraph("Test caption")
        self.assertIsNotNone(result)
        self.assertEqual(result.tag, qn("w:p"))

    def test_figure_paragraph_with_bold(self):
        """Bold text creates run with w:b element."""
        p_el = self.handler._make_figure_paragraph("Bold text", bold=True)
        r_el = p_el.find(f".//{qn('w:r')}")
        self.assertIsNotNone(r_el)
        rPr = r_el.find(qn("w:rPr"))
        self.assertIsNotNone(rPr)
        b_el = rPr.find(qn("w:b"))
        self.assertIsNotNone(b_el)

    def test_figure_paragraph_with_italic(self):
        """Italic text creates run with w:i element."""
        p_el = self.handler._make_figure_paragraph("Italic text", italic=True)
        r_el = p_el.find(f".//{qn('w:r')}")
        rPr = r_el.find(qn("w:rPr"))
        i_el = rPr.find(qn("w:i"))
        self.assertIsNotNone(i_el)

    def test_figure_paragraph_text_content(self):
        """Paragraph contains the expected text."""
        p_el = self.handler._make_figure_paragraph("Caption text")
        t_el = p_el.find(f".//{qn('w:t')}")
        self.assertEqual(t_el.text, "Caption text")

    def test_figure_paragraph_space_after(self):
        """Paragraph spacing is set correctly."""
        p_el = self.handler._make_figure_paragraph("Test", space_after="120")
        p_pr = p_el.find(qn("w:pPr"))
        sp_el = p_pr.find(qn("w:spacing"))
        self.assertEqual(sp_el.get(qn("w:after")), "120")

    def test_figure_paragraph_no_style(self):
        """Paragraph with no bold/italic has no w:b or w:i."""
        p_el = self.handler._make_figure_paragraph("Plain text")
        r_el = p_el.find(f".//{qn('w:r')}")
        rPr = r_el.find(qn("w:rPr"))
        self.assertIsNone(rPr.find(qn("w:b")))
        self.assertIsNone(rPr.find(qn("w:i")))

    def test_figure_paragraph_bold_and_italic(self):
        """Both bold and italic creates both w:b and w:i."""
        p_el = self.handler._make_figure_paragraph("Both styles", bold=True, italic=True)
        r_el = p_el.find(f".//{qn('w:r')}")
        rPr = r_el.find(qn("w:rPr"))
        self.assertIsNotNone(rPr.find(qn("w:b")))
        self.assertIsNotNone(rPr.find(qn("w:i")))


class TestFormatFigures(unittest.TestCase):
    """Tests for format_figures method (image scaling)."""

    def test_format_figures_no_images(self):
        """format_figures should handle document with no images."""
        doc = Document()
        doc.add_paragraph("Just text")
        handler = APAFiguresHandler(doc)
        handler.format_figures()
        self.assertEqual(len(doc.paragraphs), 1)

    def test_format_figures_scaling(self):
        """format_figures should scale oversized images to fit page area."""
        doc = Document()
        _add_paragraph_with_drawing(doc, cx=12000000, cy=12000000, descr="Oversized image")

        handler = APAFiguresHandler(doc)
        handler.format_figures()

        paragraphs = doc.paragraphs
        drawings = paragraphs[0]._element.findall(f".//{qn('w:drawing')}")
        self.assertEqual(len(drawings), 1)

        inline = drawings[0].find(f".//{qn('wp:inline')}")
        self.assertIsNotNone(inline)

        extent = inline.find(qn("wp:extent"))
        self.assertIsNotNone(extent)
        new_cx = int(extent.get("cx", 0))
        new_cy = int(extent.get("cy", 0))

        self.assertLess(new_cx, 12000000, "Width should be scaled down")
        self.assertLess(new_cy, 12000000, "Height should be scaled down")

    def test_format_figures_small_image_not_scaled(self):
        """format_figures should not scale images that fit within limits."""
        doc = Document()
        _add_paragraph_with_drawing(doc, cx=5000000, cy=5000000, descr="Small image")

        handler = APAFiguresHandler(doc)
        handler.format_figures()

        paragraphs = doc.paragraphs
        drawings = paragraphs[0]._element.findall(f".//{qn('w:drawing')}")
        self.assertEqual(len(drawings), 1)

        inline = drawings[0].find(f".//{qn('wp:inline')}")
        extent = inline.find(qn("wp:extent"))
        final_cx = int(extent.get("cx", 0))
        final_cy = int(extent.get("cy", 0))

        self.assertEqual(final_cx, 5000000, "Small image width should not be scaled")
        self.assertEqual(final_cy, 5000000, "Small image height should not be scaled")

    def test_format_figures_centers_image(self):
        """format_figures should center the image paragraph."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Test image")

        handler = APAFiguresHandler(doc)
        handler.format_figures()

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.assertEqual(doc.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)


class TestAddFigureCaptions(unittest.TestCase):
    """Tests for add_figure_captions method."""

    def test_add_figure_captions_empty_document(self):
        """add_figure_captions should handle empty document."""
        doc = Document()
        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()
        self.assertEqual(len(doc.paragraphs), 0)

    def test_add_figure_captions_no_images(self):
        """add_figure_captions should handle document with no images."""
        doc = Document()
        doc.add_paragraph("Just text")
        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()
        self.assertEqual(len(doc.paragraphs), 1)

    def test_caption_preserves_existing_paragraphs(self):
        """add_figure_captions should preserve existing non-image paragraphs."""
        doc = Document()
        doc.add_paragraph("Introduction")
        doc.add_paragraph("Some text")

        handler = APAFiguresHandler(doc)
        initial_count = len(doc.paragraphs)
        handler.add_figure_captions()

        self.assertGreaterEqual(len(doc.paragraphs), initial_count)

    def test_add_figure_captions_with_image_and_title(self):
        """Caption should be 'Figura N' when title is not in Nota or Figura format."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Test image")
        doc.add_paragraph("My Figure Title")

        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Figura 1", full_text)

    def test_add_figure_captions_with_nota(self):
        """Title should be extracted correctly when Nota follows title."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Test image")
        doc.add_paragraph("Title Before Nota")
        doc.add_paragraph("Nota. This is the note context")

        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Figura 1. Title Before Nota", full_text)

    def test_add_figure_captions_without_title(self):
        """Caption should be 'Figura N' when no title is found."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Test image")
        doc.add_paragraph("Some unrelated text")

        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Figura 1", full_text)

    def test_add_figure_captions_with_figura_title(self):
        """Caption should extract title from existing 'Figura N' paragraph."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Test image")
        doc.add_paragraph("Figura 1 Original Caption")

        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Figura 1. Original Caption", full_text)

    def test_add_figure_captions_multiple_images(self):
        """Multiple images should get sequential figure numbers."""
        doc = Document()
        _add_paragraph_with_drawing(doc, descr="Image 1")
        doc.add_paragraph("Figura 1 First Title")
        _add_paragraph_with_drawing(doc, descr="Image 2")
        doc.add_paragraph("Figura 2 Second Title")

        handler = APAFiguresHandler(doc)
        handler.add_figure_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Figura 1. First Title", full_text)
        self.assertIn("Figura 2. Second Title", full_text)


class TestApplyFontStyle(unittest.TestCase):
    """Tests for _apply_font_style method."""

    def test_apply_font_style_bold(self):
        """_apply_font_style should apply bold formatting."""
        doc = Document()
        handler = APAFiguresHandler(doc)
        p = doc.add_paragraph()
        run = p.add_run("Test")
        handler._apply_font_style(run, bold=True)
        self.assertTrue(run.bold)

    def test_apply_font_style_italic(self):
        """_apply_font_style should apply italic formatting."""
        doc = Document()
        handler = APAFiguresHandler(doc)
        p = doc.add_paragraph()
        run = p.add_run("Test")
        handler._apply_font_style(run, italic=True)
        self.assertTrue(run.italic)


if __name__ == "__main__":
    unittest.main()
