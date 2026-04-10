"""Tests for normadocs.formatters.base.DocumentFormatter"""

import pytest
from docx import Document

from normadocs.formatters.base import DocumentFormatter
from normadocs.models import DocumentMetadata


class ConcreteFormatter(DocumentFormatter):
    """Concrete implementation of DocumentFormatter for testing."""

    def process(self, meta: DocumentMetadata) -> None:
        pass

    def save(self, output_path: str) -> None:
        pass


@pytest.fixture
def formatter(tmp_path):
    """Create a ConcreteFormatter instance with a temporary document."""
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test content")
    doc.save(str(doc_path))
    return ConcreteFormatter(doc_path)


class TestGetConfig:
    """Tests for get_config method."""

    def test_get_config_simple(self, tmp_path):
        """Test simple key access."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.save(str(doc_path))
        formatter = ConcreteFormatter(doc_path, config={"theme": "dark"})
        assert formatter.get_config("theme") == "dark"

    def test_get_config_nested(self, tmp_path):
        """Test nested dot-notation access through multiple keys."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.save(str(doc_path))
        formatter = ConcreteFormatter(
            doc_path, config={"fonts": {"body": {"name": "Arial", "size": 12}}}
        )
        result = formatter.get_config("fonts", "body", "name")
        assert result == "Arial"

    def test_get_config_default(self, tmp_path):
        """Test default value when key not found."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.save(str(doc_path))
        formatter = ConcreteFormatter(doc_path, config={"fonts": {}})
        result = formatter.get_config("fonts", "body", "size", default="default_value")
        assert result == "default_value"

    def test_get_config_returns_default_for_non_dict(self, tmp_path):
        """Test that get_config returns default when value is not a dict."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.save(str(doc_path))
        formatter = ConcreteFormatter(doc_path, config={"key": "string_value"})
        result = formatter.get_config("key", "nested", default="fallback")
        assert result == "fallback"


class TestFormatTableCaption:
    """Tests for _format_table_caption method."""

    def test_format_table_caption_raises(self, formatter):
        """Test that _format_table_caption raises NotImplementedError."""
        table = Document().add_table(rows=2, cols=2)
        with pytest.raises(NotImplementedError) as exc_info:
            formatter._format_table_caption(table, number=1, title="Test Caption")
        assert "Caption formatting not implemented for ConcreteFormatter" in str(exc_info.value)
