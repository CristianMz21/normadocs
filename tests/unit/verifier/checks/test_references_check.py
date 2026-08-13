"""Unit tests for ReferencesCheck - APA 7th Edition references verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.references import ReferencesCheck


class TestReferencesCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant references (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_proper_references(self) -> Path:
        path = self.temp_path / "proper_refs.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        ref_heading = doc.add_paragraph("References")
        ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        ref1 = doc.add_paragraph("Smith, J. (2024). Title of the work. Publisher.")
        ref1.paragraph_format.left_indent = Inches(0.5)
        ref1.paragraph_format.first_line_indent = Inches(-0.5)  # APA 7 hanging indent

        ref2 = doc.add_paragraph("Doe, A. (2023). Another work. Journal.")
        ref2.paragraph_format.left_indent = Inches(0.5)
        ref2.paragraph_format.first_line_indent = Inches(-0.5)  # APA 7 hanging indent

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ReferencesCheck()
        return check.run(ctx)

    def test_proper_references_passes(self) -> None:
        """Document with proper APA references should pass."""
        docx_path = self._create_docx_with_proper_references()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for proper refs but got: {errors}")

    def test_every_reference_requires_hanging_indent(self) -> None:
        """A later reference without hanging indent must be reported."""
        docx_path = self._create_docx_with_proper_references()
        doc = Document(str(docx_path))
        doc.paragraphs[-1].paragraph_format.first_line_indent = None
        doc.save(str(docx_path))

        issues = self._run_check(docx_path)
        later_errors = [
            i for i in issues if "hanging_indent" in i.check and "Reference 2" in i.evidence
        ]
        self.assertGreater(len(later_errors), 0, f"Expected second-reference error: {issues}")


class TestReferencesCheckMissingSection(unittest.TestCase):
    """Tests for missing references section."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_references(self) -> Path:
        path = self.temp_path / "no_refs.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Just some text, no references section.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ReferencesCheck()
        return check.run(ctx)

    def test_no_references_section_raises_error(self) -> None:
        """No references section should raise error."""
        docx_path = self._create_docx_no_references()
        issues = self._run_check(docx_path)
        section_errors = [
            i for i in issues if "section_present" in i.check and i.severity == "error"
        ]
        self.assertGreater(len(section_errors), 0, f"Expected section error but got: {issues}")


class TestReferencesCheckHangingIndent(unittest.TestCase):
    """Tests for hanging indent violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_refs_no_hanging_indent(self) -> Path:
        path = self.temp_path / "refs_no_hanging.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("References")

        ref1 = doc.add_paragraph("Smith, J. (2024). Title without hanging indent.")
        ref1.paragraph_format.first_line_indent = None

        doc.save(str(path))
        return path

    def _create_docx_refs_wrong_indent(self) -> Path:
        path = self.temp_path / "refs_wrong_indent.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("References")

        ref1 = doc.add_paragraph("Smith, J. (2024). Title with wrong indent.")
        ref1.paragraph_format.first_line_indent = Inches(-0.25)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ReferencesCheck()
        return check.run(ctx)

    def test_no_hanging_indent_raises_error(self) -> None:
        """No hanging indent should raise error."""
        docx_path = self._create_docx_refs_no_hanging_indent()
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "hanging_indent" in i.check and i.severity == "error"]
        self.assertGreater(
            len(indent_errors), 0, f"Expected hanging indent error but got: {issues}"
        )

    def test_wrong_indent_raises_error(self) -> None:
        """Wrong hanging indent (0.25 inch) should raise error."""
        docx_path = self._create_docx_refs_wrong_indent()
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "hanging_indent" in i.check and i.severity == "error"]
        self.assertGreater(
            len(indent_errors), 0, f"Expected hanging indent error but got: {issues}"
        )


class TestReferencesCheckAlphabeticalOrder(unittest.TestCase):
    """Tests for alphabetical order violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_refs_out_of_order(self) -> Path:
        path = self.temp_path / "refs_out_of_order.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("References")

        ref1 = doc.add_paragraph("Zebra, Z. (2024). Work starting with Z.")
        ref1.paragraph_format.first_line_indent = Inches(0.5)

        ref2 = doc.add_paragraph("Alpha, A. (2023). Work starting with A.")
        ref2.paragraph_format.first_line_indent = Inches(0.5)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ReferencesCheck()
        return check.run(ctx)

    def test_out_of_order_raises_warning(self) -> None:
        """References out of alphabetical order should raise warning."""
        docx_path = self._create_docx_refs_out_of_order()
        issues = self._run_check(docx_path)
        order_warnings = [
            i for i in issues if "alphabetical_order" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(order_warnings), 0, f"Expected alphabetical order warning but got: {issues}"
        )

    def test_same_author_undated_reference_precedes_dated_reference(self) -> None:
        """APA places an undated work before dated works by the same author."""
        path = self.temp_path / "same_author_refs.docx"
        doc = Document()
        doc.add_paragraph("References")

        first = doc.add_paragraph("Servicio Nacional de Aprendizaje. (s. f.). Undated guide.")
        first.paragraph_format.first_line_indent = Inches(-0.5)
        second = doc.add_paragraph("Servicio Nacional de Aprendizaje. (2026). Dated guide.")
        second.paragraph_format.first_line_indent = Inches(-0.5)
        doc.save(str(path))

        issues = self._run_check(path)
        order_warnings = [
            i for i in issues if "alphabetical_order" in i.check and i.severity == "warning"
        ]
        self.assertEqual(order_warnings, [], f"Unexpected APA order warning: {issues}")

    def test_section_after_references_is_not_counted(self) -> None:
        """A heading after the references (e.g. appendices) must not be counted as a reference.

        Regression: a later heading that contains the word "references" in its text
        (e.g. "Slide 4 — ... references") used to trigger references mode and pollute
        the reference list. A section that follows the references (apéndices) must be
        cut at the next heading.
        """
        docx_path = self.temp_path / "refs_with_appendices.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("References")

        ref1 = doc.add_paragraph("Alpha, A. (2023). Work starting with A.")
        ref1.paragraph_format.first_line_indent = Inches(0.5)

        ref2 = doc.add_paragraph("Zebra, Z. (2024). Work starting with Z.")
        ref2.paragraph_format.first_line_indent = Inches(0.5)

        # A later heading (appendix section) with a word that contains "references"
        appendix = doc.add_paragraph("Appendix: Image references", style="Heading 1")
        appendix.paragraph_format.first_line_indent = Inches(0.5)

        appendix_body = doc.add_paragraph("Figure 1. Screenshot of the interface.")
        appendix_body.paragraph_format.first_line_indent = Inches(0.5)

        doc.save(str(docx_path))
        issues = self._run_check(docx_path)
        # Sections after the references heading must not be treated as reference
        # entries, so they must not produce alphabetical order warnings.
        order_warnings = [
            i for i in issues if "alphabetical_order" in i.check and i.severity == "warning"
        ]
        self.assertEqual(order_warnings, [], f"Unexpected order warnings: {issues}")


class TestReferencesCheckEmptySection(unittest.TestCase):
    """Tests for an empty references section (heading present, no entries)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_empty_references(self) -> Path:
        path = self.temp_path / "empty_refs.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("References")
        doc.add_paragraph("")
        doc.add_paragraph("   ")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ReferencesCheck()
        return check.run(ctx)

    def test_empty_references_section_raises_error(self) -> None:
        """A references section with no entries should raise entries_present error (L60-69)."""
        docx_path = self._create_docx_empty_references()
        issues = self._run_check(docx_path)
        entries_errors = [
            i for i in issues if "entries_present" in i.check and i.severity == "error"
        ]
        self.assertGreater(
            len(entries_errors), 0, f"Expected entries_present error but got: {issues}"
        )


if __name__ == "__main__":
    unittest.main()
