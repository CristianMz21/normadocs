"""Unit tests for FontsCheck - APA 7th Edition font verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.fonts import FontsCheck


class TestFontsCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant fonts (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
    ) -> Path:
        path = self.temp_path / f"font_{font_name}_{font_size}.docx".replace(" ", "_")
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para = doc.add_paragraph()
        run = para.add_run("Sample body text in Times New Roman 12pt.")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = False
        run.font.italic = False

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FontsCheck()
        return check.run(ctx)

    def test_times_new_roman_12pt_passes(self) -> None:
        """Times New Roman 12pt should pass with no errors."""
        docx_path = self._create_docx(font_name="Times New Roman", font_size=12)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for TNR 12pt but got: {errors}")

    def test_times_new_roman_11pt_passes_within_tolerance(self) -> None:
        """Times New Roman 11pt should pass (within 1pt tolerance)."""
        docx_path = self._create_docx(font_name="Times New Roman", font_size=11)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for TNR 11pt but got: {errors}")

    def test_times_new_roman_13pt_passes_within_tolerance(self) -> None:
        """Times New Roman 13pt should pass (within 1pt tolerance)."""
        docx_path = self._create_docx(font_name="Times New Roman", font_size=13)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for TNR 13pt but got: {errors}")


class TestFontsCheckViolation(unittest.TestCase):
    """Tests for font violations (should fail with errors)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
    ) -> Path:
        path = self.temp_path / f"font_{font_name}_{font_size}.docx".replace(" ", "_")
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para = doc.add_paragraph()
        run = para.add_run("Sample body text.")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = False
        run.font.italic = False

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FontsCheck()
        return check.run(ctx)

    def test_arial_12pt_raises_font_error(self) -> None:
        """Arial 12pt should raise error (wrong font)."""
        docx_path = self._create_docx(font_name="Arial", font_size=12)
        issues = self._run_check(docx_path)
        font_errors = [i for i in issues if i.severity == "error" and "font" in i.check]
        self.assertGreater(len(font_errors), 0, f"Expected font error for Arial but got: {issues}")

    def test_calibri_12pt_raises_font_error(self) -> None:
        """Calibri 12pt should raise error (wrong font)."""
        docx_path = self._create_docx(font_name="Calibri", font_size=12)
        issues = self._run_check(docx_path)
        font_errors = [i for i in issues if i.severity == "error" and "font" in i.check]
        self.assertGreater(
            len(font_errors), 0, f"Expected font error for Calibri but got: {issues}"
        )

    def test_times_new_roman_10pt_raises_size_error(self) -> None:
        """Times New Roman 10pt should raise size error (outside 1pt tolerance)."""
        docx_path = self._create_docx(font_name="Times New Roman", font_size=10)
        issues = self._run_check(docx_path)
        size_errors = [i for i in issues if i.severity == "error" and "size" in i.check]
        self.assertGreater(len(size_errors), 0, f"Expected size error for 10pt but got: {issues}")

    def test_times_new_roman_14pt_raises_size_error(self) -> None:
        """Times New Roman 14pt should raise size error (outside 1pt tolerance)."""
        docx_path = self._create_docx(font_name="Times New Roman", font_size=14)
        issues = self._run_check(docx_path)
        size_errors = [i for i in issues if i.severity == "error" and "size" in i.check]
        self.assertGreater(len(size_errors), 0, f"Expected size error for 14pt but got: {issues}")

    def test_comic_sans_raises_font_error(self) -> None:
        """Comic Sans should raise error (definitely not APA compliant)."""
        docx_path = self._create_docx(font_name="Comic Sans MS", font_size=12)
        issues = self._run_check(docx_path)
        font_errors = [i for i in issues if i.severity == "error" and "font" in i.check]
        self.assertGreater(
            len(font_errors), 0, f"Expected font error for Comic Sans but got: {issues}"
        )


class TestFontsCheckMixedContent(unittest.TestCase):
    """Tests for documents with mixed fonts."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_mixed_font_docx(self) -> Path:
        """Create a document where Arial is the dominant font."""
        path = self.temp_path / "mixed_fonts.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para1 = doc.add_paragraph()
        run1 = para1.add_run("Arial body text. " * 20)
        run1.font.name = "Arial"
        run1.font.size = Pt(12)

        para2 = doc.add_paragraph()
        run2 = para2.add_run("Times New Roman. " * 5)
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FontsCheck()
        return check.run(ctx)

    def test_mixed_fonts_detects_wrong_font(self) -> None:
        """Document with mostly TNR but some Arial should detect the wrong font."""
        docx_path = self._create_mixed_font_docx()
        issues = self._run_check(docx_path)
        font_errors = [i for i in issues if i.severity == "error" and "font" in i.check]
        self.assertGreater(
            len(font_errors), 0, f"Expected font error for mixed fonts but got: {issues}"
        )


if __name__ == "__main__":
    unittest.main()
