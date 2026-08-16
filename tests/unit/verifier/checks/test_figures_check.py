"""Unit tests for FiguresCheck - APA 7th Edition figure verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.figures import FiguresCheck


class TestFiguresCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant figures (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_apa_figure(self) -> Path:
        path = self.temp_path / "apa_figure.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para = doc.add_paragraph()
        run1 = para.add_run("Figure 1. ")
        run1.bold = True
        run2 = para.add_run("Title of the Figure in Italic")
        run2.italic = True

        doc.add_paragraph("Some text after the figure.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FiguresCheck()
        return check.run(ctx)

    def test_apa_figure_proper_caption_passes(self) -> None:
        """Figure with proper APA caption should pass."""
        docx_path = self._create_docx_with_apa_figure()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for APA figure but got: {errors}")


class TestFiguresCheckCaptionViolation(unittest.TestCase):
    """Tests for figure caption violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_figure_not_bold(self) -> Path:
        path = self.temp_path / "figure_not_bold.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para = doc.add_paragraph()
        run = para.add_run("Figure 1. Not Bold")
        run.bold = False

        doc.save(str(path))
        return path

    def _create_docx_figure_not_italic(self) -> Path:
        path = self.temp_path / "figure_not_italic.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        para = doc.add_paragraph()
        run = para.add_run("Figure 1. Title Not Italic")
        run.bold = True
        run.italic = False

        doc.save(str(path))
        return path

    def _create_docx_no_figure_captions(self) -> Path:
        path = self.temp_path / "no_figures.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Just some regular text, no figures here.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FiguresCheck()
        return check.run(ctx)

    def test_figure_caption_not_bold_raises_error(self) -> None:
        """Figure caption not bold should raise error."""
        docx_path = self._create_docx_figure_not_bold()
        issues = self._run_check(docx_path)
        bold_errors = [i for i in issues if "caption_bold" in i.check and i.severity == "error"]
        self.assertGreater(len(bold_errors), 0, f"Expected bold error but got: {issues}")

    def test_figure_caption_not_italic_raises_warning(self) -> None:
        """Figure caption title not italic should raise warning."""
        docx_path = self._create_docx_figure_not_italic()
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(italic_warnings), 0, f"Expected italic warning but got: {issues}")


class TestFiguresCheckNoCaptionsWithPDFEvidence(unittest.TestCase):
    """Tests for figures without captions but detected in PDF."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_figure_captions(self) -> Path:
        path = self.temp_path / "no_figure_captions.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Text content without any figure references.")

        doc.save(str(path))
        return path

    def _run_check_with_mock_pdf(self, docx_path: Path, pdf_text: str | None) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)

        mock_pdf = MagicMock()
        if pdf_text:
            mock_pdf.extract_text_by_page.return_value = {1: pdf_text}
        else:
            mock_pdf.extract_text_by_page.return_value = {}

        ctx = VerificationContext(
            pdf=mock_pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = FiguresCheck()
        return check.run(ctx)

    def test_no_captions_but_figure_in_pdf_raises_warning(self) -> None:
        """No figure captions in DOCX but 'figure' in PDF text should warn."""
        docx_path = self._create_docx_no_figure_captions()
        issues = self._run_check_with_mock_pdf(docx_path, "See figure 1 below for details")
        format_warnings = [
            i for i in issues if "caption_format" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(format_warnings), 0, f"Expected caption format warning but got: {issues}"
        )

    def test_no_captions_and_no_figure_in_pdf_no_warning(self) -> None:
        """No figure captions and no 'figure' in PDF should not warn."""
        docx_path = self._create_docx_no_figure_captions()
        issues = self._run_check_with_mock_pdf(docx_path, "Just regular text content")
        format_warnings = [i for i in issues if "caption_format" in i.check]
        self.assertEqual(format_warnings, [], f"Expected no caption warnings but got: {issues}")


def _add_image_paragraph(doc: Document) -> None:
    """Add a paragraph containing a bare w:drawing element."""
    p = doc.add_paragraph()
    p._element.append(OxmlElement("w:drawing"))


def _add_caption(doc: Document, number: int) -> None:
    """Add an APA-formatted figure caption paragraph."""
    para = doc.add_paragraph()
    run1 = para.add_run(f"Figure {number}. ")
    run1.bold = True
    run2 = para.add_run("Title of the Figure")
    run2.italic = True


class TestFiguresCheckPositionAndNumbering(unittest.TestCase):
    """Tests for caption position (above image) and numbering sequence."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _run_check(self, doc: Document) -> list:
        path = self.temp_path / "figures.docx"
        doc.save(str(path))
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=path, meta=meta)
        ctx = VerificationContext(
            pdf=MagicMock(),
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        try:
            return FiguresCheck().run(ctx)
        finally:
            verifier.close()

    def test_caption_above_image_passes(self):
        """APA 7 caption above the figure image produces no position issue."""
        doc = Document()
        _add_caption(doc, 1)
        _add_image_paragraph(doc)

        issues = self._run_check(doc)
        position = [i for i in issues if "caption_position" in i.check]
        self.assertEqual(position, [], f"Expected no position issues but got: {issues}")

    def test_caption_below_image_is_reported(self):
        """A caption placed after the image is flagged."""
        doc = Document()
        _add_image_paragraph(doc)
        _add_caption(doc, 1)

        issues = self._run_check(doc)
        position = [i for i in issues if "caption_position" in i.check]
        self.assertEqual(len(position), 1)

    def test_numbering_gap_is_reported(self):
        """Non-consecutive figure numbers are flagged."""
        doc = Document()
        _add_caption(doc, 1)
        _add_caption(doc, 3)

        issues = self._run_check(doc)
        sequence = [i for i in issues if "numbering_sequence" in i.check]
        self.assertEqual(len(sequence), 1)

    def test_sequential_numbering_passes(self):
        """Consecutive figure numbers produce no sequence issue."""
        doc = Document()
        _add_caption(doc, 1)
        _add_caption(doc, 2)

        issues = self._run_check(doc)
        sequence = [i for i in issues if "numbering_sequence" in i.check]
        self.assertEqual(sequence, [], f"Expected no sequence issues but got: {issues}")


if __name__ == "__main__":
    unittest.main()
