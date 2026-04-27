"""Unit tests for PageSetupCheck - APA 7th Edition page setup verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.page_setup import PageSetupCheck


class TestPageSetupCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant page setup (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_proper_page_setup(self) -> Path:
        path = self.temp_path / "proper_page_setup.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "PAGE"

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = PageSetupCheck()
        return check.run(ctx)

    def test_proper_page_setup_passes(self) -> None:
        """Document with proper APA page setup should pass."""
        docx_path = self._create_docx_with_proper_page_setup()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")


class TestPageSetupCheckDifferentFirstPage(unittest.TestCase):
    """Tests for different first page header/footer violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_different_first_page(self) -> Path:
        path = self.temp_path / "no_different_first_page.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = False

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = PageSetupCheck()
        return check.run(ctx)

    def test_no_different_first_page_raises_error(self) -> None:
        """Different first page not enabled should raise error."""
        docx_path = self._create_docx_no_different_first_page()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "different_first_page" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected different_first_page error but got: {issues}")


class TestPageSetupCheckPageNumbers(unittest.TestCase):
    """Tests for page number violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_page_numbers(self) -> Path:
        path = self.temp_path / "no_page_numbers.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Document Title"

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _create_docx_empty_header(self) -> Path:
        path = self.temp_path / "empty_header.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = ""

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = PageSetupCheck()
        return check.run(ctx)

    def test_no_page_numbers_raises_error(self) -> None:
        """Header without page numbers should raise error."""
        docx_path = self._create_docx_no_page_numbers()
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if "page_numbers" in i.check and i.severity == "error"]
        self.assertGreater(len(page_errors), 0, f"Expected page_numbers error but got: {issues}")

    def test_empty_header_raises_error(self) -> None:
        """Empty header should raise error."""
        docx_path = self._create_docx_empty_header()
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if "page_numbers" in i.check and i.severity == "error"]
        self.assertGreater(len(page_errors), 0, f"Expected page_numbers error but got: {issues}")


if __name__ == "__main__":
    unittest.main()
