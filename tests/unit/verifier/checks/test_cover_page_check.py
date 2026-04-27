"""Unit tests for CoverPageCheck - APA 7th Edition cover page verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.cover_page import CoverPageCheck


class TestCoverPageCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant cover page (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_ap_compliant_cover_docx(
        self,
        title: str = "The Impact of Climate Change on Agriculture",
        author: str = "John Doe",
        institution: str = "University of Example",
        year: str = "2026",
    ) -> Path:
        path = self.temp_path / "apa_cover.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(author)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        inst_para = doc.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = inst_para.add_run(institution)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(year)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = CoverPageCheck()
        return check.run(ctx)

    def test_ap_compliant_cover_no_header_passes(self) -> None:
        """APA cover with proper elements and no first-page header should pass."""
        docx_path = self._create_ap_compliant_cover_docx()
        meta = DocumentMetadata(
            title="The Impact of Climate Change on Agriculture",
            author="John Doe",
            institution="University of Example",
        )
        issues = self._run_check(docx_path, meta)
        header_errors = [i for i in issues if "no_header" in i.check and i.severity == "error"]
        self.assertEqual(header_errors, [], f"Expected no header error but got: {issues}")


class TestCoverPageCheckHeaderViolation(unittest.TestCase):
    """Tests for header on cover page violations (should fail with errors)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_first_page_header(self) -> Path:
        path = self.temp_path / "cover_with_header.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        header = section.first_page_header
        header_para = header.paragraphs[0]
        header_para.text = "Running head: this should not be here"

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Document Title")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = CoverPageCheck()
        return check.run(ctx)

    def test_first_page_header_present_raises_error(self) -> None:
        """First page header present should raise error."""
        docx_path = self._create_docx_with_first_page_header()
        meta = DocumentMetadata(title="Document Title", author="Author Name")
        issues = self._run_check(docx_path, meta)
        header_errors = [i for i in issues if "no_header" in i.check and i.severity == "error"]
        self.assertGreater(len(header_errors), 0, f"Expected header error but got: {issues}")


class TestCoverPageCheckMissingElements(unittest.TestCase):
    """Tests for missing cover page elements (should warn)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_cover_docx(self) -> Path:
        path = self.temp_path / "cover_partial.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Document Title Without Author")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = CoverPageCheck()
        return check.run(ctx)

    def test_missing_author_raises_warning(self) -> None:
        """Author not in document should raise warning."""
        docx_path = self._create_cover_docx()
        meta = DocumentMetadata(
            title="Document Title Without Author",
            author="John Doe",
        )
        issues = self._run_check(docx_path, meta)
        author_warnings = [
            i for i in issues if "author_present" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(author_warnings), 0, f"Expected author warning but got: {issues}")

    def test_missing_title_raises_warning(self) -> None:
        """Title not in document should raise warning."""
        path = self.temp_path / "no_title.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run("John Doe")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.save(str(path))

        meta = DocumentMetadata(title="Some Title That Is Not In Doc", author="John Doe")
        issues = self._run_check(path, meta)
        title_warnings = [
            i for i in issues if "title_present" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(title_warnings), 0, f"Expected title warning but got: {issues}")


class TestCoverPageCheckAlignment(unittest.TestCase):
    """Tests for title alignment warnings."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_title_not_centered(self) -> Path:
        path = self.temp_path / "title_not_centered.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.add_run("My Document Title")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = CoverPageCheck()
        return check.run(ctx)

    def test_title_not_centered_raises_warning(self) -> None:
        """Title not centered should raise warning."""
        docx_path = self._create_docx_title_not_centered()
        meta = DocumentMetadata(title="My Document Title", author="Author Name")
        issues = self._run_check(docx_path, meta)
        align_warnings = [
            i for i in issues if "title_alignment" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(align_warnings), 0, f"Expected title alignment warning but got: {issues}"
        )


if __name__ == "__main__":
    unittest.main()
