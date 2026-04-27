"""Unit tests for HeadingsCheck - APA 7th Edition heading verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.headings import HeadingsCheck


class TestHeadingsCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant headings (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_apa_headings(self) -> Path:
        path = self.temp_path / "apa_headings.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h1 = doc.add_heading("", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h1.add_run("LEVEL 1 HEADING (CENTERED, BOLD)")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True

        h2 = doc.add_heading("", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h2.add_run("Level 2 Heading (Left, Bold)")
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.bold = True

        h3 = doc.add_heading("", level=3)
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h3.add_run("Level 3 Heading (Left, Bold Italic)")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
        run.italic = True

        doc.add_paragraph("Body paragraph after headings.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = HeadingsCheck()
        return check.run(ctx)

    def test_apa_headings_all_correct_passes(self) -> None:
        """Document with correct APA heading formatting should pass."""
        docx_path = self._create_docx_with_apa_headings()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for APA headings but got: {errors}")


class TestHeadingsCheckLevel1Violation(unittest.TestCase):
    """Tests for Level 1 heading violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_h1_not_centered(self) -> Path:
        path = self.temp_path / "h1_not_centered.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h1 = doc.add_heading("", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h1.add_run("Level 1 Not Centered")
        run.font.name = "Times New Roman"
        run.bold = True

        doc.save(str(path))
        return path

    def _create_docx_h1_not_bold(self) -> Path:
        path = self.temp_path / "h1_not_bold.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h1 = doc.add_heading("", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h1.add_run("Level 1 Not Bold")
        run.font.name = "Times New Roman"
        run.bold = False

        doc.save(str(path))
        return path

    def _create_docx_h1_not_bold(self) -> Path:
        path = self.temp_path / "h1_not_bold.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h1 = doc.add_heading("Level 1 Not Bold", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
            run.font.bold = False
            run.font.name = "Times New Roman"

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = HeadingsCheck()
        return check.run(ctx)

    def test_h1_not_centered_raises_error(self) -> None:
        """Level 1 heading not centered should raise error."""
        docx_path = self._create_docx_h1_not_centered()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "level1_alignment" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected level1 alignment error but got: {issues}")

    def test_h1_not_bold_raises_error(self) -> None:
        """Level 1 heading not bold should raise error."""
        docx_path = self._create_docx_h1_not_bold()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "level1_bold" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected level1 bold error but got: {issues}")


class TestHeadingsCheckLevel2Violation(unittest.TestCase):
    """Tests for Level 2 heading violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_h2_not_left(self) -> Path:
        path = self.temp_path / "h2_not_left.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h2 = doc.add_heading("", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h2.add_run("Level 2 Not Left")
        run.font.name = "Times New Roman"
        run.bold = True

        doc.save(str(path))
        return path

    def _create_docx_h2_not_bold(self) -> Path:
        path = self.temp_path / "h2_not_bold.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h2 = doc.add_heading("", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h2.add_run("Level 2 Not Bold")
        run.font.name = "Times New Roman"
        run.bold = False

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = HeadingsCheck()
        return check.run(ctx)

    def test_h2_not_left_raises_error(self) -> None:
        """Level 2 heading not left-aligned should raise error."""
        docx_path = self._create_docx_h2_not_left()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "level2_alignment" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected level2 alignment error but got: {issues}")

    def test_h2_not_bold_raises_error(self) -> None:
        """Level 2 heading not bold should raise error."""
        docx_path = self._create_docx_h2_not_bold()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "level2_bold" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected level2 bold error but got: {issues}")


class TestHeadingsCheckLevel3Violation(unittest.TestCase):
    """Tests for Level 3 heading violations (warnings, not errors)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_h3_not_left(self) -> Path:
        path = self.temp_path / "h3_not_left.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h3 = doc.add_heading("", level=3)
        h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h3.add_run("Level 3 Not Left")
        run.font.name = "Times New Roman"
        run.bold = True
        run.italic = True

        doc.save(str(path))
        return path

    def _create_docx_h3_not_bold_italic(self) -> Path:
        path = self.temp_path / "h3_not_bold_italic.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        h3 = doc.add_heading("", level=3)
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h3.add_run("Level 3 Not Bold Italic")
        run.font.name = "Times New Roman"
        run.bold = False
        run.italic = False

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = HeadingsCheck()
        return check.run(ctx)

    def test_h3_not_left_raises_warning(self) -> None:
        """Level 3 heading not left-aligned should raise warning (not error)."""
        docx_path = self._create_docx_h3_not_left()
        issues = self._run_check(docx_path)
        warnings = [i for i in issues if "level3_alignment" in i.check and i.severity == "warning"]
        errors = [i for i in issues if "level3_alignment" in i.check and i.severity == "error"]
        self.assertGreater(len(warnings), 0, f"Expected level3 alignment warning but got: {issues}")
        self.assertEqual(errors, [], f"Level 3 should be warning, not error: {issues}")

    def test_h3_not_bold_italic_raises_warning(self) -> None:
        """Level 3 heading not bold+italic should raise warning."""
        docx_path = self._create_docx_h3_not_bold_italic()
        issues = self._run_check(docx_path)
        warnings = [
            i for i in issues if "level3_bold_italic" in i.check and i.severity == "warning"
        ]
        errors = [i for i in issues if "level3_bold_italic" in i.check and i.severity == "error"]
        self.assertGreater(
            len(warnings), 0, f"Expected level3 bold/italic warning but got: {issues}"
        )
        self.assertEqual(errors, [], f"Level 3 should be warning, not error: {issues}")


if __name__ == "__main__":
    unittest.main()
