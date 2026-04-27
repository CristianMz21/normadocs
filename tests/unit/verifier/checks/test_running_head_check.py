"""Unit tests for RunningHeadCheck - APA 7th Edition running head verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.running_head import RunningHeadCheck


class TestRunningHeadCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant running head (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_proper_running_head(self, short_title: str = "SHORT TITLE") -> Path:
        path = self.temp_path / "proper_running_head.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        default_header = section.header
        header_para = default_header.paragraphs[0]
        header_para.text = f"{short_title}                                     PAGE"

        first_header = section.first_page_header
        first_header_para = first_header.paragraphs[0]
        first_header_para.text = ""

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = RunningHeadCheck()
        return check.run(ctx)

    def test_proper_running_head_passes(self) -> None:
        """Document with proper APA running head should pass."""
        docx_path = self._create_docx_with_proper_running_head()
        meta = DocumentMetadata(title="Test Document", short_title="SHORT TITLE")
        issues = self._run_check(docx_path, meta)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")


class TestRunningHeadCheckDifferentFirstPage(unittest.TestCase):
    """Tests for different first page header/footer violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_different_first_page(self) -> Path:
        path = self.temp_path / "no_different_first_page.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = False

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document", short_title="SHORT TITLE")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = RunningHeadCheck()
        return check.run(ctx)

    def test_no_different_first_page_raises_error(self) -> None:
        """Different first page not enabled should raise error."""
        docx_path = self._create_docx_no_different_first_page()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "different_first_page" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected different_first_page error but got: {issues}")


class TestRunningHeadCheckEmptyHeader(unittest.TestCase):
    """Tests for empty default header violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_empty_default_header(self) -> Path:
        path = self.temp_path / "empty_default_header.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        default_header = section.header
        default_header.paragraphs[0].text = ""

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document", short_title="SHORT TITLE")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = RunningHeadCheck()
        return check.run(ctx)

    def test_empty_default_header_raises_error(self) -> None:
        """Empty default header should raise error."""
        docx_path = self._create_docx_empty_default_header()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "missing_on_pages_2+" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected missing_on_pages_2+ error but got: {issues}")


class TestRunningHeadCheckFirstPageHeader(unittest.TestCase):
    """Tests for first page header violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_first_page_header_present(self) -> Path:
        path = self.temp_path / "first_page_header.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        default_header = section.header
        default_header.paragraphs[0].text = "SHORT TITLE                                     PAGE"

        first_header = section.first_page_header
        first_header.paragraphs[0].text = "Should be empty!"

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document", short_title="SHORT TITLE")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = RunningHeadCheck()
        return check.run(ctx)

    def test_first_page_header_present_raises_error(self) -> None:
        """First page header present should raise error."""
        docx_path = self._create_docx_first_page_header_present()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if "present_on_cover" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected present_on_cover error but got: {issues}")


class TestRunningHeadCheckShortTitle(unittest.TestCase):
    """Tests for short title content violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_wrong_short_title(self) -> Path:
        path = self.temp_path / "wrong_short_title.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.different_first_page_header_footer = True

        default_header = section.header
        default_header.paragraphs[0].text = "WRONG TITLE                                     PAGE"

        doc.add_paragraph("Body text.")

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path, meta: DocumentMetadata) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = RunningHeadCheck()
        return check.run(ctx)

    def test_wrong_short_title_raises_error(self) -> None:
        """Header with wrong short title should raise error."""
        docx_path = self._create_docx_wrong_short_title()
        meta = DocumentMetadata(title="Test Document", short_title="CORRECT TITLE")
        issues = self._run_check(docx_path, meta)
        errors = [i for i in issues if "short_title_content" in i.check and i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected short_title_content error but got: {issues}")


if __name__ == "__main__":
    unittest.main()
