"""Unit tests for TablesCheck - APA 7th Edition table verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.tables import TablesCheck


class TestTablesCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant tables (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_apa_table(self) -> Path:
        path = self.temp_path / "apa_table.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run("Table 1")
        caption_run.bold = True

        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Title of the Table in Italic")
        title_run.italic = True

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = TablesCheck()
        return check.run(ctx)

    def test_apa_table_proper_caption_passes(self) -> None:
        """Table with separate bold caption + italic title paragraph should pass cleanly."""
        docx_path = self._create_docx_with_apa_table()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        warnings = [i for i in issues if i.severity == "warning"]
        self.assertEqual(errors, [], f"Expected no errors for APA table but got: {errors}")
        self.assertEqual(warnings, [], f"Expected no warnings for APA table but got: {warnings}")


class TestTablesCheckCaptionViolation(unittest.TestCase):
    """Tests for table caption violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_table_caption_not_bold(self) -> Path:
        path = self.temp_path / "table_not_bold.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        caption_para = doc.add_paragraph()
        run = caption_para.add_run("Table 1. Title Not Bold")
        run.bold = False
        run.italic = True

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        doc.save(str(path))
        return path

    def _create_docx_table_caption_not_italic(self) -> Path:
        path = self.temp_path / "table_not_italic.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        caption_para = doc.add_paragraph()
        run = caption_para.add_run("Table 1. Title Not Italic")
        run.bold = True
        run.italic = False

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        doc.save(str(path))
        return path

    def _create_docx_table_no_caption(self) -> Path:
        path = self.temp_path / "table_no_caption.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Just some text, no table caption.")

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = TablesCheck()
        return check.run(ctx)

    def test_table_caption_not_bold_raises_error(self) -> None:
        """Table caption not bold should raise error."""
        docx_path = self._create_docx_table_caption_not_bold()
        issues = self._run_check(docx_path)
        bold_errors = [i for i in issues if "caption_bold" in i.check and i.severity == "error"]
        self.assertGreater(len(bold_errors), 0, f"Expected bold error but got: {issues}")

    def test_table_caption_not_italic_raises_warning(self) -> None:
        """Table caption title not italic should raise warning."""
        docx_path = self._create_docx_table_caption_not_italic()
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(italic_warnings), 0, f"Expected italic warning but got: {issues}")

    def test_table_no_caption_raises_warning(self) -> None:
        """Table without caption should raise warning."""
        docx_path = self._create_docx_table_no_caption()
        issues = self._run_check(docx_path)
        caption_warnings = [
            i for i in issues if "caption_present" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(caption_warnings), 0, f"Expected caption warning but got: {issues}")


class TestTablesCheckMultipleTables(unittest.TestCase):
    """Tests for documents with multiple tables."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_multiple_tables(self) -> Path:
        path = self.temp_path / "multiple_tables.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        caption1 = doc.add_paragraph()
        run1 = caption1.add_run("Table 1")
        run1.bold = True
        run2 = caption1.add_run(". First Table")
        run2.italic = True
        doc.add_table(rows=2, cols=2)

        caption2 = doc.add_paragraph()
        run3 = caption2.add_run("Table 2")
        run3.bold = True
        run4 = caption2.add_run(". Second Table")
        run4.italic = True
        doc.add_table(rows=2, cols=2)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = TablesCheck()
        return check.run(ctx)

    def test_multiple_tables_all_proper_passes(self) -> None:
        """Document with multiple properly captioned tables should pass."""
        docx_path = self._create_docx_multiple_tables()
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")


class TestTablesCheckCaptionItalicSeparateParagraph(unittest.TestCase):
    """Tests for fix #5: italic title must live in a SEPARATE paragraph after 'Tabla N'."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        name: str,
        caption_text: str,
        caption_bold: bool,
        next_text: str | None,
        next_italic: bool = False,
        next_bold: bool = False,
    ) -> Path:
        path = self.temp_path / name
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(caption_text)
        if caption_bold:
            caption_run.bold = True

        if next_text is not None:
            next_para = doc.add_paragraph()
            next_run = next_para.add_run(next_text)
            if next_italic:
                next_run.italic = True
            if next_bold:
                next_run.bold = True

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = TablesCheck()
        return check.run(ctx)

    def test_italic_title_in_next_paragraph_passes(self) -> None:
        """Bold 'Tabla 1' + SEPARATE italic title paragraph => no caption_italic warning."""
        docx_path = self._create_docx(
            "separate_italic_passes.docx",
            caption_text="Tabla 1",
            caption_bold=True,
            next_text="Title of the table in italic",
            next_italic=True,
        )
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")
        self.assertEqual(
            italic_warnings, [], f"Expected no italic warning for separate title: {italic_warnings}"
        )

    def test_missing_italic_title_after_caption_raises_warning(self) -> None:
        """Bold 'Tabla 1' + non-italic next paragraph => caption_italic warning fires."""
        docx_path = self._create_docx(
            "missing_italic.docx",
            caption_text="Tabla 1",
            caption_bold=True,
            next_text="Title not italic",
            next_italic=False,
        )
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertGreater(len(italic_warnings), 0, f"Expected italic warning but got: {issues}")

    def test_next_paragraph_is_nota_does_not_count_as_title(self) -> None:
        """An italic 'Nota.' next paragraph must NOT be mistaken for the title (guard L95)."""
        docx_path = self._create_docx(
            "next_is_nota.docx",
            caption_text="Tabla 1",
            caption_bold=True,
            next_text="Nota. An italic note.",
            next_italic=True,
        )
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(italic_warnings),
            0,
            f"Nota must not count as title; warning should still fire: {issues}",
        )

    def test_next_paragraph_is_otra_tabla_does_not_count_as_title(self) -> None:
        """An italic 'Tabla 2' next paragraph must NOT be mistaken for the title (guard L93)."""
        docx_path = self._create_docx(
            "next_is_tabla.docx",
            caption_text="Tabla 1",
            caption_bold=True,
            next_text="Tabla 2",
            next_italic=True,
        )
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(italic_warnings),
            0,
            f"'Tabla N' must not count as title; warning should still fire: {issues}",
        )

    def test_title_starting_with_tabla_word_counts_as_title(self) -> None:
        """An italic title like 'Tabla comparativa de...' must count as the title."""
        docx_path = self._create_docx(
            "title_with_tabla_word.docx",
            caption_text="Table 5",
            caption_bold=True,
            next_text="Tabla comparativa de IDE y lenguajes",
            next_italic=True,
        )
        issues = self._run_check(docx_path)
        italic_warnings = [
            i for i in issues if "caption_italic" in i.check and i.severity == "warning"
        ]
        self.assertEqual(
            italic_warnings,
            [],
            f"'Tabla comparativa...' title must count as italic title: {issues}",
        )


class TestTablesCheckNotesAndNumbering(unittest.TestCase):
    """Tests for table note formatting and numbering sequence."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _run_check(self, doc: Document) -> list:
        path = self.temp_path / "tables.docx"
        doc.save(str(path))
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=path, meta=meta)
        ctx = VerificationContext(
            pdf=MagicMock(),
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        try:
            return TablesCheck().run(ctx)
        finally:
            verifier.close()

    @staticmethod
    def _add_table_with_caption(doc: Document, number: int) -> None:
        caption = doc.add_paragraph()
        caption.add_run(f"Table {number}").bold = True
        doc.add_paragraph("Table title").runs[0].italic = True
        doc.add_table(rows=2, cols=2)

    def test_proper_nota_passes(self):
        """An italic 'Nota.' below the table produces no note issues."""
        doc = Document()
        self._add_table_with_caption(doc, 1)
        nota = doc.add_paragraph()
        nota.add_run("Nota. ").italic = True
        nota.add_run("Elaboración propia.")

        issues = self._run_check(doc)
        note_issues = [i for i in issues if "note_" in i.check]
        self.assertEqual(note_issues, [], f"Expected no note issues but got: {note_issues}")

    def test_malformed_nota_is_reported(self):
        """A 'Nota' label without period or italics is flagged."""
        doc = Document()
        self._add_table_with_caption(doc, 1)
        doc.add_paragraph("Nota Datos propios del autor")

        issues = self._run_check(doc)
        formats = [i for i in issues if "note_format" in i.check]
        italics = [i for i in issues if "note_italic" in i.check]
        self.assertEqual(len(formats), 1)
        self.assertEqual(len(italics), 1)

    def test_numbering_gap_is_reported(self):
        """Non-consecutive table numbers are flagged."""
        doc = Document()
        self._add_table_with_caption(doc, 1)
        self._add_table_with_caption(doc, 3)

        issues = self._run_check(doc)
        sequence = [i for i in issues if "numbering_sequence" in i.check]
        self.assertEqual(len(sequence), 1)


if __name__ == "__main__":
    unittest.main()
