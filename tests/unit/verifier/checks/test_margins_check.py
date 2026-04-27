"""Unit tests for MarginsCheck - APA 7th Edition margin verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.margins import MarginsCheck


class TestMarginsCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant margins (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        top: float = 1.0,
        right: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        page_width: float = 8.5,
        page_height: float = 11.0,
    ) -> Path:
        path = self.temp_path / f"margins_{top}_{right}_{bottom}_{left}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(top)
        section.right_margin = Inches(right)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.page_width = Inches(page_width)
        section.page_height = Inches(page_height)
        doc.add_paragraph("Sample text for margin testing.")
        doc.save(str(path))
        return path

    def test_margins_check_exact_1_inch_all_sides(self) -> None:
        """Margins of exactly 1 inch on all sides should pass."""
        docx_path = self._create_docx(top=1.0, right=1.0, bottom=1.0, left=1.0)
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()

        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )

        check = MarginsCheck()
        issues = check.run(ctx)

        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")

    def test_margins_check_within_tolerance(self) -> None:
        """Margins within 0.05 inch tolerance should pass."""
        docx_path = self._create_docx(top=1.03, right=0.98, bottom=1.02, left=0.97)
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()

        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )

        check = MarginsCheck()
        issues = check.run(ctx)

        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors but got: {errors}")


class TestMarginsCheckViolation(unittest.TestCase):
    """Tests for margin violations (should fail with errors)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        top: float = 1.0,
        right: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        page_width: float = 8.5,
        page_height: float = 11.0,
    ) -> Path:
        path = self.temp_path / f"margins_{top}_{right}_{bottom}_{left}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(top)
        section.right_margin = Inches(right)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.page_width = Inches(page_width)
        section.page_height = Inches(page_height)
        doc.add_paragraph("Sample text for margin testing.")
        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = MarginsCheck()
        return check.run(ctx)

    def test_top_margin_too_small_raises_error(self) -> None:
        """Top margin of 0.8 inches should raise error."""
        docx_path = self._create_docx(top=0.8)
        issues = self._run_check(docx_path)
        margin_issues = [i for i in issues if "margins.top" in i.check and i.severity == "error"]
        self.assertGreater(len(margin_issues), 0, "Expected error for top margin 0.8 inches")

    def test_top_margin_too_large_raises_error(self) -> None:
        """Top margin of 1.5 inches should raise error."""
        docx_path = self._create_docx(top=1.5)
        issues = self._run_check(docx_path)
        margin_issues = [i for i in issues if "margins.top" in i.check and i.severity == "error"]
        self.assertGreater(len(margin_issues), 0, "Expected error for top margin 1.5 inches")

    def test_right_margin_too_small_raises_error(self) -> None:
        """Right margin of 0.8 inches should raise error."""
        docx_path = self._create_docx(right=0.8)
        issues = self._run_check(docx_path)
        margin_issues = [i for i in issues if "margins.right" in i.check and i.severity == "error"]
        self.assertGreater(len(margin_issues), 0, "Expected error for right margin 0.8 inches")

    def test_bottom_margin_too_small_raises_error(self) -> None:
        """Bottom margin of 0.8 inches should raise error."""
        docx_path = self._create_docx(bottom=0.8)
        issues = self._run_check(docx_path)
        margin_issues = [i for i in issues if "margins.bottom" in i.check and i.severity == "error"]
        self.assertGreater(len(margin_issues), 0, "Expected error for bottom margin 0.8 inches")

    def test_left_margin_too_small_raises_error(self) -> None:
        """Left margin of 0.8 inches should raise error."""
        docx_path = self._create_docx(left=0.8)
        issues = self._run_check(docx_path)
        margin_issues = [i for i in issues if "margins.left" in i.check and i.severity == "error"]
        self.assertGreater(len(margin_issues), 0, "Expected error for left margin 0.8 inches")

    def test_all_margins_wrong_raises_multiple_errors(self) -> None:
        """All margins wrong should raise errors for each."""
        docx_path = self._create_docx(top=0.8, right=0.7, bottom=1.6, left=0.9)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertGreaterEqual(
            len(errors), 4, f"Expected at least 4 errors, got {len(errors)}: {errors}"
        )


class TestMarginsCheckPageSize(unittest.TestCase):
    """Tests for page size verification."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(self, page_width: float = 8.5, page_height: float = 11.0) -> Path:
        path = self.temp_path / f"page_{page_width}x{page_height}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(page_width)
        section.page_height = Inches(page_height)
        doc.add_paragraph("Sample text for page size testing.")
        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = MarginsCheck()
        return check.run(ctx)

    def test_letter_size_passes(self) -> None:
        """Letter size (8.5 x 11) should pass."""
        docx_path = self._create_docx(page_width=8.5, page_height=11.0)
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if i.severity == "error" and "page_" in i.check]
        self.assertEqual(page_errors, [], f"Expected no page size errors but got: {page_errors}")

    def test_a4_size_fails(self) -> None:
        """A4 size (8.27 x 11.69) should fail."""
        docx_path = self._create_docx(page_width=8.27, page_height=11.69)
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if i.severity == "error" and "page_" in i.check]
        self.assertGreater(len(page_errors), 0, "Expected error for A4 page size")

    def test_legal_size_fails(self) -> None:
        """Legal size (8.5 x 14) should fail."""
        docx_path = self._create_docx(page_width=8.5, page_height=14.0)
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if i.severity == "error" and "page_" in i.check]
        self.assertGreater(len(page_errors), 0, "Expected error for legal page size")

    def test_wide_page_fails(self) -> None:
        """Wide page (11 x 8.5) should fail."""
        docx_path = self._create_docx(page_width=11.0, page_height=8.5)
        issues = self._run_check(docx_path)
        page_errors = [i for i in issues if i.severity == "error" and "page_" in i.check]
        self.assertGreater(len(page_errors), 0, "Expected error for wide page orientation")


if __name__ == "__main__":
    unittest.main()
