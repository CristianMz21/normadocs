"""Unit tests for ParagraphsCheck - APA 7th Edition paragraph formatting."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.paragraphs import ParagraphsCheck


class TestParagraphsCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant paragraph formatting (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx(
        self,
        first_line_indent_inches: float = 0.5,
        alignment: str = "justify",
        paragraph_count: int = 5,
    ) -> Path:
        path = self.temp_path / f"para_{first_line_indent_inches}_{alignment}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(paragraph_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(first_line_indent_inches)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Body paragraph {i + 1} with proper formatting.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ParagraphsCheck()
        return check.run(ctx)

    def test_exact_half_inch_indent_passes(self) -> None:
        """Exactly 0.5 inch first-line indent should pass."""
        docx_path = self._create_docx(first_line_indent_inches=0.5)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        self.assertEqual(indent_errors, [], f"Expected no indent errors but got: {indent_errors}")

    def test_indent_within_tolerance_passes(self) -> None:
        """0.55 inch indent (within 0.1 tolerance) should pass."""
        docx_path = self._create_docx(first_line_indent_inches=0.55)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        self.assertEqual(indent_errors, [], f"Expected no indent errors but got: {indent_errors}")


class TestParagraphsCheckIndentViolation(unittest.TestCase):
    """Tests for first-line indent violations."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_no_indent(self, paragraph_count: int = 5) -> Path:
        path = self.temp_path / "no_indent.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(paragraph_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = None
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Paragraph {i + 1} without indent.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _create_mixed_indent_docx(self, with_indent: int, without_indent: int) -> Path:
        """Create document with mixed indent paragraphs."""
        path = self.temp_path / "mixed_indent.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(with_indent):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Indented paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        for i in range(without_indent):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = None
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Non-indented paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ParagraphsCheck()
        return check.run(ctx)

    def _create_docx_wrong_nonzero_indent(self, indent_inches: float = 1.0) -> Path:
        """Create document where every paragraph has a non-zero indent OUTSIDE tolerance."""
        path = self.temp_path / f"wrong_nonzero_{indent_inches}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(5):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(indent_inches)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Paragraph {i + 1} with wrong indent value.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def test_no_indent_all_paragraphs_raises_error(self) -> None:
        """All paragraphs without indent should raise error."""
        docx_path = self._create_docx_no_indent(paragraph_count=5)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        self.assertGreater(len(indent_errors), 0, f"Expected indent error but got: {issues}")

    def test_wrong_nonzero_indent_counted_as_missing(self) -> None:
        """Indent present but outside tolerance (1.0in) should count as missing (branch L117)."""
        docx_path = self._create_docx_wrong_nonzero_indent(indent_inches=1.0)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        self.assertGreater(
            len(indent_errors),
            0,
            f"Expected indent error for out-of-tolerance value but got: {issues}",
        )

    def test_mixed_indent_majority_no_indent_raises_error(self) -> None:
        """4 without indent out of 5 should raise error (<50% indented, >3 lack indent)."""
        docx_path = self._create_mixed_indent_docx(with_indent=1, without_indent=4)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        self.assertGreater(
            len(indent_errors), 0, f"Expected error when majority lack indent but got: {issues}"
        )

    def test_mixed_indent_some_no_indent_raises_warning(self) -> None:
        """1 without indent out of 5 should raise warning (not >50%)."""
        docx_path = self._create_mixed_indent_docx(with_indent=4, without_indent=1)
        issues = self._run_check(docx_path)
        indent_errors = [i for i in issues if "indent" in i.check and i.severity == "error"]
        indent_warnings = [i for i in issues if "indent" in i.check and i.severity == "warning"]
        self.assertEqual(
            indent_errors, [], f"Expected no error when <50% wrong but got: {indent_errors}"
        )
        self.assertGreater(
            len(indent_warnings), 0, f"Expected warning when some lack indent but got: {issues}"
        )


class TestParagraphsCheckJustification(unittest.TestCase):
    """Tests for text justification."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_justified(self, paragraph_count: int = 5) -> Path:
        path = self.temp_path / "justified.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(paragraph_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Justified paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _create_docx_left_aligned(self, paragraph_count: int = 5) -> Path:
        path = self.temp_path / "left_aligned.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(paragraph_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(f"Left-aligned paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _create_mixed_alignment_docx(self, justified_count: int, left_count: int) -> Path:
        path = self.temp_path / "mixed_alignment.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(justified_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(f"Justified paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        for i in range(left_count):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(f"Left-aligned paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ParagraphsCheck()
        return check.run(ctx)

    def test_all_justified_raises_warning(self) -> None:
        """All paragraphs justified should raise justification warning.

        APA 7 Section 2.21: "Do not use full justification. Leave the right
        margin uneven, or 'ragged.'" Left alignment is required.
        """
        docx_path = self._create_docx_justified()
        issues = self._run_check(docx_path)
        just_warnings = [
            i for i in issues if "justification" in i.check and i.severity == "warning"
        ]
        self.assertGreater(
            len(just_warnings), 0, f"Expected justification warning but got: {issues}"
        )

    def test_all_left_aligned_no_warning(self) -> None:
        """All paragraphs left-aligned should NOT raise justification warning.

        APA 7 Section 2.21: Left alignment with ragged right margin is correct.
        """
        docx_path = self._create_docx_left_aligned()
        issues = self._run_check(docx_path)
        just_warnings = [
            i for i in issues if "justification" in i.check and i.severity == "warning"
        ]
        self.assertEqual(
            just_warnings, [], f"Expected no justification warning but got: {just_warnings}"
        )

    def test_mixed_majority_left_aligned_no_warning(self) -> None:
        """3 left-aligned out of 5 (>=50%) should not raise warning."""
        docx_path = self._create_mixed_alignment_docx(justified_count=2, left_count=3)
        issues = self._run_check(docx_path)
        just_warnings = [
            i for i in issues if "justification" in i.check and i.severity == "warning"
        ]
        self.assertEqual(
            just_warnings,
            [],
            f"Expected no warning when >=50% left-aligned but got: {just_warnings}",
        )


class TestParagraphsCheckExclusions(unittest.TestCase):
    """Tests for paragraphs that must be excluded from the first-line indent count."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _ensure_style(self, doc: Document, name: str) -> None:
        """Create a custom paragraph style if it does not already exist."""
        existing = {s.name for s in doc.styles}
        if name in existing:
            return
        from docx.enum.style import WD_STYLE_TYPE

        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        style: str | None = None,
        first_line_indent: float | None = 0.5,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        bold: bool = False,
        set_font: bool = True,
    ) -> None:
        para = doc.add_paragraph(style=style) if style is not None else doc.add_paragraph()
        run = para.add_run(text)
        if set_font:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        if bold:
            run.bold = True
        if first_line_indent is not None:
            para.paragraph_format.first_line_indent = Inches(first_line_indent)
        else:
            para.paragraph_format.first_line_indent = None
        para.alignment = alignment

    def _add_body_paragraphs(self, doc: Document, count: int = 5) -> None:
        for i in range(count):
            self._add_paragraph(
                doc,
                f"Body paragraph {i + 1} with proper left-aligned formatting.",
            )

    def _create_docx(self, name: str, build_fn) -> Path:
        path = self.temp_path / name
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        build_fn(doc)
        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = ParagraphsCheck()
        return check.run(ctx)

    def test_heading_paragraphs_excluded_from_indent(self) -> None:
        """Heading-styled paragraphs must not be counted as missing first-line indent."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(doc, "Section Heading", style="Heading 1", first_line_indent=None)
            self._add_paragraph(doc, "Another Heading", style="Heading 1", first_line_indent=None)

        docx_path = self._create_docx("headings_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(indent_issues, [], f"Headings should be excluded but got: {indent_issues}")

    def test_references_section_excluded_from_indent(self) -> None:
        """Paragraphs after a 'Referencias' heading must be excluded from indent check."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(
                doc, "Referencias", style="Heading 1", first_line_indent=None, bold=True
            )
            for j in range(3):
                self._add_paragraph(
                    doc,
                    f"Author {chr(65 + j)}, A. (2024). Work title number {j + 1}.",
                    first_line_indent=0.5,
                )

        docx_path = self._create_docx("refs_section_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"References section should be excluded but got: {indent_issues}"
        )

    def test_abstract_section_excluded_from_indent(self) -> None:
        """Paragraphs inside an abstract section must be excluded from indent check."""

        def build(doc: Document) -> None:
            self._add_paragraph(
                doc, "Abstract", style="Heading 1", first_line_indent=None, bold=True
            )
            for j in range(3):
                self._add_paragraph(
                    doc,
                    f"This is abstract paragraph {j + 1} without indent.",
                    first_line_indent=None,
                )
            self._add_body_paragraphs(doc, count=5)

        docx_path = self._create_docx("abstract_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"Abstract section should be excluded but got: {indent_issues}"
        )

    def test_toc_section_excluded_from_indent(self) -> None:
        """Table-of-contents paragraphs must be excluded from indent check."""

        def build(doc: Document) -> None:
            self._add_paragraph(
                doc, "Contenido", style="Heading 1", first_line_indent=None, bold=True
            )
            for j in range(3):
                self._add_paragraph(
                    doc, f"Chapter {j + 1} ......... page {j + 1}", first_line_indent=None
                )
            self._add_body_paragraphs(doc, count=5)

        docx_path = self._create_docx("toc_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"TOC section should be excluded but got: {indent_issues}"
        )

    def test_abstract_ends_on_keywords(self) -> None:
        """A 'Palabras clave:' paragraph ends abstract and is itself skipped."""

        def build(doc: Document) -> None:
            self._add_paragraph(
                doc, "Abstract", style="Heading 1", first_line_indent=None, bold=True
            )
            self._add_paragraph(doc, "Abstract body without indent.", first_line_indent=None)
            self._add_paragraph(
                doc, "Palabras clave: keyword one, keyword two.", first_line_indent=None
            )
            self._add_body_paragraphs(doc, count=5)

        docx_path = self._create_docx("abstract_keywords.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"Keywords paragraph should be skipped but got: {indent_issues}"
        )

    def test_list_paragraph_excluded(self) -> None:
        """Paragraphs whose style contains 'List' must be excluded from indent check."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(
                doc, "First list item.", style="List Paragraph", first_line_indent=None
            )
            self._add_paragraph(
                doc, "Second list item.", style="List Paragraph", first_line_indent=None
            )

        docx_path = self._create_docx("list_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"List paragraphs should be excluded but got: {indent_issues}"
        )

    def test_caption_style_paragraph_excluded(self) -> None:
        """Caption/Compact/Source-styled paragraphs must be excluded from indent check."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(doc, "A caption.", style="Caption", first_line_indent=None)
            self._ensure_style(doc, "Compact")
            self._add_paragraph(doc, "A compact line.", style="Compact", first_line_indent=None)
            self._ensure_style(doc, "Source")
            self._add_paragraph(doc, "Source: someone.", style="Source", first_line_indent=None)

        docx_path = self._create_docx("caption_styles_excluded.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues,
            [],
            f"Caption/Compact/Source paragraphs should be excluded but got: {indent_issues}",
        )

    def test_tabla_figura_nota_digit_excluded(self) -> None:
        """Paragraphs starting with Tabla/Figura/Nota. or all digits must be excluded."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(doc, "Tabla 1", first_line_indent=None)
            self._add_paragraph(doc, "Figura 2", first_line_indent=None)
            self._add_paragraph(doc, "Nota. This is a note.", first_line_indent=None)
            self._add_paragraph(doc, "1234", first_line_indent=None)

        docx_path = self._create_docx("tabla_figura_nota_digit.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues,
            [],
            f"Tabla/Figura/Nota/digit paragraphs should be excluded but got: {indent_issues}",
        )

    def test_empty_paragraph_skipped(self) -> None:
        """Empty paragraphs must be skipped from the indent count."""

        def build(doc: Document) -> None:
            self._add_body_paragraphs(doc, count=5)
            self._add_paragraph(doc, "", first_line_indent=None)
            self._add_paragraph(doc, "", first_line_indent=None)

        docx_path = self._create_docx("empty_skipped.docx", build)
        issues = self._run_check(docx_path)
        indent_issues = [i for i in issues if "indent" in i.check]
        self.assertEqual(
            indent_issues, [], f"Empty paragraphs should not raise indent issues: {indent_issues}"
        )


if __name__ == "__main__":
    unittest.main()
