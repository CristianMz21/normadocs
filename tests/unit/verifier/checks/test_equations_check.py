"""Unit tests for EquationsCheck - APA 7th Edition equation verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.equations import EquationsCheck


def _add_display_math(doc: Document, number: str | None = None) -> None:
    """Add a display-equation paragraph, optionally with APA number and tabs."""
    p = doc.add_paragraph()
    if number is not None:
        p.add_run("\t")
        p.add_run(number)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    o_math = OxmlElement("m:oMath")
    m_r = OxmlElement("m:r")
    m_t = OxmlElement("m:t")
    m_t.text = "x+1"
    m_r.append(m_t)
    o_math.append(m_r)
    p._element.append(o_math)


class TestEquationsCheck(unittest.TestCase):
    """Tests for the equations verifier check."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _run_check(self, doc: Document) -> list:
        path = self.temp_path / "equations.docx"
        doc.save(str(path))
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=path, meta=meta)
        ctx = VerificationContext(
            pdf=MagicMock(),
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        try:
            return EquationsCheck().run(ctx)
        finally:
            verifier.close()

    def test_numbered_equations_pass(self):
        """APA-laid-out equations produce no issues."""
        doc = Document()
        _add_display_math(doc, "(1)")
        _add_display_math(doc, "(2)")

        issues = self._run_check(doc)
        self.assertEqual(issues, [])

    def test_unnumbered_equation_is_reported(self):
        """A display equation without a number raises warnings."""
        doc = Document()
        _add_display_math(doc, "(1)")
        _add_display_math(doc, None)

        issues = self._run_check(doc)
        present = [i for i in issues if "numbering_present" in i.check]
        layout = [i for i in issues if "tab_layout" in i.check]
        self.assertEqual(len(present), 1)
        self.assertEqual(len(layout), 1)

    def test_numbering_gap_is_reported(self):
        """Non-consecutive equation numbers raise the sequence error."""
        doc = Document()
        _add_display_math(doc, "(1)")
        _add_display_math(doc, "(3)")

        issues = self._run_check(doc)
        sequence = [i for i in issues if "numbering_sequence" in i.check]
        self.assertEqual(len(sequence), 1)
        self.assertEqual(sequence[0].severity, "error")

    def test_inline_math_in_prose_is_ignored(self):
        """Inline math inside a text paragraph is not a display equation."""
        doc = Document()
        p = doc.add_paragraph("La fórmula ")
        o_math = OxmlElement("m:oMath")
        p._element.append(o_math)
        p.add_run(" aparece en el texto.")

        issues = self._run_check(doc)
        self.assertEqual(issues, [])

    def test_document_without_math_passes(self):
        """Documents without equations produce no issues."""
        doc = Document()
        doc.add_paragraph("Solo texto.")

        issues = self._run_check(doc)
        self.assertEqual(issues, [])


if __name__ == "__main__":
    unittest.main()
