"""Unit tests for SpacingCheck - APA 7th Edition line spacing verification."""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import APAVerifier, VerificationContext
from normadocs.verifier.checks.spacing import SpacingCheck
from normadocs.verifier.docx_analyzer import DOCXParagraphInfo


class TestSpacingCheckCompliant(unittest.TestCase):
    """Tests for APA-compliant double spacing (should pass)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_spacing(self, line_spacing: float) -> Path:
        path = self.temp_path / f"spacing_{line_spacing}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(5):
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = line_spacing
            run = para.add_run(f"Paragraph {i + 1} with line spacing {line_spacing}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = SpacingCheck()
        return check.run(ctx)

    def test_exact_double_spacing_passes(self) -> None:
        """Exactly 2.0 line spacing should pass with no errors."""
        docx_path = self._create_docx_with_spacing(2.0)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for 2.0 spacing but got: {errors}")

    def test_double_spacing_within_tolerance_passes(self) -> None:
        """Line spacing within 0.2 of 2.0 should pass."""
        docx_path = self._create_docx_with_spacing(2.15)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertEqual(errors, [], f"Expected no errors for 2.15 spacing but got: {errors}")

    def test_all_empty_paragraphs_no_issue(self) -> None:
        """A document of only blank paragraphs should early-return no issues (branch L40)."""
        path = self.temp_path / "all_blank.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for _ in range(5):
            doc.add_paragraph()
        doc.save(str(path))
        issues = self._run_check(path)
        self.assertEqual(issues, [], f"Expected no issues for all-blank doc but got: {issues}")


class TestSpacingCheckViolation(unittest.TestCase):
    """Tests for spacing violations (should fail with errors)."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _create_docx_with_spacing(self, line_spacing: float) -> Path:
        path = self.temp_path / f"spacing_{line_spacing}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(5):
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = line_spacing
            run = para.add_run(f"Paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _create_mixed_spacing_docx(self, correct_count: int, wrong_count: int) -> Path:
        """Create document with mixed spacing."""
        path = self.temp_path / "mixed_spacing.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        for i in range(correct_count):
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            run = para.add_run(f"Correct spacing paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        for i in range(wrong_count):
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(f"Wrong spacing paragraph {i + 1}.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        doc.save(str(path))
        return path

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = SpacingCheck()
        return check.run(ctx)

    def test_single_spacing_raises_error(self) -> None:
        """Single spacing (1.0) should raise error (>50% wrong)."""
        docx_path = self._create_docx_with_spacing(1.0)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected error for single spacing but got: {issues}")

    def test_1_5_spacing_raises_error(self) -> None:
        """1.5 spacing should raise error (outside tolerance of 2.0)."""
        docx_path = self._create_docx_with_spacing(1.5)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected error for 1.5 spacing but got: {issues}")

    def test_mixed_spacing_majority_wrong_raises_error(self) -> None:
        """3 wrong out of 5 should raise error (>50% wrong)."""
        docx_path = self._create_mixed_spacing_docx(correct_count=2, wrong_count=3)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        self.assertGreater(len(errors), 0, f"Expected error when majority wrong but got: {issues}")

    def test_mixed_spacing_some_wrong_raises_warning(self) -> None:
        """1 wrong out of 5 should raise warning (not >50% wrong)."""
        docx_path = self._create_mixed_spacing_docx(correct_count=4, wrong_count=1)
        issues = self._run_check(docx_path)
        errors = [i for i in issues if i.severity == "error"]
        warnings = [i for i in issues if i.severity == "warning"]
        self.assertEqual(errors, [], f"Expected no errors when <50% wrong but got: {errors}")
        self.assertGreater(len(warnings), 0, f"Expected warning when some wrong but got: {issues}")

    def _run_check_with_info(self, docx_path: Path, info: DOCXParagraphInfo) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = SpacingCheck()
        with patch.object(ctx.docx, "get_paragraphs_info", return_value=[info]):
            return check.run(ctx)

    def test_int_line_spacing_handled(self) -> None:
        """An int line_spacing value must be cast to float and treated as double (branch L65-66)."""
        docx_path = self._create_docx_with_spacing(2.0)
        info = DOCXParagraphInfo(
            text="Synthetic paragraph with int spacing.",
            style_name="Normal",
            alignment="left",
            first_line_indent=None,
            space_before=None,
            space_after=None,
            line_spacing=2,
            runs=[],
        )
        issues = self._run_check_with_info(docx_path, info)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"int line_spacing=2 should be treated as double: {spacing_issues}"
        )

    def test_nonnumeric_line_spacing_skipped(self) -> None:
        """A non-int/non-float line_spacing must be silently skipped (branch L67-68)."""
        docx_path = self._create_docx_with_spacing(2.0)
        info = DOCXParagraphInfo(
            text="Synthetic paragraph with bogus spacing.",
            style_name="Normal",
            alignment="left",
            first_line_indent=None,
            space_before=None,
            space_after=None,
            line_spacing="double",
            runs=[],
        )
        issues = self._run_check_with_info(docx_path, info)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Non-numeric line_spacing should be skipped: {spacing_issues}"
        )


class TestSpacingCheckExclusions(unittest.TestCase):
    """Tests for spacing exclusions: table captions/titles/Nota/Heading must be skipped."""

    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_dir = TemporaryDirectory()
        cls.temp_path = Path(cls.temp_dir.name)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temp_dir.cleanup()

    def _ensure_style(self, doc: Document, name: str) -> None:
        existing = {s.name for s in doc.styles}
        if name in existing:
            return
        from docx.enum.style import WD_STYLE_TYPE

        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    def _add_body(self, doc: Document, count: int = 5) -> None:
        for i in range(count):
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            run = para.add_run(f"Body paragraph {i + 1} with double spacing.")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def _create_docx(self, name: str, build_fn) -> Path:
        path = self.temp_path / name
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        build_fn(doc)
        doc.save(str(path))
        return path

    def _add_single_spaced(self, doc: Document, text: str, style: str | None = None) -> None:
        para = doc.add_paragraph(style=style) if style is not None else doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.0

    def _run_check(self, docx_path: Path) -> list:
        pdf_path = self.temp_path / "output.pdf"
        pdf_path.touch()
        meta = DocumentMetadata(title="Test Document")
        verifier = APAVerifier(pdf_path=pdf_path, docx_path=docx_path, meta=meta)
        ctx = VerificationContext(
            pdf=verifier.pdf,
            docx=verifier.docx,
            meta=meta,
            strict=False,
        )
        check = SpacingCheck()
        return check.run(ctx)

    def test_tabla_caption_single_spaced_not_flagged(self) -> None:
        """A single-spaced 'Tabla 1' paragraph must not be flagged (filter L47)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Tabla 1")

        docx_path = self._create_docx("tabla_caption.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Tabla caption should be excluded from spacing: {spacing_issues}"
        )

    def test_figura_caption_single_spaced_not_flagged(self) -> None:
        """A single-spaced 'Figura 1' paragraph must not be flagged (filter L48)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Figura 1")

        docx_path = self._create_docx("figura_caption.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Figura caption should be excluded from spacing: {spacing_issues}"
        )

    def test_table_english_caption_not_flagged(self) -> None:
        """A single-spaced 'Table 1' paragraph must not be flagged (English prefix)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Table 1")

        docx_path = self._create_docx("table_english_caption.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Table caption should be excluded from spacing: {spacing_issues}"
        )

    def test_figure_english_caption_not_flagged(self) -> None:
        """A single-spaced 'Figure 1' paragraph must not be flagged (English prefix)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Figure 1")

        docx_path = self._create_docx("figure_english_caption.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Figure caption should be excluded from spacing: {spacing_issues}"
        )

    def test_table_title_line_not_flagged(self) -> None:
        """A single-spaced table-title line ('. Process inputs') must not be flagged."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, ". Process inputs")

        docx_path = self._create_docx("table_title_line.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues,
            [],
            f"Table title line should be excluded from spacing: {spacing_issues}",
        )

    def test_nota_paragraph_single_spaced_not_flagged(self) -> None:
        """A single-spaced 'Nota. ...' paragraph must not be flagged (filter L49)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Nota. This is a single-spaced note.")

        docx_path = self._create_docx("nota_para.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"Nota paragraph should be excluded from spacing: {spacing_issues}"
        )

    def test_all_digit_paragraph_single_spaced_not_flagged(self) -> None:
        """A single-spaced all-digit paragraph must not be flagged (filter L50)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "2024")

        docx_path = self._create_docx("digit_para.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues, [], f"All-digit paragraph should be excluded: {spacing_issues}"
        )

    def test_heading_and_caption_styles_not_flagged(self) -> None:
        """Single-spaced Heading- and Caption-styled paragraphs must not be flagged (L52-54)."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            self._add_single_spaced(doc, "Heading text", style="Heading 1")
            self._add_single_spaced(doc, "A caption", style="Caption")

        docx_path = self._create_docx("heading_caption_styles.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues,
            [],
            f"Heading/Caption-styled paragraphs should be excluded: {spacing_issues}",
        )

    def test_title_line_after_caption_not_flagged(self) -> None:
        """An italic single-spaced title right after a 'Table N' caption must be excluded."""

        def build(doc: Document) -> None:
            self._add_body(doc, count=5)
            cap = doc.add_paragraph()
            cap_run = cap.add_run("Table 1")
            cap_run.bold = True
            cap_run.font.name = "Times New Roman"
            cap_run.font.size = Pt(12)
            cap.paragraph_format.line_spacing = 1.0
            title = doc.add_paragraph()
            title_run = title.add_run("Tabla comparativa de herramientas")
            title_run.italic = True
            title_run.font.name = "Times New Roman"
            title_run.font.size = Pt(12)
            title.paragraph_format.line_spacing = 1.0

        docx_path = self._create_docx("caption_with_title.docx", build)
        issues = self._run_check(docx_path)
        spacing_issues = [i for i in issues if "line_spacing" in i.check]
        self.assertEqual(
            spacing_issues,
            [],
            f"Caption+title block should be excluded from spacing: {spacing_issues}",
        )

    def test_only_excluded_paragraphs_no_issue(self) -> None:
        """A doc with ONLY excluded paragraphs should early-return no issues (branch L57)."""

        def build(doc: Document) -> None:
            self._add_single_spaced(doc, "Tabla 1")
            self._add_single_spaced(doc, "Figura 1")
            self._add_single_spaced(doc, "Nota. note here.", style="Normal")
            self._add_single_spaced(doc, "Heading one", style="Heading 1")

        docx_path = self._create_docx("only_excluded.docx", build)
        issues = self._run_check(docx_path)
        self.assertEqual(
            issues, [], f"Only-excluded doc should produce no spacing issues: {issues}"
        )


if __name__ == "__main__":
    unittest.main()
