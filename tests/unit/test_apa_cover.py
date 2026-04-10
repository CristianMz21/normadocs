"""
Unit Tests for APA Formatter Cover Page.

Tests _add_cover_page method.
"""

import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from normadocs.formatters.apa import APADocxFormatter
from normadocs.models import DocumentMetadata


class TestAddCoverPage(unittest.TestCase):
    """Tests for _add_cover_page method."""

    def _create_formatter_with_doc(
        self, paragraphs_data: list[dict] | None = None
    ) -> tuple[APADocxFormatter, str]:
        """Create a formatter with optional custom paragraphs. Returns (formatter, temp_path)."""
        doc = Document()
        if paragraphs_data:
            for para_data in paragraphs_data:
                style_name = para_data.get("style", "Normal")
                p = doc.add_paragraph(para_data["text"], style=style_name)
                if "runs" in para_data:
                    for run in list(p.runs):
                        run._r.getparent().remove(run._r)
                    for run_data in para_data["runs"]:
                        r = p.add_run(run_data["text"])
                        if run_data.get("bold"):
                            r.bold = True
                        if run_data.get("italic"):
                            r.italic = True

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        formatter = APADocxFormatter(temp_path)
        formatter.doc = doc
        return formatter, temp_path

    def test_title_appears_centered(self):
        """Title should appear in the document and be centered."""
        meta = DocumentMetadata(
            title="El Efecto del Aprendizaje Automático en la Educación Superior",
            author="Juan Pérez García",
        )
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Contenido de la introducción...", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            # Collect all paragraph texts
            all_texts = [p.text for p in formatter.doc.paragraphs]

            # Title should be present somewhere
            self.assertIn(
                meta.title,
                all_texts,
                f"Title '{meta.title}' not found in document paragraphs",
            )

            # Find the paragraph with the title and check it's centered
            title_para = None
            for p in formatter.doc.paragraphs:
                if p.text.strip() == meta.title.strip():
                    title_para = p
                    break

            self.assertIsNotNone(title_para, "Title paragraph not found")
            self.assertEqual(
                title_para.alignment,
                WD_ALIGN_PARAGRAPH.CENTER,
                "Title should be centered",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_author_appears(self):
        """Author name should appear in the cover page."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Juan Pérez García",
        )
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Contenido...", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            all_texts = [p.text for p in formatter.doc.paragraphs]
            self.assertIn(
                meta.author,
                all_texts,
                f"Author '{meta.author}' not found in document",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_institution_appears(self):
        """Institution should appear in the cover page."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Juan Pérez García",
            institution="SENA Centro de Tecnologías",
        )
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Contenido...", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            all_texts = [p.text for p in formatter.doc.paragraphs]
            self.assertIn(
                meta.institution,
                all_texts,
                f"Institution '{meta.institution}' not found in document",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_page_break_after_cover(self):
        """Page break should appear after the cover page (before first heading or content)."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Autor Test",
        )
        paragraphs = [
            {"text": "Párrafo de contenido.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            # Check that at least one paragraph after the cover has page_break_before
            # The first non-centered paragraph with text should trigger the break
            found_page_break = False
            for p in formatter.doc.paragraphs:
                if p.paragraph_format.page_break_before:
                    found_page_break = True
                    break

            self.assertTrue(
                found_page_break,
                "A paragraph should have page_break_before=True after cover page",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_cover_with_subtitle(self):
        """Cover page should handle subtitle if present in metadata."""
        meta = DocumentMetadata(
            title="Título del Documento",
            subtitle="Un Subtítulo Importante",
            author="Autor Test",
        )
        paragraphs = [
            {"text": "Contenido", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            all_texts = [p.text for p in formatter.doc.paragraphs]
            self.assertIn(
                "Un Subtítulo Importante",
                all_texts,
                "Subtitle should appear in cover page",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_cover_with_affiliation(self):
        """Cover page should handle affiliation if present in metadata."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Autor Test",
            affiliation="Departamento de Ingeniería",
        )
        paragraphs = [
            {"text": "Contenido", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            all_texts = [p.text for p in formatter.doc.paragraphs]
            self.assertIn(
                "Departamento de Ingeniería",
                all_texts,
                "Affiliation should appear in cover page",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_cover_with_date(self):
        """Cover page should display date if present in metadata."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Autor Test",
            date="2026-04-09",
        )
        paragraphs = [
            {"text": "Contenido", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            all_texts = [p.text for p in formatter.doc.paragraphs]
            self.assertIn(
                "2026-04-09",
                all_texts,
                "Date should appear in cover page",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_title_is_bold(self):
        """Title in cover page should be bold."""
        meta = DocumentMetadata(
            title="Título en Negrita",
            author="Autor Test",
        )
        paragraphs = [
            {"text": "Contenido", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            # Find the title paragraph and check bold on its runs
            for p in formatter.doc.paragraphs:
                if p.text.strip() == meta.title.strip():
                    # At least one run should be bold
                    bold_runs = [r for r in p.runs if r.bold and r.text.strip()]
                    self.assertGreater(
                        len(bold_runs),
                        0,
                        f"Title run should be bold but no bold runs found in '{p.text}'",
                    )
                    break
            else:
                self.fail(f"Title paragraph '{meta.title}' not found")
        finally:
            import os

            os.unlink(temp_path)

    def test_cover_uses_double_line_spacing(self):
        """Cover page paragraphs should use double line spacing."""
        meta = DocumentMetadata(
            title="Título del Documento",
            author="Autor Test",
        )
        paragraphs = [
            {"text": "Contenido", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._add_cover_page(meta)

            from docx.enum.text import WD_LINE_SPACING

            found_spacing = False
            for p in formatter.doc.paragraphs[:10]:
                if p.text.strip():
                    spacing = p.paragraph_format.line_spacing_rule
                    if spacing == WD_LINE_SPACING.DOUBLE:
                        found_spacing = True
                        break

            self.assertTrue(
                found_spacing, "At least one cover page paragraph should have double spacing"
            )
        finally:
            import os

            os.unlink(temp_path)


class TestBackwardCompatibilityModule(unittest.TestCase):
    """Tests for the backward compatibility apa.py module."""

    def _load_apa_module_directly(self):
        """Load the apa.py module file directly, bypassing package import.

        The apa.py module is shadowed by the apa/ package in normal imports.
        This method pre-loads the package first, then loads apa.py which can
        then import from the already-loaded package.
        """
        import importlib.util
        import sys

        # Clear apa-related modules from cache to allow re-import
        modules_to_clear = [k for k in list(sys.modules.keys()) if "apa" in k.lower()]
        for mod in modules_to_clear:
            del sys.modules[mod]

        # First load the actual apa subpackage so it's in sys.modules
        # This allows apa.py's import statement to find it

        # Now load apa.py with a distinct name to avoid sys.modules conflict
        spec = importlib.util.spec_from_file_location(
            "apa_compat_module", "src/normadocs/formatters/apa.py"
        )
        module = importlib.util.module_from_spec(spec)
        sys.modules["apa_compat_module"] = module
        spec.loader.exec_module(module)
        return module

    def test_apa_module_exports_apadocxformatter(self):
        """Test backward compatibility import from apa.py module.

        The apa.py module exists for backward compatibility and re-exports
        APADocxFormatter from the apa/ subpackage. This test forces the
        module to be loaded via importlib to ensure coverage.
        """

        module = self._load_apa_module_directly()

        self.assertTrue(hasattr(module, "APADocxFormatter"))
        self.assertEqual(module.APADocxFormatter.__name__, "APADocxFormatter")
        self.assertTrue(callable(module.APADocxFormatter))

    def test_apa_module_in_apa_all(self):
        """Test that APADocxFormatter is in __all__ of backward compat module."""
        module = self._load_apa_module_directly()

        self.assertIn("APADocxFormatter", module.__all__)


if __name__ == "__main__":
    unittest.main()
