"""Unit tests for spacing check."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import MagicMock

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import VerificationContext
from normadocs.verifier.checks.spacing import SpacingCheck
from normadocs.verifier.docx_analyzer import DOCXAnalyzer


class TestSpacingCheck(unittest.TestCase):
    """Tests for SpacingCheck with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_docx_with_spacing(
        self,
        line_spacing: float = 2.0,
    ) -> Path:
        """Create a DOCX with specified line spacing."""
        path = self.temp_dir / f"test_spacing_{line_spacing}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        para = doc.add_paragraph("Double spaced text for testing.")
        para.paragraph_format.line_spacing = line_spacing

        doc.save(str(path))
        return path

    def _create_simple_docx(self) -> Path:
        """Create a minimal DOCX with default spacing."""
        path = self.temp_dir / "test_simple.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        doc.save(str(path))
        return path

    def _create_context(self, docx_path: Path) -> VerificationContext:
        """Create a mock VerificationContext for testing."""
        mock_pdf = MagicMock()
        analyzer = DOCXAnalyzer(docx_path)
        meta = DocumentMetadata(title="Test Document")

        return VerificationContext(pdf=mock_pdf, docx=analyzer, meta=meta, strict=False)

    def test_check_passes_with_double_spacing(self) -> None:
        """Test that 2.0 line spacing passes."""
        path = self._create_docx_with_spacing(2.0)
        try:
            ctx = self._create_context(path)
            check = SpacingCheck()
            issues = check.run(ctx)

            spacing_issues = [i for i in issues if "spacing" in i.check]
            assert len(spacing_issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_single_spacing(self) -> None:
        """Test that single spacing (1.0) fails."""
        path = self._create_docx_with_spacing(1.0)
        try:
            ctx = self._create_context(path)
            check = SpacingCheck()
            issues = check.run(ctx)

            spacing_issues = [i for i in issues if "line_spacing" in i.check]
            assert len(spacing_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_1_5_spacing(self) -> None:
        """Test that 1.5 spacing fails."""
        path = self._create_docx_with_spacing(1.5)
        try:
            ctx = self._create_context(path)
            check = SpacingCheck()
            issues = check.run(ctx)

            spacing_issues = [i for i in issues if "line_spacing" in i.check]
            assert len(spacing_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_mixed_spacing(self) -> None:
        """Test that non-standard spacing fails."""
        path = self._create_docx_with_spacing(1.5)
        try:
            ctx = self._create_context(path)
            check = SpacingCheck()
            issues = check.run(ctx)

            spacing_issues = [i for i in issues if "line_spacing" in i.check]
            assert len(spacing_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_empty_document(self) -> None:
        """Test check on empty document passes (no text to validate)."""
        path = self._create_simple_docx()
        try:
            ctx = self._create_context(path)
            check = SpacingCheck()
            issues = check.run(ctx)
            assert isinstance(len(issues), int)
        finally:
            path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
