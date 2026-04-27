"""Unit tests for headings check."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import MagicMock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import VerificationContext
from normadocs.verifier.checks.headings import HeadingsCheck
from normadocs.verifier.docx_analyzer import DOCXAnalyzer


class TestHeadingsCheck(unittest.TestCase):
    """Tests for HeadingsCheck with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_docx_with_headings(self) -> Path:
        """Create a DOCX with properly formatted headings."""
        path = self.temp_dir / "test_headings.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        h1 = doc.add_heading("Heading 1", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
            run.bold = True

        h2 = doc.add_heading("Heading 2", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h2.runs:
            run.bold = True
            run.italic = True

        h3 = doc.add_heading("Heading 3", level=3)
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h3.runs:
            run.bold = True
            run.italic = True

        doc.save(str(path))
        return path

    def _create_docx_with_wrong_headings(self) -> Path:
        """Create a DOCX with improperly formatted headings."""
        path = self.temp_dir / "test_wrong_headings.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        h1 = doc.add_heading("Heading 1", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        h2 = doc.add_heading("Heading 2", level=2)
        for run in h2.runs:
            run.bold = False

        doc.save(str(path))
        return path

    def _create_context(self, docx_path: Path) -> VerificationContext:
        """Create a mock VerificationContext for testing."""
        mock_pdf = MagicMock()
        analyzer = DOCXAnalyzer(docx_path)
        meta = DocumentMetadata(title="Test Document")

        return VerificationContext(pdf=mock_pdf, docx=analyzer, meta=meta, strict=False)

    def test_check_passes_with_correct_headings(self) -> None:
        """Test that properly formatted headings pass."""
        path = self._create_docx_with_headings()
        try:
            ctx = self._create_context(path)
            check = HeadingsCheck()
            issues = check.run(ctx)

            assert len(issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_wrong_alignment(self) -> None:
        """Test that left-aligned Level 1 heading fails."""
        path = self._create_docx_with_wrong_headings()
        try:
            ctx = self._create_context(path)
            check = HeadingsCheck()
            issues = check.run(ctx)

            alignment_issues = [i for i in issues if "alignment" in i.check]
            assert len(alignment_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_not_bold(self) -> None:
        """Test that non-bold Level 2 heading fails."""
        path = self._create_docx_with_wrong_headings()
        try:
            ctx = self._create_context(path)
            check = HeadingsCheck()
            issues = check.run(ctx)

            bold_issues = [i for i in issues if "bold" in i.check]
            assert len(bold_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_no_headings_document(self) -> None:
        """Test check on document without headings passes (no headings is valid)."""
        path = self.temp_dir / "test_no_headings.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        doc.add_paragraph("Just a paragraph.")
        doc.save(str(path))

        try:
            ctx = self._create_context(path)
            check = HeadingsCheck()
            issues = check.run(ctx)

            heading_issues = [i for i in issues if "heading" in i.check]
            assert len(heading_issues) == 0
        finally:
            path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
