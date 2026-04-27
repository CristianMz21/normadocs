"""Unit tests for fonts check."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import MagicMock

from docx import Document
from docx.shared import Inches, Pt

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import VerificationContext
from normadocs.verifier.checks.fonts import FontsCheck
from normadocs.verifier.docx_analyzer import DOCXAnalyzer


class TestFontsCheck(unittest.TestCase):
    """Tests for FontsCheck with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_docx_with_fonts(
        self,
        body_font: str = "Times New Roman",
        body_size: int = 12,
        heading_font: str = "Times New Roman",
    ) -> Path:
        """Create a DOCX with specified fonts."""
        path = self.temp_dir / f"test_fonts_{body_font.replace(' ', '_')}_{body_size}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        body_para = doc.add_paragraph("This is body text.")
        for run in body_para.runs:
            run.font.name = body_font
            run.font.size = Pt(body_size)

        heading = doc.add_heading("Test Heading", level=1)
        for run in heading.runs:
            run.font.name = heading_font
            run.font.size = Pt(14)

        doc.save(str(path))
        return path

    def _create_simple_docx(self) -> Path:
        """Create a minimal DOCX."""
        path = self.temp_dir / "test_simple.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        doc.save(str(path))
        return path

    def _create_context(self, docx_path: Path) -> VerificationContext:
        """Create a mock VerificationContext for testing."""
        mock_pdf = MagicMock()
        analyzer = DOCXAnalyzer(docx_path)
        meta = DocumentMetadata(title="Test Document")

        return VerificationContext(pdf=mock_pdf, docx=analyzer, meta=meta, strict=False)

    def test_check_passes_with_times_new_roman_12(self) -> None:
        """Test that Times New Roman 12pt passes."""
        path = self._create_docx_with_fonts("Times New Roman", 12)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            font_issues = [i for i in issues if "font" in i.check]
            assert len(font_issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_wrong_font(self) -> None:
        """Test that non-Times New Roman font fails."""
        path = self._create_docx_with_fonts("Arial", 12)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            assert len(issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_wrong_size(self) -> None:
        """Test that wrong font size fails (10pt is outside tolerance)."""
        path = self._create_docx_with_fonts("Times New Roman", 10)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            size_issues = [i for i in issues if "size" in i.check]
            assert len(size_issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_correct_size(self) -> None:
        """Test that correct font size passes."""
        path = self._create_docx_with_fonts("Arial", 12)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            size_issues = [i for i in issues if "size" in i.check]
            assert len(size_issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_calibri(self) -> None:
        """Test that Calibri is rejected."""
        path = self._create_docx_with_fonts("Calibri", 12)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            assert len(issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_georgia(self) -> None:
        """Test that Georgia is rejected."""
        path = self._create_docx_with_fonts("Georgia", 12)
        try:
            ctx = self._create_context(path)
            check = FontsCheck()
            issues = check.run(ctx)

            assert len(issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_pt_from_emu_conversion(self) -> None:
        """Test EMU to point conversion."""
        check = FontsCheck()
        assert abs(check._pt_from_emu(12700) - 1.0) < 0.01
        assert abs(check._pt_from_emu(12700 * 12) - 12.0) < 0.01


if __name__ == "__main__":
    unittest.main()
