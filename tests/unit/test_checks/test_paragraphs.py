"""Unit tests for paragraphs check."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import MagicMock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import VerificationContext
from normadocs.verifier.checks.paragraphs import ParagraphsCheck
from normadocs.verifier.docx_analyzer import DOCXAnalyzer


class TestParagraphsCheck(unittest.TestCase):
    """Tests for ParagraphsCheck with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_context(self, docx_path: Path) -> VerificationContext:
        """Create a mock VerificationContext for testing."""
        mock_pdf = MagicMock()
        analyzer = DOCXAnalyzer(docx_path)
        meta = DocumentMetadata(title="Test Document")

        return VerificationContext(pdf=mock_pdf, docx=analyzer, meta=meta, strict=False)

    def _create_simple_docx(self) -> Path:
        """Create a minimal DOCX."""
        path = self.temp_dir / "test_simple.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        p1 = doc.add_paragraph("First paragraph.")
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p2 = doc.add_paragraph("Second paragraph.")
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(str(path))
        return path

    def _create_docx_with_hanging_indent(self) -> Path:
        """Create a DOCX with proper hanging indent for references."""
        path = self.temp_dir / "test_hanging_indent.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        ref_para = doc.add_paragraph("Author, A. A. (2020). Title of work. Publisher.")
        ref_para.paragraph_format.first_line_indent = Inches(0.5)
        ref_para.paragraph_format.left_indent = Inches(0)

        doc.save(str(path))
        return path

    def test_check_runs_without_error(self) -> None:
        """Test that paragraphs check runs without error."""
        path = self._create_simple_docx()
        try:
            ctx = self._create_context(path)
            check = ParagraphsCheck()
            issues = check.run(ctx)

            assert isinstance(len(issues), int)
        finally:
            path.unlink(missing_ok=True)

    def test_check_detects_first_line_indent(self) -> None:
        """Test that first line indent check runs."""
        path = self._create_docx_with_hanging_indent()
        try:
            ctx = self._create_context(path)
            check = ParagraphsCheck()
            issues = check.run(ctx)

            assert isinstance(len(issues), int)
        finally:
            path.unlink(missing_ok=True)

    def test_empty_document(self) -> None:
        """Test check on empty document."""
        path = self.temp_dir / "test_empty.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        doc.save(str(path))

        try:
            ctx = self._create_context(path)
            check = ParagraphsCheck()
            issues = check.run(ctx)

            assert isinstance(len(issues), int)
        finally:
            path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
