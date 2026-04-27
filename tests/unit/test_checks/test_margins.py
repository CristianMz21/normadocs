"""Unit tests for margins check."""

import unittest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import MagicMock

from docx import Document
from docx.shared import Inches

from normadocs.models import DocumentMetadata
from normadocs.verifier.apa_verifier import VerificationContext
from normadocs.verifier.checks.margins import MarginsCheck
from normadocs.verifier.docx_analyzer import DOCXAnalyzer


class TestMarginsCheck(unittest.TestCase):
    """Tests for MarginsCheck with controlled DOCX creation."""

    @classmethod
    def setUpClass(cls) -> None:
        """Create test DOCX fixtures."""
        with NamedTemporaryFile(delete=False, suffix=".docx", mode="w") as f:
            cls.temp_dir = Path(f.name).parent
        cls.temp_dir.mkdir(exist_ok=True)

    def _create_docx_with_margins(
        self,
        top: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        right: float = 1.0,
    ) -> Path:
        """Create a DOCX with specified margins."""
        path = self.temp_dir / f"test_margins_{top}_{bottom}_{left}_{right}.docx"
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        doc.add_paragraph("Test content.")
        doc.save(str(path))
        return path

    def _create_context(self, docx_path: Path) -> VerificationContext:
        """Create a mock VerificationContext for testing."""
        mock_pdf = MagicMock()
        mock_pdf.get_page_info.return_value = MagicMock()

        analyzer = DOCXAnalyzer(docx_path)
        meta = DocumentMetadata(title="Test Document")

        return VerificationContext(pdf=mock_pdf, docx=analyzer, meta=meta, strict=False)

    def test_check_passes_with_standard_margins(self) -> None:
        """Test that 1-inch margins pass."""
        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 1.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            assert len(issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_with_wrong_margins(self) -> None:
        """Test that non-1-inch margins fail."""
        path = self._create_docx_with_margins(1.5, 1.5, 1.5, 1.5)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            assert len(issues) > 0
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_top_margin(self) -> None:
        """Test that wrong top margin is detected."""
        path = self._create_docx_with_margins(2.0, 1.0, 1.0, 1.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("margin_top" in t or "margins.top" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_bottom_margin(self) -> None:
        """Test that wrong bottom margin is detected."""
        path = self._create_docx_with_margins(1.0, 0.5, 1.0, 1.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("margin_bottom" in t or "margins.bottom" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_left_margin(self) -> None:
        """Test that wrong left margin is detected."""
        path = self._create_docx_with_margins(1.0, 1.0, 1.25, 1.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("margin_left" in t or "margins.left" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)

    def test_check_fails_right_margin(self) -> None:
        """Test that wrong right margin is detected."""
        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 2.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("margin_right" in t or "margins.right" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)

    def test_tolerance_handles_slight_variations(self) -> None:
        """Test that slight margin variations within tolerance pass."""
        path = self._create_docx_with_margins(1.02, 1.0, 1.0, 1.0)
        try:
            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            assert len(issues) == 0
        finally:
            path.unlink(missing_ok=True)

    def test_page_width_check(self) -> None:
        """Test that wrong page width is detected."""
        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 1.0)
        try:
            doc = Document(str(path))
            section = doc.sections[0]
            section.page_width = Inches(7.5)
            doc.save(str(path))

            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("page_width" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)

    def test_page_height_check(self) -> None:
        """Test that wrong page height is detected."""
        path = self._create_docx_with_margins(1.0, 1.0, 1.0, 1.0)
        try:
            doc = Document(str(path))
            section = doc.sections[0]
            section.page_height = Inches(10.0)
            doc.save(str(path))

            ctx = self._create_context(path)
            check = MarginsCheck()
            issues = check.run(ctx)

            issue_types = [i.check for i in issues]
            assert any("page_height" in t for t in issue_types)
        finally:
            path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
