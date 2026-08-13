"""
Unit Tests for APA Page Handler.

Tests setup_page_layout, _add_page_number, add_section_page_breaks methods.
"""

import os
import tempfile
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from normadocs.formatters.apa.apa_page import APAPageHandler


class TestSetupPageLayout(unittest.TestCase):
    """Tests for setup_page_layout method."""

    def _create_doc_with_config(self, config: dict | None = None) -> tuple[Document, str]:
        """Create a document with optional config. Returns (doc, temp_path)."""
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Content here.", style="Normal")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        # Reload doc to ensure we're testing actual formatting
        doc = Document(temp_path)
        handler = APAPageHandler(doc, config)
        return doc, handler, temp_path

    def test_margins_one_inch_all_sides(self):
        """All margins should be 1 inch (default config)."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                self.assertAlmostEqual(section.top_margin.inches, 1.0, places=2)
                self.assertAlmostEqual(section.bottom_margin.inches, 1.0, places=2)
                self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
                self.assertAlmostEqual(section.right_margin.inches, 1.0, places=2)
        finally:
            os.unlink(temp_path)

    def test_margins_custom_config(self):
        """Custom margins should be applied when specified in config."""
        config = {
            "margins": {
                "top": 1.5,
                "bottom": 1.5,
                "left": 1.0,
                "right": 1.0,
                "unit": "inches",
            }
        }
        doc, handler, temp_path = self._create_doc_with_config(config)

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                self.assertAlmostEqual(section.top_margin.inches, 1.5, places=2)
                self.assertAlmostEqual(section.bottom_margin.inches, 1.5, places=2)
                self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
                self.assertAlmostEqual(section.right_margin.inches, 1.0, places=2)
        finally:
            os.unlink(temp_path)

    def test_page_size_8_5x11(self):
        """Page size should be 8.5 x 11 inches (Letter)."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                self.assertAlmostEqual(section.page_width.inches, 8.5, places=2)
                self.assertAlmostEqual(section.page_height.inches, 11.0, places=2)
        finally:
            os.unlink(temp_path)

    def test_different_first_page_header_footer_enabled(self):
        """different_first_page_header_footer should be True for running head support."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                # different_first_page_header_footer must be True to support
                # running head on pages 2+ while keeping cover page header empty
                self.assertTrue(section.different_first_page_header_footer)
        finally:
            os.unlink(temp_path)

    def test_footer_cleared(self):
        """Footer should be cleared and empty after setup_page_layout."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                footer_text = "".join(p.text for p in section.footer.paragraphs)
                self.assertEqual(footer_text.strip(), "")
        finally:
            os.unlink(temp_path)


class TestAddPageNumber(unittest.TestCase):
    """Tests for _add_page_number method."""

    def _create_doc_with_header(
        self, config: dict | None = None
    ) -> tuple[Document, APAPageHandler, str]:
        """Create a document for header testing."""
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Content here.", style="Normal")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        doc = Document(temp_path)
        handler = APAPageHandler(doc, config)
        return doc, handler, temp_path

    def test_page_number_in_header(self):
        """Page number should appear in header after _add_page_number."""
        doc, handler, temp_path = self._create_doc_with_header()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                header_text = "".join(p.text for p in section.header.paragraphs)
                # Page number placeholder "1" should be in header
                self.assertIn("1", header_text)
        finally:
            os.unlink(temp_path)

    def test_page_number_is_present_on_cover_header(self):
        """Student APA cover should have page number 1 without running-head text."""
        doc, handler, temp_path = self._create_doc_with_header()

        try:
            handler.setup_page_layout()
            first_header_text = "".join(
                p.text for p in doc.sections[0].first_page_header.paragraphs
            )
            self.assertIn("1", first_header_text)
        finally:
            os.unlink(temp_path)

    def test_student_profile_keeps_page_number_without_running_head(self):
        """Student APA profile must suppress optional running-head text."""
        doc, handler, temp_path = self._create_doc_with_header({"running_head": {"enabled": False}})

        try:
            handler.setup_page_layout()
            handler.setup_running_head("OPTIONAL RUNNING HEAD")

            header_text = " ".join(p.text for p in doc.sections[0].header.paragraphs).strip()
            self.assertNotIn("OPTIONAL RUNNING HEAD", header_text)
            self.assertIn("1", header_text)
        finally:
            os.unlink(temp_path)

    def test_page_number_uses_times_new_roman(self):
        """Page number should use Times New Roman font."""
        doc, handler, temp_path = self._create_doc_with_header()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                header = section.header
                if header.paragraphs:
                    para = header.paragraphs[0]
                    for run in para.runs:
                        if run.text.strip():
                            self.assertEqual(run.font.name, "Times New Roman")
        finally:
            os.unlink(temp_path)

    def test_page_number_uses_12pt(self):
        """Page number should be 12pt font size."""
        doc, handler, temp_path = self._create_doc_with_header()

        try:
            handler.setup_page_layout()

            for section in doc.sections:
                header = section.header
                if header.paragraphs:
                    para = header.paragraphs[0]
                    for run in para.runs:
                        if run.text.strip():
                            self.assertEqual(run.font.size, Pt(12))
        finally:
            os.unlink(temp_path)


class TestMarginToInches(unittest.TestCase):
    """Tests for _margin_to_inches helper method."""

    def _create_handler(self, config: dict | None = None) -> APAPageHandler:
        doc = Document()
        return APAPageHandler(doc, config)

    def test_cm_conversion(self):
        """CM values should be converted to inches."""
        handler = self._create_handler({"margins": {"unit": "cm"}})

        result = handler._margin_to_inches(2.54, "cm")
        self.assertAlmostEqual(result.inches, 1.0, places=2)

    def test_inches_passthrough(self):
        """INCHES values should pass through unchanged."""
        handler = self._create_handler({"margins": {"unit": "inches"}})

        result = handler._margin_to_inches(1.5, "inches")
        self.assertAlmostEqual(result.inches, 1.5, places=2)

    def test_default_unit_inches(self):
        """Default unit should be inches."""
        handler = self._create_handler()
        result = handler._margin_to_inches(1.0, "inches")
        self.assertIsInstance(result, Inches)


class TestAddSectionPageBreaks(unittest.TestCase):
    """Tests for add_section_page_breaks method."""

    def _create_doc_with_headings(
        self, headings: list[str]
    ) -> tuple[Document, APAPageHandler, str]:
        """Create a document with specified headings."""
        doc = Document()
        for heading in headings:
            doc.add_paragraph(heading, style="Heading 1")
            doc.add_paragraph("Content under this heading.", style="Normal")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        doc = Document(temp_path)
        handler = APAPageHandler(doc)
        return doc, handler, temp_path

    def test_page_break_before_conclusiones(self):
        """Page break should be added before 'Conclusiones' heading."""
        headings = ["Introducción", "Métodos", "Resultados", "Conclusiones"]
        doc, handler, temp_path = self._create_doc_with_headings(headings)

        try:
            # First setup page layout to initialize headers
            handler.setup_page_layout()
            handler.add_section_page_breaks()

            # Find the Conclusiones paragraph
            conclusiones_para = None
            for p in doc.paragraphs:
                if p.text.strip() == "Conclusiones":
                    conclusiones_para = p
                    break

            self.assertIsNotNone(conclusiones_para)

            # Check if there's a page break element before it
            prev_elem = conclusiones_para._element.getprevious()
            has_page_break = False
            while prev_elem is not None:
                if any(br.get(qn("w:type")) == "page" for br in prev_elem.iter(qn("w:br"))):
                    has_page_break = True
                    break
                prev_elem = prev_elem.getprevious()

            self.assertTrue(has_page_break, "Conclusiones should have a page break before it")
        finally:
            os.unlink(temp_path)

    def test_page_break_before_referencias(self):
        """Page break should be added before 'Referencias' heading."""
        headings = ["Introducción", "Referencias"]
        doc, handler, temp_path = self._create_doc_with_headings(headings)

        try:
            handler.setup_page_layout()
            handler.add_section_page_breaks()

            # Find the Referencias paragraph
            referencias_para = None
            for p in doc.paragraphs:
                if p.text.strip() == "Referencias":
                    referencias_para = p
                    break

            self.assertIsNotNone(referencias_para)

            # Check if there's a page break element before it
            prev_elem = referencias_para._element.getprevious()
            has_page_break = False
            while prev_elem is not None:
                if any(br.get(qn("w:type")) == "page" for br in prev_elem.iter(qn("w:br"))):
                    has_page_break = True
                    break
                prev_elem = prev_elem.getprevious()

            self.assertTrue(has_page_break, "Referencias should have a page break before it")
        finally:
            os.unlink(temp_path)

    def test_page_break_before_english_references(self):
        """English References should also start on a new page."""
        headings = ["Introduction", "References"]
        doc, handler, temp_path = self._create_doc_with_headings(headings)

        try:
            handler.add_section_page_breaks()
            references = next(p for p in doc.paragraphs if p.text.strip() == "References")
            self.assertTrue(handler._has_page_break_before(references))
        finally:
            os.unlink(temp_path)

    def test_no_page_break_for_other_sections(self):
        """No page break should be added before sections not in new_page_sections."""
        headings = ["Introducción", "Métodos", "Discusión"]
        doc, handler, temp_path = self._create_doc_with_headings(headings)

        try:
            handler.setup_page_layout()
            handler.add_section_page_breaks()

            # Find the Discusión paragraph
            diskusion_para = None
            for p in doc.paragraphs:
                if p.text.strip() == "Discusión":
                    diskusion_para = p
                    break

            self.assertIsNotNone(diskusion_para)

            # Check if there's a page break element before it
            prev_elem = diskusion_para._element.getprevious()
            has_page_break = False
            while prev_elem is not None:
                if any(br.get(qn("w:type")) == "page" for br in prev_elem.iter(qn("w:br"))):
                    has_page_break = True
                    break
                prev_elem = prev_elem.getprevious()

            self.assertFalse(has_page_break, "Discusión should NOT have a page break before it")
        finally:
            os.unlink(temp_path)

    def test_existing_break_does_not_suppress_later_section_break(self):
        """An existing break before Referencias must not suppress Conclusiones."""
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Introduction content.", style="Normal")
        references = doc.add_paragraph("References", style="Heading 1")
        break_paragraph = OxmlElement("w:p")
        break_run = OxmlElement("w:r")
        explicit_break = OxmlElement("w:br")
        explicit_break.set(qn("w:type"), "page")
        break_run.append(explicit_break)
        break_paragraph.append(break_run)
        references._element.addprevious(break_paragraph)
        doc.add_paragraph("Reference entry.", style="Normal")
        conclusions = doc.add_paragraph("Conclusiones", style="Heading 1")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        doc = Document(temp_path)
        handler = APAPageHandler(doc)

        try:
            handler.add_section_page_breaks()

            references = next(p for p in doc.paragraphs if p.text == "References")
            conclusions = next(p for p in doc.paragraphs if p.text == "Conclusiones")
            self.assertTrue(handler._has_page_break_before(references))
            self.assertTrue(handler._has_page_break_before(conclusions))
        finally:
            os.unlink(temp_path)

    def test_page_break_case_insensitive(self):
        """Page break should work regardless of case (conclusiones vs Conclusiones)."""
        headings = ["Introducción", "conclusiones"]
        doc, handler, temp_path = self._create_doc_with_headings(headings)

        try:
            handler.setup_page_layout()
            handler.add_section_page_breaks()

            # Find the conclusiones paragraph (lowercase)
            conclusiones_para = None
            for p in doc.paragraphs:
                if p.text.strip() == "conclusiones":
                    conclusiones_para = p
                    break

            self.assertIsNotNone(conclusiones_para)

            # Check if there's a page break element before it
            prev_elem = conclusiones_para._element.getprevious()
            has_page_break = False
            while prev_elem is not None:
                if any(br.get(qn("w:type")) == "page" for br in prev_elem.iter(qn("w:br"))):
                    has_page_break = True
                    break
                prev_elem = prev_elem.getprevious()

            self.assertTrue(has_page_break, "conclusiones (lowercase) should have a page break")
        finally:
            os.unlink(temp_path)


if __name__ == "__main__":
    unittest.main()
