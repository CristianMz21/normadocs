"""Unit tests for the APA citations handler (in-text citations and references)."""

import unittest

from docx import Document

from normadocs.formatters.apa.apa_citations import APACitationsHandler


class TestFixCitations(unittest.TestCase):
    """Tests for fix_citations (in-text citation normalization)."""

    def _doc_with_paragraph(self, text: str, style: str = "Normal") -> Document:
        doc = Document()
        doc.add_paragraph(text, style=style)
        return doc

    def test_three_authors_become_et_al(self):
        """Parenthetical citations with 3+ authors are truncated to 'et al.'."""
        doc = self._doc_with_paragraph(
            "Los estudios previos coinciden (García, López y Martínez, 2020) en el efecto."
        )
        APACitationsHandler(doc).fix_citations()
        self.assertIn("(García et al., 2020)", doc.paragraphs[0].text)

    def test_two_authors_use_ampersand(self):
        """Two-author citations replace 'y' with '&'."""
        doc = self._doc_with_paragraph("Otro estudio (Pérez y Rojas, 2021) confirma el hallazgo.")
        APACitationsHandler(doc).fix_citations()
        self.assertIn("(Pérez & Rojas, 2021)", doc.paragraphs[0].text)

    def test_multiple_citations_in_one_parenthesis(self):
        """Each ';'-separated citation is fixed independently."""
        doc = self._doc_with_paragraph(
            "Varias fuentes (García y López, 2023; Martínez, Ruiz y Silva, 2024) avalan esto."
        )
        APACitationsHandler(doc).fix_citations()
        text = doc.paragraphs[0].text
        self.assertIn("(García & López, 2023; Martínez et al., 2024)", text)

    def test_existing_et_al_is_preserved(self):
        """Citations already using 'et al.' are left untouched."""
        doc = self._doc_with_paragraph("Ya truncado (Martínez et al., 2024) previamente.")
        APACitationsHandler(doc).fix_citations()
        self.assertIn("(Martínez et al., 2024)", doc.paragraphs[0].text)

    def test_narrative_citation_becomes_et_al(self):
        """Narrative citations with 3+ authors are truncated."""
        doc = self._doc_with_paragraph(
            "García, López y Martínez (2020) encontraron resultados similares."
        )
        APACitationsHandler(doc).fix_citations()
        self.assertIn("García et al. (2020)", doc.paragraphs[0].text)

    def test_non_citation_parentheses_untouched(self):
        """Statistical parentheses are not mangled."""
        doc = self._doc_with_paragraph("La mejora fue significativa (p < 0.05; n = 30).")
        APACitationsHandler(doc).fix_citations()
        self.assertIn("(p < 0.05; n = 30)", doc.paragraphs[0].text)

    def test_references_section_is_skipped(self):
        """Reference entries keep their full author lists."""
        doc = Document()
        doc.add_paragraph("Cuerpo del texto (García, López y Silva, 2020).", style="Normal")
        doc.add_paragraph("Referencias", style="Heading 1")
        doc.add_paragraph(
            "García, A., López, B. y Silva, C. (2020). Estudio completo. Journal, 1(1), 1-2.",
            style="Normal",
        )
        APACitationsHandler(doc).fix_citations()
        self.assertIn("García, A., López, B. y Silva, C. (2020)", doc.paragraphs[2].text)


class TestFormatReferences(unittest.TestCase):
    """Tests for format_references (reference-list formatting)."""

    def _references_doc(self, entries: list[str]) -> Document:
        doc = Document()
        doc.add_paragraph("Referencias", style="Heading 1")
        for entry in entries:
            doc.add_paragraph(entry, style="Normal")
        return doc

    def test_spanish_conjunction_becomes_ampersand(self):
        """'y' before the last author becomes ', &'."""
        doc = self._references_doc(["García, A. y López, B. (2023). Machine learning."])
        APACitationsHandler(doc).format_references()
        self.assertIn("García, A., & López, B. (2023)", doc.paragraphs[1].text)

    def test_references_are_sorted_alphabetically(self):
        """Entries are reordered alphabetically below the heading."""
        doc = self._references_doc(
            [
                "Martínez, C. (2024). Sistemas adaptativos.",
                "García, A. (2023). Machine learning.",
            ]
        )
        APACitationsHandler(doc).format_references()
        texts = [p.text for p in doc.paragraphs[1:]]
        self.assertEqual(texts[0], "García, A. (2023). Machine learning.")
        self.assertEqual(texts[1], "Martínez, C. (2024). Sistemas adaptativos.")

    def test_retrieved_from_is_removed(self):
        """APA 6 'Recuperado de' prefixes are dropped."""
        doc = self._references_doc(["Sena, A. (2020). Guía. Recuperado de https://example.com/doc"])
        APACitationsHandler(doc).format_references()
        self.assertNotIn("Recuperado de", doc.paragraphs[1].text)
        self.assertIn("https://example.com/doc", doc.paragraphs[1].text)

    def test_journal_and_volume_are_italicized(self):
        """Plain-text journal entries get italic 'Journal, Volume' runs."""
        doc = self._references_doc(
            ["García, A. (2023). Machine learning en educación. Revista Educación, 45(2), 112-130."]
        )
        APACitationsHandler(doc).format_references()
        para = doc.paragraphs[1]
        italic_text = "".join(r.text for r in para.runs if r.italic)
        self.assertIn("Revista Educación, 45", italic_text)
        self.assertNotIn("Machine learning", italic_text)

    def test_already_italic_entries_are_not_modified(self):
        """Entries with markdown italics keep their original runs."""
        doc = self._references_doc(["García, A. (2023). Título."])
        run = doc.paragraphs[1].runs[0]
        run.text = "García, A. (2023). Título."
        run.italic = True
        APACitationsHandler(doc).format_references()
        self.assertTrue(doc.paragraphs[1].runs[0].italic)

    def test_no_references_section_is_a_noop(self):
        """Documents without a references heading are left alone."""
        doc = Document()
        doc.add_paragraph("Solo texto del cuerpo.")
        APACitationsHandler(doc).format_references()
        self.assertEqual(len(doc.paragraphs), 1)


if __name__ == "__main__":
    unittest.main()
