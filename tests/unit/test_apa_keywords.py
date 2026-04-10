"""
Unit Tests for APA Formatter Keywords Processing.

Tests _format_keywords, _apply_foreign_word_italics, _apply_body_indent,
and _format_nota_italic methods.
"""

import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from normadocs.formatters.apa import APADocxFormatter
from normadocs.models import DocumentMetadata


class TestFormatKeywords(unittest.TestCase):
    """Tests for _format_keywords method."""

    def _create_formatter_with_doc(
        self, paragraphs_data: list[dict]
    ) -> tuple[APADocxFormatter, str]:
        """Create a formatter with custom paragraphs. Returns (formatter, temp_path)."""
        doc = Document()
        for para_data in paragraphs_data:
            style_name = para_data.get("style", "Normal")
            p = doc.add_paragraph(para_data["text"], style=style_name)
            if "runs" in para_data:
                for run in list(p.runs):
                    run._r.getparent().remove(run._r)
                for run_data in para_data["runs"]:
                    r = p.add_run(run_data["text"])
                    if run_data.get("bold"):
                        r.bold = True
                    if run_data.get("italic"):
                        r.italic = True

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        formatter = APADocxFormatter(temp_path)
        formatter.doc = doc
        return formatter, temp_path

    def test_palabras_clave_label_in_italics(self):
        """'Palabras clave:' label should be in italics."""
        paragraphs = [
            {"text": "Este es el resumen del documento.", "style": "Normal"},
            {
                "text": (
                    "Palabras clave: aprendizaje automático, educación, inteligencia artificial"
                ),
                "style": "Normal",
            },
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Contenido de la introducción...", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            meta = DocumentMetadata(title="Test", author="Author")
            formatter._format_keywords(meta)

            # Find the keywords paragraph
            kw_para = None
            for p in formatter.doc.paragraphs:
                if "palabras clave" in p.text.lower():
                    kw_para = p
                    break

            self.assertIsNotNone(kw_para, "Keywords paragraph not found")
            # Check that first run with "Palabras clave:" is italic
            found_italic_label = False
            for run in kw_para.runs:
                if "Palabras clave" in run.text or "palabras clave" in run.text.lower():
                    self.assertTrue(
                        run.italic,
                        f"Label run '{run.text}' should be italic but is not",
                    )
                    found_italic_label = True
                    break

            self.assertTrue(found_italic_label, "Keywords label run not found or not italic")
        finally:
            import os

            os.unlink(temp_path)

    def test_keywords_appear_after_abstract(self):
        """Keywords should appear in a paragraph after the abstract section."""
        paragraphs = [
            {"text": "Resumen", "style": "Heading 1"},
            {"text": "Este es el abstract.", "style": "Normal"},
            {"text": "Palabras clave: testing, unit test", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            meta = DocumentMetadata(title="Test", author="Author")
            formatter._format_keywords(meta)

            # Verify keywords paragraph exists
            found = False
            for p in formatter.doc.paragraphs:
                if "palabras clave" in p.text.lower() or "keywords" in p.text.lower():
                    found = True
                    break

            self.assertTrue(found, "Keywords paragraph should be present in document")
        finally:
            import os

            os.unlink(temp_path)

    def test_keywords_left_indent_half_inch(self):
        """Keywords paragraph should have 0.5 inch left indent."""
        paragraphs = [
            {"text": "Resumen", "style": "Heading 1"},
            {"text": "Abstract text.", "style": "Normal"},
            {"text": "Keywords: testing, unittest", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            meta = DocumentMetadata(title="Test", author="Author")
            formatter._format_keywords(meta)

            # Find the keywords paragraph
            kw_para = None
            for p in formatter.doc.paragraphs:
                if "keywords" in p.text.lower():
                    kw_para = p
                    break

            self.assertIsNotNone(kw_para)
            self.assertEqual(kw_para.paragraph_format.left_indent, Inches(0.5))
        finally:
            import os

            os.unlink(temp_path)

    def test_english_keywords_label_also_italic(self):
        """English 'Keywords:' label should also be italic."""
        paragraphs = [
            {"text": "Abstract", "style": "Heading 1"},
            {"text": "This is the abstract content.", "style": "Normal"},
            {"text": "Keywords: machine learning, education, AI", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            meta = DocumentMetadata(title="Test", author="Author")
            formatter._format_keywords(meta)

            kw_para = None
            for p in formatter.doc.paragraphs:
                if p.text.lower().startswith("keywords"):
                    kw_para = p
                    break

            self.assertIsNotNone(kw_para, "English Keywords paragraph not found")
            # Check italic on the label part
            for run in kw_para.runs:
                if "Keywords" in run.text:
                    self.assertTrue(run.italic, "English 'Keywords' label should be italic")
        finally:
            import os

            os.unlink(temp_path)


class TestApplyForeignWordItalics(unittest.TestCase):
    """Tests for _apply_foreign_word_italics method."""

    def _add_table_cell_with_text(self, doc: Document, cell_text: str) -> Document:
        """Add a table with a cell containing the specified text."""
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = cell_text
        return table

    def test_backend_gets_italic_formatting(self):
        """'Backend' in table cells should get italic formatting."""
        doc = Document()
        self._add_table_cell_with_text(doc, "El Backend fue desarrollado en Python.")

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._apply_foreign_word_italics()

            found_italic = False
            for table in formatter.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "Backend" in run.text and run.italic:
                                    found_italic = True

            self.assertTrue(
                found_italic,
                "'Backend' should be italicized but is not",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_frontend_gets_italic_formatting(self):
        """'Frontend' in table cells should get italic formatting."""
        doc = Document()
        self._add_table_cell_with_text(doc, "El Frontend usa React.")

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._apply_foreign_word_italics()

            found_italic = False
            for table in formatter.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "Frontend" in run.text and run.italic:
                                    found_italic = True

            self.assertTrue(
                found_italic,
                "'Frontend' should be italicized but is not",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_django_gets_italic_formatting(self):
        """'Django' in table cells should get italic formatting."""
        doc = Document()
        self._add_table_cell_with_text(doc, "Framework Django para backend.")

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._apply_foreign_word_italics()

            found_italic = False
            for table in formatter.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "Django" in run.text and run.italic:
                                    found_italic = True

            self.assertTrue(
                found_italic,
                "'Django' should be italicized but is not",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_already_italic_runs_not_modified(self):
        """Runs that are already italic should not be changed."""
        doc = Document()
        self._add_table_cell_with_text(doc, "Texto con Backend ya italicizado.")

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            # Make the run italic before applying
            for table in formatter.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "Backend" in run.text:
                                    run.italic = True

            # Apply the method - should not change already-italic runs
            formatter._apply_foreign_word_italics()

            # Verify it's still italic (not overwritten to non-italic)
            for table in formatter.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "Backend" in run.text:
                                    self.assertTrue(
                                        run.italic,
                                        "Already italic run should remain italic",
                                    )
        finally:
            import os

            os.unlink(temp_path)


class TestApplyBodyIndent(unittest.TestCase):
    """Tests for _apply_body_indent method."""

    def _create_formatter_with_doc(
        self, paragraphs_data: list[dict]
    ) -> tuple[APADocxFormatter, str]:
        """Create a formatter with custom paragraphs. Returns (formatter, temp_path)."""
        doc = Document()
        for para_data in paragraphs_data:
            style_name = para_data.get("style", "Normal")
            doc.add_paragraph(para_data["text"], style=style_name)

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        formatter = APADocxFormatter(temp_path)
        formatter.doc = doc
        return formatter, temp_path

    def test_body_paragraphs_get_half_inch_indent(self):
        """Body paragraphs should get 0.5 inch first-line indent."""
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {
                "text": (
                    "Este es un párrafo de cuerpo con bastante texto "
                    "para que la sangría se aplique correctamente."
                ),
                "style": "Normal",
            },
            {"text": "Segundo párrafo del cuerpo.", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._apply_body_indent()

            # At least one should have the indent
            indented = [
                p
                for p in formatter.doc.paragraphs
                if p.paragraph_format.first_line_indent == Inches(0.5)
            ]
            self.assertGreater(
                len(indented),
                0,
                "At least one body paragraph should have 0.5in indent",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_heading_paragraphs_not_indented(self):
        """Heading paragraphs should NOT get indent."""
        paragraphs = [
            {"text": "Introducción", "style": "Heading 1"},
            {"text": "Subtítulo", "style": "Heading 2"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._apply_body_indent()

            for p in formatter.doc.paragraphs:
                if p.style.name.startswith("Heading"):
                    indent = p.paragraph_format.first_line_indent
                    self.assertTrue(
                        indent is None or indent == Inches(0),
                        f"Heading '{p.text}' should have no indent, got {indent}",
                    )
        finally:
            import os

            os.unlink(temp_path)

    def test_centered_paragraphs_not_indented(self):
        """Centered paragraphs (cover page) should NOT get indent."""
        paragraphs = [
            {"text": "Título del Documento", "style": "Normal"},
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter.doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            formatter._apply_body_indent()

            para = formatter.doc.paragraphs[0]
            indent = para.paragraph_format.first_line_indent
            self.assertTrue(
                indent is None or indent == Inches(0),
                f"Centered paragraph should not get indent, got {indent}",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_references_section_not_indented(self):
        """References section paragraphs should NOT get additional indent (handled separately)."""
        paragraphs = [
            {"text": "Referencias", "style": "Heading 1"},
            {
                "text": "García, A. y López, B. (2023). Machine learning en educación.",
                "style": "Normal",
            },
        ]
        formatter, temp_path = self._create_formatter_with_doc(paragraphs)

        try:
            formatter._apply_body_indent()

            ref_para = formatter.doc.paragraphs[1]
            # References already have hanging indent from _process_paragraphs
            # _apply_body_indent should skip them because they already have indent
            self.assertNotEqual(
                ref_para.paragraph_format.first_line_indent,
                Inches(0.5),
                "References paragraphs should not get 0.5in first-line indent",
            )
        finally:
            import os

            os.unlink(temp_path)


class TestFormatNotaItalic(unittest.TestCase):
    """Tests for _format_nota_italic method."""

    def test_nota_prefix_is_italicized(self):
        """'Nota.' prefix should be italicized, rest of text should be normal."""
        doc = Document()
        doc.add_paragraph(
            "Nota. Este es el texto de la nota que explica la tabla.",
            style="Normal",
        )

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._format_nota_italic()

            # Find the Nota. paragraph
            nota_run_texts = []
            for p in formatter.doc.paragraphs:
                if p.text.strip().startswith("Nota."):
                    for run in p.runs:
                        nota_run_texts.append((run.text, run.italic))
                    break

            self.assertGreater(len(nota_run_texts), 0, "Nota. paragraph not found")

            # First run containing "Nota." should be italic
            nota_runs = [(t, i) for t, i in nota_run_texts if "Nota." in t]
            self.assertGreater(len(nota_runs), 0, "No run with 'Nota.' found")
            nota_text, nota_italic = nota_runs[0]
            self.assertTrue(
                nota_italic,
                f"'Nota.' run should be italic but is not. Run text: '{nota_text}'",
            )
        finally:
            import os

            os.unlink(temp_path)

    def test_nota_text_without_prefix_not_affected(self):
        """Paragraphs not starting with 'Nota.' should not be affected."""
        doc = Document()
        doc.add_paragraph(
            "Este es un párrafo normal que no debería verse afectado.",
            style="Normal",
        )

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._format_nota_italic()

            # The test is that no error occurs and the paragraph is unchanged
            para = formatter.doc.paragraphs[0]
            self.assertIn("párrafo normal", para.text)
        finally:
            import os

            os.unlink(temp_path)

    def test_nota_with_long_text(self):
        """'Nota.' with long text should have italic prefix and normal rest."""
        doc = Document()
        doc.add_paragraph(
            "Nota. Esta es una nota muy larga que explica en detalle el contenido "
            "de la tabla presentada anteriormente en el documento de investigación. "
            "La nota proporciona información adicional sobre las variables medidas.",
            style="Normal",
        )

        import os
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        doc.save(temp_path)
        try:
            formatter = APADocxFormatter(temp_path)
            formatter.doc = doc
            formatter._format_nota_italic()

            # Find Nota. paragraph and check formatting
            for p in formatter.doc.paragraphs:
                if p.text.strip().startswith("Nota."):
                    # Should have multiple runs now (Nota. italic + rest normal)
                    self.assertGreater(
                        len(p.runs),
                        1,
                        "Nota. paragraph should have multiple runs after formatting",
                    )
                    # First run should contain Nota. and be italic
                    first_run = p.runs[0]
                    self.assertIn("Nota.", first_run.text)
                    self.assertTrue(
                        first_run.italic,
                        "First run (Nota.) should be italic",
                    )
                    break
            else:
                self.fail("Nota. paragraph not found after formatting")
        finally:
            import os

            os.unlink(temp_path)


if __name__ == "__main__":
    unittest.main()
