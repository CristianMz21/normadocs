"""
Unit Tests for APA Styles Handler.

Tests create_styles, _neutralize_table_style, _apply_font_style methods.
"""

import unittest
import tempfile
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

from normadocs.formatters.apa.apa_styles import APAStylesHandler


class TestCreateStyles(unittest.TestCase):
    """Tests for create_styles method."""

    def _create_doc_with_config(self, config: dict | None = None) -> tuple[Document, APAStylesHandler, str]:
        """Create a document with optional config. Returns (doc, handler, temp_path)."""
        doc = Document()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Content here.", style="Normal")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        doc = Document(temp_path)
        handler = APAStylesHandler(doc, config)
        return doc, handler, temp_path

    def test_normal_style_times_new_roman_12pt(self):
        """Normal style should use Times New Roman 12pt."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            normal = doc.styles["Normal"]
            self.assertEqual(normal.font.name, "Times New Roman")
            self.assertEqual(normal.font.size, Pt(12))
        finally:
            os.unlink(temp_path)

    def test_normal_style_double_spacing(self):
        """Normal style should use double line spacing."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            normal = doc.styles["Normal"]
            self.assertEqual(normal.paragraph_format.line_spacing_rule, WD_LINE_SPACING.DOUBLE)
        finally:
            os.unlink(temp_path)

    def test_normal_style_left_aligned(self):
        """Normal style should be left aligned."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            normal = doc.styles["Normal"]
            self.assertEqual(normal.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        finally:
            os.unlink(temp_path)

    def test_heading_1_centered_bold(self):
        """Heading 1 should be centered and bold."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            h1 = doc.styles["Heading 1"]
            self.assertEqual(h1.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(h1.font.bold)
        finally:
            os.unlink(temp_path)

    def test_heading_2_left_bold(self):
        """Heading 2 should be left-aligned and bold."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            h2 = doc.styles["Heading 2"]
            self.assertEqual(h2.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(h2.font.bold)
        finally:
            os.unlink(temp_path)

    def test_heading_3_left_bold_italic(self):
        """Heading 3 should be left-aligned, bold, and italic."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            h3 = doc.styles["Heading 3"]
            self.assertEqual(h3.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(h3.font.bold)
            self.assertTrue(h3.font.italic)
        finally:
            os.unlink(temp_path)

    def test_heading_4_indented(self):
        """Heading 4 should be indented 0.5in."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            h4 = doc.styles["Heading 4"]
            self.assertIsNotNone(h4.paragraph_format.first_line_indent)
            self.assertAlmostEqual(h4.paragraph_format.first_line_indent.inches, 0.5, places=2)
        finally:
            os.unlink(temp_path)

    def test_heading_5_indented_italic(self):
        """Heading 5 should be indented 0.5in and italic."""
        doc, handler, temp_path = self._create_doc_with_config()

        try:
            handler.create_styles()

            h5 = doc.styles["Heading 5"]
            self.assertIsNotNone(h5.paragraph_format.first_line_indent)
            self.assertAlmostEqual(h5.paragraph_format.first_line_indent.inches, 0.5, places=2)
            self.assertTrue(h5.font.italic)
        finally:
            os.unlink(temp_path)

    def test_custom_font_config(self):
        """Custom font should be applied when specified in config."""
        config = {
            "fonts": {
                "body": {
                    "name": "Arial",
                    "size": 11,
                }
            }
        }
        doc, handler, temp_path = self._create_doc_with_config(config)

        try:
            handler.create_styles()

            normal = doc.styles["Normal"]
            self.assertEqual(normal.font.name, "Arial")
            self.assertEqual(normal.font.size, Pt(11))
        finally:
            os.unlink(temp_path)

    def test_custom_spacing_config(self):
        """Custom spacing should be applied when specified in config."""
        config = {
            "spacing": {"line": "1.5"}
        }
        doc, handler, temp_path = self._create_doc_with_config(config)

        try:
            handler.create_styles()

            normal = doc.styles["Normal"]
            self.assertEqual(normal.paragraph_format.line_spacing_rule, WD_LINE_SPACING.ONE_POINT_FIVE)
        finally:
            os.unlink(temp_path)


class TestNeutralizeTableStyle(unittest.TestCase):
    """Tests for _neutralize_table_style method."""

    def _create_doc_with_table(self) -> tuple[Document, APAStylesHandler, str]:
        """Create a document with a table for table style testing."""
        doc = Document()
        # Add a table - this should trigger creation of Table style
        table = doc.add_table(rows=2, cols=2)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        doc = Document(temp_path)
        handler = APAStylesHandler(doc)
        return doc, handler, temp_path

    def _get_table_style_element(self, doc: Document):
        """Helper to find the Table style element in document."""
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for style_el in doc.styles.element.findall(f"{{{ns}}}style"):
            style_id = style_el.get(f"{{{ns}}}styleId", "")
            if style_id == "Table":
                return style_el
        return None

    def test_table_style_no_borders(self):
        """Table style should have no borders after _neutralize_table_style."""
        doc, handler, temp_path = self._create_doc_with_table()

        try:
            handler.create_styles()

            table_style_el = self._get_table_style_element(doc)
            if table_style_el is None:
                self.skipTest("Table style not present in document")

            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            tblPr = table_style_el.find(f"{{{ns}}}tblPr")
            self.assertIsNotNone(tblPr, "Table should have tblPr")
            tblBorders = tblPr.find(f"{{{ns}}}tblBorders")
            self.assertIsNotNone(tblBorders, "Table should have tblBorders")

            # Check that all borders are set to "none"
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = tblBorders.find(f"{{{ns}}}{edge}")
                self.assertIsNotNone(border_el, f"Border element for {edge} should exist")
                self.assertEqual(border_el.get(f"{{{ns}}}val"), "none")
        finally:
            os.unlink(temp_path)

    def test_table_style_left_alignment(self):
        """Table style should be left-aligned after _neutralize_table_style."""
        doc, handler, temp_path = self._create_doc_with_table()

        try:
            handler.create_styles()

            table_style_el = self._get_table_style_element(doc)
            if table_style_el is None:
                self.skipTest("Table style not present in document")

            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            pPr = table_style_el.find(f"{{{ns}}}pPr")
            self.assertIsNotNone(pPr, "Table style should have pPr")
            jc = pPr.find(f"{{{ns}}}jc")
            self.assertIsNotNone(jc, "Table style should have jc (justification)")
            self.assertEqual(jc.get(f"{{{ns}}}val"), "left")
        finally:
            os.unlink(temp_path)

    def test_table_style_single_spacing(self):
        """Table style should have single spacing after _neutralize_table_style."""
        doc, handler, temp_path = self._create_doc_with_table()

        try:
            handler.create_styles()

            table_style_el = self._get_table_style_element(doc)
            if table_style_el is None:
                self.skipTest("Table style not present in document")

            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            pPr = table_style_el.find(f"{{{ns}}}pPr")
            self.assertIsNotNone(pPr, "Table style should have pPr")
            spacing = pPr.find(f"{{{ns}}}spacing")
            self.assertIsNotNone(spacing, "Table style should have spacing")
            # Single spacing is 240 twips (12pt at 20 twips/point)
            self.assertEqual(spacing.get(f"{{{ns}}}line"), "240")
        finally:
            os.unlink(temp_path)


class TestApplyFontStyle(unittest.TestCase):
    """Tests for _apply_font_style helper method."""

    def _create_handler(self) -> APAStylesHandler:
        doc = Document()
        return APAStylesHandler(doc)

    def test_applies_font_name(self):
        """_apply_font_style should set font name."""
        handler = self._create_handler()
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]

        handler._apply_font_style(run, font_name="Arial")

        self.assertEqual(run.font.name, "Arial")

    def test_applies_size(self):
        """_apply_font_style should set font size."""
        handler = self._create_handler()
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]

        handler._apply_font_style(run, size=14)

        self.assertEqual(run.font.size, Pt(14))

    def test_applies_bold(self):
        """_apply_font_style should set bold."""
        handler = self._create_handler()
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]

        handler._apply_font_style(run, bold=True)

        self.assertTrue(run.font.bold)

    def test_applies_italic(self):
        """_apply_font_style should set italic."""
        handler = self._create_handler()
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]

        handler._apply_font_style(run, italic=True)

        self.assertTrue(run.font.italic)

    def test_applies_multiple_attributes(self):
        """_apply_font_style should apply multiple attributes at once."""
        handler = self._create_handler()
        doc = Document()
        para = doc.add_paragraph("Test")
        run = para.runs[0]

        handler._apply_font_style(run, font_name="Times New Roman", size=12, bold=True, italic=True)

        self.assertEqual(run.font.name, "Times New Roman")
        self.assertEqual(run.font.size, Pt(12))
        self.assertTrue(run.font.bold)
        self.assertTrue(run.font.italic)


class TestGetFontHelpers(unittest.TestCase):
    """Tests for _get_font_name and _get_font_size helper methods."""

    def _create_handler_with_config(self, config: dict | None = None) -> APAStylesHandler:
        doc = Document()
        return APAStylesHandler(doc, config)

    def test_get_font_name_default(self):
        """_get_font_name should return Times New Roman by default."""
        handler = self._create_handler_with_config()
        self.assertEqual(handler._get_font_name(), "Times New Roman")

    def test_get_font_name_custom(self):
        """_get_font_name should return custom font when configured."""
        config = {"fonts": {"body": {"name": "Arial"}}}
        handler = self._create_handler_with_config(config)
        self.assertEqual(handler._get_font_name(), "Arial")

    def test_get_font_size_default(self):
        """_get_font_size should return 12 by default."""
        handler = self._create_handler_with_config()
        self.assertEqual(handler._get_font_size(), 12)

    def test_get_font_size_custom(self):
        """_get_font_size should return custom size when configured."""
        config = {"fonts": {"body": {"size": 14}}}
        handler = self._create_handler_with_config(config)
        self.assertEqual(handler._get_font_size(), 14)

    def test_get_spacing_line_default(self):
        """_get_spacing_line should return 'double' by default."""
        handler = self._create_handler_with_config()
        self.assertEqual(handler._get_spacing_line(), "double")

    def test_get_spacing_line_custom(self):
        """_get_spacing_line should return custom spacing when configured."""
        config = {"spacing": {"line": "single"}}
        handler = self._create_handler_with_config(config)
        self.assertEqual(handler._get_spacing_line(), "single")


if __name__ == "__main__":
    unittest.main()
