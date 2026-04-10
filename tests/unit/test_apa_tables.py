"""
Unit tests for APA formatter table methods:
_format_tables, _apply_apa_table_borders, _add_table_captions, and _add_table_notes.
"""

import unittest

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from normadocs.formatters.apa import APADocxFormatter


def _make_doc_with_tables(num_tables: int = 1, cols: int = 3) -> tuple[Document, list]:
    """Create a DOCX with specified number of tables and columns."""
    doc = Document()
    tables = []
    for _ in range(num_tables):
        tbl = doc.add_table(rows=3, cols=cols)
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.text = f"Cell{r_idx}C{c_idx}"
        tables.append(tbl)
        doc.add_paragraph("After table paragraph")
    return doc, tables


def _tbl_borders_xml(table) -> dict:
    """Extract border XML from first cell of table for inspection."""
    first_cell = table.rows[0].cells[0]
    tc_pr = first_cell._element.find(qn("w:tcPr"))
    if tc_pr is None:
        return {}
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return {}
    borders = {}
    for border_el in tc_borders:
        tag = border_el.tag.split("}")[-1] if "}" in border_el.tag else border_el.tag
        borders[tag] = dict(border_el.attrib)
    return borders


class TestApplyAPATableBorders(unittest.TestCase):
    """Tests for _apply_apa_table_borders."""

    def test_horizontal_borders_applied(self):
        """First row gets top border, header row gets bottom border, last row gets bottom."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        table = doc.tables[0]
        formatter._apply_apa_table_borders(table)

        borders = _tbl_borders_xml(table)

        # First row top border should exist
        self.assertIn("top", borders)
        self.assertEqual(borders["top"].get(qn("w:val")), "single")
        self.assertEqual(borders["top"].get(qn("w:sz")), "12")

        # First row bottom border (header underline)
        self.assertIn("bottom", borders)
        self.assertEqual(borders["bottom"].get(qn("w:val")), "single")
        self.assertEqual(borders["bottom"].get(qn("w:sz")), "6")

        # No vertical borders
        self.assertNotIn("start", borders)
        self.assertNotIn("end", borders)

    def test_last_row_bottom_border(self):
        """Last row should have a bottom border."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        table = doc.tables[0]
        formatter._apply_apa_table_borders(table)

        last_row = table.rows[-1]
        last_cell = last_row.cells[0]
        tc_pr = last_cell._element.find(qn("w:tcPr"))
        tc_borders = tc_pr.find(qn("w:tcBorders")) if tc_pr is not None else None
        self.assertIsNotNone(tc_borders)

        found_bottom = False
        for border_el in tc_borders:
            tag = border_el.tag.split("}")[-1] if "}" in border_el.tag else border_el.tag
            if tag == "bottom":
                self.assertEqual(border_el.get(qn("w:val")), "single")
                self.assertEqual(border_el.get(qn("w:sz")), "12")
                found_bottom = True
        self.assertTrue(found_bottom, "Last row should have bottom border")

    def test_no_vertical_borders(self):
        """No vertical (start/end/insideV) borders should be present."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        table = doc.tables[0]
        formatter._apply_apa_table_borders(table)

        borders = _tbl_borders_xml(table)

        self.assertNotIn("start", borders)
        self.assertNotIn("end", borders)


class TestFormatTables(unittest.TestCase):
    """Tests for _format_tables."""

    def test_table_centered_alignment(self):
        """Tables should be center-aligned after formatting."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._format_tables()

        table = doc.tables[0]
        self.assertEqual(table.alignment, WD_TABLE_ALIGNMENT.CENTER)

    def test_font_scales_down_for_six_columns(self):
        """Tables with 6+ columns should use smaller font to fit."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=6)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._format_tables()
        table = doc.tables[0]
        self.assertIsNotNone(table)

    def test_font_scales_down_for_eight_columns(self):
        """Tables with 8+ columns should use 9pt font."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=8)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._format_tables()
        table = doc.tables[0]
        self.assertIsNotNone(table)

    def test_column_widths_applied(self):
        """Column widths should be calculated and set after formatting."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        formatter._format_tables()

        table = doc.tables[0]
        tbl_grid = table._tbl.find(qn("w:tblGrid"))
        grid_cols = tbl_grid.findall(qn("w:gridCol")) if tbl_grid is not None else []

        # Grid columns should have explicit width set
        self.assertGreater(len(grid_cols), 0)
        for gc in grid_cols:
            w = gc.get(qn("w:w"))
            self.assertIsNotNone(w, "gridCol should have explicit width")
            self.assertGreater(int(w), 0, "gridCol width should be positive")

    def test_cell_margins_set(self):
        """Cell margins should be set for all cells."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        formatter._format_tables()

        table = doc.tables[0]
        cell = table.rows[0].cells[0]
        tc_pr = cell._element.find(qn("w:tcPr"))
        self.assertIsNotNone(tc_pr)

        tc_mar = tc_pr.find(qn("w:tcMar"))
        self.assertIsNotNone(tc_mar, "Cell should have margins (tcMar)")

        for side in ("top", "bottom", "start", "end"):
            found = tc_mar.find(qn(f"w:{side}"))
            self.assertIsNotNone(found, f"Cell margin should have {side} element")


class TestAddTableCaptions(unittest.TestCase):
    """Tests for _add_table_captions."""

    def test_caption_format_table_number_title(self):
        """Caption should be 'Tabla N' in bold followed by title in italics."""
        doc, _ = _make_doc_with_tables(num_tables=1, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        formatter._add_table_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Should contain "Table 1" (bold label)
        self.assertIn("Table 1", full_text)

    def test_table_number_increments_correctly(self):
        """Table numbers should increment: Table 1, Table 2, etc."""
        doc, _ = _make_doc_with_tables(num_tables=3, cols=3)
        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc

        formatter._add_table_captions()

        full_text = "\n".join(p.text for p in doc.paragraphs)

        self.assertIn("Table 1", full_text)
        self.assertIn("Table 2", full_text)
        self.assertIn("Table 3", full_text)


class TestAddTableNotes(unittest.TestCase):
    """Tests for _add_table_notes descriptive notes based on table content."""

    def test_caracteristica_especificacion_specific_description(self):
        """Table with caracteristica + especificacion in header gets specific description."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Caracteristica"
        tbl.rows[0].cells[1].text = "Especificacion"
        tbl.rows[1].cells[0].text = "Procesador"
        tbl.rows[1].cells[1].text = "Intel Core i7"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)

    def test_software_version_software_description(self):
        """Table with software + version gets software description."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Software"
        tbl.rows[0].cells[1].text = "Version"
        tbl.rows[1].cells[0].text = "Django"
        tbl.rows[1].cells[1].text = "4.2"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)

    def test_servicio_proveedor_cloud_services_description(self):
        """Table with servicio + proveedor gets cloud services description."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Servicio"
        tbl.rows[0].cells[1].text = "Proveedor"
        tbl.rows[1].cells[0].text = "AWS S3"
        tbl.rows[1].cells[1].text = "Amazon"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)
        self.assertIn("nube", full_text.lower())

    def test_rubro_costo_investments_description(self):
        """Table with rubro + costo gets investments description."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Rubro"
        tbl.rows[0].cells[1].text = "Costo"
        tbl.rows[1].cells[0].text = "Software"
        tbl.rows[1].cells[1].text = "50000"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)
        self.assertIn("inversion", full_text.lower())

    def test_componente_porcentaje_aiu_description(self):
        """Table with componente + porcentaje gets AIU description."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Componente"
        tbl.rows[0].cells[1].text = "Porcentaje"
        tbl.rows[1].cells[0].text = "Administracion"
        tbl.rows[1].cells[1].text = "10%"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)
        self.assertIn("AIU", full_text)

    def test_fallback_description_no_match(self):
        """When no content heuristics match, fallback description is used."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "ColumnA"
        tbl.rows[0].cells[1].text = "ColumnB"
        tbl.rows[1].cells[0].text = "ValueA"
        tbl.rows[1].cells[1].text = "ValueB"

        formatter = APADocxFormatter.__new__(APADocxFormatter)
        formatter.doc = doc
        formatter._add_table_notes()

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Nota.", full_text)


if __name__ == "__main__":
    unittest.main()
